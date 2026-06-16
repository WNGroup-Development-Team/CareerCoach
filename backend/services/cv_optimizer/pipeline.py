from __future__ import annotations

import io
import json
import re
import shutil
import subprocess
import tempfile
import textwrap
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional

from docx.oxml import OxmlElement
from docx.shared import Pt

from .section_catalog import SECTION_ALIASES as SHARED_SECTION_ALIASES
from .section_catalog import additional_field_section_key, canonical_section_key, normalize_section_title

SECTION_ALIASES = {
    "profilo": SHARED_SECTION_ALIASES["profile"],
    "esperienze": SHARED_SECTION_ALIASES["experience"],
    "formazione": SHARED_SECTION_ALIASES["education"],
    "competenze": SHARED_SECTION_ALIASES["hard_skills"],
    "hard_skills": SHARED_SECTION_ALIASES["hard_skills"],
    "soft_skills": SHARED_SECTION_ALIASES["soft_skills"],
    "lingue": SHARED_SECTION_ALIASES["languages"],
    "certificazioni": SHARED_SECTION_ALIASES["certifications"],
    "pubblicazioni": SHARED_SECTION_ALIASES["publications"],
    "progetti": SHARED_SECTION_ALIASES["projects"],
    "contatti": SHARED_SECTION_ALIASES["contacts"],
}


def _cv_rewrite_llm_enabled() -> bool:
    try:
        from main import CV_REWRITE_LLM_ENABLED as enabled
        return bool(enabled)
    except Exception:
        return False

SECTION_HEADINGS = {heading for values in SECTION_ALIASES.values() for heading in values}
SECTION_LEAK_MARKERS = {
    "contatti", "lingue", "hard skills", "soft skills", "formazione",
    "istruzione", "esperienze professionali", "esperienza professionale",
    "esperienze", "progetti", "certificazioni", "pubblicazioni",
}
GENERIC_KEYWORDS = {
    "data", "analyst", "analysis", "business", "project", "manager", "team", "office",
    "azienda", "ruolo", "junior", "senior", "stage", "internship", "lavoro",
}

# Tokens that are generic role nouns and should not be proposed alone as keywords
GENERIC_ROLE_TOKENS = {
    "scientist", "analyst", "engineer", "developer", "manager", "specialist",
    "consultant", "designer", "researcher", "assistant",
}
SYNONYMS = {
    "analisi dati": {"data analysis", "data analytics", "analisi dati"},
    "excel": {"excel", "microsoft excel", "fogli di calcolo"},
    "power bi": {"power bi", "powerbi", "business intelligence"},
    "sql": {"sql", "database", "query"},
    "python": {"python", "pandas", "numpy"},
    "ai": {"ai", "artificial intelligence", "intelligenza artificiale"},
    "machine learning": {"machine learning", "ml", "apprendimento automatico"},
    "nlp": {"nlp", "natural language processing", "elaborazione del linguaggio naturale"},
    "llm": {"llm", "large language models", "large language model", "modelli llm"},
    "clustering testuale": {"clustering testuale", "clustering", "text clustering"},
    "accuratezza": {"accuratezza", "accuracy"},
    "tempi di risposta": {"tempi di risposta", "response time", "response times"},
    "preparazione dati": {"preparazione dati", "data preparation", "preprocessing", "preparazione dei dati"},
    "confronto modelli": {"confronto modelli", "model comparison", "confronto delle prestazioni"},
    "comunicazione": {"comunicazione", "communication"},
    "problem solving": {"problem solving", "risoluzione problemi"},
}

CONTROLLED_PM_SKILLS = {
    "gestione budget": ("budget", "costi", "spese"),
    "budget management": ("budget", "costi", "spese"),
    "leadership strategica": ("leadership strategica",),
    "negoziazione": ("negoziazione", "negoziato"),
    "negoziazione avanzata": ("negoziazione avanzata",),
    "microsoft project": ("microsoft project",),
    "monday.com": ("monday.com", "monday"),
    "notion": ("notion",),
    "risk management": ("risk management", "gestione rischi", "analisi dei rischi"),
    "gestione rischi": ("gestione rischi", "analisi dei rischi"),
}

INFORMAL_CV_MARKERS = (
    "l'ho usato", "l ho usato", "l'ho utilizzato", "l ho utilizzato",
    "ho usato", "ho utilizzato", "ho applicato", "visto all'esame",
    "vista all'esame", "usata durante i progetti", "usato durante i progetti",
    "usata in progetto", "usato in progetto",
)


def normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").lower()).strip()


def tokenize(value: str) -> List[str]:
    return re.findall(r"[a-zA-ZÀ-ÖØ-öø-ÿ0-9][a-zA-ZÀ-ÖØ-öø-ÿ0-9+#.-]{1,}", normalize_text(value))


def is_section_heading(line: str) -> bool:
    clean = normalize_text(line).strip(":")
    stripped = (line or "").strip().strip(":")
    if any(marker in stripped for marker in {"●", "○", "•", "◦"}):
        return False
    if canonical_section_key(stripped):
        return True
    words = clean.split()
    if (line or "").strip().endswith(":"):
        return (
            1 <= len(words) <= 7
            and len(clean) <= 64
            and not re.search(r"[.!?]", stripped)
        )
    return False


def canonical_section(line: str) -> str:
    shared_key = canonical_section_key(line)
    if shared_key:
        return {
            "profile": "profilo",
            "experience": "esperienze",
            "education": "formazione",
            "hard_skills": "hard_skills",
            "soft_skills": "soft_skills",
            "languages": "lingue",
            "certifications": "certificazioni",
            "publications": "pubblicazioni",
            "projects": "progetti",
            "contacts": "contatti",
            "header": "intestazione",
        }.get(shared_key, shared_key)
    clean = normalize_section_title(line)
    return clean or "contenuto"


@dataclass
class ResumeSection:
    name: str
    heading: str
    lines: List[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n".join(self.lines).strip()


@dataclass
class RewriteInstruction:
    section: str
    original: str
    replacement: str
    reason: str = ""
    category: str = ""
    source_id: str = ""


@dataclass
class ParagraphContext:
    paragraph: Any
    section: str


class ResumeParser:
    def _prepared_lines(self, cv_text: str) -> List[str]:
        prepared = re.sub(r"\r\n?", "\n", cv_text or "")
        prepared = re.sub(r"[ \t]+", " ", prepared)
        heading_pattern = "|".join(
            re.escape(heading)
            for heading in sorted(SECTION_HEADINGS, key=len, reverse=True)
        )
        if heading_pattern:
            prepared = re.sub(
                rf"(?<![\wÀ-ÖØ-öø-ÿ])({heading_pattern})\s*:?(?=\s|$)",
                lambda match: f"\n{match.group(1).strip()}\n",
                prepared,
                flags=re.IGNORECASE,
            )
        return [line.strip() for line in prepared.splitlines() if line.strip()]

    def parse_text(self, cv_text: str) -> List[ResumeSection]:
        sections: List[ResumeSection] = []
        current = ResumeSection(name="intestazione", heading="", lines=[])

        for line in self._prepared_lines(cv_text):
            if is_section_heading(line):
                if current.heading or current.lines:
                    sections.append(current)
                current = ResumeSection(name=canonical_section(line), heading=line, lines=[])
            else:
                current.lines.append(line)

        if current.heading or current.lines:
            sections.append(current)
        return sections

    def section_text_map(self, cv_text: str) -> Dict[str, str]:
        result: Dict[str, str] = {}
        for section in self.parse_text(cv_text):
            result.setdefault(section.name, section.text)
        return result


class JobAnalyzer:
    def __init__(self, normalizer: Callable[[str], str] = normalize_text):
        self.normalizer = normalizer

    def extract_keywords(self, role: str, description: str = "", required_skills: str = "") -> List[str]:
        explicit = [
            item.strip()
            for item in re.split(r"[,;\n]+", required_skills or "")
            if item.strip()
        ]
        source = description.strip() or role.strip()
        # If role looks like a multi-word title (e.g., "Data Scientist"), prefer keeping it whole
        candidates = explicit + self._phrases(source) + tokenize(source)
        role_title = role.strip()
        if role_title and " " in role_title and len(role_title) >= 4:
            candidates = [role_title] + candidates

        keywords: List[str] = []
        seen = set()
        for candidate in candidates:
            clean = self.normalizer(candidate).strip(" .,:;")
            if not clean or clean in seen:
                continue
            # Skip generic keywords and very short tokens (except common acronyms like SQL)
            if clean in GENERIC_KEYWORDS or (len(clean) < 4 and clean not in {"sql"}):
                continue
            # Avoid proposing single-word generic role tokens (e.g. "scientist", "engineer")
            if " " not in clean and clean in GENERIC_ROLE_TOKENS:
                continue
            # Also avoid single-word fragments that are just "data" or similar generic terms
            if " " not in clean and clean in {"data", "analyst"}:
                continue
            seen.add(clean)
            keywords.append(candidate.strip())
            if len(keywords) >= 20:
                break
        return keywords

    def _phrases(self, text: str) -> List[str]:
        normalized = self.normalizer(text)
        phrases = []
        for phrase in sorted(SYNONYMS, key=len, reverse=True):
            if phrase in normalized or any(alias in normalized for alias in SYNONYMS[phrase]):
                phrases.append(phrase)
        words = [word for word in tokenize(text) if word not in GENERIC_KEYWORDS and word not in GENERIC_ROLE_TOKENS]
        for size in (3, 2):
            for index in range(0, max(len(words) - size + 1, 0)):
                phrase = " ".join(words[index:index + size])
                if phrase and phrase not in GENERIC_KEYWORDS:
                    phrases.append(phrase)
        quoted = re.findall(r"['\"]([^'\"]{4,40})['\"]", text or "")
        return phrases + quoted


class MatchingEngine:
    def __init__(self, normalizer: Callable[[str], str] = normalize_text):
        self.normalizer = normalizer

    def split_present_missing(self, cv_text: str, keywords: Iterable[str]) -> Dict[str, List[str]]:
        cv_plain = self.normalizer(cv_text)
        present: List[str] = []
        missing: List[str] = []
        for keyword in keywords:
            variants = SYNONYMS.get(self.normalizer(keyword), {self.normalizer(keyword)})
            if any(variant and variant in cv_plain for variant in variants):
                present.append(keyword)
            else:
                missing.append(keyword)
        return {"present": present, "missing": missing}


class CoachSuggestionEngine:
    def accepted_only(self, suggestions: Any) -> List[Dict[str, Any]]:
        if not isinstance(suggestions, list):
            return []
        accepted = []
        for item in suggestions:
            if isinstance(item, dict) and item.get("type") == "actionableEdit":
                description = str(item.get("description") or item.get("action") or item.get("coach_tip") or "").strip()
                proposed_text = str(item.get("proposed_text") or item.get("replacement") or "").strip()
                original_text = str(item.get("original_text") or item.get("original") or "").strip()
                section = str(item.get("section") or "").strip()
                if proposed_text and original_text and section:
                    accepted.append({
                        **item,
                        "description": description,
                        "section": section,
                        "original_text": original_text,
                        "proposed_text": proposed_text,
                    })
        return accepted[:30]


class ResumeRewriter:
    FORBIDDEN_OUTPUT_MARKERS = (
        "CV ottimizzato - bozza guidata",
        "Ottimizzazioni accettate dall'utente",
        "CV originale da rifinire",
        "Punteggio",
        "ATS simulato",
        "Suggerimenti coach",
    )

    def __init__(self, parser: Optional[ResumeParser] = None):
        self.parser = parser or ResumeParser()

    def build_prompt(
        self,
        cv_text: str,
        target: Dict[str, str],
        accepted_suggestions: List[Dict[str, Any]],
        user_data: Dict[str, Any],
        sections: Optional[List[ResumeSection]] = None,
    ) -> str:
        section_payload = [
            {"name": section.name, "heading": section.heading, "text": section.text}
            for section in (sections or self.parser.parse_text(cv_text))
        ]
        return f"""
Sei un resume editor. Devi restituire SOLO JSON valido.

Obiettivo: crea SOLO istruzioni di riscrittura puntuali per piccoli blocchi del CV.
Non generare mai un CV completo e non restituire testo completo da incollare nel DOCX.

Target:
{target}

Sezioni CV originali:
{section_payload}

Modifiche accettate dall'utente:
{accepted_suggestions}

Dati confermati dall'utente:
{user_data}

Schema JSON:
{{
  "instructions": [
    {{
      "section": "nome sezione canonico",
      "original": "testo esatto o frase del CV da sostituire. Lascia vuoto '' SOLO se stai aggiungendo una entry del tutto nuova.",
      "replacement": "nuovo testo naturale o lista",
      "reason": "breve motivo interno"
    }}
  ]
}}

Regole:
- DEVI ASSOLUTAMENTE includere e integrare nel CV TUTTE le modifiche accettate dall'utente e TUTTI i dati confermati (es. confirmed_skills). Se non lo fai, l'utente perderà i dati inseriti!
- Per AGGIUNGERE nuove competenze/skill (es. confermate dall'utente) usa SEMPRE una istruzione dedicata con `section`: "competenze", `original`: "" e `replacement`: "Nome skill 1, Nome skill 2". NON copiare le vecchie skill nel replacement di questa istruzione, il sistema appenderà quelle nuove automaticamente alla sezione corretta.
- ATTENZIONE CRITICA: Se invece stai modificando o migliorando un'esperienza, formazione o profilo GIÀ PRESENTE nel CV, DEVI OBBLIGATORIAMENTE inserire in `original` il frammento esatto di testo originale da sostituire.
- Non duplicare mai sezioni già presenti. Usa `original: ""` SOLO ED ESCLUSIVAMENTE per aggiungere skill, progetti o esperienze completamente nuove non menzionate nel CV.
- Non aggiungere nuove pagine o sezioni finali artificiali solo per skill, keyword o suggerimenti.
- Riscrivi in modo intelligente e naturale.
- Mantieni separazione tra Hard Skills e Soft Skills se necessario, oppure uniscile logicamente.
- Ogni replacement con original NON vuoto deve sostituire solo il blocco originale della stessa sezione.
- Non inventare esperienze o competenze non confermate.
- Restituisci da 1 a 15 istruzioni concrete.
"""

    def apply_to_text(self, cv_text: str, instructions: List[RewriteInstruction]) -> str:
        updated = cv_text or ""
        for instruction in instructions:
            original = (instruction.original or "").strip()
            replacement = (instruction.replacement or "").strip()
            if not replacement:
                continue
            if not original:
                updated = self._append_replacement_to_text_section(updated, instruction)
                continue
            if original in updated:
                updated = updated.replace(original, replacement, 1)
                continue
            pattern = re.compile(re.escape(original), flags=re.IGNORECASE)
            updated, count = pattern.subn(replacement, updated, count=1)
            if count:
                continue
            original_words = re.findall(r"\S+", original)
            if len(original_words) < 4:
                continue
            flexible_pattern = r"\s+".join(re.escape(word) for word in original_words)
            updated, count = re.subn(
                flexible_pattern,
                lambda _match: replacement,
                updated,
                count=1,
                flags=re.IGNORECASE,
            )
            if not count:
                updated = self._append_replacement_to_text_section(updated, instruction)
        return self.clean_final_text(updated)

    def _append_replacement_to_text_section(self, cv_text: str, instruction: RewriteInstruction) -> str:
        replacement = (instruction.replacement or "").strip()
        if not replacement:
            return cv_text
        if normalize_text(replacement) in normalize_text(cv_text):
            return cv_text

        target = canonical_section(instruction.section or instruction.category or "")
        heading = (instruction.section or target or "COMPETENZE").strip().upper()
        lines = (cv_text or "").splitlines()
        if not lines:
            return f"{heading}\n{replacement}".strip()

        heading_index = None
        for index, line in enumerate(lines):
            if is_section_heading(line) and canonical_section(line) == target:
                heading_index = index
                break

        if heading_index is None:
            return (cv_text.rstrip() + f"\n\n{heading}\n{replacement}").strip()

        insert_at = len(lines)
        for index in range(heading_index + 1, len(lines)):
            if is_section_heading(lines[index]):
                insert_at = index
                break
        return "\n".join([*lines[:insert_at], replacement, *lines[insert_at:]]).strip()

    def fallback_text(self, cv_text: str, accepted_suggestions: List[Dict[str, Any]], user_data: Dict[str, Any]) -> str:
        instructions = self.instructions_from_suggestions(accepted_suggestions or [])
        confirmed = (user_data or {}).get("confirmed_skills", [])
        if isinstance(confirmed, list) and confirmed:
            skill_names = []
            for item in confirmed:
                name = str(item.get("name") or item.get("skill") or item if isinstance(item, dict) else item).strip()
                if name and normalize_text(name) not in {normalize_text(existing) for existing in skill_names}:
                    skill_names.append(name)
            if skill_names:
                instructions.append(RewriteInstruction(
                    section="COMPETENZE",
                    original="",
                    replacement=", ".join(skill_names),
                    reason="Skill confermate dall'utente integrate con fallback locale.",
                    category="skills",
                    source_id="fallback_confirmed_skills",
                ))
        return self.apply_to_text(cv_text, instructions) if instructions else self.clean_final_text(cv_text)

    def instructions_from_result(self, result: Dict[str, Any]) -> List[RewriteInstruction]:
        instructions = []
        for item in result.get("instructions") or []:
            if not isinstance(item, dict):
                continue
            original = str(item.get("original") or "").strip()
            raw_replacement = str(item.get("replacement") or "").strip()
            section = canonical_section(str(item.get("section") or ""))
            if self._is_bad_cv_instruction(section, original, raw_replacement, item):
                print(
                    "Istruzione LLM scartata da quality gate: "
                    f"id={item.get('id') or item.get('title') or 'llm_instruction'}, "
                    f"section={section}, replacement={raw_replacement[:180]}"
                )
                continue
            if not self.is_safe_replacement(section, raw_replacement):
                print(
                    "Suggerimento riscrittura scartato: "
                    f"id={item.get('id') or item.get('title') or 'llm_instruction'}, "
                    f"section={section}, replacement={raw_replacement[:180]}"
                )
                continue
            replacement = self.clean_replacement(raw_replacement)
            if not replacement or original == replacement:
                continue
            instructions.append(RewriteInstruction(
                section=section,
                original=original,
                replacement=replacement,
                reason=str(item.get("reason") or "").strip(),
                category=str(item.get("category") or "").strip(),
                source_id=str(item.get("id") or item.get("title") or "llm_instruction").strip(),
            ))
        return instructions[:40]

    def instructions_from_suggestions(self, suggestions: List[Dict[str, Any]]) -> List[RewriteInstruction]:
        instructions = []
        for item in suggestions:
            original = str(item.get("original_text") or item.get("original") or "").strip()
            raw_proposed = str(item.get("proposed_text") or item.get("replacement") or "").strip()
            section = canonical_section(str(item.get("section") or item.get("category") or ""))
            if not raw_proposed:
                continue
            if self._is_bad_cv_instruction(section, original, raw_proposed, item):
                print(
                    "Suggerimento coach scartato da quality gate: "
                    f"id={item.get('id') or item.get('title') or 'coach_suggestion'}, "
                    f"section={section}, proposed_text={raw_proposed[:180]}"
                )
                continue
            if not self.is_safe_replacement(section, raw_proposed):
                print(
                    "Suggerimento coach scartato: "
                    f"id={item.get('id') or item.get('title') or 'coach_suggestion'}, "
                    f"section={section}, category={item.get('category') or ''}, "
                    f"proposed_text={raw_proposed[:180]}"
                )
                continue
            proposed = self.clean_replacement(raw_proposed)
            instructions.append(RewriteInstruction(
                section=section,
                original=original,
                replacement=proposed,
                reason=str(item.get("reason") or item.get("description") or "").strip(),
                category=str(item.get("category") or "").strip(),
                source_id=str(item.get("id") or item.get("title") or "coach_suggestion").strip(),
            ))
        return instructions[:40]

    def _is_bad_cv_instruction(self, section: str, original: str, proposed: str, item: Any = None) -> bool:
        try:
            from services.cv_optimizer.safe_cv_guard import is_bad_suggestion
            payload = dict(item or {})
            payload.setdefault("section", section)
            payload.setdefault("original_text", original)
            payload.setdefault("proposed_text", proposed)
            return is_bad_suggestion(payload)
        except Exception:
            combined = normalize_text(f"{original} {proposed}")
            section_name = canonical_section(section)
            if section_name == "formazione":
                return True
            if section_name == "profilo" and any(marker in combined for marker in ["captive portal", "identity provider", "tesi di laurea", "corso di laurea"]):
                return True
            if any(marker in combined for marker in ["percorso formativo valorizzato", "percorso formativo coerente"]):
                return True
            return False

    def is_safe_replacement(self, section: str, replacement: str) -> bool:
        raw = replacement or ""
        normalized = normalize_text(raw)
        if not normalized:
            return False
        section_name = canonical_section(section)
        lines = [line.strip() for line in raw.splitlines() if line.strip()]
        exact_heading_lines = {
            normalize_text(marker).strip(":")
            for marker in SECTION_LEAK_MARKERS
        }
        if any(normalize_text(line).strip(":") in exact_heading_lines for line in lines):
            return False
        # No length limits, allowing arbitrary CV sizes and details.
        blocked_profile_lines = {"formazione", "esperienze professionali", "esperienza professionale", "contatti", "lingue"}
        if section_name == "competenze" and any(normalize_text(line).strip(":") in blocked_profile_lines for line in lines):
            return False
        if section_name == "formazione" and any(normalize_text(line).strip(":") in {"contatti", "lingue", "hard skills", "soft skills", "esperienze professionali", "esperienza professionale"} for line in lines):
            return False
        if section_name == "esperienze" and any(normalize_text(line).strip(":") in {"contatti", "lingue", "hard skills", "soft skills", "formazione"} for line in lines):
            return False
        return True

    def clean_final_text(self, text: str) -> str:
        lines = []
        for raw_line in (text or "").splitlines():
            line = self.clean_line(raw_line)
            if not line:
                continue
            if any(marker.lower() in line.lower() for marker in self.FORBIDDEN_OUTPUT_MARKERS):
                continue
            lines.append(line)
        return "\n".join(lines).strip()

    def clean_line(self, line: str) -> str:
        line = re.sub(r"\s+", " ", line or "").strip()
        for heading in SECTION_HEADINGS:
            upper = heading.upper()
            line = re.sub(rf"(?<!^)\b{re.escape(upper)}\b(?!$)", "", line)
        return re.sub(r"\s{2,}", " ", line).strip()

    def clean_replacement(self, value: str) -> str:
        lines = [self.clean_line(line) for line in (value or "").splitlines()]
        return "\n".join(line for line in lines if line).strip()

class DocxPreserver:
    def __init__(self) -> None:
        self.applied_source_ids: List[str] = []
        self.skipped_source_ids: List[str] = []

    def dedupe_contact_line(self, line: str, existing_contact_ids: Optional[set[str]] = None) -> str:
        cleaned = re.sub(r"\s+", " ", line or "").strip()
        if not cleaned or not self._contact_identifiers(cleaned):
            return cleaned

        contact_value_pattern = r"(?:https?://|www\.)?[a-z0-9.-]+\.[a-z]{2,}(?:/[^\s,;|]*)?|[\w.+-]+@[\w.-]+\.\w+|\+?\d[\d\s().-]{7,}\d"
        fragment_pattern = re.compile(contact_value_pattern, flags=re.IGNORECASE)
        label_pattern = r"linkedin|github|portfolio|website|sito web|email|e-mail|telefono|phone|mobile|cellulare"
        labeled_item_pattern = re.compile(
            rf"(?P<item>(?:(?P<label>{label_pattern})\s*:\s*)?(?P<value>{contact_value_pattern}))",
            flags=re.IGNORECASE,
        )

        def collapse_repeated_fragment(fragment: str) -> str:
            token = fragment.strip()
            for size in range(len(token) // 2, 7, -1):
                if len(token) % size != 0:
                    continue
                unit = token[:size]
                repeats = len(token) // size
                if repeats <= 1:
                    continue
                if token == unit * repeats and self._contact_identifiers(unit):
                    return unit
            return token

        rebuilt: List[str] = []
        seen_contacts: set[str] = set()
        cursor = 0
        for match in fragment_pattern.finditer(cleaned):
            between = cleaned[cursor:match.start()]
            fragment = collapse_repeated_fragment(match.group(0))
            identifiers = self._contact_identifiers(fragment)
            is_duplicate = bool(identifiers) and identifiers.issubset(seen_contacts)
            if not is_duplicate or not re.fullmatch(r"[\s,;|/-]*", between or ""):
                rebuilt.append(between)
            if not is_duplicate:
                rebuilt.append(fragment)
                seen_contacts.update(identifiers)
            cursor = match.end()
        rebuilt.append(cleaned[cursor:])

        cleaned = "".join(rebuilt)
        seen_contacts = set(existing_contact_ids or set())
        matches = list(labeled_item_pattern.finditer(cleaned))
        if matches:
            unique_items: List[str] = []
            duplicates_found = False
            for match in matches:
                value = collapse_repeated_fragment(match.group("value"))
                label = (match.group("label") or "").strip()
                item_text = f"{label}: {value}" if label else value
                identifiers = self._contact_identifiers(item_text)
                if identifiers and identifiers.issubset(seen_contacts):
                    duplicates_found = True
                    continue
                if identifiers:
                    seen_contacts.update(identifiers)
                unique_items.append(item_text)
            if duplicates_found and unique_items:
                separator = " | " if "|" in cleaned else (" • " if "•" in cleaned else ("; " if ";" in cleaned else " "))
                cleaned = separator.join(unique_items)
            elif duplicates_found and not unique_items:
                cleaned = ""

        cleaned = re.sub(
            rf"(?:\s*[|,;•/-]?\s*)(?:{label_pattern})\s*:\s*$",
            "",
            cleaned,
            flags=re.IGNORECASE,
        )
        cleaned = re.sub(r"\s{2,}", " ", cleaned).strip(" ,;|/-")
        if cleaned and existing_contact_ids and self._contact_identifiers(cleaned).issubset(existing_contact_ids):
            return ""
        return cleaned

    def apply(self, original_file_bytes: bytes, instructions: List[RewriteInstruction]) -> tuple[bytes, int]:
        from docx import Document

        # The specialized pipeline owns the section-aware replacement helpers.
        # Keep the public base class compatible for callers and tests that use it directly.
        if type(self) is DocxPreserver:
            delegate = ResumeDocxOptimizationPipeline()
            result = delegate.apply(original_file_bytes, instructions)
            self.applied_source_ids = list(delegate.applied_source_ids)
            self.skipped_source_ids = list(delegate.skipped_source_ids)
            return result

        self.applied_source_ids = []
        self.skipped_source_ids = []
        document = Document(io.BytesIO(original_file_bytes))
        applied = 0
        for i, instruction in enumerate(instructions):
            section_name = instruction.section or instruction.category or "unknown"
            original_preview = (instruction.original or "")[:80].replace('\n', '\\n') if instruction.original else "(no original)"
            if not self._valid_instruction(instruction):
                self.skipped_source_ids.append(instruction.source_id or instruction.section)
                continue
            contexts = self._paragraph_contexts(document)
            if self._should_append(instruction):
                self._append_section(document, instruction)
                applied += 1
                self.applied_source_ids.append(instruction.source_id or instruction.section)
                continue
            if instruction.original and self._replace_exact(contexts, instruction):
                applied += 1
                self.applied_source_ids.append(instruction.source_id or instruction.section)
                continue
            if instruction.original and self._replace_section_block(contexts, instruction):
                applied += 1
                self.applied_source_ids.append(instruction.source_id or instruction.section)
                continue
            if instruction.original and self._replace_similar(contexts, instruction):
                applied += 1
                self.applied_source_ids.append(instruction.source_id or instruction.section)
                continue
            if instruction.original and self._replace_section_block(contexts, instruction):
                applied += 1
                self.applied_source_ids.append(instruction.source_id or instruction.section)
                continue
            if self._replace_in_section(contexts, instruction):
                applied += 1
                self.applied_source_ids.append(instruction.source_id or instruction.section)
                continue
            if self._append_to_target_section(document, instruction):
                applied += 1
                self.applied_source_ids.append(instruction.source_id or instruction.section)
                continue
            self.skipped_source_ids.append(instruction.source_id or instruction.section)

        output = io.BytesIO()
        document.save(output)
        result_bytes = output.getvalue()
        return result_bytes, applied


@dataclass
class StructuredRewriteInstruction:
    suggestion_id: str
    target_section: str
    action: str
    old_text_hint: str
    new_text: str
    items: List[str] = field(default_factory=list)
    reason: str = ""
    confidence: float = 0.0
    source_field: str = ""
    llm_target_section: str = ""
    section_override_reason: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return {
            "suggestion_id": self.suggestion_id,
            "target_section": self.target_section,
            "action": self.action,
            "old_text_hint": self.old_text_hint,
            "new_text": self.new_text,
            "items": list(self.items),
            "reason": self.reason,
            "confidence": self.confidence,
            "source_field": self.source_field,
            "llm_target_section": self.llm_target_section,
            "section_override_reason": self.section_override_reason,
        }


@dataclass
class DocxApplyResult:
    file_bytes: bytes
    sections_detected: List[str]
    applied_ids: List[str]
    partially_applied_ids: List[str]
    failed_ids: List[str]
    duplicate_warnings: List[str]
    validation_report: Dict[str, Any]


class ResumeDocxOptimizationPipeline(DocxPreserver):
    def __init__(self) -> None:
        self.rewriter = ResumeRewriter()
        self.parser = ResumeParser()
        self._source_cv_text = ""
        self._user_additional_data: Dict[str, Any] = {}

    def generate_structured_instructions(
        self,
        cv_text: str,
        role: str,
        company: str,
        goal: str,
        accepted_suggestions: Optional[List[Dict[str, Any]]] = None,
        user_additional_data: Optional[Dict[str, Any]] = None,
        use_llm: bool = True,
    ) -> List[StructuredRewriteInstruction]:
        accepted_suggestions = accepted_suggestions or []
        self._source_cv_text = cv_text or ""
        self._user_additional_data = dict(user_additional_data or {})
        normalized_suggestions: List[Dict[str, Any]] = []
        strong_section_by_id: Dict[str, Dict[str, str]] = {}
        for index, suggestion in enumerate(accepted_suggestions):
            if not isinstance(suggestion, dict):
                continue
            suggestion_id = str(
                suggestion.get("suggestion_id")
                or suggestion.get("id")
                or suggestion.get("source_id")
                or f"suggestion_{index + 1}"
            )
            target_section = str(suggestion.get("section") or suggestion.get("target_section") or "").strip()
            if not target_section:
                continue
            old_text_hint = str(
                suggestion.get("original_text")
                or suggestion.get("original")
                or suggestion.get("old_text_hint")
                or ""
            ).strip()
            new_text = str(
                suggestion.get("proposed_text")
                or suggestion.get("replacement")
                or suggestion.get("new_text")
                or ""
            ).strip()
            if not new_text:
                continue
            source_field = self._source_field_from_id(suggestion_id)
            strong_target = self._strong_section_for_source(source_field, target_section)
            items = suggestion.get("items") if isinstance(suggestion.get("items"), list) else []
            if source_field:
                strong_section_by_id[suggestion_id] = {
                    "source_field": source_field,
                    "target_section": strong_target,
                }
            normalized_suggestions.append({
                "suggestion_id": suggestion_id,
                "target_section": strong_target,
                "action": str(
                    suggestion.get("action")
                    or ("append" if not old_text_hint else "replace")
                ).strip().lower(),
                "old_text_hint": old_text_hint,
                "new_text": new_text,
                "items": [str(item) for item in items if str(item).strip()],
                "reason": str(suggestion.get("description") or suggestion.get("reason") or "").strip(),
                "source_field": source_field,
            })

        confirmed_skill_groups: Dict[str, List[str]] = {}
        existing_skill_text = normalize_text(
            " ".join(str(item.get("new_text") or "") for item in normalized_suggestions)
        )
        confirmed_skills = (user_additional_data or {}).get("confirmed_skills") or []
        if isinstance(confirmed_skills, list):
            for item in confirmed_skills:
                if isinstance(item, dict):
                    name = str(item.get("name") or item.get("skill") or "").strip()
                    category = normalize_text(str(item.get("category") or "hard_skill"))
                else:
                    name = str(item or "").strip()
                    category = "hard_skill"
                if not name:
                    continue
                if normalize_text(name) in existing_skill_text:
                    continue
                target_section = "SOFT SKILLS" if category in {"soft_skill", "soft skill"} else "COMPETENZE TECNICHE"
                group = confirmed_skill_groups.setdefault(target_section, [])
                if normalize_text(name) not in {normalize_text(existing) for existing in group}:
                    group.append(name)

        for target_section, skill_names in confirmed_skill_groups.items():
            normalized_suggestions.append({
                "suggestion_id": f"confirmed_{normalize_text(target_section).replace(' ', '_')}",
                "target_section": target_section,
                "action": "append",
                "old_text_hint": "",
                "new_text": " | ".join(skill_names),
                "items": list(skill_names),
                "reason": "Skill confermate dall'utente integrate nella sezione corretta.",
                "source_field": "confirmed_skills",
            })

        prompt = f"""
Restituisci SOLO JSON valido con una chiave `instructions` che contenga una lista di istruzioni operative.
Ogni istruzione deve avere: suggestion_id, target_section, action, old_text_hint, new_text, items.

CV originale:
{cv_text[:6000]}

Ruolo: {role or "Non specificato"}
Azienda: {company or "Non specificata"}
Obiettivo: {goal or "Non specificato"}

Suggerimenti accettati:
{json.dumps(normalized_suggestions, ensure_ascii=False)}

Dati aggiuntivi utente:
{json.dumps(user_additional_data or {}, ensure_ascii=False)}
"""
        instructions: List[StructuredRewriteInstruction] = []
        if use_llm and _cv_rewrite_llm_enabled():
            try:
                result = extract_json(call_groq(prompt, temperature=0.1, max_tokens=1400, json_mode=True), context="structured_docx_instructions")
                raw_instructions = result.get("instructions") if isinstance(result, dict) else []
                if isinstance(raw_instructions, list):
                    for item in raw_instructions:
                        if not isinstance(item, dict):
                            continue
                        suggestion_id = str(item.get("suggestion_id") or "")
                        llm_target = str(item.get("target_section") or "")
                        strong = strong_section_by_id.get(suggestion_id)
                        final_target = strong["target_section"] if strong else llm_target
                        source_field = strong.get("source_field", "") if strong else ""
                        override_reason = ""
                        if strong and canonical_section(llm_target) != canonical_section(final_target):
                            override_reason = "hint forte dal frontend: la sezione proposta dall'LLM e stata ignorata"
                            print(
                                "[CV-OPT DEBUG] structured section override: "
                                f"source_field={source_field}, final_target={final_target}, "
                                f"llm_target={llm_target}, reason={override_reason}"
                            )
                        instructions.append(StructuredRewriteInstruction(
                            suggestion_id=suggestion_id,
                            target_section=final_target,
                            action=str(item.get("action") or "replace").strip() or "replace",
                            old_text_hint=str(item.get("old_text_hint") or ""),
                            new_text=str(item.get("new_text") or ""),
                            items=[str(x) for x in (item.get("items") if isinstance(item.get("items"), list) else []) if str(x).strip()],
                            reason=str(item.get("reason") or ""),
                            confidence=float(item.get("confidence") or 0.0),
                            source_field=source_field,
                            llm_target_section=llm_target,
                            section_override_reason=override_reason,
                        ))
            except Exception as exc:
                print(f"Generazione istruzioni strutturate non disponibile: {exc}")

        if not instructions:
            for suggestion in normalized_suggestions:
                instructions.append(StructuredRewriteInstruction(
                    suggestion_id=suggestion["suggestion_id"],
                    target_section=suggestion["target_section"],
                    action=suggestion["action"],
                    old_text_hint=suggestion["old_text_hint"],
                    new_text=suggestion["new_text"],
                    items=suggestion["items"],
                    reason=suggestion["reason"],
                    confidence=0.5,
                    source_field=suggestion.get("source_field", ""),
                    llm_target_section="",
                    section_override_reason="",
                ))
        return self._sanitize_structured_instructions(
            instructions,
            cv_text,
            user_additional_data or {},
        )

    def _sanitize_structured_instructions(
        self,
        instructions: List[StructuredRewriteInstruction],
        cv_text: str,
        user_additional_data: Dict[str, Any],
    ) -> List[StructuredRewriteInstruction]:
        source_sections = self.parser.section_text_map(cv_text)
        support_text = normalize_text(
            " ".join([
                cv_text or "",
                *self._confirmed_skill_names(user_additional_data),
                *self._user_note_texts(user_additional_data),
            ])
        )
        note_texts = [
            normalize_text(note)
            for note in self._user_note_texts(user_additional_data)
            if len(normalize_text(note)) >= 14
        ]
        cleaned: List[StructuredRewriteInstruction] = []
        seen = set()

        # Le info extra confermate dall'utente possono creare nuove sezioni anche se assenti nell'originale.
        user_provided_sections: set[str] = set()
        if isinstance(user_additional_data, dict):
            for field_name, raw_value in user_additional_data.items():
                if field_name in {"adaptation_answers", "confirmed_skills"}:
                    continue
                if not isinstance(raw_value, str) or not raw_value.strip():
                    continue
                section_key = additional_field_section_key(field_name)
                section_name = self._pipeline_section_from_key(section_key or "")
                if section_name:
                    user_provided_sections.add(section_name)
            answers = user_additional_data.get("adaptation_answers") or []
            if isinstance(answers, list):
                for ans in answers:
                    if not isinstance(ans, dict):
                        continue
                    if not str(ans.get("answer") or "").strip():
                        continue
                    section_key = additional_field_section_key(str(ans.get("category") or ""))
                    section_name = self._pipeline_section_from_key(section_key or "")
                    if section_name:
                        user_provided_sections.add(section_name)
            confirmed_skills = user_additional_data.get("confirmed_skills") or []
            if isinstance(confirmed_skills, list):
                for item in confirmed_skills:
                    if not isinstance(item, dict):
                        continue
                    detail = normalize_text(
                        str(item.get("user_example") or item.get("detail") or "")
                    )
                    if not detail:
                        continue
                    if any(term in detail for term in ["azienda", "cliente", "lavoro", "tirocinio", "stage", "impiego"]):
                        user_provided_sections.add("esperienze")
                    if any(term in detail for term in ["progetto", "dashboard", "prototipo", "dataset", "portfolio"]):
                        user_provided_sections.add("progetti")
                    if any(term in detail for term in ["certificazione", "certificato", "attestato", "licenza"]):
                        user_provided_sections.add("certificazioni")
                    if any(term in detail for term in ["laurea", "universita", "corso", "esame", "formazione"]):
                        user_provided_sections.add("formazione")

        for instruction in instructions:
            target = canonical_section(instruction.target_section)
            if target == "esperienze" and not source_sections.get("esperienze") and "esperienze" not in user_provided_sections:
                continue
            if target == "certificazioni" and not source_sections.get("certificazioni") and "certificazioni" not in user_provided_sections:
                explicit_certification = normalize_text(
                    " ".join([instruction.new_text, *self._user_note_texts(user_additional_data)])
                )
                if not any(term in explicit_certification for term in [
                    "certificazione", "certificato", "attestato", "licenza", "corso", "master",
                ]):
                    continue

            text = self._clean_instruction_text(instruction.new_text, note_texts)
            text = self._format_section_text(target, text, support_text)
            text = self._fix_common_grammar(text)
            key = (target, normalize_text(text))
            if not text or key in seen:
                continue
            seen.add(key)
            cleaned.append(StructuredRewriteInstruction(
                suggestion_id=instruction.suggestion_id,
                target_section=instruction.target_section,
                action="append" if str(instruction.action).strip().lower() == "append" else "replace",
                old_text_hint=instruction.old_text_hint,
                new_text=text,
                items=list(instruction.items),
                reason=instruction.reason,
                confidence=instruction.confidence,
                source_field=instruction.source_field,
                llm_target_section=instruction.llm_target_section,
                section_override_reason=instruction.section_override_reason,
            ))
            print(
                "[CV-OPT DEBUG] structured instruction sanitized: "
                f"source_field={instruction.source_field or '-'}, "
                f"target_final={instruction.target_section}, "
                f"llm_target={instruction.llm_target_section or '-'}, "
                f"override_reason={instruction.section_override_reason or '-'}"
            )
        return cleaned

    def _confirmed_skill_names(self, user_data: Dict[str, Any]) -> List[str]:
        names: List[str] = []
        for item in user_data.get("confirmed_skills", []) if isinstance(user_data.get("confirmed_skills"), list) else []:
            name = str(item.get("name") or item.get("skill") or "") if isinstance(item, dict) else str(item or "")
            if name.strip():
                names.append(name.strip())
        return names

    def _source_field_from_id(self, suggestion_id: str) -> str:
        raw = str(suggestion_id or "")
        patterns = (
            r"user_box_([a-z_]+)_\d+",
            r"user_additional_answer_\d+_\d+_([a-z_]+)",
            r"confirmed_([a-z_]+)",
        )
        for pattern in patterns:
            match = re.search(pattern, raw)
            if match:
                return (match.group(1) or "").strip("_")
        return ""

    def _strong_section_for_source(self, source_field: str, fallback_section: str) -> str:
        section_key = additional_field_section_key(source_field)
        if not section_key:
            return fallback_section
        return self._display_section_from_key(section_key, fallback_section)

    def _pipeline_section_from_key(self, section_key: str, fallback: str = "") -> str:
        return {
            "profile": "profilo",
            "experience": "esperienze",
            "education": "formazione",
            "hard_skills": "hard_skills",
            "soft_skills": "soft_skills",
            "languages": "lingue",
            "projects": "progetti",
            "certifications": "certificazioni",
            "contacts": "contatti",
        }.get(section_key, fallback)

    def _display_section_from_key(self, section_key: str, fallback: str = "") -> str:
        return {
            "profile": "PROFILO",
            "experience": "ESPERIENZE PROFESSIONALI",
            "education": "FORMAZIONE",
            "hard_skills": "COMPETENZE TECNICHE",
            "soft_skills": "SOFT SKILLS",
            "languages": "LINGUE",
            "projects": "PROGETTI",
            "certifications": "CERTIFICAZIONI",
            "contacts": "CONTATTI",
        }.get(section_key, fallback)

    def _user_note_texts(self, user_data: Dict[str, Any]) -> List[str]:
        notes: List[str] = []
        for field_name in (
            "experiences", "projects", "certifications", "languages", "education",
            "technical_skills", "tools", "soft_skills", "measurable_results",
            "company_role_notes", "additional_notes",
        ):
            value = str(user_data.get(field_name) or "").strip()
            if value:
                notes.append(value)
        for item in user_data.get("adaptation_answers", []) if isinstance(user_data.get("adaptation_answers"), list) else []:
            if isinstance(item, dict) and str(item.get("answer") or "").strip():
                notes.append(str(item.get("answer")).strip())
        additional = str(user_data.get("additional_notes") or "").strip()
        if additional:
            notes.append(additional)
        for item in user_data.get("confirmed_skills", []) if isinstance(user_data.get("confirmed_skills"), list) else []:
            if isinstance(item, dict):
                detail = str(item.get("user_example") or item.get("detail") or "").strip()
                if detail:
                    notes.append(detail)
        return notes

    def _format_section_text(self, target: str, text: str, support_text: str) -> str:
        if not text:
            return ""
        if target in {"hard_skills", "soft_skills", "competenze"}:
            return self._format_skill_lines(
                self._filter_supported_skill_text(text, support_text)
            )
        if target == "progetti":
            return self._format_project_entries(text)
        if target in {"lingue", "certificazioni", "formazione"}:
            return self._format_fact_lines(text)
        if target in {"esperienze", "profilo", "attivita rilevanti"}:
            return self._format_prose_lines(text)
        return self._format_prose_lines(text)

    def _format_skill_lines(self, text: str) -> str:
        values: List[str] = []
        seen = set()
        for part in re.split(r"[,;|Ã‚Â·Ã¢â‚¬Â¢\n]+", text or ""):
            value = re.sub(r"\s+", " ", part).strip(" -.")
            key = normalize_text(value)
            if value and key and key not in seen:
                seen.add(key)
                values.append(value)
        return " | ".join(values)

    def _format_fact_lines(self, text: str) -> str:
        rows: List[str] = []
        seen = set()
        for raw_line in re.split(r"\n+|(?<=[.!?])\s+", text or ""):
            line = re.sub(r"\s+", " ", raw_line).strip(" -")
            line = self._normalize_certification_line(line)
            key = normalize_text(line)
            if not line or not key or key in seen:
                continue
            seen.add(key)
            rows.append(line.rstrip("."))
        return "\n".join(rows)

    def _normalize_certification_line(self, line: str) -> str:
        cleaned = line or ""
        language_typos = {
            r"\bfrance[sc]e\b": "Francese",
            r"\bfrancesce\b": "Francese",
            r"\bfrancesee\b": "Francese",
            r"\bingles[eai]\b": "Inglese",
            r"\bspagnollo\b": "Spagnolo",
            r"\btedescoo\b": "Tedesco",
        }
        for pattern, replacement in language_typos.items():
            cleaned = re.sub(pattern, replacement, cleaned, flags=re.IGNORECASE)
        plain = normalize_text(cleaned)
        level_match = re.search(r"\b([abc][12])\b", plain)
        provider_aliases = {
            "cambridge": "Cambridge",
            "ielts": "IELTS",
            "toefl": "TOEFL",
            "dele": "DELE",
            "delf": "DELF",
            "dalf": "DALF",
            "goethe": "Goethe",
        }
        provider = next((label for marker, label in provider_aliases.items() if marker in plain), "")
        if level_match and provider:
            language = next(
                (
                    label
                    for marker, label in {
                        "francese": "Francese",
                        "inglese": "Inglese",
                        "spagnolo": "Spagnolo",
                        "tedesco": "Tedesco",
                        "italiano": "Italiano",
                    }.items()
                    if marker in plain
                ),
                "linguistica",
            )
            return f"Certificazione {language} {level_match.group(1).upper()} {provider}"
        if level_match and "certific" in plain:
            return f"Certificazione linguistica {level_match.group(1).upper()}"
        return cleaned

    def _format_prose_lines(self, text: str) -> str:
        rows: List[str] = []
        seen = set()
        for raw_line in (text or "").splitlines():
            line = re.sub(r"\s+", " ", raw_line).strip(" -")
            key = normalize_text(line)
            if not line or not key or key in seen:
                continue
            seen.add(key)
            rows.append(line if line.endswith((".", ":", ";")) else f"{line}.")
        return "\n".join(rows)

    def _normalize_sentence_case(self, text: str) -> str:
        cleaned = re.sub(r"\s+", " ", text or "").strip()
        if not cleaned:
            return ""
        if cleaned.isupper() and len(cleaned.split()) > 1:
            cleaned = cleaned.lower()
        return cleaned[:1].upper() + cleaned[1:]

    def _clean_instruction_text(self, text: str, note_texts: List[str]) -> str:
        lines: List[str] = []
        for raw_line in re.sub(r"\r\n?", "\n", text or "").splitlines():
            original_line = re.sub(r"[ \t]+", " ", raw_line).strip()
            line = self._professionalize_informal_line(original_line)
            plain = normalize_text(line)
            if not line:
                continue
            if any(marker in plain for marker in INFORMAL_CV_MARKERS):
                continue
            lines.append(line)
        return "\n".join(lines).strip()

    def _professionalize_informal_line(self, line: str) -> str:
        plain = normalize_text(line)
        if "allineare requisiti e aspettative" in plain:
            return (
                "Collaborazione con il team per l'analisi dei requisiti e "
                "l'allineamento delle attivita progettuali."
            )
        replacements = [
            (r"^ho\s+applicato\s+", "Applicazione di "),
            (r"^ho\s+(usato|utilizzato)\s+", "Utilizzo di "),
            (r"^(usata?|utilizzata?)\s+in\s+", "Applicazione in "),
            (r"^(usata?|utilizzata?)\s+durante\s+", "Applicazione durante "),
            (r"\b(visto|vista)\s+all'?esame\b", "approfondito in ambito universitario"),
        ]
        cleaned = line
        for pattern, replacement in replacements:
            cleaned = re.sub(pattern, replacement, cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"[ \t]{2,}", " ", cleaned).strip()
        return cleaned[:1].upper() + cleaned[1:] if cleaned else ""

    def _filter_supported_skill_text(self, text: str, support_text: str) -> str:
        filtered = text
        for skill, evidence_terms in CONTROLLED_PM_SKILLS.items():
            if skill not in normalize_text(filtered):
                continue
            if not any(term in support_text for term in evidence_terms):
                filtered = re.sub(
                    rf"(?i)\b{re.escape(skill)}\b",
                    "",
                    filtered,
                )
        rows: List[str] = []
        seen = set()
        for raw_line in filtered.splitlines():
            line = re.sub(r"\s*[,;|Â·â€¢]\s*", " Â· ", raw_line).strip(" Â·,;|")
            parts = [part.strip() for part in line.split(" Â· ") if part.strip()]
            unique_parts = []
            for part in parts:
                key = normalize_text(part)
                if key and key not in seen:
                    seen.add(key)
                    unique_parts.append(part)
            if unique_parts:
                rows.append(" Â· ".join(unique_parts))
        return "\n".join(rows).strip()

    def _format_project_entries(self, text: str) -> str:
        lines = [line.strip(" -Â·â€¢") for line in (text or "").splitlines() if line.strip()]
        if not lines:
            return ""
        if len(lines) == 1:
            single = lines[0].rstrip(".")
            return f"{self._project_title_from_text(single)}\n{self._normalize_sentence_case(single)}."

        rows: List[str] = []
        index = 0
        while index < len(lines):
            title = lines[index]
            plain = normalize_text(title)
            if not plain.startswith(("progetto ", "project ", "tesi ")):
                if len(plain.split()) <= 7 and len(title) <= 70:
                    rows.append(
                        f"{self._project_title_from_text(title)}\n"
                        f"{self._normalize_sentence_case(title).rstrip('.')}."
                    )
                index += 1
                continue
            if index + 1 >= len(lines):
                rows.append(f"{self._normalize_sentence_case(title).rstrip('.')}.")
                break
            description = lines[index + 1].strip()
            if (
                normalize_text(description).startswith(("progetto ", "project ", "tesi "))
                or any(marker in normalize_text(description) for marker in INFORMAL_CV_MARKERS)
            ):
                index += 1
                continue
            rows.extend([
                self._project_title_from_text(title),
                self._normalize_sentence_case(description).rstrip(".") + ".",
            ])
            index += 2
        return "\n".join(rows).strip()

    def _project_title_from_text(self, text: str) -> str:
        cleaned = re.sub(
            r"(?i)^\s*(progetto|project)\s+(personale|accademico|universitario|professionale)?\s*[:/-]?\s*",
            "",
            text or "",
        ).strip()
        cleaned = re.sub(r"(?i)^\s*sviluppo\s+di\s+", "", cleaned).strip()
        words = cleaned.split()
        title = " ".join(words[:5]).strip(" .,:;-") or "Progetto"
        return title[:1].upper() + title[1:]

    def _fix_common_grammar(self, text: str) -> str:
        text = self._repair_mojibake(text)
        cleaned = re.sub(
            r"\brischiare\s+(attivit[aÃ ]|progetti?)\b",
            r"gestire i rischi delle \1",
            text or "",
            flags=re.IGNORECASE,
        )
        cleaned = re.sub(r"\bcon applicand[oa]?\b", "applicando", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\battraverso di esam(?:e|es|i)\b", "attraverso esami", cleaned, flags=re.IGNORECASE)
        return re.sub(r"[ \t]{2,}", " ", cleaned).strip()

    def _repair_mojibake(self, text: str) -> str:
        repaired = text or ""
        for broken, replacement in {
            "\u00c2\u00b7": "|",
            "\u00c3\u201a\u00c2\u00b7": "|",
            "\u00e2\u20ac\u00a2": "|",
            "\u00c3\u00a0": "a",
            "\u00c3\u00a8": "e",
            "\u00c3\u00a9": "e",
            "\u00c3\u00ac": "i",
            "\u00c3\u00b2": "o",
            "\u00c3\u00b9": "u",
        }.items():
            repaired = repaired.replace(broken, replacement)
        return repaired

    def apply_instructions_to_docx(
        self,
        original_file_bytes: bytes,
        instructions: List[StructuredRewriteInstruction],
    ) -> DocxApplyResult:
        from docx import Document

        document = Document(io.BytesIO(original_file_bytes))
        textbox_mirror_groups = self._textbox_mirror_groups(document)
        sections_detected = self._detect_sections(document)
        applied_ids: List[str] = []
        partially_applied_ids: List[str] = []
        failed_ids: List[str] = []

        for instruction in instructions:
            outcome = self._apply_single_instruction(document, instruction, sections_detected)
            if outcome == "applied":
                applied_ids.append(instruction.suggestion_id)
            elif outcome == "partial":
                partially_applied_ids.append(instruction.suggestion_id)
            else:
                failed_ids.append(instruction.suggestion_id)

        self._clean_corrupted_document_text(document)
        self._synchronize_textbox_mirrors(textbox_mirror_groups)
        self._polish_document_typography(document)
        self._remove_oversized_table_row_minimums(document)
        self._collapse_trailing_blank_paragraphs(document)
        output = io.BytesIO()
        document.save(output)
        file_bytes = output.getvalue()
        validation_report = self.validate_generated_docx(file_bytes, original_file_bytes, instructions)
        duplicate_warnings = validation_report.get("duplicate_warnings", [])
        return DocxApplyResult(
            file_bytes=file_bytes,
            sections_detected=sections_detected,
            applied_ids=applied_ids,
            partially_applied_ids=partially_applied_ids,
            failed_ids=failed_ids,
            duplicate_warnings=duplicate_warnings,
            validation_report=validation_report,
        )

    def _remove_empty_section_headings(self, document) -> None:
        """Remove orphan headings without touching sections backed by a table."""
        paragraphs = list(document.paragraphs)
        for index, paragraph in enumerate(paragraphs):
            if not is_section_heading(paragraph.text or ""):
                continue
            has_content = False
            for following in paragraphs[index + 1:]:
                text = (following.text or "").strip()
                if not text:
                    continue
                if is_section_heading(text):
                    break
                has_content = True
                break
            if has_content:
                continue

            sibling = paragraph._p.getnext()
            while sibling is not None:
                if sibling.tag.endswith("}tbl"):
                    has_content = True
                    break
                if sibling.tag.endswith("}p"):
                    sibling_text = "".join(sibling.itertext()).strip()
                    if sibling_text:
                        break
                sibling = sibling.getnext()
            if not has_content:
                self._replace_paragraph_preserving_style(paragraph, "")

    def _polish_document_typography(self, document) -> None:
        """Harmonize equivalent font aliases while preserving the original design."""
        from docx.oxml.ns import qn

        font_aliases = {
            "arial mt": "Arial",
            "arialmt": "Arial",
        }

        def normalize_run_font(run) -> None:
            current = str(run.font.name or "").strip()
            replacement = font_aliases.get(current.lower())
            if not replacement:
                return
            run.font.name = replacement
            r_pr = run._r.get_or_add_rPr()
            r_fonts = r_pr.rFonts
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.insert(0, r_fonts)
            for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                r_fonts.set(qn(f"w:{attribute}"), replacement)

        for style in document.styles:
            style_font = getattr(style, "font", None)
            if style_font is None:
                continue
            current = str(style_font.name or "").strip()
            replacement = font_aliases.get(current.lower())
            if replacement:
                style_font.name = replacement

        seen = set()
        for context in self._paragraph_contexts(document):
            paragraph = context.paragraph
            paragraph_id = id(paragraph._p)
            if paragraph_id in seen:
                continue
            seen.add(paragraph_id)
            for run in paragraph.runs:
                normalize_run_font(run)
            if is_section_heading(paragraph.text or ""):
                paragraph.paragraph_format.keep_with_next = True

    def _remove_oversized_table_row_minimums(self, document) -> None:
        """Let full-page layout tables size to content instead of forcing overflow."""
        for table in document.tables:
            for row in table.rows:
                tr_pr = row._tr.get_or_add_trPr()
                for child in list(tr_pr):
                    if child.tag.endswith("}trHeight"):
                        tr_pr.remove(child)

    def _collapse_trailing_blank_paragraphs(self, document) -> None:
        """Keep only one hidden trailing blank paragraph and minimize its footprint."""
        from docx.shared import Pt

        body = document._element.body
        children = list(body)
        if not children:
            return
        trailing_paragraphs = []
        for paragraph in reversed(document.paragraphs):
            if (paragraph.text or "").strip():
                break
            trailing_paragraphs.append(paragraph)
        if not trailing_paragraphs:
            return
        for index, paragraph in enumerate(trailing_paragraphs):
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(1)
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
            run.font.size = Pt(1)
            run.font.hidden = True
            if index > 0:
                paragraph._element.getparent().remove(paragraph._element)

    def validate_generated_docx(
        self,
        generated_file_bytes: bytes,
        original_file_bytes: bytes,
        instructions: List[StructuredRewriteInstruction],
    ) -> Dict[str, Any]:
        final_text = self._extract_docx_text_bytes(generated_file_bytes)
        original_text = self._extract_docx_text_bytes(original_file_bytes)
        normalized_final = normalize_text(final_text)
        normalized_original = normalize_text(original_text)

        applied = []
        partial = []
        failed = []
        for instruction in instructions:
            needle = normalize_text(instruction.new_text)
            needle_lines = [
                normalize_text(line)
                for line in (instruction.new_text or "").splitlines()
                if normalize_text(line)
            ]
            hint = normalize_text(instruction.old_text_hint)
            matched_lines = sum(
                1 for line in needle_lines
                if self._semantic_match(normalized_final, line)
            )
            if needle_lines and matched_lines == len(needle_lines):
                applied.append(instruction.suggestion_id)
            elif needle and self._semantic_match(normalized_final, needle):
                applied.append(instruction.suggestion_id)
            elif needle_lines and matched_lines:
                partial.append(instruction.suggestion_id)
            elif hint and self._semantic_match(normalized_final, hint):
                partial.append(instruction.suggestion_id)
            else:
                failed.append(instruction.suggestion_id)

        duplicate_warnings = []
        for marker in self._duplicate_markers(final_text):
            duplicate_warnings.append(marker)

        contamination_warnings = self._skill_section_contamination_warnings(
            generated_file_bytes
        )
        heading_warnings = self._duplicate_section_heading_warnings(
            generated_file_bytes
        )
        added_sections = self._unexpected_sections(normalized_original, normalized_final)
        content_warnings = self._final_content_warnings(final_text, original_text)
        status = "applied"
        if failed:
            status = "partially_applied" if (applied or partial) else "failed"
        if duplicate_warnings and len(failed) > 0:
            status = "partially_applied" if (applied or partial) else "failed"
        if applied or partial:
            status = "applied" if not failed else "partially_applied"
        objective_targets = {"profilo", "objective", "obiettivo"}
        education_targets = {"formazione", "education"}
        if any(canonical_section(inst.target_section) in objective_targets for inst in instructions):
            original_objective = self._extract_section_text(normalized_original, objective_targets)
            final_objective = self._extract_section_text(normalized_final, objective_targets)
            if original_objective and final_objective and SequenceMatcher(None, original_objective, final_objective).ratio() >= 0.88:
                status = "failed"
                failed = [inst.suggestion_id for inst in instructions if canonical_section(inst.target_section) in objective_targets]
        if self._objective_followed_by_education(normalized_final):
            status = "failed"
            failed = [inst.suggestion_id for inst in instructions if canonical_section(inst.target_section) in education_targets] or failed
        if contamination_warnings:
            status = "failed"
            failed = list(dict.fromkeys([
                *failed,
                *[
                    inst.suggestion_id
                    for inst in instructions
                    if canonical_section(inst.target_section) in {
                        "hard_skills", "soft_skills", "competenze",
                    }
                ],
            ]))
        if heading_warnings:
            status = "failed"
        if content_warnings and status == "applied":
            status = "partially_applied"

        return {
            "status": status,
            "applied": applied,
            "partially_applied": partial,
            "failed": failed,
            "duplicate_warnings": duplicate_warnings,
            "heading_warnings": heading_warnings,
            "contamination_warnings": contamination_warnings,
            "content_warnings": content_warnings,
            "unexpected_sections": added_sections,
            "final_text": final_text,
        }

    def _final_content_warnings(self, final_text: str, original_text: str) -> List[str]:
        warnings: List[str] = []
        final_plain = normalize_text(final_text)
        original_sections = self.parser.section_text_map(original_text)
        final_sections = self.parser.section_text_map(final_text)

        for marker in INFORMAL_CV_MARKERS:
            if marker in final_plain:
                warnings.append(f"Formula informale ancora presente: {marker}")

        for note in self._user_note_texts(self._user_additional_data):
            note_plain = normalize_text(note)
            if len(note_plain) >= 14 and note_plain in final_plain:
                warnings.append("Una nota utente risulta copiata letteralmente nel CV.")
                break

        support_text = normalize_text(
            " ".join([
                original_text,
                *self._confirmed_skill_names(self._user_additional_data),
                *self._user_note_texts(self._user_additional_data),
            ])
        )
        for skill, evidence_terms in CONTROLLED_PM_SKILLS.items():
            if skill in final_plain and not any(term in support_text for term in evidence_terms):
                warnings.append(f"Competenza non supportata rilevata: {skill}")

        project_text = final_sections.get("progetti", "")
        if project_text:
            project_lines = [line.strip() for line in project_text.splitlines() if line.strip()]
            for index, line in enumerate(project_lines):
                if index % 2 == 0 and not normalize_text(line).startswith(("progetto ", "project ", "tesi ")):
                    warnings.append("La sezione PROGETTI non rispetta il formato titolo e descrizione.")
                    break
        return list(dict.fromkeys(warnings))[:10]

    def _skill_section_contamination_warnings(self, file_bytes: bytes) -> List[str]:
        from docx import Document

        try:
            document = Document(io.BytesIO(file_bytes))
        except Exception:
            return ["Impossibile verificare le sezioni skill del DOCX."]

        warnings: List[str] = []
        blocked_terms = {
            "universita", "università", "laurea", "diploma", "liceo",
            "tirocinio", "curriculare", "formazione", "esperienza professionale",
        }
        for context in self._paragraph_contexts(document):
            if context.section not in {"hard_skills", "soft_skills", "competenze"}:
                continue
            text = (context.paragraph.text or "").strip()
            plain = normalize_text(text)
            if not text or is_section_heading(text):
                continue
            if (
                any(term in plain for term in blocked_terms)
                or bool(re.search(r"\b(?:19|20)\d{2}\b", text))
                or plain.startswith(("e ", "in ", "con "))
                or text.endswith((".", "!", "?"))
                or bool(re.search(r"\b[A-ZÀ-ÖØ-Þ]{2,}\s+[A-ZÀ-ÖØ-Þ]{2,}\b", text))
            ):
                warnings.append(
                    f"{context.section}: contenuto estraneo rilevato: {text[:140]}"
                )
        return warnings[:8]

    def _detect_sections(self, document) -> List[str]:
        sections: List[str] = []
        for context in self._paragraph_contexts(document):
            text = (context.paragraph.text or "").strip()
            if is_section_heading(text):
                sections.append(canonical_section(text))
        seen: List[str] = []
        for section in sections:
            if section and section not in seen:
                seen.append(section)
        return seen

    def _detect_sections_from_table(self, table) -> List[str]:
        sections: List[str] = []
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = (paragraph.text or "").strip()
                    if is_section_heading(text):
                        sections.append(canonical_section(text))
                for nested in cell.tables:
                    sections.extend(self._detect_sections_from_table(nested))
        return sections

    def _apply_single_instruction(self, document, instruction: StructuredRewriteInstruction, sections_detected: List[str]) -> str:
        target = canonical_section(instruction.target_section)
        if not target:
            return "failed"
        if not self._is_safe_target_text(instruction):
            print(f"DOCX APPLY - blocked wrong section injection: suggestion_id={instruction.suggestion_id}, target={instruction.target_section}")
            return "failed"
        contexts = self._paragraph_contexts(document)
        matching = [
            context
            for context in contexts
            if context.section == target
            and self._is_editable_paragraph(context.paragraph)
            and self._is_safe_text_destination(context.paragraph)
        ]
        if not matching:
            return self._insert_missing_section(document, instruction, sections_detected)

        replacement_lines = [
            (self.dedupe_contact_line(line) if target == "contatti" else line.strip())
            for line in instruction.new_text.splitlines()
            if line.strip()
        ]
        if not replacement_lines:
            return "failed"
        if instruction.action == "append":
            if target in {"hard_skills", "soft_skills", "competenze"}:
                return self._rewrite_skill_section_canonically(
                    matching,
                    replacement_lines,
                )
            self._append_to_existing_section(document, matching, replacement_lines)
            return "applied"
        if self._is_high_confidence_section_instruction(instruction):
            self._rewrite_section_block(document, matching, replacement_lines, instruction)
        else:
            self._rewrite_single_anchor(document, matching, replacement_lines, instruction)
        return "applied"

    def _insert_missing_section(self, document, instruction: StructuredRewriteInstruction, sections_detected: List[str]) -> str:
        target = canonical_section(instruction.target_section)
        replacement = (instruction.new_text or "").strip()
        if not replacement:
            return "failed"
        if target in sections_detected:
            section_contexts = [
                context
                for context in self._paragraph_contexts(document)
                if context.section == target
                and self._is_safe_text_destination(context.paragraph)
            ]
            if section_contexts:
                replacement_lines = [
                    (self.dedupe_contact_line(line) if target == "contatti" else line.strip())
                    for line in replacement.splitlines()
                    if line.strip()
                ]
                self._append_to_existing_section(document, section_contexts, replacement_lines)
                return "applied"
            heading_anchor = self._find_section_heading_paragraph(document, target)
            if heading_anchor is not None:
                self._append_content_after_anchor(
                    heading_anchor,
                    target,
                    replacement.splitlines(),
                )
                return "applied"
        heading = self._section_display_name(target)
        shape_host = next(
            (
                paragraph
                for paragraph in document.paragraphs
                if self._paragraph_hosts_textbox(paragraph)
            ),
            None,
        )
        if shape_host is not None:
            self._insert_section_block_before(
                shape_host,
                heading,
                replacement,
            )
            return "applied"
        if not document.paragraphs:
            self._append_section_block(document, heading, replacement, None)
            return "applied"

        anchor = self._find_best_section_anchor(document, target, sections_detected)
        if anchor is None:
            self._append_section_block(document, heading, replacement, None)
            return "applied"
        self._append_section_block(document, heading, replacement, anchor)
        return "applied"

    def _find_section_heading_paragraph(self, document, target: str):
        for context in reversed(self._paragraph_contexts(document)):
            text = (context.paragraph.text or "").strip()
            if is_section_heading(text) and canonical_section(text) == target:
                return context.paragraph
        return None

    def _append_content_after_anchor(
        self,
        anchor_paragraph,
        section_name: str,
        replacement_lines: List[str],
    ) -> None:
        previous = anchor_paragraph
        lines = [line.strip() for line in replacement_lines if line.strip()]
        for line_index, line in enumerate(lines):
            new_paragraph = self._insert_paragraph_after(
                previous,
                line,
                source_format=anchor_paragraph,
            )
            self._style_section_line(new_paragraph, section_name, line_index)
            previous = new_paragraph

    def _is_safe_target_text(self, instruction: StructuredRewriteInstruction) -> bool:
        section = canonical_section(instruction.target_section)
        heading_lines = {
            canonical_section(line)
            for line in (instruction.new_text or "").splitlines()
            if is_section_heading(line)
        }
        if section in {"profilo", "objective", "obiettivo"} and heading_lines.intersection(
            {"formazione", "esperienze", "progetti", "hard_skills", "soft_skills", "competenze", "contatti", "lingue", "certificazioni"}
        ):
            return False
        if section in {"formazione", "education"} and heading_lines.intersection(
            {"esperienze", "obiettivo", "profilo", "progetti"}
        ):
            return False
        if section in {"esperienze", "experience"} and heading_lines.intersection(
            {"formazione", "obiettivo", "profilo"}
        ):
            return False
        return True

    def _append_to_existing_section(
        self,
        document,
        matching_contexts: List[ParagraphContext],
        replacement_lines: List[str],
    ) -> None:
        existing_lines = {
            normalize_text(context.paragraph.text or "")
            for context in matching_contexts
            if normalize_text(context.paragraph.text or "")
        }
        existing_contact_ids = self._document_contact_identifiers(document)
        lines_to_append: List[str] = []
        for line in replacement_lines:
            candidate_line = self.dedupe_contact_line(line, existing_contact_ids)
            normalized_line = normalize_text(candidate_line)
            contact_ids = self._contact_identifiers(candidate_line)
            if (
                not normalized_line
                or normalized_line in existing_lines
                or bool(contact_ids.intersection(existing_contact_ids))
            ):
                continue
            existing_lines.add(normalized_line)
            existing_contact_ids.update(contact_ids)
            lines_to_append.append(candidate_line)

        if not lines_to_append:
            return

        anchor = matching_contexts[-1]
        section_name = anchor.section
        previous = anchor.paragraph
        for line_index, line in enumerate(lines_to_append):
            new_paragraph = self._insert_paragraph_after(previous, line)
            self._copy_paragraph_format(previous, new_paragraph)
            self._style_section_line(new_paragraph, section_name, line_index)
            previous = new_paragraph

    def _contact_identifiers(self, text: str) -> set[str]:
        identifiers = {
            f"email:{match.lower()}"
            for match in re.findall(r"[\w.+-]+@[\w.-]+\.\w+", text or "")
        }
        for raw_phone in re.findall(r"\+?\d[\d\s().-]{7,}\d", text or ""):
            digits = re.sub(r"\D+", "", raw_phone)
            if len(digits) >= 8:
                identifiers.add(f"phone:{digits}")
                identifiers.add(f"phone-suffix:{digits[-9:]}")
        for raw_url in re.findall(
            r"(?:https?://|www\.)?[a-z0-9.-]+\.[a-z]{2,}(?:/[^\s,;|]*)?",
            text or "",
            flags=re.IGNORECASE,
        ):
            normalized_url = raw_url.strip().lower().rstrip("./")
            normalized_url = re.sub(r"^https?://", "", normalized_url)
            normalized_url = re.sub(r"^www\.", "", normalized_url)
            normalized_url = normalized_url.split("?", 1)[0].split("#", 1)[0].rstrip("/")
            if normalized_url:
                identifiers.add(f"url:{normalized_url}")
        return identifiers

    def _document_contact_identifiers(self, document) -> set[str]:
        texts = [
            context.paragraph.text or ""
            for context in self._paragraph_contexts(document)
        ]
        for section in getattr(document, "sections", []):
            for container in (section.header, section.footer):
                texts.extend(
                    paragraph.text or ""
                    for paragraph in container.paragraphs
                )
                for table in container.tables:
                    texts.extend(self._collect_table_text(table))
        return {
            contact_id
            for text in texts
            for contact_id in self._contact_identifiers(text)
        }

    def _merge_skill_lines_into_existing_section(
        self,
        matching_contexts: List[ParagraphContext],
        replacement_lines: List[str],
    ) -> str:
        existing_text = "\n".join(
            context.paragraph.text or ""
            for context in matching_contexts
        )
        existing_keys = {
            normalize_text(item)
            for item in re.split(r"[,;|Â·\n]+", existing_text)
            if normalize_text(item)
        }
        missing = [
            item.strip()
            for line in replacement_lines
            for item in re.split(r"[,;|Â·\n]+", line)
            if item.strip() and normalize_text(item) not in existing_keys
        ]
        if not missing:
            return "applied"

        anchor = matching_contexts[-1].paragraph
        separator = self._skill_separator(anchor.text or "")
        merged = (anchor.text or "").strip()
        addition = separator.join(dict.fromkeys(missing))
        merged = f"{merged}{separator if merged else ''}{addition}"
        self._replace_paragraph_preserving_style(anchor, merged)
        return "applied"

    def _skill_separator(self, text: str) -> str:
        if "Â·" in text:
            return " Â· "
        if "|" in text:
            return " | "
        if ";" in text:
            return "; "
        return ", "

    def _rewrite_skill_section_canonically(
        self,
        matching_contexts: List[ParagraphContext],
        replacement_lines: List[str],
    ) -> str:
        if not matching_contexts:
            return "failed"

        anchor = self._best_skill_anchor_paragraph(matching_contexts)
        if anchor is None:
            return "failed"

        existing_items: List[str] = []
        for context in matching_contexts:
            existing_items.extend(self._split_skill_items(context.paragraph.text or ""))

        new_items: List[str] = []
        for line in replacement_lines:
            new_items.extend(self._split_skill_items(line))

        canonical_items = self._dedupe_skill_items(existing_items + new_items)
        separator = self._skill_separator(anchor.text or "")
        canonical_text = separator.join(canonical_items)
        self._replace_paragraph_preserving_style(anchor, canonical_text)
        self._remove_duplicate_skill_paragraphs(matching_contexts, anchor)
        return "applied"

    def _split_skill_items(self, text: str) -> List[str]:
        raw_items = re.split(r"[,;|·•\n]+", text or "")
        items: List[str] = []
        for item in raw_items:
            clean = re.sub(r"\s+", " ", item).strip(" \t\r\n-–—•·;:,.")
            if clean:
                items.append(clean)
        return items

    def _dedupe_skill_items(self, items: List[str]) -> List[str]:
        deduped: List[str] = []
        seen = set()
        for item in items:
            normalized = re.sub(r"\s+", " ", normalize_text(item)).strip(" \t\r\n-–—•·;:,.")
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            deduped.append(re.sub(r"\s+", " ", item).strip())
        filtered: List[str] = []
        normalized_items = [
            re.sub(r"\s+", " ", normalize_text(item)).strip(" \t\r\n-â€“â€”â€¢Â·;:,.")
            for item in deduped
        ]
        for index, item in enumerate(deduped):
            normalized = normalized_items[index]
            contained_items = [
                other
                for other_index, other in enumerate(normalized_items)
                if other_index != index and other and len(other) >= 4 and other in normalized
            ]
            if len(normalized.split()) >= 4 and len(contained_items) >= 2:
                continue
            filtered.append(item)
        return filtered

    def _best_skill_anchor_paragraph(self, matching_contexts: List[ParagraphContext]):
        preferred = [
            context
            for context in matching_contexts
            if not self._paragraph_hosts_textbox(context.paragraph)
            and not self._paragraph_is_in_table(context.paragraph)
        ]
        if preferred:
            return preferred[-1].paragraph
        return matching_contexts[-1].paragraph if matching_contexts else None

    def _remove_duplicate_skill_paragraphs(self, matching_contexts: List[ParagraphContext], anchor_paragraph) -> None:
        anchor_id = id(getattr(anchor_paragraph, "_p", anchor_paragraph))
        for context in matching_contexts:
            paragraph = context.paragraph
            if id(paragraph._p) == anchor_id:
                continue
            self._replace_paragraph_preserving_style(paragraph, "")

    def _rewrite_section_block(
        self,
        document,
        matching_contexts: List[ParagraphContext],
        replacement_lines: List[str],
        instruction: StructuredRewriteInstruction,
    ) -> None:
        anchor = None
        if instruction.old_text_hint:
            hint_norm = normalize_text(instruction.old_text_hint)
            for context in matching_contexts:
                if hint_norm and hint_norm in normalize_text(context.paragraph.text or ""):
                    anchor = context
                    break
        if anchor is None:
            anchor = matching_contexts[0]

        section_name = anchor.section
        # Replace only the matched anchor. Other existing entries in the section
        # must remain untouched.
        self._replace_paragraph_preserving_style(anchor.paragraph, replacement_lines[0])
        self._style_section_line(anchor.paragraph, section_name, 0)
        previous = anchor.paragraph
        for line_index, extra_line in enumerate(replacement_lines[1:], start=1):
            new_paragraph = self._insert_paragraph_after(previous, extra_line)
            self._copy_paragraph_format(previous, new_paragraph)
            self._style_section_line(new_paragraph, section_name, line_index)
            previous = new_paragraph

    def _rewrite_single_anchor(
        self,
        document,
        matching_contexts: List[ParagraphContext],
        replacement_lines: List[str],
        instruction: StructuredRewriteInstruction,
    ) -> None:
        anchor = None
        if instruction.old_text_hint:
            hint_norm = normalize_text(instruction.old_text_hint)
            for context in matching_contexts:
                if hint_norm and hint_norm in normalize_text(context.paragraph.text or ""):
                    anchor = context
                    break
        if anchor is None:
            anchor = matching_contexts[0]

        section_name = anchor.section
        self._replace_paragraph_preserving_style(anchor.paragraph, replacement_lines[0])
        self._style_section_line(anchor.paragraph, section_name, 0)
        previous = anchor.paragraph
        for line_index, extra_line in enumerate(replacement_lines[1:], start=1):
            new_paragraph = self._insert_paragraph_after(previous, extra_line)
            self._copy_paragraph_format(previous, new_paragraph)
            self._style_section_line(new_paragraph, section_name, line_index)
            previous = new_paragraph

    def _is_high_confidence_section_instruction(self, instruction: StructuredRewriteInstruction) -> bool:
        target = canonical_section(instruction.target_section)
        if target in {"formazione", "esperienze"}:
            return True
        if target in {"hard_skills", "soft_skills"}:
            return False
        return len(tokenize(instruction.old_text_hint)) >= 18 or len((instruction.old_text_hint or "").strip()) >= 220

    def _append_section_block(self, document, heading: str, replacement: str, anchor_paragraph) -> None:
        lines = [line.strip() for line in replacement.splitlines() if line.strip()]
        if not lines:
            lines = [replacement.strip()]

        heading_reference = self._last_heading_paragraph(document)
        body_reference = self._last_styled_paragraph(document)
        if anchor_paragraph is None:
            heading_paragraph = document.add_paragraph()
        else:
            heading_paragraph = self._insert_paragraph_after(anchor_paragraph, "")
        if heading_reference is not None:
            self._copy_paragraph_format(heading_reference, heading_paragraph)
        self._replace_paragraph_preserving_style(heading_paragraph, heading.upper())
        self._make_heading_like(heading_paragraph)

        previous = heading_paragraph
        section_name = canonical_section(heading)
        for line_index, line in enumerate(lines):
            content_paragraph = self._insert_paragraph_after(previous, line, source_format=body_reference or previous)
            self._style_section_line(content_paragraph, section_name, line_index)
            previous = content_paragraph

    def _insert_section_block_before(
        self,
        anchor_paragraph,
        heading: str,
        replacement: str,
    ) -> None:
        from docx.text.paragraph import Paragraph

        lines = [line.strip() for line in replacement.splitlines() if line.strip()]
        elements = []
        for text in [heading.upper(), *lines]:
            new_p = OxmlElement("w:p")
            anchor_paragraph._p.addprevious(new_p)
            paragraph = Paragraph(new_p, anchor_paragraph._parent)
            self._copy_paragraph_format(anchor_paragraph, paragraph)
            self._replace_paragraph_preserving_style(paragraph, text)
            elements.append(paragraph)
        if not elements:
            return
        self._make_heading_like(elements[0])
        section_name = canonical_section(heading)
        for index, paragraph in enumerate(elements[1:]):
            self._style_section_line(paragraph, section_name, index)

    def _style_section_line(self, paragraph, section_name: str, line_index: int) -> None:
        if canonical_section(section_name) == "progetti":
            if line_index % 2 == 0:
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.space_after = Pt(1)
                if paragraph.runs:
                    paragraph.runs[0].bold = True
            elif paragraph.runs:
                paragraph.runs[0].bold = False
                paragraph.paragraph_format.space_after = Pt(6)
        elif paragraph.runs:
            paragraph.runs[0].bold = False

    def _find_best_section_anchor(self, document, target: str, sections_detected: List[str]):
        contexts = self._paragraph_contexts(document)
        target_index = self._section_index(sections_detected, target)
        prefer_body = target not in {"contatti", "lingue", "hard_skills", "soft_skills", "competenze"}
        if target_index > 0:
            previous_section = sections_detected[target_index - 1]
            body_candidates = []
            for context in reversed(contexts):
                if context.section == previous_section and (context.paragraph.text or '').strip():
                    if (
                        not self._paragraph_is_in_textbox(context.paragraph)
                        and not self._paragraph_hosts_textbox(context.paragraph)
                        and (
                            not prefer_body
                            or not self._paragraph_is_in_table(context.paragraph)
                        )
                    ):
                        return context.paragraph
                    body_candidates.append(context.paragraph)
            if body_candidates:
                return body_candidates[0]
        if prefer_body:
            for context in reversed(contexts):
                if (context.paragraph.text or '').strip() and not self._paragraph_is_in_table(context.paragraph):
                    return context.paragraph
        for context in reversed(contexts):
            if (
                (context.paragraph.text or '').strip()
                and not self._paragraph_is_in_textbox(context.paragraph)
                and not self._paragraph_hosts_textbox(context.paragraph)
            ):
                return context.paragraph
        return None

    def _is_safe_text_destination(self, paragraph) -> bool:
        if self._paragraph_hosts_textbox(paragraph):
            return False
        if not self._paragraph_is_in_textbox(paragraph):
            return True
        textbox = self._textbox_ancestor(paragraph)
        if textbox is None:
            return False
        from docx.text.paragraph import Paragraph

        texts = [
            (Paragraph(child, paragraph._parent).text or "").strip()
            for child in textbox
            if child.tag.endswith("}p")
        ]
        return any(is_section_heading(text) for text in texts if text)

    def _textbox_ancestor(self, paragraph):
        marked_textbox = getattr(paragraph, "_careercoach_textbox", None)
        if marked_textbox is not None:
            return marked_textbox
        element = getattr(paragraph, "_p", None)
        if element is not None:
            ancestors = element.xpath("ancestor::w:txbxContent")
            if ancestors:
                return ancestors[0]
        while element is not None:
            if element.tag.endswith("}txbxContent"):
                return element
            element = element.getparent()
        return None

    def _paragraph_is_in_textbox(self, paragraph) -> bool:
        return bool(
            getattr(paragraph, "_careercoach_in_textbox", False)
            or self._textbox_ancestor(paragraph) is not None
        )

    def _paragraph_hosts_textbox(self, paragraph) -> bool:
        element = getattr(paragraph, "_p", None)
        return bool(
            element is not None
            and element.xpath(".//w:txbxContent")
        )

    def _section_index(self, sections_detected: List[str], target: str) -> int:
        if target in sections_detected:
            return sections_detected.index(target)
        if target in {"hard_skills", "soft_skills"} and "competenze" in sections_detected:
            return sections_detected.index("competenze")
        return -1

    def _section_display_name(self, section: str) -> str:
        return {
            "profilo": "Profilo professionale",
            "esperienze": "Esperienze professionali",
            "formazione": "Formazione",
            "competenze": "Competenze",
            "hard_skills": "Hard Skills",
            "soft_skills": "Soft Skills",
            "lingue": "Lingue",
            "certificazioni": "Certificazioni",
            "pubblicazioni": "Pubblicazioni",
            "progetti": "Progetti",
            "contatti": "Contatti",
        }.get(section, section.capitalize())

    def _supplemental_section_heading(self, section: str) -> str:
        section_name = canonical_section(section)
        base_label = self._section_display_name(section_name).strip()
        if not base_label:
            return "Ulteriori informazioni"

        gendered_labels = {
            "Esperienze professionali": "Ulteriori esperienze professionali",
            "Competenze": "Ulteriori competenze",
            "Lingue": "Ulteriori lingue",
            "Certificazioni": "Ulteriori certificazioni",
            "Pubblicazioni": "Ulteriori pubblicazioni",
        }
        if base_label in gendered_labels:
            return gendered_labels[base_label]
        return f"Ulteriori {base_label.lower()}"

    def _make_heading_like(self, paragraph) -> None:
        if not paragraph.runs:
            paragraph.add_run("")
        paragraph.paragraph_format.keep_with_next = True
        if paragraph.paragraph_format.space_before is None:
            paragraph.paragraph_format.space_before = Pt(6)
        if paragraph.paragraph_format.space_after is None:
            paragraph.paragraph_format.space_after = Pt(4)
        first_run = paragraph.runs[0]
        if first_run.bold is None:
            first_run.bold = True

    def _make_subheading_like(self, paragraph) -> None:
        if not paragraph.runs:
            paragraph.add_run("")
        paragraph.paragraph_format.keep_with_next = True
        if paragraph.paragraph_format.space_before is None:
            paragraph.paragraph_format.space_before = Pt(4)
        if paragraph.paragraph_format.space_after is None:
            paragraph.paragraph_format.space_after = Pt(2)
        first_run = paragraph.runs[0]
        first_run.bold = True
        if first_run.italic is None:
            first_run.italic = False

    def _replace_paragraph_preserving_style(self, paragraph, replacement: str) -> None:
        # Gestione hyperlink: se il replacement contiene il testo gia' presente
        # dentro un <w:hyperlink> (es. "linkedin.com/in/..."), lo lasciamo dentro
        # l'hyperlink (per non perdere la cliccabilita') e lo rimuoviamo dal
        # testo che verra' scritto nei run normali. Cosi' evitiamo duplicati
        # senza rompere il link.
        ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        remaining = replacement
        try:
            for hyperlink in paragraph._p.iter(f"{ns_w}hyperlink"):
                t_elements = list(hyperlink.iter(f"{ns_w}t"))
                link_text = "".join((t.text or "") for t in t_elements)
                stripped = link_text.strip()
                if stripped and stripped in remaining:
                    # Rimuovo dalla stringa da scrivere; il link resta cliccabile.
                    remaining = remaining.replace(stripped, "", 1)
                else:
                    # Hyperlink non piu' citato nel nuovo testo: svuotalo per
                    # evitare doppioni residui.
                    for t in t_elements:
                        t.text = ""
            # Normalizza spazi residui (es. "LinkedIn: " senza URL accanto).
            import re as _re
            remaining = _re.sub(r"[ \t]{2,}", " ", remaining).strip()
        except Exception:
            remaining = replacement
        runs = list(paragraph.runs)
        if not runs:
            paragraph.text = remaining
            return
        first_run = runs[0]
        first_run.text = remaining
        for run in runs[1:]:
            run.text = ""

    def _copy_paragraph_format(self, source, target) -> None:
        try:
            target.style = source.style
        except Exception:
            pass
        if source._p.pPr is not None:
            if target._p.pPr is not None:
                target._p.remove(target._p.pPr)
            target._p.insert(0, deepcopy(source._p.pPr))

        try:
            target.paragraph_format.alignment = source.paragraph_format.alignment
            target.paragraph_format.left_indent = source.paragraph_format.left_indent
            target.paragraph_format.right_indent = source.paragraph_format.right_indent
            target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
            target.paragraph_format.space_before = source.paragraph_format.space_before
            target.paragraph_format.space_after = source.paragraph_format.space_after
            target.paragraph_format.line_spacing = source.paragraph_format.line_spacing
            target.paragraph_format.keep_together = source.paragraph_format.keep_together
            target.paragraph_format.keep_with_next = source.paragraph_format.keep_with_next
            target.paragraph_format.widow_control = source.paragraph_format.widow_control
        except Exception:
            pass

        source_run = next((run for run in source.runs if (run.text or "").strip()), None)
        target_run = target.runs[0] if target.runs else target.add_run()
        if source_run is not None and source_run._r.rPr is not None:
            if target_run._r.rPr is not None:
                target_run._r.remove(target_run._r.rPr)
            target_run._r.insert(0, deepcopy(source_run._r.rPr))
        try:
            target.paragraph_format.alignment = source.paragraph_format.alignment
            target.paragraph_format.left_indent = source.paragraph_format.left_indent
            target.paragraph_format.right_indent = source.paragraph_format.right_indent
            target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
            target.paragraph_format.space_before = source.paragraph_format.space_before
            target.paragraph_format.space_after = source.paragraph_format.space_after
            target.paragraph_format.line_spacing = source.paragraph_format.line_spacing
            target.paragraph_format.keep_together = source.paragraph_format.keep_together
            target.paragraph_format.keep_with_next = source.paragraph_format.keep_with_next
            target.paragraph_format.widow_control = source.paragraph_format.widow_control
        except Exception:
            pass
        try:
            target_run.font.name = source_run.font.name if source_run is not None else target_run.font.name
            target_run.font.size = source_run.font.size if source_run is not None else target_run.font.size
            target_run.font.bold = source_run.font.bold if source_run is not None else target_run.font.bold
            target_run.font.italic = source_run.font.italic if source_run is not None else target_run.font.italic
            target_run.font.underline = source_run.font.underline if source_run is not None else target_run.font.underline
        except Exception:
            pass
        try:
            target_run.font.name = source_run.font.name if source_run is not None else target_run.font.name
            target_run.font.size = source_run.font.size if source_run is not None else target_run.font.size
            target_run.font.bold = source_run.font.bold if source_run is not None else target_run.font.bold
            target_run.font.italic = source_run.font.italic if source_run is not None else target_run.font.italic
            target_run.font.underline = source_run.font.underline if source_run is not None else target_run.font.underline
        except Exception:
            pass

    def _duplicate_markers(self, final_text: str) -> List[str]:
        warnings: List[str] = []
        lines = [line.strip() for line in (final_text or "").splitlines() if line.strip()]
        counts: Dict[str, int] = {}
        for line in lines:
            normalized = normalize_text(line)
            counts[normalized] = counts.get(normalized, 0) + 1
        for line, count in counts.items():
            if count > 2 and len(line) > 25:
                warnings.append(line[:120])
        return warnings

    def _duplicate_section_heading_warnings(self, file_bytes: bytes) -> List[str]:
        from docx import Document

        try:
            document = Document(io.BytesIO(file_bytes))
        except Exception:
            return ["Impossibile verificare i titoli delle sezioni."]

        counts: Dict[str, int] = {}
        for context in self._paragraph_contexts(document):
            text = (context.paragraph.text or "").strip()
            if not is_section_heading(text):
                continue
            section = canonical_section(text)
            if section:
                counts[section] = counts.get(section, 0) + 1
        return [
            f"Titolo di sezione duplicato: {section}"
            for section, count in counts.items()
            if count > 1
        ]

    def _clean_corrupted_document_text(self, document) -> None:
        for context in self._paragraph_contexts(document):
            paragraph = context.paragraph
            text = (paragraph.text or "").strip()
            if not text or is_section_heading(text):
                continue
            cleaned = self._remove_obviously_corrupted_tokens(text)
            if cleaned != text:
                self._replace_paragraph_preserving_style(paragraph, cleaned)

    def _remove_obviously_corrupted_tokens(self, text: str) -> str:
        known_corrupted = {"fhaurehds"}
        kept: List[str] = []
        for token in (text or "").split():
            if re.search(r"https?://|www\.|[\w.+-]+@[\w.-]+\.\w+", token, re.IGNORECASE):
                kept.append(token)
                continue
            bare = re.sub(r"^[^\w]+|[^\w]+$", "", token)
            normalized = normalize_text(bare)
            letters = re.sub(r"[^a-z]", "", normalized)
            has_long_consonant_run = bool(
                re.search(r"[bcdfghjklmnpqrstvwxyz]{6,}", letters)
            )
            if normalized in known_corrupted or (
                len(letters) >= 9 and has_long_consonant_run
            ):
                continue
            kept.append(token)
        return re.sub(r"\s+", " ", " ".join(kept)).strip()

    def _unexpected_sections(self, original_text: str, final_text: str) -> List[str]:
        original_sections = {
            canonical_section(section)
            for section in self.parser.section_text_map(original_text).keys()
        }
        final_sections = {
            canonical_section(section)
            for section in self.parser.section_text_map(final_text).keys()
        }
        return sorted(section for section in final_sections - original_sections if section not in {"intestazione", "contenuto"})

    def _insert_paragraph_after(self, paragraph, text: str, source_format=None):
        from docx.text.paragraph import Paragraph

        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_paragraph = Paragraph(new_p, paragraph._parent)
        if source_format is None:
            source_format = paragraph
        if source_format is not None:
            self._copy_paragraph_format(source_format, new_paragraph)
        self._replace_paragraph_preserving_style(new_paragraph, text)
        return new_paragraph

    def _extract_docx_text_bytes(self, file_bytes: bytes) -> str:
        try:
            from docx import Document

            document = Document(io.BytesIO(file_bytes))
            parts = []
            seen_paragraphs = set()
            for context in self._paragraph_contexts(document):
                paragraph = context.paragraph
                paragraph_id = id(paragraph._p)
                if paragraph_id in seen_paragraphs or not paragraph.text:
                    continue
                seen_paragraphs.add(paragraph_id)
                parts.append(paragraph.text)
            for section in getattr(document, "sections", []):
                header = getattr(section, "header", None)
                footer = getattr(section, "footer", None)
                if header:
                    parts.extend(paragraph.text for paragraph in header.paragraphs if paragraph.text)
                    for table in getattr(header, "tables", []):
                        parts.extend(self._collect_table_text(table))
                if footer:
                    parts.extend(paragraph.text for paragraph in footer.paragraphs if paragraph.text)
                    for table in getattr(footer, "tables", []):
                        parts.extend(self._collect_table_text(table))
            seen_cells = set()
            for table in document.tables:
                parts.extend(self._collect_table_text(table, seen_cells))
            return "\n".join(parts)
        except Exception:
            return ""

    def _collect_table_text(self, table, seen_cells: Optional[set] = None) -> List[str]:
        texts: List[str] = []
        seen_cells = seen_cells or set()
        for row in table.rows:
            for cell in row.cells:
                if id(cell._tc) in seen_cells:
                    continue
                seen_cells.add(id(cell._tc))
                texts.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text)
                for nested_table in cell.tables:
                    texts.extend(self._collect_table_text(nested_table, seen_cells))
        return texts

    def _semantic_match(self, normalized_final: str, normalized_target: str) -> bool:
        if not normalized_target:
            return False
        if normalized_target in normalized_final:
            return True
        target_tokens = [token for token in tokenize(normalized_target) if len(token) > 2]
        if len(target_tokens) < 3:
            return False
        final_tokens = set(token for token in tokenize(normalized_final) if len(token) > 2)
        window = " ".join(target_tokens[:6])
        return window in normalized_final or len(set(target_tokens).intersection(final_tokens)) >= max(3, len(target_tokens) // 2)

    def _objective_followed_by_education(self, normalized_final: str) -> bool:
        lines = [line.strip() for line in normalized_final.splitlines() if line.strip()]
        objective_markers = {"obiettivo", "profilo", "profilo professionale", "chi sono", "objective"}
        education_markers = {"formazione", "istruzione", "education"}
        for index, line in enumerate(lines[:-1]):
            if line in objective_markers and lines[index + 1] in education_markers:
                return True
        return False

    def _extract_section_text(self, normalized_text: str, section_names: set[str]) -> str:
        lines = [line.strip() for line in normalized_text.splitlines() if line.strip()]
        captured: List[str] = []
        capture = False
        stop_markers = {"obiettivo", "profilo", "profilo professionale", "chi sono", "objective", "formazione", "istruzione", "education", "esperienze", "esperienze professionali", "progetti", "hard skills", "soft skills", "competenze", "contatti", "lingue", "certificazioni", "pubblicazioni", "publications"}
        for line in lines:
            if line in section_names:
                capture = True
                continue
            if capture and line in stop_markers:
                break
            if capture:
                captured.append(line)
        return " ".join(captured).strip()

    def _paragraph_contexts(self, document) -> List[ParagraphContext]:
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        contexts: List[ParagraphContext] = []
        current_section = "intestazione"
        for child in document._element.body.iterchildren():
            if child.tag.endswith("}p"):
                paragraph = Paragraph(child, document._body)
                text = (paragraph.text or "").strip()
                if is_section_heading(text):
                    current_section = canonical_section(text)
                contexts.append(ParagraphContext(
                    paragraph=paragraph,
                    section=current_section,
                ))
            elif child.tag.endswith("}tbl"):
                table = Table(child, document._body)
                table_contexts = self._table_paragraph_contexts(
                    table,
                    current_section,
                )
                contexts.extend(table_contexts)
                if table_contexts:
                    current_section = table_contexts[-1].section
        contexts.extend(self._textbox_paragraph_contexts(document))
        return contexts

    def _textbox_paragraph_contexts(self, document) -> List[ParagraphContext]:
        from docx.text.paragraph import Paragraph

        contexts: List[ParagraphContext] = []
        seen_signatures = set()
        for textbox in document._element.xpath(".//w:txbxContent"):
            paragraphs = [
                child for child in textbox
                if child.tag.endswith("}p")
            ]
            signature = tuple(
                normalize_text("".join(paragraph.itertext()))
                for paragraph in paragraphs
                if normalize_text("".join(paragraph.itertext()))
            )
            if signature in seen_signatures:
                continue
            seen_signatures.add(signature)
            current_section = "intestazione"
            for paragraph_element in paragraphs:
                paragraph = Paragraph(paragraph_element, document._body)
                paragraph._careercoach_in_textbox = True
                paragraph._careercoach_textbox = textbox
                text = (paragraph.text or "").strip()
                if is_section_heading(text):
                    current_section = canonical_section(text)
                contexts.append(ParagraphContext(
                    paragraph=paragraph,
                    section=current_section,
                ))
        return contexts

    def _textbox_mirror_groups(self, document) -> List[List[Any]]:
        grouped: Dict[tuple, List[Any]] = {}
        for textbox in document._element.xpath(".//w:txbxContent"):
            paragraphs = [
                child for child in textbox
                if child.tag.endswith("}p")
            ]
            signature = tuple(
                normalize_text("".join(paragraph.itertext()))
                for paragraph in paragraphs
                if normalize_text("".join(paragraph.itertext()))
            )
            if signature:
                grouped.setdefault(signature, []).append(textbox)
        return [elements for elements in grouped.values() if len(elements) > 1]

    def _synchronize_textbox_mirrors(self, mirror_groups: List[List[Any]]) -> None:
        for group in mirror_groups:
            source = group[0]
            for mirror in group[1:]:
                for child in list(mirror):
                    mirror.remove(child)
                for child in source:
                    mirror.append(deepcopy(child))

    def _table_paragraph_contexts(self, table, inherited_section: str = "intestazione") -> List[ParagraphContext]:
        contexts: List[ParagraphContext] = []
        current_section = inherited_section or "intestazione"
        seen_tc = set()
        for row in table.rows:
            row_section = current_section
            for cell in row.cells:
                if id(cell._tc) in seen_tc:
                    continue
                seen_tc.add(id(cell._tc))
                cell_has_heading = any(
                    is_section_heading(paragraph.text or "")
                    for paragraph in cell.paragraphs
                    if (paragraph.text or "").strip()
                )
                # A cell with its own headings starts an independent column.
                # Heading-only cells still pass their section to an adjacent body cell.
                cell_section = inherited_section if cell_has_heading else row_section
                cell_text = " ".join(
                    paragraph.text or ""
                    for paragraph in cell.paragraphs
                    if (paragraph.text or "").strip()
                )
                if self._looks_like_contact_block(cell_text):
                    cell_section = "contatti"
                    row_section = cell_section
                    current_section = cell_section
                for paragraph in cell.paragraphs:
                    text = (paragraph.text or "").strip()
                    if is_section_heading(text):
                        cell_section = canonical_section(text)
                        row_section = cell_section
                        current_section = cell_section
                    contexts.append(ParagraphContext(paragraph=paragraph, section=cell_section))
                for nested_table in cell.tables:
                    nested_contexts = self._table_paragraph_contexts(nested_table, cell_section)
                    contexts.extend(nested_contexts)
                    if nested_contexts:
                        cell_section = nested_contexts[-1].section
                        row_section = cell_section
                        current_section = cell_section
        return contexts

    def _looks_like_contact_block(self, text: str) -> bool:
        value = text or ""
        plain = normalize_text(value)
        return bool(
            re.search(r"[\w.+-]+@[\w.-]+\.\w+", value)
            or re.search(r"\+?\d[\d\s().-]{7,}", value)
            or re.search(r"https?://|www\.|linkedin\.com", value, flags=re.IGNORECASE)
            or plain.startswith(("via ", "viale ", "piazza ", "address "))
        )

    def _valid_instruction(self, instruction: RewriteInstruction) -> bool:
        replacement = (instruction.replacement or "").strip()
        if not replacement:
            return False

        target = canonical_section(instruction.section or instruction.category or "")
        is_append = not (instruction.original or "").strip()

        # Profilo/esperienze/formazione sono sezioni narrative: le sostituzioni
        # sicure devono essere consentite, altrimenti il DOCX finale resta quasi
        # invariato. Blocchiamo solo gli append liberi/non tracciati.
        if is_append and target == "profilo":
            self._log_blocked(
                instruction,
                "append vietato nel profilo: serve un testo originale da sostituire",
            )
            return False

        if is_append and target in {"formazione", "esperienze"}:
            source_id = instruction.source_id or ""
            is_user_confirmed = "user_additional_info" in source_id
            if not is_user_confirmed:
                self._log_blocked(
                    instruction,
                    "append narrativo bloccato: serve informazione confermata dall'utente",
                )
                return False

        if not ResumeRewriter().is_safe_replacement(instruction.section, replacement):
            self._log_blocked(instruction, "replacement non valido per la sezione")
            return False
        return True

    def _replace_exact(self, contexts: List[ParagraphContext], instruction: RewriteInstruction) -> bool:
        original_norm = normalize_text(instruction.original)
        for context in contexts:
            paragraph = context.paragraph
            paragraph_text = paragraph.text or ""
            if not self._is_editable_paragraph(paragraph):
                continue
            if normalize_text(paragraph_text) == original_norm or instruction.original in paragraph_text:
                section = context.section
                if not self._section_matches_instruction(section, instruction):
                    self._log_blocked(instruction, f"testo originale trovato fuori sezione: {section or 'sconosciuta'}")
                    continue
                self._replace_and_cleanup(contexts, context, instruction)
                return True
        return False

    def _replace_similar(self, contexts: List[ParagraphContext], instruction: RewriteInstruction) -> bool:
        original_tokens = set(tokenize(instruction.original))
        if len(original_tokens) < 4:
            return False
        best = None
        best_score = 0
        for context in contexts:
            paragraph = context.paragraph
            if not self._is_editable_paragraph(paragraph):
                continue
            current_section = context.section
            if not self._section_matches_instruction(current_section, instruction):
                continue
            paragraph_tokens = set(tokenize(paragraph.text or ""))
            if not paragraph_tokens:
                continue
            score = len(original_tokens.intersection(paragraph_tokens)) / max(len(original_tokens), 1)
            if score > best_score:
                best = context
                best_score = score
        if best is not None and best_score >= 0.62:
            self._replace_and_cleanup(contexts, best, instruction)
            return True
        return False

    def _replace_in_section(self, contexts: List[ParagraphContext], instruction: RewriteInstruction) -> bool:
        section_names = self._section_candidates(instruction)
        if not section_names:
            return False

        start_index = None
        for index, context in enumerate(contexts):
            text = normalize_text(context.paragraph.text or "").strip(":")
            if self._is_heading_for_candidates(text, section_names):
                start_index = index
                break
        if start_index is None:
            self._log_blocked(instruction, "sezione non trovata in modo sicuro")
            return False

        for context in contexts[start_index + 1:start_index + 10]:
            paragraph = context.paragraph
            text = (paragraph.text or "").strip()
            if not text:
                continue
            if is_section_heading(text):
                break
            if not self._section_matches_instruction(context.section, instruction):
                continue
            if not self._is_editable_paragraph(paragraph):
                continue
            self._replace_and_cleanup(contexts, context, instruction)
            return True
        self._log_blocked(instruction, "nessun paragrafo sostituibile trovato sotto la sezione")
        return False

    def _replace_section_block(
        self,
        contexts: List[ParagraphContext],
        instruction: RewriteInstruction,
    ) -> bool:
        expected = canonical_section(instruction.section or instruction.category or "")
        if not expected or expected == "contenuto":
            return False

        candidates = [
            context
            for context in contexts
            if context.section == expected and self._is_editable_paragraph(context.paragraph)
        ]
        if not candidates:
            return False

        original_tokens = set(tokenize(instruction.original))
        if not original_tokens:
            return False

        best_contexts: List[ParagraphContext] = []
        best_score = 0.0
        for start in range(len(candidates)):
            block: List[ParagraphContext] = []
            block_tokens = set()
            for candidate in candidates[start:start + 8]:
                block.append(candidate)
                block_tokens.update(tokenize(candidate.paragraph.text or ""))
                score = len(original_tokens.intersection(block_tokens)) / max(len(original_tokens), 1)
                if score > best_score:
                    best_score = score
                    best_contexts = list(block)
                if score >= 0.9:
                    break

        if not best_contexts or best_score < 0.55:
            return False

        self._replace_paragraph_block(best_contexts, instruction.replacement)
        return True

    def _replace_paragraph_block(
        self,
        contexts: List[ParagraphContext],
        replacement: str,
    ) -> None:
        lines = [line.strip() for line in (replacement or "").splitlines() if line.strip()]
        if not lines:
            lines = [replacement.strip()]

        for index, context in enumerate(contexts):
            value = lines[index] if index < len(lines) else ""
            self._replace_preserving_first_run(context.paragraph, value)

        # Clear any remaining paragraphs when the replacement is shorter than the matched block.
        if len(lines) < len(contexts):
            for context in contexts[len(lines):]:
                self._replace_preserving_first_run(context.paragraph, "")

        if len(lines) <= len(contexts):
            return

        anchor = contexts[-1].paragraph
        format_reference = contexts[-1].paragraph
        for value in lines[len(contexts):]:
            parent = anchor._parent
            new_paragraph = parent.add_paragraph(value)
            new_paragraph._p.getparent().remove(new_paragraph._p)
            anchor._p.addnext(new_paragraph._p)
            self._copy_paragraph_format(format_reference, new_paragraph)
            anchor = new_paragraph

    def _append_to_target_section(self, document, instruction: RewriteInstruction) -> bool:
        if not instruction.replacement.strip():
            return False
        self._append_section(document, RewriteInstruction(
            section=instruction.section,
            original="",
            replacement=instruction.replacement,
            reason=instruction.reason,
            category=instruction.category,
            source_id=instruction.source_id,
        ))
        return True

    def _should_append(self, instruction: RewriteInstruction) -> bool:
        if instruction.original.strip():
            return False
            
        section_name = canonical_section(instruction.section)
        category = normalize_text(instruction.category)
        
        return (
            section_name in {
                "progetti", "esperienze", "formazione", 
                "certificazioni", "competenze", "lingue"
            }
            or category in {
                "project", "extra_page", "skills", "soft_skills",
                "experience", "education", "certification", "hard_skill", "soft_skill"
            }
        )

    def _append_section(self, document, instruction: RewriteInstruction) -> None:
        from docx.enum.text import WD_BREAK

        heading = (instruction.section or "PROGETTI").strip().upper()
        if heading == "PAGINA AGGIUNTIVA":
            heading = "PROGETTI"
        section_name = canonical_section(heading)
        replacement_lines = [line.strip() for line in str(instruction.replacement or "").splitlines() if line.strip()]
        contexts = self._paragraph_contexts(document)
        preferred_order = [
            "profilo",
            "esperienze",
            "competenze",
            "formazione",
            "certificazioni",
            "lingue",
            "progetti",
            "contatti",
        ]
        matching_heading_index = next(
            (
                index
                for index, context in enumerate(contexts)
                if is_section_heading(context.paragraph.text or "")
                and canonical_section(context.paragraph.text or "") == section_name
            ),
            None,
        )

        if matching_heading_index is not None:
            anchor = contexts[matching_heading_index].paragraph
            body_reference = None
            existing_section_parts = []
            for context in contexts[matching_heading_index + 1:]:
                paragraph = context.paragraph
                if is_section_heading(paragraph.text or "") or context.section != section_name:
                    break
                anchor = paragraph
                current_text = (paragraph.text or "").strip()
                if current_text:
                    existing_section_parts.append(current_text)
                    body_reference = paragraph
            existing_section_text = "\n".join(existing_section_parts)
            if normalize_text(instruction.replacement) in normalize_text(existing_section_text):
                return
            insert_subheading = bool(
                existing_section_parts
                and section_name in {"progetti", "esperienze", "formazione", "certificazioni", "lingue", "competenze"}
                and len(replacement_lines) > 1
            )
            if insert_subheading:
                subheading_paragraph = document.add_paragraph(self._supplemental_section_heading(section_name))
                subheading_paragraph._p.getparent().remove(subheading_paragraph._p)
                anchor._p.addnext(subheading_paragraph._p)
                if body_reference is not None:
                    self._copy_paragraph_format(body_reference, subheading_paragraph)
                self._make_subheading_like(subheading_paragraph)
                anchor = subheading_paragraph
            body_paragraph = document.add_paragraph(instruction.replacement)
            body_paragraph._p.getparent().remove(body_paragraph._p)
            anchor._p.addnext(body_paragraph._p)
            if body_reference is not None:
                self._copy_paragraph_format(body_reference, body_paragraph)
            return

        heading_reference = self._last_heading_paragraph(document)
        body_reference = self._last_styled_paragraph(document)
        anchor = body_reference or heading_reference
        insertion_index = len(contexts)
        section_rank = preferred_order.index(section_name) if section_name in preferred_order else len(preferred_order)
        for index, context in enumerate(contexts):
            current_section = canonical_section(context.section or "")
            if current_section == section_name:
                insertion_index = index
                break
            current_rank = preferred_order.index(current_section) if current_section in preferred_order else len(preferred_order)
            if current_rank > section_rank:
                insertion_index = index
                break

        if insertion_index < len(contexts) and insertion_index > 0:
            anchor = contexts[insertion_index - 1].paragraph

        insert_page_break = section_name == "progetti" and len(instruction.replacement.splitlines()) > 2
        if insert_page_break and anchor is not None:
            page_break = document.add_paragraph()
            page_break.add_run().add_break(WD_BREAK.PAGE)
            page_break._p.getparent().remove(page_break._p)
            anchor._p.addnext(page_break._p)
            anchor = page_break

        heading_paragraph = document.add_paragraph(heading)
        body_paragraph = document.add_paragraph(instruction.replacement)
        if heading_reference is not None:
            self._copy_paragraph_format(heading_reference, heading_paragraph)
        if body_reference is not None:
            self._copy_paragraph_format(body_reference, body_paragraph)
        if anchor is not None:
            anchor._p.addnext(heading_paragraph._p)
            heading_paragraph._p.addnext(body_paragraph._p)
        elif heading_reference is not None:
            heading_reference._p.addnext(heading_paragraph._p)
            heading_paragraph._p.addnext(body_paragraph._p)
        # Leave the new section in document flow without forcing a page break.

    def _replace_and_cleanup(
        self,
        contexts: List[ParagraphContext],
        context: ParagraphContext,
        instruction: RewriteInstruction,
    ) -> None:
        self._replace_preserving_first_run(context.paragraph, instruction.replacement)
        if canonical_section(instruction.section) != "profilo":
            return

        replacement_norm = normalize_text(instruction.replacement)
        try:
            start_index = next(
                index for index, candidate in enumerate(contexts)
                if candidate.paragraph is context.paragraph
            )
        except StopIteration:
            return

        replacement_tokens = set(tokenize(instruction.replacement))
        for candidate in contexts[start_index + 1:]:
            if candidate.section != context.section:
                break
            paragraph = candidate.paragraph
            text = (paragraph.text or "").strip()
            if not text:
                continue
            if is_section_heading(text):
                break
            text_norm = normalize_text(text)
            text_tokens = set(tokenize(text))
            overlap = (
                len(text_tokens.intersection(replacement_tokens)) / max(len(text_tokens), 1)
                if text_tokens
                else 0
            )
            if text_norm in replacement_norm or overlap >= 0.85:
                self._replace_preserving_first_run(paragraph, "")
                continue
            break

    def _last_heading_paragraph(self, document):
        for context in reversed(self._paragraph_contexts(document)):
            if is_section_heading(context.paragraph.text or ""):
                return context.paragraph
        return None

    def _last_styled_paragraph(self, document):
        for context in reversed(self._paragraph_contexts(document)):
            paragraph = context.paragraph
            text = (paragraph.text or "").strip()
            if text and not is_section_heading(text):
                return paragraph
        return None

    def _copy_paragraph_format(self, source, target) -> None:
        try:
            target.style = source.style
        except Exception:
            pass
        if source._p.pPr is not None:
            if target._p.pPr is not None:
                target._p.remove(target._p.pPr)
            target._p.insert(0, deepcopy(source._p.pPr))

        source_run = next((run for run in source.runs if (run.text or "").strip()), None)
        target_run = target.runs[0] if target.runs else target.add_run()
        if source_run is not None and source_run._r.rPr is not None:
            if target_run._r.rPr is not None:
                target_run._r.remove(target_run._r.rPr)
            target_run._r.insert(0, deepcopy(source_run._r.rPr))

    def _section_candidates(self, instruction: RewriteInstruction) -> List[str]:
        raw_values = [instruction.section, instruction.category]
        candidates = set()
        for value in raw_values:
            clean = canonical_section(value or "")
            if clean and clean != "contenuto":
                candidates.add(clean)
                candidates.update(SECTION_ALIASES.get(clean, set()))
        if instruction.category == "profile":
            candidates.update(SECTION_ALIASES["profilo"])
        elif instruction.category == "experience":
            candidates.update(SECTION_ALIASES["esperienze"])
        elif instruction.category == "skills":
            candidates.update(SECTION_ALIASES["competenze"])
        elif instruction.category == "education":
            candidates.update(SECTION_ALIASES["formazione"])
        return [normalize_text(candidate) for candidate in candidates if candidate]

    def _is_editable_paragraph(self, paragraph) -> bool:
        text = (paragraph.text or "").strip()
        return bool(text) and not is_section_heading(text)

    def _paragraph_is_in_table(self, paragraph) -> bool:
        parent = getattr(paragraph, "_parent", None)
        while parent is not None:
            if parent.__class__.__name__ == "Table":
                return True
            parent = getattr(parent, "_parent", None)
        return False

    def _is_heading_for_candidates(self, normalized_text: str, section_names: List[str]) -> bool:
        clean = normalize_text(normalized_text).strip(":")
        return clean in section_names

    def _section_matches_instruction(self, current_section: str, instruction: RewriteInstruction) -> bool:
        expected = canonical_section(instruction.section or instruction.category or "")
        if not expected or expected == "contenuto":
            return True
        if current_section == expected:
            return True
        if instruction.category == "profile" and current_section == "profilo":
            return True
        if instruction.category == "experience" and current_section == "esperienze":
            return True
        if instruction.category == "skills" and current_section == "competenze":
            return True
        if instruction.category == "education" and current_section == "formazione":
            return True
        return False

    def _log_blocked(self, instruction: RewriteInstruction, reason: str) -> None:
        print(
            "Inserimento CV bloccato: "
            f"id={instruction.source_id or 'unknown'}, "
            f"section={instruction.section}, category={instruction.category}, "
            f"reason={reason}, proposed_text={instruction.replacement[:220]}"
        )

    def _replace_preserving_first_run(self, paragraph, replacement: str) -> None:
        runs = list(paragraph.runs)
        original_text = paragraph.text
        if not runs:
            paragraph.add_run(replacement)
            return

        runs[0].text = replacement
        
        # Verifica se il testo è stato salvato
        runs_after = list(paragraph.runs)
        for run in runs[1:]:
            run.text = ""


class ExportService:
    DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def docx_to_pdf(self, docx_bytes: bytes) -> Optional[bytes]:
        executable = shutil.which("soffice") or shutil.which("libreoffice")
        if not executable:
            return None
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            docx_path = tmp_path / "cv-ottimizzato.docx"
            docx_path.write_bytes(docx_bytes)
            subprocess.run(
                [executable, "--headless", "--convert-to", "pdf", "--outdir", str(tmp_path), str(docx_path)],
                check=False,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=30,
            )
            pdf_path = tmp_path / "cv-ottimizzato.pdf"
            return pdf_path.read_bytes() if pdf_path.exists() else None

    def plain_pdf(self, text: str) -> tuple[bytes, str, str]:
        try:
            import fitz

            doc = fitz.open()
            page = doc.new_page()
            y = 46
            for paragraph in (text or "").splitlines():
                for line in textwrap.wrap(paragraph.strip(), width=92) or [""]:
                    if y > 790:
                        page = doc.new_page()
                        y = 46
                    page.insert_text((48, y), line, fontsize=10.2, fontname="helv", color=(0.16, 0.20, 0.25))
                    y += 13.5
                y += 5
            pdf_bytes = doc.write()
            doc.close()
            return pdf_bytes, "application/pdf", "pdf"
        except Exception:
            return text.encode("utf-8"), "text/plain; charset=utf-8", "txt"
