from docx import Document
import sys

# Accept optional path argument, otherwise default to backend path
path_arg = sys.argv[1] if len(sys.argv) > 1 else r"C:\Users\silvi\OneDrive\Desktop\CareerCoach\backend\optimized_55.docx"
try:
    doc=Document(path_arg)
    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
    seen_cells = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if id(cell._tc) not in seen_cells:
                    seen_cells.add(id(cell._tc))
                    parts.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text)

    full = "\n".join(parts)
    if not full.strip():
        print('DOCX vuoto o non leggibile')
    else:
        keywords = ['Linguaggi', 'Python', 'SQL', 'Java', 'COMPETENZE', 'HARD SKILLS', 'SOFT SKILLS', 'GAME']
        found = {k: [] for k in keywords}
        for line in full.splitlines():
            for k in keywords:
                if k.lower() in line.lower():
                    found[k].append(line.strip())
        for k,v in found.items():
            print(f"{k}: {len(v)}")
            for item in v[:3]:
                print('  -', item)
        print('\n--- Text excerpt (first 4000 chars) ---')
        print(full[:4000])
except Exception as e:
    print('ERROR',e)
