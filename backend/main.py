import os
import json
import re
import unicodedata
import sqlite3
import requests
import secrets
import hashlib
import hmac
import smtplib
import base64
import binascii
import urllib.parse
import io
import ipaddress
import socket
import textwrap
import zipfile
import xml.etree.ElementTree as ET
from difflib import SequenceMatcher
from datetime import datetime, timedelta
from email.message import EmailMessage
from functools import lru_cache
from html.parser import HTMLParser
from pathlib import Path
from typing import Optional, List, Dict, Any
from urllib3.exceptions import MaxRetryError, NewConnectionError
from typing import Iterable, List

from dotenv import dotenv_values, load_dotenv
from fastapi import FastAPI, HTTPException, Header, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import RedirectResponse, Response
from openai import OpenAI
from pydantic import BaseModel
from starlette.concurrency import run_in_threadpool
from services.cv_optimizer import (
    CoachSuggestionEngine,
    DocxApplyResult,
    DocxPreserver,
    ExportService,
    JobAnalyzer,
    MatchingEngine,
    ResumeDocxOptimizationPipeline,
    ResumeParser,
    ResumeRewriter,
    RewriteInstruction,
    StructuredRewriteInstruction,
)
from services.cv_optimizer.section_catalog import (
    SECTION_ALIASES as SHARED_CV_SECTION_ALIASES,
    additional_field_section_key,
    canonical_section_key,
    normalize_section_title,
)
from services.cv_image_safety import validate_cv_images


# =========================
# CONFIGURAZIONE AMBIENTE
# =========================

BACKEND_DIR = os.path.dirname(os.path.abspath(__file__))
#ENV_FILE_PATH = os.path.join(BACKEND_DIR, ".env")
#load_dotenv(ENV_FILE_PATH, override=False)
#ENV_FILE_VALUES = dotenv_values(ENV_FILE_PATH)

# DOPO
ENV_FILE_PATH = os.path.join(BACKEND_DIR, ".env")
load_dotenv(ENV_FILE_PATH, override=False)
ENV_FILE_VALUES = os.environ  # ← usa le env vars di sistema/Render

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
#GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.1-8b-instant")
GROQ_MODEL = os.getenv("GROQ_MODEL", "openai/gpt-oss-20b")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "").strip()
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-3.1-flash-lite").strip()
GEMINI_API_URL = os.getenv(
    "GEMINI_API_URL",
    "https://generativelanguage.googleapis.com/v1beta",
).rstrip("/")
LLM_PROVIDER = ENV_FILE_VALUES.get("LLM_PROVIDER", "groq").strip().lower()
OLLAMA_TEXT_MODEL = ENV_FILE_VALUES.get("OLLAMA_TEXT_MODEL", "gemma3:4b")
OLLAMA_TEXT_TIMEOUT = int(ENV_FILE_VALUES.get("OLLAMA_TEXT_TIMEOUT", "120"))
OLLAMA_TEXT_NUM_CTX = int(ENV_FILE_VALUES.get("OLLAMA_TEXT_NUM_CTX", "8192"))
OPENAI_API_KEY = ENV_FILE_VALUES.get("OPENAI_API_KEY")
VISION_PROVIDER = ENV_FILE_VALUES.get("VISION_PROVIDER", "ollama").strip().lower()
OLLAMA_URL = ENV_FILE_VALUES.get("OLLAMA_URL", "http://127.0.0.1:11434").rstrip("/")
OLLAMA_VISION_MODEL = ENV_FILE_VALUES.get("OLLAMA_VISION_MODEL", "moondream")
OLLAMA_OCR_MODEL = ENV_FILE_VALUES.get("OLLAMA_OCR_MODEL", "qwen2.5vl:3b")
OPENAI_VISION_MODEL = ENV_FILE_VALUES.get("OPENAI_VISION_MODEL", "gpt-4o-mini")

TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "").strip() or None
FRONTEND_URL = os.getenv("FRONTEND_URL", "http://127.0.0.1:5173")
SMTP_HOST = os.getenv("SMTP_HOST")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD")
SMTP_FROM = os.getenv("SMTP_FROM", SMTP_USER or "noreply@careercoach.local")
SESSION_DAYS = int(os.getenv("SESSION_DAYS", "30"))
DEBUG_MODE = os.getenv("DEBUG", "false").strip().lower() in {"1", "true", "yes", "on"}

# CV AI policy
# Default: the CV pipeline is deterministic/local, inspired by Resume Matcher/OpenResume.
# Set these flags to true only if you explicitly want to use an external/local LLM
# for the corresponding CV step. This prevents Groq/Ollama timeouts from blocking CV analysis.
# When the chosen provider is Ollama local, enable the CV LLM flow by default unless
# the user explicitly disabled it.
def _env_bool(name: str, default: bool = False) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return str(raw).strip().lower() in {"1", "true", "yes", "on"}

CV_LLM_ENABLED = _env_bool("CV_LLM_ENABLED", default=True)
CV_EVALUATION_LLM_ENABLED = _env_bool("CV_EVALUATION_LLM_ENABLED", default=False)
CV_REWRITE_LLM_ENABLED = _env_bool("CV_REWRITE_LLM_ENABLED", default=True)
CV_QUALITY_LLM_ENABLED = _env_bool("CV_QUALITY_LLM_ENABLED", default=False)
OAUTH_REDIRECT_BASE_URL = os.getenv("OAUTH_REDIRECT_BASE_URL", "http://localhost:8000")
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID")
GOOGLE_CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET")
GOOGLE_REDIRECT_URI = os.getenv("GOOGLE_REDIRECT_URI")
LINKEDIN_CLIENT_ID = os.getenv("LINKEDIN_CLIENT_ID")
LINKEDIN_CLIENT_SECRET = os.getenv("LINKEDIN_CLIENT_SECRET")
LINKEDIN_REDIRECT_URI = os.getenv("LINKEDIN_REDIRECT_URI")

if not GROQ_API_KEY:
    raise RuntimeError("Manca GROQ_API_KEY nel file .env")

if not TAVILY_API_KEY:
    print("Avviso: TAVILY_API_KEY non configurata. Le ricerche web saranno disabilitate.")


groq_client = OpenAI(
    api_key=GROQ_API_KEY,
    base_url="https://api.groq.com/openai/v1",
    timeout=30.0
)

openai_moderation_client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
openai_visual_rate_limited_until = None


class GroqRateLimitError(Exception):
    """Errore specifico per limite di rate limit Groq."""
    pass


groq_rate_limited_until: Optional[datetime] = None


def groq_is_temporarily_blocked() -> bool:
    return groq_rate_limited_until is not None and datetime.utcnow() < groq_rate_limited_until


def mark_groq_rate_limit(error_text: str) -> None:
    global groq_rate_limited_until
    wait_seconds = 30 * 60
    minutes_match = re.search(r"try again in (\d+)m([0-9.]+)s", error_text or "", re.IGNORECASE)
    seconds_match = re.search(r"try again in ([0-9.]+)s", error_text or "", re.IGNORECASE)
    if minutes_match:
        wait_seconds = int(minutes_match.group(1)) * 60 + int(float(minutes_match.group(2))) + 5
    elif seconds_match:
        wait_seconds = int(float(seconds_match.group(1))) + 5
    groq_rate_limited_until = datetime.utcnow() + timedelta(seconds=max(wait_seconds, 60))
    print(f"Groq disattivato temporaneamente per rate limit fino a {groq_rate_limited_until.isoformat()} UTC: uso fallback locale.")


# =========================
# APP FASTAPI
# =========================

app = FastAPI(title="CareerCoach API - Interview Gym")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# =========================
# DATABASE SQLITE
# =========================

DB_NAME = str(Path(__file__).resolve().parent / "careercoach.db")


def get_connection():
    return sqlite3.connect(DB_NAME)


def add_column_if_not_exists(cursor, table_name: str, column_name: str, column_sql: str):
    cursor.execute(f"PRAGMA table_info({table_name})")
    existing_columns = [row[1] for row in cursor.fetchall()]

    if column_name not in existing_columns:
        cursor.execute(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {column_sql}")


def init_db():
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        email TEXT,
        education TEXT,
        target_role TEXT,
        sector TEXT,
        experience_level TEXT,
        interview_language TEXT
    )
    """)

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS interview_sessions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        interview_type TEXT,
        difficulty TEXT,
        company TEXT,
        question_mode TEXT,
        total_score INTEGER,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(user_id) REFERENCES users(id)
    )
    """)

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS questions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        session_id INTEGER NOT NULL,
        question_text TEXT NOT NULL,
        category TEXT,
        difficulty TEXT,
        company TEXT,
        question_mode TEXT,
        sources_json TEXT,
        FOREIGN KEY(session_id) REFERENCES interview_sessions(id)
    )
    """)

    cursor.execute("PRAGMA table_info(interview_sessions)")
    existing_columns = [row[1] for row in cursor.fetchall()]
    if "role" not in existing_columns:
        cursor.execute("ALTER TABLE interview_sessions ADD COLUMN role TEXT")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS answers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        question_id INTEGER NOT NULL,
        user_answer TEXT NOT NULL,
        clarity_score INTEGER,
        completeness_score INTEGER,
        relevance_score INTEGER,
        professionalism_score INTEGER,
        synthesis_score INTEGER,
        speaking_score INTEGER,
        total_score INTEGER,
        feedback TEXT,
        improved_answer TEXT,
        speaking_feedback TEXT,
        solution_explanation TEXT,
        speech_metrics_json TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(question_id) REFERENCES questions(id)
    )
    """)

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS web_sources (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT,
        url TEXT UNIQUE,
        content TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    """)

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS question_web_sources (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        question_id INTEGER NOT NULL,
        source_id INTEGER NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(question_id) REFERENCES questions(id),
        FOREIGN KEY(source_id) REFERENCES web_sources(id)
    )
    """)

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS user_sessions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        token TEXT NOT NULL UNIQUE,
        expires_at TIMESTAMP NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(user_id) REFERENCES users(id)
    )
    """)

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS oauth_states (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        provider TEXT NOT NULL,
        state TEXT NOT NULL UNIQUE,
        expires_at TIMESTAMP NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    """)

    add_column_if_not_exists(cursor, "users", "password_hash", "TEXT")
    add_column_if_not_exists(cursor, "users", "phone", "TEXT")
    add_column_if_not_exists(cursor, "users", "email_verified", "INTEGER DEFAULT 0")
    add_column_if_not_exists(cursor, "users", "email_verification_token", "TEXT")
    add_column_if_not_exists(cursor, "users", "email_verification_expires_at", "TIMESTAMP")
    add_column_if_not_exists(cursor, "users", "password_reset_token", "TEXT")
    add_column_if_not_exists(cursor, "users", "password_reset_expires_at", "TIMESTAMP")
    add_column_if_not_exists(cursor, "users", "created_at", "TIMESTAMP")
    add_column_if_not_exists(cursor, "users", "last_login_at", "TIMESTAMP")
    add_column_if_not_exists(cursor, "users", "auth_provider", "TEXT")
    add_column_if_not_exists(cursor, "users", "provider_user_id", "TEXT")
    add_column_if_not_exists(cursor, "users", "cv_filename", "TEXT")
    add_column_if_not_exists(cursor, "users", "cv_content_type", "TEXT")
    add_column_if_not_exists(cursor, "users", "cv_size", "INTEGER")
    add_column_if_not_exists(cursor, "users", "cv_text", "TEXT")
    add_column_if_not_exists(cursor, "users", "cv_file_base64", "TEXT")
    add_column_if_not_exists(cursor, "users", "cv_uploaded_at", "TIMESTAMP")
    add_column_if_not_exists(cursor, "users", "optimized_cv_filename", "TEXT")
    add_column_if_not_exists(cursor, "users", "optimized_cv_content_type", "TEXT")
    add_column_if_not_exists(cursor, "users", "optimized_cv_text", "TEXT")
    add_column_if_not_exists(cursor, "users", "optimized_cv_file_base64", "TEXT")
    add_column_if_not_exists(cursor, "users", "optimized_cv_generated_at", "TIMESTAMP")
    add_column_if_not_exists(cursor, "users", "linkedin_url", "TEXT")
    add_column_if_not_exists(cursor, "users", "portfolio_url", "TEXT")
    add_column_if_not_exists(cursor, "users", "instagram_handle", "TEXT")
    add_column_if_not_exists(cursor, "users", "digital_analysis_json", "TEXT")
    add_column_if_not_exists(cursor, "users", "linkedin_profile_filename", "TEXT")
    add_column_if_not_exists(cursor, "users", "linkedin_profile_text", "TEXT")
    add_column_if_not_exists(cursor, "users", "linkedin_oauth_profile_json", "TEXT")
    add_column_if_not_exists(cursor, "users", "profile_image_data_url", "TEXT")

    add_column_if_not_exists(cursor, "interview_sessions", "company", "TEXT")
    add_column_if_not_exists(cursor, "interview_sessions", "question_mode", "TEXT")

    add_column_if_not_exists(cursor, "questions", "company", "TEXT")
    add_column_if_not_exists(cursor, "questions", "question_mode", "TEXT")
    add_column_if_not_exists(cursor, "questions", "sources_json", "TEXT")

    add_column_if_not_exists(cursor, "answers", "speaking_score", "INTEGER")
    add_column_if_not_exists(cursor, "answers", "speaking_feedback", "TEXT")
    add_column_if_not_exists(cursor, "answers", "solution_explanation", "TEXT")
    add_column_if_not_exists(cursor, "answers", "speech_metrics_json", "TEXT")

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS optimized_cvs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        filename TEXT NOT NULL,
        content_type TEXT,
        text TEXT,
        file_base64 TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(user_id) REFERENCES users(id)
    )
    """)
    add_column_if_not_exists(cursor, "optimized_cvs", "target_role", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "target_company", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "job_description", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "analysis_score", "INTEGER")
    add_column_if_not_exists(cursor, "optimized_cvs", "selected_suggestions_json", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "rejected_suggestions_json", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "additional_info_json", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "confirmed_skills_json", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "generation_status", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "applied_changes_count", "INTEGER")
    add_column_if_not_exists(cursor, "optimized_cvs", "skipped_changes_json", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "docx_filename", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "docx_content_type", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "docx_file_base64", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "pdf_filename", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "pdf_content_type", "TEXT")
    add_column_if_not_exists(cursor, "optimized_cvs", "pdf_file_base64", "TEXT")

    conn.commit()
    conn.close()


init_db()


# =========================
# MODELLI PYDANTIC
# =========================

class UserCreate(BaseModel):
    name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    education: str
    target_role: str
    sector: str
    experience_level: str
    interview_language: str = "Italiano"


class UserUpdate(BaseModel):
    name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    education: str
    target_role: str
    sector: str
    experience_level: str
    interview_language: str = "Italiano"


class UserCvUpload(BaseModel):
    filename: str
    content_type: Optional[str] = None
    size: int
    text: Optional[str] = None
    file_base64: Optional[str] = None


class DigitalPresenceUpdate(BaseModel):
    linkedin_url: Optional[str] = None
    portfolio_url: Optional[str] = None
    instagram_handle: Optional[str] = None
    target_role: Optional[str] = None
    linkedin_connected: bool = False


class ProfileImageUpdate(BaseModel):
    image_data_url: str


class CvOptimizationAnalysisRequest(BaseModel):
    company: Optional[str] = None
    role: Optional[str] = None
    role_level: Optional[str] = None
    goal: Optional[str] = None
    job_link: Optional[str] = None
    original_cv_text: Optional[str] = None
    cv_fingerprint: Optional[str] = None
    job_data: Optional[Dict[str, Any]] = None
    cv_evaluation: Optional[Any] = None
    strategic_analysis: Optional[Any] = None
    recommended_adaptations: Optional[Any] = None
    acceptedSuggestionIds: Optional[List[str]] = None
    selected_suggestion_ids: Optional[List[str]] = None
    rejected_suggestion_ids: Optional[List[str]] = None
    accepted_suggestions: Optional[Any] = None
    rejected_suggestions: Optional[Any] = None
    user_additional_data: Optional[Dict[str, Any]] = None
    additionalInfo: Optional[Any] = None
    answers: Optional[Any] = None
    extraAnswers: Optional[Any] = None
    confirmedSkills: Optional[Any] = None
    confirmedKeywords: Optional[Any] = None
    acceptedSkillConfirmations: Optional[Any] = None
    acceptedKeywordConfirmations: Optional[Any] = None
    rejectedSkillConfirmations: Optional[Any] = None
    rejectedKeywordConfirmations: Optional[Any] = None


class JobValidationRequest(BaseModel):
    description: Optional[str] = None
    company: Optional[str] = None
    role: Optional[str] = None
    role_level: Optional[str] = None
    link: Optional[str] = None
    sector: Optional[str] = None
    required_skills: Optional[str] = None


class RegisterRequest(BaseModel):
    name: str
    email: str
    password: str
    phone: Optional[str] = None


class LoginRequest(BaseModel):
    identifier: str
    password: str


class TokenRequest(BaseModel):
    token: str


class ForgotPasswordRequest(BaseModel):
    identifier: str


class ResetPasswordRequest(BaseModel):
    token: str
    password: str


class GenerateQuestionRequest(BaseModel):
    user_id: int
    interview_type: str
    difficulty: str = "intermedio"
    company: Optional[str] = "Azienda Generica"
    goal: Optional[str] = ""
    role: Optional[str] = ""
    job_link: Optional[str] = ""
    question_mode: Optional[str] = "web"


class SpeechMetrics(BaseModel):
    duration_seconds: Optional[float] = None
    words_count: Optional[int] = None
    words_per_minute: Optional[float] = None
    filler_words_count: Optional[int] = None
    filler_words: Optional[List[str]] = None


class EvaluateAnswerRequest(BaseModel):
    question_id: int
    answer: str
    speech_metrics: Optional[SpeechMetrics] = None


class SearchInterviewQuestionsRequest(BaseModel):
    company: str
    role: str
    interview_type: str
    language: str = "Italiano"


# =========================
# FUNZIONI AUTH / EMAIL
# =========================

EMAIL_PATTERN = re.compile(r"^[A-Za-z0-9.!#$%&'*+/=?^_`{|}~-]+@[A-Za-z0-9-]+(?:\.[A-Za-z0-9-]+)+$")
BLOCKED_EMAIL_DOMAINS = {
    "example.com",
    "example.it",
    "test.com",
    "test.it",
    "email.com",
    "mailinator.com",
    "tempmail.com",
    "10minutemail.com",
}


def utc_now() -> datetime:
    return datetime.utcnow()


def normalize_email(email: str) -> str:
    return email.strip().lower()


def normalize_phone(phone: Optional[str]) -> Optional[str]:
    if not phone:
        return None

    cleaned = re.sub(r"[^\d+]", "", phone.strip())
    if cleaned.startswith("00"):
        cleaned = "+" + cleaned[2:]
    return cleaned or None


def validate_email_address(email: str) -> str:
    normalized = normalize_email(email)
    domain = normalized.split("@")[-1] if "@" in normalized else ""
    tld = domain.rsplit(".", 1)[-1] if "." in domain else ""

    if (
        not EMAIL_PATTERN.match(normalized)
        or len(tld) < 2
        or domain in BLOCKED_EMAIL_DOMAINS
        or ".." in normalized
    ):
        raise HTTPException(status_code=400, detail="Inserisci un indirizzo email reale e valido.")

    return normalized


def validate_password(password: str):
    if len(password) < 8:
        raise HTTPException(status_code=400, detail="La password deve avere almeno 8 caratteri.")

    if not re.search(r"[A-Za-z]", password) or not re.search(r"\d", password):
        raise HTTPException(status_code=400, detail="La password deve includere almeno una lettera e un numero.")


def validate_phone(phone: Optional[str]) -> Optional[str]:
    normalized = normalize_phone(phone)
    if not normalized:
        return None

    digits = re.sub(r"\D", "", normalized)
    if len(digits) < 8 or len(digits) > 15:
        raise HTTPException(status_code=400, detail="Inserisci un numero di cellulare valido.")

    return normalized


def hash_password(password: str) -> str:
    salt = secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), 120000)
    return f"pbkdf2_sha256$120000${salt}${digest.hex()}"


def verify_password(password: str, stored_hash: Optional[str]) -> bool:
    if not stored_hash:
        return False

    try:
        algorithm, iterations, salt, expected = stored_hash.split("$", 3)
        if algorithm != "pbkdf2_sha256":
            return False
        digest = hashlib.pbkdf2_hmac(
            "sha256",
            password.encode("utf-8"),
            salt.encode("utf-8"),
            int(iterations),
        ).hex()
        return hmac.compare_digest(digest, expected)
    except ValueError:
        return False


def make_token() -> str:
    return secrets.token_urlsafe(32)


def make_frontend_link(kind: str, token: str) -> str:
    return f"{FRONTEND_URL}/?{kind}={token}"


def normalize_oauth_provider(provider: str) -> str:
    return "google" if provider == "gmail" else provider.strip().lower()


def make_oauth_callback_url(provider: str) -> str:
    provider = normalize_oauth_provider(provider)
    provider_redirects = {
        "google": GOOGLE_REDIRECT_URI,
        "linkedin": LINKEDIN_REDIRECT_URI,
    }

    return provider_redirects.get(provider) or f"{OAUTH_REDIRECT_BASE_URL}/auth/oauth/{provider}/callback"


def make_frontend_oauth_link(token: str) -> str:
    return f"{FRONTEND_URL}/?oauth_token={token}"


def decode_jwt_payload(token: str) -> Dict:
    try:
        payload = token.split(".")[1]
        payload += "=" * (-len(payload) % 4)
        return json.loads(base64.urlsafe_b64decode(payload.encode("utf-8")))
    except Exception:
        return {}


def send_email(to_email: str, subject: str, body: str) -> bool:
    if not SMTP_HOST or not SMTP_USER or not SMTP_PASSWORD:
        return False

    try:
        message = EmailMessage()
        message["Subject"] = subject
        message["From"] = SMTP_FROM
        message["To"] = to_email
        message.set_content(body)

        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=15) as smtp:
            smtp.starttls()
            smtp.login(SMTP_USER, SMTP_PASSWORD)
            smtp.send_message(message)

        return True
    except Exception as exc:
        print(f"Invio email non riuscito: {exc}")
        return False


def user_to_response(row):
    digital_analysis_json = row[19] if len(row) > 19 else None
    linkedin_profile_filename = row[20] if len(row) > 20 else None
    linkedin_profile_text = row[21] if len(row) > 21 else None
    linkedin_oauth_profile_json = row[22] if len(row) > 22 else None
    auth_provider = row[23] if len(row) > 23 else None
    profile_image_data_url = row[24] if len(row) > 24 else None
    return {
        "id": row[0],
        "name": row[1],
        "email": row[2],
        "education": row[3],
        "target_role": row[4],
        "sector": row[5],
        "experience_level": row[6],
        "interview_language": row[7],
        "phone": row[8],
        "email_verified": bool(row[9]),
        "cv_filename": row[10],
        "cv_uploaded": bool(row[10]),
        "cv_text": row[13] or "",
        "cv_uploaded_at": row[15],
        "linkedin_url": row[16],
        "portfolio_url": row[17],
        "instagram_handle": row[18],
        "digital_analysis": json.loads(digital_analysis_json) if digital_analysis_json else None,
        "linkedin_profile_filename": linkedin_profile_filename,
        "linkedin_profile_uploaded": bool(linkedin_profile_filename),
        "linkedin_oauth_profile": json.loads(linkedin_oauth_profile_json) if linkedin_oauth_profile_json else None,
        "auth_provider": auth_provider,
        "profile_image_data_url": profile_image_data_url,
    }


def create_session(cursor, user_id: int) -> str:
    token = make_token()
    expires_at = utc_now() + timedelta(days=SESSION_DAYS)
    cursor.execute("""
    INSERT INTO user_sessions (user_id, token, expires_at)
    VALUES (?, ?, ?)
    """, (user_id, token, expires_at.isoformat()))
    cursor.execute("UPDATE users SET last_login_at = CURRENT_TIMESTAMP WHERE id = ?", (user_id,))
    return token


def fetch_user_by_id(cursor, user_id: int):
    cursor.execute("""
    SELECT
        id,
        name,
        email,
        education,
        target_role,
        sector,
        experience_level,
        interview_language,
        phone,
        email_verified,
        cv_filename,
        cv_content_type,
        cv_size,
        cv_text,
        cv_file_base64,
        cv_uploaded_at,
        linkedin_url,
        portfolio_url,
        instagram_handle,
        digital_analysis_json,
        linkedin_profile_filename,
        linkedin_profile_text,
        linkedin_oauth_profile_json,
        auth_provider,
        profile_image_data_url
    FROM users
    WHERE id = ?
    """, (user_id,))
    return cursor.fetchone()


def get_oauth_config(provider: str) -> Dict:
    provider = normalize_oauth_provider(provider)

    configs = {
        "google": {
            "client_id": GOOGLE_CLIENT_ID,
            "client_secret": GOOGLE_CLIENT_SECRET,
            "redirect_uri": make_oauth_callback_url("google"),
            "auth_url": "https://accounts.google.com/o/oauth2/v2/auth",
            "token_url": "https://oauth2.googleapis.com/token",
            "userinfo_url": "https://openidconnect.googleapis.com/v1/userinfo",
            "scope": "openid email profile",
        },
    
        "linkedin": {
            "client_id": LINKEDIN_CLIENT_ID,
            "client_secret": LINKEDIN_CLIENT_SECRET,
            "redirect_uri": make_oauth_callback_url("linkedin"),
            "auth_url": "https://www.linkedin.com/oauth/v2/authorization",
            "token_url": "https://www.linkedin.com/oauth/v2/accessToken",
            "userinfo_url": "https://api.linkedin.com/v2/userinfo",
            "scope": "openid profile email",
            "setup_hint": (
                "In LinkedIn Developers abilita il prodotto 'Sign In with LinkedIn using OpenID Connect' "
                "e registra esattamente il redirect URI configurato nel backend."
            ),
        },
    }

    if provider not in configs:
        raise HTTPException(status_code=404, detail="Provider non supportato.")

    config = configs[provider]
    if not config["client_id"] or not config["client_secret"]:
        raise HTTPException(
            status_code=400,
            detail=f"Configura {provider.upper()}_CLIENT_ID e {provider.upper()}_CLIENT_SECRET nel file .env.",
        )

    return {"provider": provider, **config}


def create_oauth_state(cursor, provider: str) -> str:
    state = make_token()
    expires_at = utc_now() + timedelta(minutes=10)
    cursor.execute("""
    INSERT INTO oauth_states (provider, state, expires_at)
    VALUES (?, ?, ?)
    """, (provider, state, expires_at.isoformat()))
    return state


def build_oauth_authorization_url(provider: str, state: str) -> str:
    config = get_oauth_config(provider)
    params = {
        "client_id": config["client_id"],
        "redirect_uri": config["redirect_uri"],
        "response_type": "code",
        "scope": config["scope"],
        "state": state,
    }

    return f"{config['auth_url']}?{urllib.parse.urlencode(params)}"


def consume_oauth_state(cursor, provider: str, state: str):
    cursor.execute("""
    SELECT expires_at
    FROM oauth_states
    WHERE provider = ? AND state = ?
    """, (provider, state))
    row = cursor.fetchone()

    if not row:
        raise HTTPException(status_code=400, detail="Sessione OAuth non valida.")

    cursor.execute("DELETE FROM oauth_states WHERE state = ?", (state,))

    if datetime.fromisoformat(row[0]) < utc_now():
        raise HTTPException(status_code=400, detail="Sessione OAuth scaduta.")


def fetch_oauth_profile(provider: str, code: str) -> Dict:
    config = get_oauth_config(provider)
    provider = config["provider"]
    token_response = requests.post(
        config["token_url"],
        data={
            "grant_type": "authorization_code",
            "code": code,
            "redirect_uri": config["redirect_uri"],
            "client_id": config["client_id"],
            "client_secret": config["client_secret"],
        },
        headers={"Accept": "application/json"},
        timeout=15,
    )

    if token_response.status_code >= 400:
        provider_hint = config.get("setup_hint", "Controlla client id, client secret e redirect URI del provider.")
        raise HTTPException(
            status_code=400,
            detail=(
                f"Accesso {provider} non riuscito durante lo scambio del codice. "
                f"{provider_hint} Dettaglio provider: {token_response.text[:300]}"
            ),
        )

    token_data = token_response.json()
    accesstoken = token_data.get("access_token")
    if not accesstoken:
        raise HTTPException(status_code=400, detail=f"Token OAuth non ricevuto dal provider.")

# Recupera il profilo utente    
    userinforesponse = requests.get(
        config["userinfo_url"],
        headers={"Authorization": f"Bearer {accesstoken}"},
        timeout=15,
)   
    if userinforesponse.status_code != 200:
        raise HTTPException(status_code=400, detail=f"Impossibile recuperare il profilo da {provider}.")

    profile = userinforesponse.json()

    email = profile.get("email")
    if not email:
        raise HTTPException(status_code=400, detail="Il provider non ha restituito un indirizzo email.")

    name = " ".join(part for part in [profile.get("given_name"), profile.get("family_name")] if part and part.strip()) or email.split("@")[0]

    return {
        "provider_user_id": str(profile.get("sub") or profile.get("id") or email),
        "email": validate_email_address(email),
        "name": name,
        "oauth_profile": {
            "name": name,
            "picture": profile.get("picture"),
            "locale": profile.get("locale"),
            "email_verified": bool(profile.get("email_verified")),
    },
}

def find_or_create_oauth_user(cursor, provider: str, profile: Dict) -> int:
    oauth_profile_json = json.dumps(profile.get("oauth_profile") or {}, ensure_ascii=False)
    cursor.execute("""
    SELECT id
    FROM users
    WHERE auth_provider = ? AND provider_user_id = ?
    """, (provider, profile["provider_user_id"]))
    provider_row = cursor.fetchone()
    if provider_row:
        if provider == "linkedin":
            cursor.execute("""
            UPDATE users
            SET linkedin_oauth_profile_json = ?
            WHERE id = ?
            """, (oauth_profile_json, provider_row[0]))
        return provider_row[0]

    cursor.execute("SELECT id FROM users WHERE lower(email) = lower(?)", (profile["email"],))
    email_row = cursor.fetchone()
    if email_row:
        cursor.execute("""
        UPDATE users
        SET auth_provider = ?,
            provider_user_id = ?,
            email_verified = 1,
            linkedin_oauth_profile_json = CASE WHEN ? = 'linkedin' THEN ? ELSE linkedin_oauth_profile_json END
        WHERE id = ?
        """, (provider, profile["provider_user_id"], provider, oauth_profile_json, email_row[0]))
        return email_row[0]

    cursor.execute("""
    INSERT INTO users (
        name,
        email,
        email_verified,
        auth_provider,
        provider_user_id,
        linkedin_oauth_profile_json,
        education,
        target_role,
        sector,
        experience_level,
        interview_language
    )
    VALUES (?, ?, 1, ?, ?, ?, '', '', '', 'Junior', 'Italiano')
    """, (
        profile["name"],
        profile["email"],
        provider,
        profile["provider_user_id"],
        oauth_profile_json if provider == "linkedin" else None,
    ))
    return cursor.lastrowid


# =========================
# FUNZIONI AI / WEB SEARCH
# =========================

def _debug_print(*args: object, **kwargs: object) -> None:
    if DEBUG_MODE:
        print(*args, **kwargs)


def call_ollama(
    prompt: str,
    temperature: float = 0.7,
    max_tokens: int = 1000,
    timeout: Optional[int] = None,
    json_mode: bool = False,
) -> str:
    timeout = OLLAMA_TEXT_TIMEOUT if timeout is None else max(1, timeout)
    timeout = min(timeout, 120)
    print(f"Ollama timeout impostato a {timeout}s, model={OLLAMA_TEXT_MODEL}")
    _debug_print(f"Ollama debug: model={OLLAMA_TEXT_MODEL}, max_tokens={max_tokens}, timeout={timeout}")
    _debug_print(f"Ollama debug prompt:\n{prompt[:500]}...")
    payload: Dict[str, Any] = {
        "model": OLLAMA_TEXT_MODEL,
        "stream": False,
        "messages": [
            {
                "role": "system",
                "content": (
                    "Sei un career coach e resume editor esperto in CV, candidature, "
                    "compatibilità ATS stimata e preparazione ai colloqui. "
                    "Lavori esclusivamente sui dati forniti nella richiesta corrente, "
                    "senza ricordare o riutilizzare persone, ruoli o contenuti precedenti."
                ),
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    if json_mode:
        payload["format"] = "json"

    try:
        print(f"Ollama richiesta POST a {OLLAMA_URL}/api/chat in avvio (timeout={timeout}s)...")
        import time
        start = time.time()
        response = requests.post(
            f"{OLLAMA_URL}/api/chat",
            json=payload,
            timeout=timeout,
        )
        elapsed = time.time() - start
        print(f"Ollama richiesta POST completata in {elapsed:.1f}s, status={response.status_code}")
        
        if not response.ok:
            try:
                error_message = response.json().get("error", response.text)
            except ValueError:
                error_message = response.text
            raise RuntimeError(f"Ollama HTTP {response.status_code}: {error_message}")

        content = response.json().get("message", {}).get("content", "")
        print(f"Ollama risposta ricevuta: {len(content)} char, tempo totale={elapsed:.1f}s")
        return (content or "").strip()
    except requests.exceptions.Timeout as exc:
        print(f"Ollama TIMEOUT dopo {timeout}s: {exc}")
        raise RuntimeError(f"Ollama timeout dopo {timeout} secondi - modello troppo lento")
    except requests.exceptions.ConnectionError as exc:
        print(f"Ollama CONNECTION ERROR: {exc}")
        raise RuntimeError(f"Ollama non disponibile: {OLLAMA_URL}")
    except Exception as exc:
        print(f"Errore Ollama locale: {exc}")
        raise RuntimeError(str(exc))


def call_gemini(
    prompt: str,
    temperature: float = 0.1,
    max_tokens: int = 1100,
    timeout: int = 60,
    json_mode: bool = True,
) -> str:
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY non configurata.")

    generation_config: Dict[str, Any] = {
        "temperature": temperature,
        "maxOutputTokens": max_tokens,
    }
    if json_mode:
        generation_config["responseMimeType"] = "application/json"

    payload = {
        "systemInstruction": {
            "parts": [{
                "text": (
                    "Sei un resume editor senior. Riscrivi in italiano naturale e professionale, "
                    "usando esclusivamente i fatti forniti. Non inventare competenze, risultati, "
                    "aziende, ruoli, date o responsabilita."
                ),
            }],
        },
        "contents": [{
            "role": "user",
            "parts": [{"text": prompt}],
        }],
        "generationConfig": generation_config,
    }
    url = f"{GEMINI_API_URL}/models/{urllib.parse.quote(GEMINI_MODEL, safe='')}:generateContent"
    try:
        print(f"Gemini richiesta avviata: model={GEMINI_MODEL}, timeout={timeout}s")
        response = requests.post(
            url,
            params={"key": GEMINI_API_KEY},
            json=payload,
            timeout=max(1, min(timeout, 120)),
        )
        if not response.ok:
            try:
                error_payload = response.json()
                error_message = error_payload.get("error", {}).get("message") or response.text
            except ValueError:
                error_message = response.text
            raise RuntimeError(f"Gemini HTTP {response.status_code}: {error_message}")

        data = response.json()
        candidates = data.get("candidates") if isinstance(data, dict) else []
        if not isinstance(candidates, list) or not candidates:
            raise RuntimeError("Gemini non ha restituito candidati.")
        parts = candidates[0].get("content", {}).get("parts", [])
        text = "".join(
            str(part.get("text") or "")
            for part in parts
            if isinstance(part, dict)
        ).strip()
        if not text:
            raise RuntimeError("Gemini ha restituito una risposta vuota.")
        print(f"Gemini risposta ricevuta: {len(text)} char")
        return text
    except requests.exceptions.Timeout as exc:
        raise RuntimeError(f"Gemini timeout dopo {timeout} secondi") from exc
    except requests.exceptions.ConnectionError as exc:
        raise RuntimeError("Gemini non raggiungibile.") from exc


def _call_groq_impl(
    prompt: str,
    temperature: float = 0.7,
    max_tokens: int = 1000,
    timeout: int = 25,
    json_mode: bool = False,
) -> str:
    if groq_is_temporarily_blocked():
        raise GroqRateLimitError("Groq temporaneamente disattivato per rate limit: uso fallback locale.")

    def _completion_kwargs(include_json_mode: bool) -> Dict[str, Any]:
        kwargs: Dict[str, Any] = {
            "model": GROQ_MODEL,
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "Sei un career coach e resume editor esperto in CV, candidature, "
                        "compatibilità ATS stimata e preparazione ai colloqui. "
                        "Lavori esclusivamente sui dati forniti nella richiesta corrente, "
                        "senza ricordare o riutilizzare persone, ruoli o contenuti precedenti."
                    )
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": temperature,
            "max_tokens": max_tokens,
            "timeout": timeout,
        }
        if include_json_mode:
            kwargs["response_format"] = {"type": "json_object"}
        return kwargs

    try:
        print("Chiamata Groq avviata...")
        try:
            response = groq_client.chat.completions.create(**_completion_kwargs(json_mode))
        except Exception as first_exc:
            first_error = str(first_exc).lower()
            unsupported_json_mode = (
                json_mode
                and "response_format" in first_error
                and any(marker in first_error for marker in ["unsupported", "not supported", "invalid", "unknown"])
            )
            if not unsupported_json_mode:
                raise
            print("Groq non supporta response_format json_object per questo modello: ritento senza json_mode.")
            response = groq_client.chat.completions.create(**_completion_kwargs(False))

        print("Chiamata Groq completata.")
        return (response.choices[0].message.content or "").strip()

    except Exception as e:
        error_text = str(e)
        if (
            "429" in error_text
            or "rate_limit" in error_text.lower()
            or "too many requests" in error_text.lower()
            or "rate_limit_exceeded" in error_text.lower()
        ):
            print("Errore Groq: rate limit rilevato, uso fallback locale deterministico.")
            mark_groq_rate_limit(error_text)
            raise GroqRateLimitError(error_text)
        if "invalid api key" in error_text.lower() or "invalid_api_key" in error_text.lower():
            print("Errore Groq: chiave API non valida. Aggiorna GROQ_API_KEY nel file .env e riavvia il backend.")
            detail = "Chiave GroqCloud non valida. Aggiorna GROQ_API_KEY nel file .env e riavvia il backend."
        else:
            print(f"Errore Groq: {e}")
            detail = f"Errore durante la chiamata a GroqCloud: {str(e)}"
        raise HTTPException(
            status_code=500,
            detail=detail
        )


def call_groq(
    prompt: str,
    temperature: float = 0.7,
    max_tokens: int = 1000,
    timeout: int = 25,
    json_mode: bool = False,
) -> str:
    if LLM_PROVIDER == "ollama":
        return call_ollama(prompt, temperature, max_tokens, timeout, json_mode)

    if LLM_PROVIDER == "auto":
        try:
            return _call_groq_impl(prompt, temperature, max_tokens, timeout, json_mode)
        except GroqRateLimitError:
            print("Groq non disponibile, uso Ollama locale come fallback.")
            return call_ollama(prompt, temperature, max_tokens, timeout, json_mode)

    return _call_groq_impl(prompt, temperature, max_tokens, timeout, json_mode)


def call_structured_llm(
    prompt: str,
    context: str = "",
    temperature: float = 0.2,
    max_tokens: int = 1200,
    timeout: int = 25,
    preferred_order: Optional[List[str]] = None,
) -> Dict[str, Any]:
    """
    Structured JSON call with local-first fallback.
    Groq can reject strict JSON on some prompts; this helper retries with Ollama
    and always routes the result through extract_json for normalization.
    """
    last_error: Optional[Exception] = None
    order = [
        item
        for item in (preferred_order or ["groq", "ollama"])
        if item in {"gemini", "groq", "ollama"}
    ]
    attempts = []
    for source in order:
        if source == "gemini":
            attempts.append((
                "gemini",
                lambda: call_gemini(
                    prompt,
                    temperature=temperature,
                    max_tokens=max_tokens,
                    timeout=timeout,
                    json_mode=True,
                ),
            ))
        elif source == "groq":
            attempts.append(("groq-json", lambda: call_groq(prompt, temperature=temperature, max_tokens=max_tokens, timeout=timeout, json_mode=True)))
            attempts.append(("groq-plain", lambda: call_groq(prompt, temperature=temperature, max_tokens=max_tokens, timeout=timeout, json_mode=False)))
        elif source == "ollama":
            attempts.append(("ollama", lambda: call_ollama(prompt, temperature=temperature, max_tokens=max_tokens, timeout=timeout, json_mode=True)))
    for source, runner in attempts:
        try:
            return extract_json(runner(), context=context or source)
        except Exception as exc:
            last_error = exc
            print(f"Structured LLM fallback from {source} for {context or 'generic'}: {exc}")
    raise RuntimeError(f"Nessun modello strutturato disponibile per {context or 'generic'}: {last_error}")


def call_analysis_llm(
    prompt: str,
    context: str = "",
    temperature: float = 0.15,
    max_tokens: int = 1100,
    timeout: int = 45,
) -> Dict[str, Any]:
    return call_structured_llm(
        prompt,
        context=context,
        temperature=temperature,
        max_tokens=max_tokens,
        timeout=timeout,
        preferred_order=["gemini"],
    )


def call_lightweight_analysis_llm(
    prompt: str,
    context: str = "",
    temperature: float = 0.1,
    max_tokens: int = 700,
    timeout: int = 35,
) -> Dict[str, Any]:
    return call_structured_llm(
        prompt,
        context=context,
        temperature=temperature,
        max_tokens=max_tokens,
        timeout=timeout,
        preferred_order=["gemini"],
    )


def build_cv_analysis_prompt(
    company: str,
    role: str,
    description: str,
    required_skills: str,
    link: str,
    sources_prompt: str,
    ats_analysis: Dict[str, Any],
    cv_text: str,
    compact: bool = False,
) -> str:
    sources_excerpt = sources_prompt[:500] if compact else sources_prompt[:900]
    ats_excerpt = (
        json.dumps(ats_analysis, ensure_ascii=False, separators=(",", ":"))[:450]
        if compact
        else json.dumps(ats_analysis, ensure_ascii=False, indent=2)[:900]
    )
    cv_excerpt = cv_text[:900] if compact else cv_text[:1400]
    if compact:
        return f"""
Sei un recruiter senior. Restituisci SOLO JSON valido.

Valuta il CV per il ruolo {role or "non specificato"} presso {company or "azienda non specificata"}.

Descrizione: {description[:450]}
Competenze richieste: {required_skills or "Non specificate"}
Link: {link or "Non inserito"}

ATS:
{ats_excerpt}

CV:
{cv_excerpt}

Schema:
{{
  "overall_score": 0,
  "ats_score": 0,
  "job_match_score": 0,
  "role_match_score": 0,
  "company_fit_score": 0,
  "clarity_score": 0,
  "completeness_score": 0,
  "strengths": ["..."],
  "weaknesses": ["..."],
  "present_keywords": ["..."],
  "missing_keywords": ["..."],
  "sections_to_improve": [{{"section":"...","suggestion":"..."}}],
  "questions_for_user": [{{"question":"...","reason":"...","category":"..."}}],
  "summary": "..."
}}

Regole:
- Italiano.
- Solo JSON.
- Frasi brevi.
"""
    return f"""
Sei un recruiter senior. Restituisci SOLO JSON valido.

Valuta il CV rispetto alla candidatura.

Contesto:
- Azienda: {company}
- Ruolo: {role}
- Descrizione: {description}
- Competenze richieste: {required_skills or "Non specificate"}
- Link: {link or "Non inserito"}

Fonti:
{sources_excerpt}

ATS preliminare:
{ats_excerpt}

CV:
{cv_excerpt}

Schema minimo:
{{
  "overall_score": 0,
  "ats_score": 0,
  "job_match_score": 0,
  "role_match_score": 0,
  "company_fit_score": 0,
  "clarity_score": 0,
  "completeness_score": 0,
  "strengths": ["..."],
  "weaknesses": ["..."],
  "present_keywords": ["..."],
  "missing_keywords": ["..."],
  "sections_to_improve": [{{"section": "...", "suggestion": "..."}}],
  "questions_for_user": [{{"question": "...", "reason": "...", "category": "..."}}],
  "summary": "..."
}}

Regole:
- Italiano, frasi brevi, niente testo fuori dal JSON.
- Usa punteggi conservativi.
"""

def build_cv_rewrite_prompt(
    cv_text: str,
    company: str,
    role: str,
    goal: str,
    job_link: str,
    sources: List[Dict[str, str]],
    cv_evaluation: Optional[Any] = None,
    strategic_analysis: Optional[Any] = None,
    recommended_adaptations: Optional[Any] = None,
    accepted_coach_suggestions: Optional[List[Dict[str, Any]]] = None,
    clean_additional_data: Optional[Dict[str, Any]] = None,
) -> str:
    compact_suggestions = []
    for item in (accepted_coach_suggestions or [])[:8]:
        if not isinstance(item, dict):
            continue
        compact_suggestions.append({
            "section": str(item.get("section") or item.get("category") or "").strip(),
            "title": str(item.get("title") or item.get("message") or "").strip()[:90],
            "proposed_text": str(item.get("proposed_text") or item.get("replacement") or "").strip()[:280],
            "keywords_added": [str(x).strip() for x in (item.get("keywords_added") or [])[:5] if str(x).strip()],
        })
    return f"""
Sei un resume editor. Restituisci SOLO JSON valido.

Obiettivo: proponi modifiche puntuali al CV per il ruolo indicato usando solo contenuti presenti.

Contesto:
- Azienda: {company or "Non specificata"}
- Ruolo: {role or "Non specificato"}
- Obiettivo: {goal or "Non specificato"}

Modifiche accettate:
{json.dumps(compact_suggestions, ensure_ascii=False)[:1100]}

Dati confermati:
{json.dumps(clean_additional_data or {}, ensure_ascii=False)[:700]}

CV:
{cv_text[:2200]}

Schema:
{{
  "instructions": [
    {{
      "id": "identificativo breve",
      "section": "nome della sezione esistente",
      "original": "testo esatto da sostituire",
      "replacement": "testo professionale riscritto",
      "reason": "motivazione breve",
      "category": "categoria"
    }}
  ]
}}

Regole:
- Mantieni tono naturale e professionale.
- Integra le modifiche senza perderle.
- Genera al massimo 8 istruzioni brevi.
- Usa original copiato esattamente dal CV; se devi aggiungere una voce usa original vuoto.
- Non restituire il CV completo.
- Non inventare fatti, competenze, risultati, aziende o date.
- Non aggiungere spiegazioni.
"""


def call_rewrite_llm(
    prompt: str,
    context: str = "",
    temperature: float = 0.1,
    max_tokens: int = 1100,
    timeout: int = 60,
) -> Dict[str, Any]:
    return call_structured_llm(
        prompt,
        context=context,
        temperature=temperature,
        max_tokens=max_tokens,
        timeout=timeout,
        preferred_order=["gemini"],
    )


def extract_json(text: str, context: str = ""):
    raw_text = text or ""
    if not raw_text.strip():
        raise ValueError(f"JSON vuoto restituito dal modello{f' in {context}' if context else ''}.")

    text = raw_text.strip().replace("﻿", "").replace(" ", " ")
    fenced = re.search(r"```(?:json)?\s*(.*?)```", text, flags=re.IGNORECASE | re.DOTALL)
    if fenced:
        text = fenced.group(1).strip()

    def _cleanup(candidate: str) -> str:
        candidate = candidate.strip()
        candidate = candidate.replace("“", '"').replace("”", '"').replace("’", "'")
        candidate = re.sub(r",\s*([}\]])", r"\1", candidate)
        return candidate

    def _loads(candidate: str):
        return json.loads(_cleanup(candidate))

    try:
        return _loads(text)
    except json.JSONDecodeError:
        pass

    decoder = json.JSONDecoder()
    for index, char in enumerate(text):
        if char not in "[{":
            continue
        try:
            value, _end = decoder.raw_decode(_cleanup(text[index:]))
            return value
        except json.JSONDecodeError:
            continue

    candidates = []
    for opener, closer in (("{", "}"), ("[", "]")):
        start = text.find(opener)
        end = text.rfind(closer)
        if start != -1 and end != -1 and end > start:
            candidates.append(text[start:end + 1])
    for candidate in candidates:
        try:
            return _loads(candidate)
        except json.JSONDecodeError:
            continue

    preview = re.sub(r"\s+", " ", raw_text)[:500]
    print(
        "JSON non valido restituito dal modello"
        + (f" in {context}" if context else "")
        + f": len={len(raw_text)}, preview={preview!r}"
    )
    raise ValueError(f"JSON non valido restituito dal modello{f' in {context}' if context else ''}: {preview}")


CV_SECTION_KEYWORDS = {
    "contatti": [
        "contatti", "telefono", "email", "e-mail", "linkedin", "indirizzo",
        "numero", "cellulare", "mobile", "phone"
    ],
    "profilo personale": [
        "profilo personale", "profilo professionale", "personal profile",
        "chi sono", "summary", "about me", "presentazione", "obiettivo",
        "career objective"
    ],
    "esperienze professionali": [
        "esperienze professionali", "esperienza professionale", "esperienza lavorativa", "esperienza",
        "esperienze", "experience", "work experience", "employment", "lavoro",
        "azienda", "tirocinio", "stage", "internship", "ruolo", "position"
    ],
    "formazione": [
        "formazione", "formazione accademica", "istruzione", "studi", "percorso di studi", "education",
        "universita", "università", "university", "laurea", "laurea magistrale",
        "laurea triennale", "diploma", "liceo", "master", "degree"
    ],
    "competenze": [
        "competenze", "competenze tecniche", "skills", "technical skills",
        "hard skills", "soft skills", "capacita", "capacità", "tecnologie",
        "linguaggi", "software", "tools", "python", "sql", "java", "c++",
        "machine learning", " ai ", " ml "
    ],
    "lingue": ["lingue", "languages", "italiano", "inglese", "francese", "spagnolo"],
    "certificazioni": [
        "certificazioni", "certifications", "certificate", "attestati",
        "abilitazioni"
    ],
    "progetti": ["progetti", "projects", "project work", "tesi", "tirocinio"],
    "linkedin": ["linkedin", "linkedin.com"],
    "github": ["github", "github.com"],
    "portfolio": ["portfolio", "website", "sito personale", "behance", "dribbble"],
}


def decode_text_bytes(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def clean_extracted_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def cv_content_fingerprint(text: str) -> str:
    normalized = clean_extracted_text(text)
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def cv_analysis_target_fingerprint(
    role: str,
    company: str = "",
    description: str = "",
    required_skills: str = "",
) -> str:
    normalized_company = normalize_plain_text(company)
    if normalized_company in {"generica", "azienda generica", "non specificata"}:
        normalized_company = ""
    normalized = "|".join([
        normalize_plain_text(role),
        normalized_company,
        normalize_plain_text(description),
        normalize_plain_text(required_skills),
    ])
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def normalize_for_cv_detection(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text.lower())
    normalized = "".join(char for char in normalized if not unicodedata.combining(char))
    return f" {normalized} "


def extract_pdf_with_pymupdf(file_bytes: bytes) -> str:
    text_parts = []

    try:
        import fitz

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            page_text = page.get_text("text")
            if page_text:
                text_parts.append(page_text)
        doc.close()
        return "\n".join(text_parts).strip()
    except Exception as exc:
        print("Errore PyMuPDF:", exc)
        return ""


def extract_pdf_with_pypdf(file_bytes: bytes) -> str:
    text_parts = []

    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts).strip()
    except Exception as exc:
        print("Errore pypdf:", exc)
        return ""


def extract_text_from_file_bytes(file_bytes: bytes, filename: str) -> tuple[str, str]:
    lower_filename = filename.lower()

    if lower_filename.endswith(".pdf"):
        text = extract_pdf_with_pymupdf(file_bytes)
        if len(text.strip()) >= 50:
            return text, "pymupdf"

        text = extract_pdf_with_pypdf(file_bytes)
        if len(text.strip()) >= 50:
            return text, "pypdf"

        return text, "pdf_failed_or_scanned"

    if lower_filename.endswith(".docx"):
        try:
            from docx import Document

            document = Document(io.BytesIO(file_bytes))
            text_parts = []

            def append_paragraphs(paragraphs):
                for paragraph in paragraphs:
                    if paragraph.text:
                        text_parts.append(paragraph.text)

            append_paragraphs(document.paragraphs)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)

            for section in document.sections:
                append_paragraphs(section.header.paragraphs)
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:
                                text_parts.append(cell.text)
                append_paragraphs(section.footer.paragraphs)
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:
                                text_parts.append(cell.text)

            docx_text = "\n".join(text_parts).strip()
            xml_text = _extract_docx_text_from_xml(file_bytes)
            
            docx_alnum = len(re.sub(r'\s+', '', docx_text))
            xml_alnum = len(re.sub(r'\s+', '', xml_text))
            
            if xml_alnum > docx_alnum:
                return xml_text, "docx_xml"
            return docx_text or xml_text, "docx"
        except Exception as exc:
            print("Errore DOCX:", exc)
            xml_text = _extract_docx_text_from_xml(file_bytes)
            return xml_text, "docx_xml"

    return "", "failed"


def _extract_text_from_word_ml(xml_bytes: bytes) -> str:
    try:
        root = ET.fromstring(xml_bytes)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs = []
        for paragraph in root.findall(".//w:p", ns):
            text_segments = [node.text for node in paragraph.findall('.//w:t', ns) if node.text]
            if text_segments:
                paragraphs.append("".join(text_segments))
        return "\n".join(paragraphs).strip()
    except Exception:
        return ""


def _extract_docx_text_from_xml(file_bytes: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
            content_files = sorted(
                name
                for name in archive.namelist()
                if name == "word/document.xml" or name.startswith("word/header") or name.startswith("word/footer")
            )
            text_parts = []
            for name in content_files:
                xml_text = _extract_text_from_word_ml(archive.read(name))
                if xml_text:
                    text_parts.append(xml_text)
            return "\n".join(text_parts).strip()
    except Exception as exc:
        print("Errore DOCX XML parsing:", exc)
        return ""


def extract_text_with_method(file_bytes: bytes, filename: str) -> Dict:
    text, method = extract_text_from_file_bytes(file_bytes, filename)
    return {"text": text, "method": method, "errors": []}


def extract_text_from_cv_upload(filename: str, content: bytes) -> str:
    text, _method = extract_text_from_file_bytes(content, filename)
    return text


def recover_saved_cv_text(cursor, user_row) -> str:
    existing_text = clean_extracted_text(user_row[13] or "")
    if existing_text:
        return existing_text

    filename = user_row[10] or ""
    file_base64 = user_row[14] or ""
    if not filename or not file_base64:
        return ""

    try:
        file_bytes = base64.b64decode(file_base64, validate=True)
        extracted_text, _method = extract_text_from_file_bytes(file_bytes, filename)
        normalized_text = clean_extracted_text(extracted_text)[:20000]
    except Exception as exc:
        print(f"Recupero testo CV salvato non riuscito: {exc}")
        return ""

    if normalized_text:
        cursor.execute("""
        UPDATE users
        SET cv_text = ?
        WHERE id = ?
        """, (normalized_text, user_row[0]))

    return normalized_text


def analyze_cv_heuristics(text: str) -> Dict:
    normalized = normalize_for_cv_detection(text)
    detected_sections = []
    signals = []

    has_email = bool(re.search(r"[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}", normalized))
    has_phone = bool(re.search(r"(?:\+?\d[\s().-]*){8,}", normalized))
    has_name_like_line = any(
        2 <= len(re.findall(r"\b[A-ZÀ-ÖØ-Ý][A-Za-zÀ-ÖØ-öø-ÿ'`-]{1,}\b", line)) <= 5
        for line in text.splitlines()[:12]
    )

    if has_email:
        detected_sections.append("email")
        detected_sections.append("contatti")
        signals.append("email")
        signals.append("contatti")
    if has_phone or "numero" in normalized:
        detected_sections.append("telefono")
        detected_sections.append("contatti")
        signals.append("telefono")
        signals.append("contatti")
    if has_name_like_line:
        signals.append("nome e cognome")

    for section, keywords in CV_SECTION_KEYWORDS.items():
        if any(keyword in normalized for keyword in keywords):
            detected_sections.append(section)
            signals.append(section)

    unique_sections = sorted(set(detected_sections))
    unique_signals = sorted(set(signals))

    score = 0
    score += 15 if "email" in unique_signals else 0
    score += 15 if "telefono" in unique_signals else 0
    score += 10 if "linkedin" in unique_signals else 0
    score += 10 if "nome e cognome" in unique_signals else 0
    score += 20 if "esperienze professionali" in unique_signals else 0
    score += 20 if "formazione" in unique_signals else 0
    score += 20 if "competenze" in unique_signals else 0
    score += 10 if "lingue" in unique_signals else 0

    optional_sections = {"certificazioni", "progetti", "github", "portfolio"}
    score += 10 if optional_sections.intersection(unique_signals) else 0
    score += 10 if "profilo personale" in unique_signals else 0
    score = min(score, 100)

    reason = "Il documento e leggibile, ma non contiene abbastanza elementi tipici di un curriculum."
    if score >= 45:
        reason = (
            "Il documento contiene elementi tipici di un CV: "
            f"{', '.join(unique_sections[:8])}."
        )
    elif 35 <= score <= 49 and ("email" in unique_signals or "telefono" in unique_signals):
        reason = (
            "Il documento contiene alcuni elementi tipici di un CV e almeno un contatto: "
            f"{', '.join(unique_sections[:8])}."
        )

    return {
        "score": score,
        "confidence": score,
        "detected_sections": unique_sections,
        "signals": unique_signals,
        "reason": reason,
    }


def classify_cv_with_llm(text: str) -> Optional[Dict]:
    prompt = f"""
Devi stabilire se il documento seguente e un curriculum vitae.

Classifica il documento come:
- CV valido
- probabilmente CV
- non CV

Restituisci SOLO JSON valido con questa struttura:
{{
  "is_cv": true,
  "confidence": 0,
  "reason": "spiegazione breve",
  "detected_sections": []
}}

Testo estratto dal documento:
{text[:6000]}
"""

    try:
        raw_output = call_groq(prompt, temperature=0.1, max_tokens=500)
        result = extract_json(raw_output)
        return {
            "is_cv": bool(result.get("is_cv")),
            "confidence": clamp_score(result.get("confidence", 0)),
            "reason": str(result.get("reason", "")).strip(),
            "detected_sections": result.get("detected_sections", []),
        }
    except Exception as exc:
        print(f"Classificazione LLM CV non disponibile: {exc}")
        return None


def extract_questions_list(text: str) -> List[str]:
    """
    Estrae una lista pulita di domande dal testo restituito dal modello.
    Gestisce:
    - JSON corretto {"questions": [...]}
    - lista JSON [...]
    - testo numerato
    - risposte con frase introduttiva
    """

    text = text.strip()

    if text.startswith("```json"):
        text = text.replace("```json", "").replace("```", "").strip()
    elif text.startswith("```"):
        text = text.replace("```", "").strip()

    questions = []

    try:
        data = json.loads(text)

        if isinstance(data, dict) and "questions" in data:
            questions = data["questions"]
        elif isinstance(data, list):
            questions = data

    except Exception:
        questions = []

    if not questions:
        lines = text.split("\n")

        for line in lines:
            line = line.strip()

            if not line:
                continue

            cleaned = line.lstrip("-• ")
            cleaned = cleaned.strip()

            parts = cleaned.split(maxsplit=1)
            if parts:
                first = parts[0].replace(".", "").replace(")", "")
                if first.isdigit() and len(parts) > 1:
                    cleaned = parts[1].strip()

            lower = cleaned.lower()

            intro_phrases = [
                "ecco le 10 domande",
                "ecco dieci domande",
                "di seguito",
                "queste sono",
                "certamente",
                "ecco una lista",
                "ecco le domande"
            ]

            if any(phrase in lower for phrase in intro_phrases):
                continue

            if "?" in cleaned and len(cleaned) > 15:
                questions.append(cleaned)

    clean_questions = []

    for question in questions:
        if not isinstance(question, str):
            continue

        q = question.strip()
        lower_q = q.lower()

        if any(phrase in lower_q for phrase in [
            "ecco le 10 domande",
            "ecco dieci domande",
            "di seguito",
            "queste sono",
            "ecco una lista",
            "ecco le domande"
        ]):
            continue

        if len(q) > 15 and "?" in q:
            clean_questions.append(q)

    return clean_questions[:10]


def clamp_score(value):
    try:
        value = int(value)
    except Exception:
        value = 0

    return max(0, min(100, value))


def compute_total_score(
    clarity_score,
    completeness_score,
    relevance_score,
    professionalism_score,
    synthesis_score,
    speaking_score
):
    scores = [
        clarity_score,
        completeness_score,
        relevance_score,
        professionalism_score,
        synthesis_score
    ]

    if speaking_score > 0:
        scores.append(speaking_score)

    return round(sum(scores) / len(scores))


def compute_weighted_cv_job_score(payload: Dict[str, Any]) -> int:
    components = [
        (clamp_score(payload.get("role_match_score", 0)), 0.30),
        (clamp_score(payload.get("completeness_score", 0)), 0.20),
        (clamp_score(payload.get("ats_score", 0)), 0.17),
        (clamp_score(payload.get("keyword_score", payload.get("ats_score", 0))), 0.13),
        (clamp_score(payload.get("format_score", 0)), 0.08),
        (clamp_score(payload.get("clarity_score", 0)), 0.06),
        (clamp_score(payload.get("professionalism_score", 0)), 0.06),
    ]
    if payload.get("company_provided"):
        components.append((clamp_score(payload.get("company_fit_score", 0)), 0.10))

    total_weight = sum(weight for _score, weight in components)
    if total_weight <= 0:
        return 0
    return clamp_score(round(
        sum(score * weight for score, weight in components) / total_weight
    ))


def build_cv_score_explanation(payload: Dict[str, Any]) -> Dict[str, Any]:
    role_match = clamp_score(payload.get("role_match_score", 0))
    company_fit = clamp_score(payload.get("company_fit_score", 0))
    completeness = clamp_score(payload.get("completeness_score", 0))
    ats_score = clamp_score(payload.get("ats_score", 0))
    format_score = clamp_score(payload.get("format_score", 0))
    clarity = clamp_score(payload.get("clarity_score", 0))
    professionalism = clamp_score(payload.get("professionalism_score", 0))
    keyword_score = clamp_score(payload.get("keyword_score", ats_score))
    ats_quality = round((ats_score * 0.7) + (format_score * 0.3))
    overall = clamp_score(payload.get("overall_score", compute_weighted_cv_job_score(payload)))

    strongest = sorted(
        [
            ("Coerenza col ruolo", role_match),
            ("Struttura e completezza", completeness),
            ("Compatibilità ATS", ats_score),
            ("Copertura competenze", keyword_score),
            ("Formato", format_score),
            *(([("Adattamento all'azienda", company_fit)]) if payload.get("company_provided") else []),
            ("Chiarezza", clarity),
            ("Professionalità", professionalism),
        ],
        key=lambda item: item[1],
        reverse=True,
    )
    weakest = sorted(
        [
            ("Coerenza col ruolo", role_match),
            ("Struttura e completezza", completeness),
            ("Compatibilità ATS", ats_score),
            ("Copertura competenze", keyword_score),
            ("Formato", format_score),
            *(([("Adattamento all'azienda", company_fit)]) if payload.get("company_provided") else []),
            ("Chiarezza", clarity),
            ("Professionalità", professionalism),
        ],
        key=lambda item: item[1],
    )

    summary_parts = []
    if strongest:
        top_label, top_value = strongest[0]
        summary_parts.append(f"punto più forte: {top_label.lower()} ({top_value}/100)")
    if weakest:
        low_label, low_value = weakest[0]
        summary_parts.append(f"area più debole: {low_label.lower()} ({low_value}/100)")

    return {
        "overall_score": overall,
        "weighted_components": {
            "role_match_score": role_match,
            "company_fit_score": company_fit,
            "completeness_score": completeness,
            "ats_score": ats_score,
            "format_score": format_score,
            "ats_quality_score": ats_quality,
            "keyword_score": keyword_score,
            "clarity_score": clarity,
            "professionalism_score": professionalism,
        },
        "top_strengths": [
            {"label": label, "score": value}
            for label, value in strongest[:3]
        ],
        "top_gaps": [
            {"label": label, "score": value}
            for label, value in weakest[:3]
        ],
        "summary": (
            f"Punteggio complessivo {overall}/100. "
            + "; ".join(summary_parts)
            + "."
        ).strip(),
        "explanation": [
            f"Il punteggio finale pesa soprattutto la coerenza col ruolo, la completezza e la qualità ATS, evitando di contare due volte keyword e struttura.",
            f"Coerenza col ruolo: {role_match}/100.",
            f"Compatibilità ATS: {ats_score}/100.",
            f"Copertura competenze e termini target: {keyword_score}/100.",
            f"Formato testuale e sezioni riconoscibili: {format_score}/100.",
            f"Completezza: {completeness}/100.",
            *(
                [f"Adattamento all'azienda: {company_fit}/100."]
                if payload.get("company_provided")
                else ["Adattamento all'azienda non incluso: azienda non specificata."]
            ),
            f"Chiarezza: {clarity}/100, professionalità: {professionalism}/100.",
        ],
    }


def build_speaking_feedback_from_metrics(metrics: SpeechMetrics) -> str:
    """
    Costruisce un feedback testuale sul parlato usando le metriche ricevute dal frontend.
    Non serve mostrarle a schermo: servono solo per produrre una valutazione più utile.
    """

    duration = metrics.duration_seconds or 0
    words = metrics.words_count or 0
    wpm = metrics.words_per_minute or 0
    fillers = metrics.filler_words_count or 0

    if words <= 0:
        return (
            "Il modo di parlare non è valutabile perché la risposta non contiene parole trascritte "
            "in modo sufficiente."
        )

    if wpm < 80:
        ritmo = (
            "Il ritmo risulta piuttosto lento: in un colloquio può trasmettere esitazione "
            "o poca sicurezza, quindi conviene allenarsi a parlare in modo leggermente più fluido."
        )
    elif wpm <= 160:
        ritmo = (
            "Il ritmo è complessivamente adeguato per un colloquio: permette a chi ascolta "
            "di seguire il discorso senza difficoltà."
        )
    else:
        ritmo = (
            "Il ritmo è abbastanza veloce: conviene rallentare leggermente per risultare più chiara, "
            "ordinata e sicura."
        )

    if fillers == 0:
        riempitivi = (
            "Non emergono parole riempitive rilevanti, quindi il discorso appare abbastanza pulito."
        )
    elif fillers <= 2:
        riempitivi = (
            "Sono presenti poche parole riempitive: non compromettono la risposta, ma si può ancora "
            "migliorare la fluidità."
        )
    else:
        riempitivi = (
            "Sono presenti diverse parole riempitive: questo può rendere il discorso meno sicuro "
            "e meno professionale."
        )

    if duration < 10 or words < 15:
        struttura = (
            "La risposta è breve: per un colloquio sarebbe utile sviluppare meglio il discorso "
            "con un esempio concreto o una motivazione più chiara."
        )
    else:
        struttura = (
            "La risposta ha una durata sufficiente per essere valutata anche dal punto di vista espositivo."
        )

    return f"{ritmo} {riempitivi} {struttura}"


# =========================
# FILTRO RISPOSTE NON VALIDE
# =========================

def normalize_answer_text(text: str) -> str:
    """
    Normalizza il testo per riconoscere meglio risposte non valide:
    - minuscole;
    - rimozione accenti;
    - rimozione punteggiatura;
    - spazi uniformi.
    """

    if not text:
        return ""

    text = text.lower().strip()
    text = unicodedata.normalize("NFD", text)
    text = "".join(char for char in text if unicodedata.category(char) != "Mn")
    text = re.sub(r"[^\w\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()

    return text


def get_zero_answer_reason(answer: str, question: str = "") -> Optional[str]:
    # Restituisce il motivo per cui una risposta deve prendere 0.
    # Se la risposta e valutabile, restituisce None.
    # Serve per bloccare prima dell'LLM risposte come:
    # - "che criterio usi per valutarmi?";
    # - "quanto mi dai?";
    # - "qual e la risposta giusta?";
    # - "boh", "non lo so";
    # - testo casuale o troppo breve.
    if not answer:
        return "Risposta vuota"

    original_text = answer.strip()
    text = normalize_answer_text(original_text)

    if len(text) < 3:
        return "Risposta troppo breve"

    zero_phrases = {
        "boh",
        "bo",
        "mah",
        "non lo so",
        "non so",
        "chi lo sa",
        "non saprei",
        "non ne ho idea",
        "nessuna idea",
        "non mi viene",
        "non mi viene in mente",
        "non voglio rispondere",
        "skip",
        "passo",
        "idk",
        "i don t know",
        "dont know",
        "no idea"
    }

    if text in zero_phrases:
        return "Risposta di rinuncia o assenza di contenuto"

    words = text.split()

    if len(words) < 3:
        return "Risposta troppo breve per essere valutata"

    for phrase in zero_phrases:
        if phrase in text and len(words) <= 8:
            return "Risposta di rinuncia o forte incertezza"

    meta_patterns = [
        r"\bche criterio\b.*\bvalut",
        r"\bche criteri\b.*\bvalut",
        r"\bquale criterio\b.*\bvalut",
        r"\bquali criteri\b.*\bvalut",
        r"\bcome\b.*\bvalut",
        r"\bcome mi valuti\b",
        r"\bquanto mi dai\b",
        r"\bquanto mi daresti\b",
        r"\bche punteggio\b",
        r"\bche voto\b",
        r"\bche giudizio\b",
        r"\bqual e la risposta corretta\b",
        r"\bqual e la risposta giusta\b",
        r"\bcosa dovrei rispondere\b",
        r"\bcosa devo rispondere\b",
        r"\bmi aiuti\b.*\brispondere\b",
        r"\bpuoi aiutarmi\b.*\brispondere\b",
        r"\bpuoi farmi un esempio\b",
        r"\bfammi un esempio\b",
        r"\bdammi la risposta\b",
        r"\bsuggeriscimi\b.*\brisposta\b",
        r"\bpuoi spiegarmi\b.*\bdomanda\b",
        r"\bnon ho capito\b.*\bdomanda\b",
        r"\bcome funziona\b.*\bvalut",
        r"\bcome funziona\b.*\bsistema\b",
        r"\bcome funziona\b.*\bcodice\b",
        r"\bcome e costruito\b.*\bcodice\b",
        r"\blogica\b.*\bprogetto\b",
        r"\bsistema di valutazione\b",
        r"\bcriteri di valutazione\b"
    ]

    for pattern in meta_patterns:
        if re.search(pattern, text):
            return (
                "Risposta non valida: il candidato fa una domanda sul sistema, "
                "sulla valutazione o su cosa rispondere"
            )

    question_words = [
        "che",
        "come",
        "quanto",
        "quale",
        "quali",
        "cosa",
        "perche",
        "puoi",
        "potresti",
        "sapresti",
        "mi"
    ]

    if original_text.endswith("?") and words and words[0] in question_words:
        return "Risposta non valida: il candidato fa una domanda invece di rispondere"

    irrelevant_short_answers = {
        "ciao",
        "ok",
        "okay",
        "va bene",
        "bene",
        "male",
        "si",
        "sì",
        "no",
        "forse",
        "niente",
        "nulla",
        "tutto bene",
        "mi piace",
        "non mi piace",
        "dipende"
    }

    if text in irrelevant_short_answers:
        return "Risposta troppo generica o non pertinente"

    vowels = "aeiou"
    weird_words = 0

    for word in words:
        clean_word = "".join(ch for ch in word if ch.isalpha())

        if not clean_word:
            weird_words += 1
            continue

        if len(clean_word) >= 5 and not any(v in clean_word for v in vowels):
            weird_words += 1

        consecutive_consonants = 0
        max_consecutive_consonants = 0

        for ch in clean_word:
            if ch.isalpha() and ch not in vowels:
                consecutive_consonants += 1
                max_consecutive_consonants = max(max_consecutive_consonants, consecutive_consonants)
            else:
                consecutive_consonants = 0

        if max_consecutive_consonants >= 5:
            weird_words += 1

    weird_ratio = weird_words / len(words)

    if weird_ratio >= 0.5:
        return "Risposta indecifrabile o composta da testo casuale"

    letters = sum(1 for ch in text if ch.isalpha())
    total = len(text)

    if total > 0 and letters / total < 0.45:
        return "Risposta non valutabile perché contiene troppo poco testo significativo"

    return None


def is_zero_answer(answer: str, question: str = "") -> bool:
    return get_zero_answer_reason(answer, question) is not None


def build_zero_feedback(
    reason: str = "Risposta non valutabile",
    question: str = "",
    interview_type: str = ""
):
    """
    Feedback standard per risposte indecifrabili, non pertinenti o prive di contenuto.
    """

    if interview_type == "logica":
        improved_answer = (
            "La risposta non è valutabile. Per una domanda di logica, una risposta corretta dovrebbe "
            "spiegare il ragionamento passo dopo passo, chiarendo le ipotesi fatte, i passaggi intermedi "
            "e la conclusione. Non basta dare un numero o dire 'non lo so'."
        )
        solution_explanation = (
            f"Per risolvere questa domanda bisogna partire dal testo: '{question}'. "
            "Individua prima cosa viene chiesto, poi elenca le ipotesi, costruisci un ragionamento "
            "progressivo e arriva a una conclusione motivata. Nelle domande di stima non conta il numero "
            "esatto, ma la qualità del ragionamento."
        )
    else:
        improved_answer = (
            f"Una risposta efficace alla domanda '{question}' dovrebbe essere chiara, pertinente e strutturata. "
            "Puoi iniziare rispondendo direttamente alla domanda, poi aggiungere un esempio concreto o una "
            "motivazione collegata al ruolo e concludere spiegando cosa potresti portare all'azienda."
        )
        solution_explanation = ""

    return {
        "clarity_score": 0,
        "completeness_score": 0,
        "relevance_score": 0,
        "professionalism_score": 0,
        "synthesis_score": 0,
        "speaking_score": 0,
        "total_score": 0,
        "feedback": (
            f"{reason}. La risposta non può essere considerata valida in un colloquio, "
            "perché non fornisce contenuto utile, non risponde alla domanda oppure risulta "
            "troppo vaga/indecifrabile."
        ),
        "improved_answer": improved_answer,
        "speaking_feedback": (
            "Il modo di parlare non è valutabile perché il contenuto della risposta non è valido."
        ),
        "solution_explanation": solution_explanation
    }


def get_speech_metrics_json(metrics: Optional[SpeechMetrics]) -> Optional[str]:
    if not metrics:
        return None

    try:
        return metrics.model_dump_json()
    except Exception:
        return metrics.json()


def save_answer_result(
    cursor,
    question_id: int,
    user_answer: str,
    result: Dict,
    speech_metrics_json: Optional[str]
):
    cursor.execute("""
    INSERT INTO answers (
        question_id,
        user_answer,
        clarity_score,
        completeness_score,
        relevance_score,
        professionalism_score,
        synthesis_score,
        speaking_score,
        total_score,
        feedback,
        improved_answer,
        speaking_feedback,
        solution_explanation,
        speech_metrics_json
    )
    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        question_id,
        user_answer,
        result["clarity_score"],
        result["completeness_score"],
        result["relevance_score"],
        result["professionalism_score"],
        result["synthesis_score"],
        result["speaking_score"],
        result["total_score"],
        result["feedback"],
        result["improved_answer"],
        result["speaking_feedback"],
        result["solution_explanation"],
        speech_metrics_json
    ))


def build_search_query(company: str, role: str, interview_type: str, language: str) -> str:
    company = company.strip()
    role = role.strip()
    interview_type = interview_type.strip()

    if language.lower().startswith("ingles"):
        if interview_type == "conoscitive_motivazionali":
            return (
                f"{company} interview questions motivation behavioral cultural fit "
                f"tell me about yourself why do you want to work here goals teamwork"
            )

        if interview_type == "tecniche":
            return (
                f"{company} {role} technical interview questions skills assessment "
                f"role specific interview questions"
            )

        if interview_type == "logica":
            return (
                f"{company} logic interview questions brain teasers reasoning puzzles "
                f"problem solving interview questions"
            )

        return f"{company} {role} interview questions"

    if interview_type == "conoscitive_motivazionali":
        return (
            f"{company} domande colloquio motivazionale conoscitivo "
            f"parlami di te obiettivi perché vuoi lavorare qui lavoro di gruppo "
            f"cosa sai dell'azienda"
        )

    if interview_type == "tecniche":
        return (
            f"{company} {role} domande colloquio tecnico competenze "
            f"domande tecniche ruolo selezione candidato"
        )

    if interview_type == "logica":
        return (
            f"{company} domande colloquio logica ragionamento indovinelli "
            f"brain teaser problem solving trabocchetto stima"
        )

    return f"{company} {role} domande colloquio"


def search_web_interview_questions(
    company: str,
    role: str,
    interview_type: str,
    language: str = "Italiano"
) -> List[Dict[str, str]]:
    if not TAVILY_API_KEY:
        print("Ricerca Tavily saltata: TAVILY_API_KEY non configurata.")
        return []

    query = build_search_query(company, role, interview_type, language)

    try:
        print(f"Ricerca Tavily avviata: {query}")

        response = requests.post(
            "https://api.tavily.com/search",
            json={
                "api_key": TAVILY_API_KEY,
                "query": query,
                "search_depth": "basic",
                "max_results": 3,
                "include_answer": False,
                "include_raw_content": False
            },
            timeout=10
        )

        response.raise_for_status()

        data = response.json()
        results = data.get("results", [])

        clean_results = []

        for item in results:
            clean_results.append({
                "title": item.get("title", ""),
                "url": item.get("url", ""),
                "content": item.get("content", "")
            })

        print(f"Ricerca Tavily completata. Fonti trovate: {len(clean_results)}")
        return clean_results

    except Exception as e:
        print(f"Errore Tavily, continuo senza fonti web: {e}")
        return []


JOB_SOURCE_KEYWORDS = (
    "job", "jobs", "careers", "career", "apply", "hiring", "position",
    "role", "vacancy", "offerta", "lavoro", "candidatura", "lavora con noi",
    "work with us", "join us",
)

RELIABLE_JOB_SOURCE_DOMAINS = (
    "indeed.",
    "glassdoor.",
    "lever.co",
    "greenhouse.io",
    "workdayjobs.com",
    "smartrecruiters.com",
)

EXCLUDED_JOB_SOURCE_TERMS = (
    "certificate", "certification", "certificato", "certificazione",
    "course", "corso", "training", "blog", "forum", "quora", "reddit",
)


def is_reliable_job_source(source: Dict[str, str]) -> bool:
    title = str(source.get("title") or "")
    url = str(source.get("url") or "")
    if not url.strip():
        return False

    parsed = urllib.parse.urlparse(url)
    domain = parsed.netloc.lower()
    path = parsed.path.lower()
    normalized_title = title.lower()
    normalized_url = url.lower().replace("-", " ").replace("_", " ")
    haystack = f"{normalized_title} {normalized_url}"

    if "reddit.com" in domain or "quora.com" in domain:
        return False
    if path.endswith(".pdf") or normalized_url.split("?", 1)[0].endswith(".pdf"):
        return False
    if any(term in haystack for term in EXCLUDED_JOB_SOURCE_TERMS):
        return False

    is_linkedin_jobs = "linkedin.com" in domain and path.startswith("/jobs")
    is_google_jobs = domain in {"careers.google.com", "jobs.google.com"}
    is_known_job_board = any(domain.endswith(item) or item in domain for item in RELIABLE_JOB_SOURCE_DOMAINS)
    has_job_keyword = any(keyword in haystack for keyword in JOB_SOURCE_KEYWORDS)

    return is_linkedin_jobs or is_google_jobs or is_known_job_board or has_job_keyword


def filter_candidate_sources(sources: List[Dict[str, str]]) -> List[Dict[str, str]]:
    filtered = []
    seen_urls = set()
    for source in sources or []:
        if not isinstance(source, dict):
            continue
        url = str(source.get("url") or "").strip()
        if not url or url in seen_urls:
            continue
        if not is_reliable_job_source(source):
            print(
                "Fonte candidatura esclusa: "
                f"{source.get('title', '')} | {url}"
            )
            continue
        seen_urls.add(url)
        filtered.append({
            "title": str(source.get("title") or url).strip(),
            "url": url,
            "content": str(source.get("content") or "").strip(),
        })
    return filtered[:4]


def search_job_context(company: str, role: str, job_link: str) -> List[Dict[str, str]]:
    parts = [value.strip() for value in [job_link, company, role, "requisiti competenze job description"] if value and value.strip()]
    query = " ".join(parts)

    if not query:
        return []

    try:
        print(f"Ricerca contesto candidatura avviata: {query}")
        response = requests.post(
            "https://api.tavily.com/search",
            json={
                "api_key": TAVILY_API_KEY,
                "query": query,
                "search_depth": "basic",
                "max_results": 4,
                "include_answer": False,
                "include_raw_content": False
            },
            timeout=10
        )
        response.raise_for_status()
        results = response.json().get("results", [])
        raw_sources = []
        if job_link:
            raw_sources.append({
                "title": f"Annuncio o pagina candidatura: {role or company}".strip(),
                "url": job_link,
                "content": "",
            })
        raw_sources.extend([
            {
                "title": item.get("title", ""),
                "url": item.get("url", ""),
                "content": item.get("content", "")
            }
            for item in results
        ])
        filtered_sources = filter_candidate_sources(raw_sources)
        print(
            "Ricerca contesto candidatura completata. "
            f"Fonti affidabili: {len(filtered_sources)}/{len(raw_sources)}"
        )
        return filtered_sources
    except Exception as e:
        print(f"Errore ricerca contesto candidatura, continuo senza fonti web: {e}")
        return []


def sources_to_prompt(sources: List[Dict[str, str]]) -> str:
    if not sources:
        return "Nessuna fonte web trovata."

    text = ""

    for index, source in enumerate(sources, start=1):
        text += f"""
Fonte {index}
Titolo: {source.get("title", "")}
URL: {source.get("url", "")}
Estratto: {source.get("content", "")}
"""

    return text


def normalize_instagram_handle(handle: Optional[str]) -> str:
    if not handle:
        return ""

    cleaned = handle.strip()
    cleaned = cleaned.replace("https://www.instagram.com/", "")
    cleaned = cleaned.replace("https://instagram.com/", "")
    cleaned = cleaned.split("?")[0].strip("/")
    cleaned = cleaned.lstrip("@")
    return cleaned


def normalize_public_profile_url(url: Optional[str]) -> str:
    if not url:
        return ""

    cleaned = url.strip()
    if cleaned and not cleaned.startswith(("http://", "https://")):
        cleaned = f"https://{cleaned}"
    return cleaned.rstrip("/")


def profile_path(url: str) -> str:
    try:
        parsed = urllib.parse.urlparse(normalize_public_profile_url(url))
        return parsed.path.strip("/").lower()
    except Exception:
        return ""


def canonical_public_url(url: Optional[str]) -> str:
    normalized = normalize_public_profile_url(url)
    if not normalized:
        return ""

    parsed = urllib.parse.urlparse(normalized)
    hostname = (parsed.hostname or "").lower()
    path = parsed.path.rstrip("/").lower()
    return f"{hostname}{path}"


def url_belongs_to_public_profile(profile_url: str, result_url: str) -> bool:
    profile = urllib.parse.urlparse(normalize_public_profile_url(profile_url))
    result = urllib.parse.urlparse(normalize_public_profile_url(result_url))
    profile_host = (profile.hostname or "").lower().removeprefix("www.")
    result_host = (result.hostname or "").lower().removeprefix("www.")
    profile_path = profile.path.rstrip("/").lower()
    result_path = result.path.rstrip("/").lower()

    if not profile_host or profile_host != result_host:
        return False

    return not profile_path or result_path == profile_path or result_path.startswith(f"{profile_path}/")


def normalize_identity_tokens(value: Optional[str]) -> set[str]:
    if not value:
        return set()

    normalized = unicodedata.normalize("NFKD", value.lower())
    normalized = "".join(char for char in normalized if not unicodedata.combining(char))
    return {
        token
        for token in re.findall(r"[a-z0-9]+", normalized)
        if len(token) >= 2
    }


def normalize_linkedin_profile_url(url: Optional[str]) -> str:
    normalized = normalize_public_profile_url(url)
    if not normalized:
        return ""

    parsed = urllib.parse.urlparse(normalized)
    hostname = (parsed.hostname or "").lower()
    path_parts = [part for part in parsed.path.strip("/").split("/") if part]

    if hostname not in {"linkedin.com", "www.linkedin.com"}:
        raise HTTPException(status_code=400, detail="Inserisci un link LinkedIn valido.")

    if len(path_parts) != 2 or path_parts[0].lower() != "in":
        raise HTTPException(
            status_code=400,
            detail="Inserisci il link pubblico del profilo LinkedIn nel formato https://www.linkedin.com/in/tuo-profilo.",
        )

    return f"https://www.linkedin.com/in/{path_parts[1]}"


def build_linkedin_basic_info(linkedin_url: str) -> Optional[Dict[str, str]]:
    if not linkedin_url:
        return None

    public_identifier = profile_path(linkedin_url).split("/", 1)[-1]
    return {
        "profile_url": linkedin_url,
        "public_identifier": public_identifier,
        "access_level": "basic_public",
        "message": (
            "CareerCoach puo verificare il link e usare solo dati base o snippet pubblici. "
            "Esperienze, competenze e formazione complete richiedono autorizzazioni LinkedIn."
        ),
    }


OFFICIAL_PROFILE_CAPABILITIES = {
    "linkedin": {
        "api": "LinkedIn OpenID Connect",
        "source_quality": "official_oauth",
        "can_verify_identity": True,
        "can_read_basic_profile": True,
        "can_read_professional_sections": False,
        "can_read_posts": False,
        "can_read_media": False,
        "limitations": (
            "Con OpenID Connect sono disponibili dati base di identita. "
            "Esperienze, formazione, competenze, post e media richiedono permessi/prodotti LinkedIn non inclusi nel login base."
        ),
    },
    "instagram": {
        "api": "Instagram Graph API / Instagram API with Login",
        "source_quality": "official_oauth_available",
        "can_verify_identity": "limited",
        "can_read_basic_profile": "with_permissions",
        "can_read_professional_sections": False,
        "can_read_posts": "business_or_creator_with_permissions",
        "can_read_media": "business_or_creator_with_permissions",
        "limitations": (
            "Per account personali l'accesso ufficiale e limitato. Media e insight sono affidabili solo con account professionali "
            "e permessi approvati da Meta."
        ),
    },
    "facebook": {
        "api": "Facebook Graph API",
        "source_quality": "official_oauth_available",
        "can_verify_identity": True,
        "can_read_basic_profile": True,
        "can_read_professional_sections": False,
        "can_read_posts": "requires_permissions_and_app_review",
        "can_read_media": "requires_permissions_and_app_review",
        "limitations": (
            "I contenuti personali sono fortemente limitati. Pagine e asset gestiti sono piu adatti a un controllo ufficiale."
        ),
    },
}


def build_standardized_official_profile_source(platform: str, oauth_profile: Dict, profile_url: str = "") -> Dict:
    capabilities = OFFICIAL_PROFILE_CAPABILITIES.get(platform, {})
    display_name = str(oauth_profile.get("name") or "").strip()
    profile_image_available = bool(oauth_profile.get("picture"))
    normalized_source = {
        "platform": platform,
        "source_quality": capabilities.get("source_quality", "official_oauth"),
        "api": capabilities.get("api", "OAuth API"),
        "profile_url": profile_url,
        "identity": {
            "display_name": display_name,
            "email_verified": bool(oauth_profile.get("email_verified")),
            "profile_image_available": profile_image_available,
            "locale": oauth_profile.get("locale") or "",
        },
        "professional_profile": {
            "headline": "",
            "bio": "",
            "work_experience": [],
            "education": [],
            "skills": [],
        },
        "content": {
            "posts_summary": [],
            "media_summary": [],
            "media_risk": {
                "analyzed": False,
                "sensitive_count": 0,
                "reason": "I media non sono disponibili tramite il profilo OAuth base.",
            },
        },
        "capabilities": capabilities,
        "limitations": capabilities.get("limitations", ""),
    }
    return normalized_source


def build_official_profile_sources(user: Dict) -> List[Dict]:
    sources = []
    linkedin_oauth_profile = user.get("linkedin_oauth_profile") or {}
    if linkedin_oauth_profile:
        sources.append(
            build_standardized_official_profile_source(
                "linkedin",
                linkedin_oauth_profile,
                user.get("linkedin_url", ""),
            )
        )
    return sources


def official_profile_source_to_text(source: Dict) -> str:
    identity = source.get("identity", {})
    professional_profile = source.get("professional_profile", {})
    capabilities = source.get("capabilities", {})
    return (
        f"Piattaforma: {source.get('platform', '')}. "
        f"Qualita fonte: {source.get('source_quality', '')}. "
        f"API: {source.get('api', '')}. "
        f"Nome verificato via OAuth: {identity.get('display_name', '')}. "
        f"Email verificata dal provider: {bool(identity.get('email_verified'))}. "
        f"Foto profilo disponibile: {bool(identity.get('profile_image_available'))}. "
        f"Headline disponibile: {bool(professional_profile.get('headline'))}. "
        f"Esperienze disponibili: {len(professional_profile.get('work_experience') or [])}. "
        f"Competenze disponibili: {len(professional_profile.get('skills') or [])}. "
        f"Media disponibili: {capabilities.get('can_read_media')}. "
        f"Limiti: {source.get('limitations', '')}"
    )


def search_public_profile_signals(user: Dict, digital_presence: DigitalPresenceUpdate) -> List[Dict[str, str]]:
    search_targets = []
    linkedin_url = normalize_linkedin_profile_url(digital_presence.linkedin_url)
    portfolio_url = normalize_public_profile_url(digital_presence.portfolio_url)
    linkedin_path = profile_path(linkedin_url)
    instagram_handle = normalize_instagram_handle(digital_presence.instagram_handle)

    if linkedin_url:
        search_targets.append((linkedin_url, "linkedin"))
    if portfolio_url:
        search_targets.append((portfolio_url, "additional_link"))
    if instagram_handle:
        search_targets.append((f"https://www.instagram.com/{instagram_handle}/", "instagram"))

    sources = []
    seen_snippet_urls = set()

    if user.get("linkedin_profile_text"):
        sources.append({
            "title": "Esportazione profilo LinkedIn caricata dal candidato",
            "url": linkedin_url,
            "content": user["linkedin_profile_text"][:6000],
            "kind": "linkedin_export",
        })

    for official_source in build_official_profile_sources(user):
        sources.append({
            "title": f"Dati ufficiali standardizzati: {official_source['platform']}",
            "url": linkedin_url,
            "content": official_profile_source_to_text(official_source),
            "kind": f"official_oauth_{official_source['platform']}",
        })

    if linkedin_url:
        linkedin_identifier = linkedin_path.split("/", 1)[-1]
        sources.append({
            "title": f"LinkedIn pubblico: {linkedin_identifier}",
            "url": linkedin_url,
            "content": (
                "Profilo LinkedIn indicato direttamente dal candidato. "
                "Sono disponibili il link pubblico e l'identificativo del profilo; "
                "eventuali altri dati possono essere usati solo se presenti negli snippet pubblici."
            ),
            "kind": "linkedin_reference",
        })

    if instagram_handle:
        instagram_url = f"https://www.instagram.com/{instagram_handle}/"
        sources.append({
            "title": f"Instagram @{instagram_handle}",
            "url": instagram_url,
            "content": (
                "Profilo Instagram indicato direttamente dal candidato. "
                "Bio, foto e contenuti possono essere valutati solo se risultano pubblicamente accessibili."
            ),
            "kind": "instagram_reference",
        })

    if portfolio_url:
        sources.append({
            "title": "Link aggiuntivo indicato dal candidato",
            "url": portfolio_url,
            "content": (
                "Link pubblico indicato direttamente dal candidato. "
                "L'analisi usa eventuali snippet pubblici trovati e non inventa contenuti non accessibili."
            ),
            "kind": "other_profile_reference",
        })

    if not TAVILY_API_KEY:
        print("Ricerca presenza digitale Tavily saltata: TAVILY_API_KEY non configurata.")
        return sources[:8]

    for query, target_kind in search_targets[:3]:
        try:
            response = requests.post(
                "https://api.tavily.com/search",
                json={
                    "api_key": TAVILY_API_KEY,
                    "query": query,
                    "search_depth": "basic",
                    "max_results": 3,
                    "include_answer": False,
                    "include_raw_content": False,
                },
                timeout=10,
            )
            response.raise_for_status()
            for item in response.json().get("results", []):
                url = item.get("url", "")
                title = item.get("title", "")
                content = item.get("content", "")
                normalized_url = normalize_public_profile_url(url)
                if target_kind == "linkedin":
                    if not linkedin_path:
                        continue
                    result_path = profile_path(normalized_url)
                    if result_path != linkedin_path:
                        continue
                if target_kind == "instagram":
                    if not instagram_handle:
                        continue
                    instagram_path = urllib.parse.urlparse(normalized_url).path.strip("/").split("/")
                    if not instagram_path or instagram_path[0].lower() != instagram_handle.lower():
                        continue
                if (
                    target_kind == "additional_link"
                    and not url_belongs_to_public_profile(portfolio_url, normalized_url)
                ):
                    continue
                canonical_url = canonical_public_url(normalized_url)
                if canonical_url and canonical_url in seen_snippet_urls:
                    continue
                if canonical_url:
                    seen_snippet_urls.add(canonical_url)
                sources.append({
                    "title": title,
                    "url": normalized_url or url,
                    "content": content,
                    "kind": (
                        "instagram_public_metadata"
                        if target_kind == "instagram"
                        else "linkedin_public_snippet"
                        if target_kind == "linkedin"
                        else "other_profile_public_snippet"
                    ),
                })
        except Exception as exc:
            print(f"Ricerca presenza digitale non riuscita per '{query}': {exc}")

    return sources[:8]


def has_public_instagram_metadata(sources: List[Dict[str, str]]) -> bool:
    return any(source.get("kind") == "instagram_public_metadata" for source in sources)


def has_public_other_profile_signals(sources: List[Dict[str, str]]) -> bool:
    return any(source.get("kind") == "other_profile_public_snippet" for source in sources)


class PublicImageParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.image_urls = []

    def handle_starttag(self, tag: str, attrs):
        attributes = dict(attrs)
        if tag == "meta" and attributes.get("property") in {"og:image", "og:image:url"}:
            self.image_urls.append({"url": attributes.get("content", ""), "kind": "public_preview"})
        elif tag == "meta" and attributes.get("name") in {"twitter:image", "twitter:image:src"}:
            self.image_urls.append({"url": attributes.get("content", ""), "kind": "public_preview"})
        elif tag == "img":
            self.image_urls.append({"url": attributes.get("src", ""), "kind": "page_image"})


def is_safe_public_url(url: str) -> bool:
    try:
        parsed = urllib.parse.urlparse(url)
        if parsed.scheme not in {"http", "https"} or not parsed.hostname:
            return False
        hostname = parsed.hostname.lower()
        if hostname == "localhost" or hostname.endswith(".local"):
            return False
        for info in socket.getaddrinfo(hostname, None):
            address = ipaddress.ip_address(info[4][0])
            if address.is_private or address.is_loopback or address.is_link_local or address.is_reserved:
                return False
        return True
    except Exception:
        return False


def collect_public_image_urls(user: Dict) -> Dict:
    instagram_handle = normalize_instagram_handle(user.get("instagram_handle"))
    profile_urls = []
    if instagram_handle:
        profile_urls.append(f"https://www.instagram.com/{instagram_handle}/")
    if user.get("portfolio_url"):
        profile_urls.append(normalize_public_profile_url(user["portfolio_url"]))

    image_urls = []
    pages_checked = []
    seen_urls = set()
    headers = {
        "User-Agent": "Mozilla/5.0 (compatible; CareerCoach/1.0; public-profile-check)",
        "Accept": "text/html,application/xhtml+xml",
    }
    for profile_url in profile_urls[:3]:
        if not is_safe_public_url(profile_url):
            continue
        try:
            response = requests.get(profile_url, headers=headers, timeout=8)
            response.raise_for_status()
            content_type = response.headers.get("content-type", "")
            pages_checked.append(profile_url)
            if "html" not in content_type.lower():
                continue
            parser = PublicImageParser()
            parser.feed(response.text[:1_000_000])
            for candidate in parser.image_urls:
                candidate_url = candidate["url"]
                image_url = urllib.parse.urljoin(profile_url, candidate_url)
                canonical_url = canonical_public_url(image_url)
                if (
                    not candidate_url
                    or not canonical_url
                    or canonical_url in seen_urls
                    or not is_safe_public_url(image_url)
                ):
                    continue
                seen_urls.add(canonical_url)
                image_urls.append({
                    "url": image_url,
                    "kind": candidate["kind"],
                    "profile_url": profile_url,
                })
                if len(image_urls) >= 8:
                    break
        except Exception as exc:
            print(f"Recupero media pubblici non riuscito per '{profile_url}': {exc}")
        if len(image_urls) >= 8:
            break

    return {"pages_checked": pages_checked, "image_urls": image_urls}


def extract_image_base64(image_input: Dict) -> str:
    data_url = image_input.get("image_url", {}).get("url", "")
    if not data_url.startswith("data:image/") or "," not in data_url:
        return ""
    return data_url.split(",", 1)[1]


VISUAL_SENSITIVE_CATEGORIES = {
    "nudita",
    "contenuto sessuale esplicito",
    "contenuto intimo o non professionale",
    "sexual",
    "sexual/minors",
}


def is_sensitive_visual_category(category: str) -> bool:
    normalized = str(category or "").strip().lower()
    return (
        normalized in VISUAL_SENSITIVE_CATEGORIES
        or "sexual" in normalized
        or "nudit" in normalized
        or "nudita" in normalized
        or "nude" in normalized
        or "intim" in normalized
    )


def build_visual_analysis_result(
    source: str,
    discovered_count: int,
    analyzed: List[Dict],
    failed_count: int,
    content_count: Optional[int] = None,
) -> Dict:
    flagged_results = [result for result in analyzed if result.get("flagged")]
    sensitive_results = [
        result
        for result in flagged_results
        if any(is_sensitive_visual_category(category) for category in result.get("categories", []))
    ]
    categories = sorted({
        category
        for result in flagged_results
        for category in result.get("categories", [])
    })
    incomplete_message = f" {failed_count} media non sono risultati leggibili." if failed_count else ""
    analyzed_content_count = len(analyzed) if content_count is None else content_count
    preview_count = max(0, len(analyzed) - analyzed_content_count)
    analyzed_label = "contenuto caricato" if len(analyzed) == 1 else "contenuti caricati"
    preview_label = "anteprima pubblica" if len(analyzed) == 1 else "anteprime pubbliche"
    scope_message = (
        f"Controllo visuale preliminare locale completato su {len(analyzed)} {analyzed_label}"
        if source == "uploaded_screenshots"
        else f"Controllo visuale preliminare locale completato su {len(analyzed)} {preview_label}"
    )
    return {
        "status": "completed",
        "provider": VISION_PROVIDER,
        "source": source,
        "discovered_count": discovered_count,
        "analyzed_count": len(analyzed),
        "analyzed_content_count": analyzed_content_count,
        "analyzed_preview_count": preview_count,
        "flagged_count": len(flagged_results),
        "sensitive_flagged_count": len(sensitive_results),
        "failed_count": failed_count,
        "flagged_categories": categories,
        "message": (
            f"{scope_message}: {len(flagged_results)} richiedono una verifica manuale.{incomplete_message}"
            if flagged_results
            else (
                f"{scope_message}: "
                f"non sono emersi contenuti sensibili evidenti.{incomplete_message}"
            )
        ),
    }


def analyze_image_with_ollama(image_input: Dict) -> Dict:
    encoded_image = extract_image_base64(image_input)
    if not encoded_image:
        raise ValueError("Immagine Base64 non disponibile.")

    uses_lightweight_description = OLLAMA_VISION_MODEL.split(":", 1)[0] == "moondream"
    if uses_lightweight_description:
        prompt = "Describe only what is visibly present in this image in one short factual sentence."
    else:
        prompt = (
            "Review the image. Return JSON only: "
            '{"flagged": boolean, "categories": string[], "summary": string}. '
            "Use categories only from: nudita, contenuto sessuale esplicito, "
            "contenuto intimo o non professionale, violenza, armi, droghe, "
            "linguaggio offensivo visibile. Flag intimate, revealing, underwear, "
            "topless, nude, or sexually suggestive images even when they are not explicit. "
            "Do not flag pets, landscapes, food, hobbies, or neutral casual photos unless "
            "another listed risk is visible. Keep summary brief and factual."
        )
    payload = {
        "model": OLLAMA_VISION_MODEL,
        "stream": False,
        "messages": [{
            "role": "user",
            "content": prompt,
            "images": [encoded_image],
        }],
    }
    if not uses_lightweight_description:
        payload["format"] = "json"

    response = requests.post(
        f"{OLLAMA_URL}/api/chat",
        json=payload,
        timeout=240,
    )
    if not response.ok:
        try:
            error_message = response.json().get("error", response.text)
        except ValueError:
            error_message = response.text
        raise RuntimeError(f"Ollama: {error_message}")
    content = response.json().get("message", {}).get("content", "").strip()
    if uses_lightweight_description:
        description = content.lower()
        category_keywords = {
            "nudita": ("nude", "nudity", "naked", "topless", "bare chest"),
            "contenuto sessuale esplicito": ("explicit sexual", "sexual act", "sexually explicit"),
            "contenuto intimo o non professionale": (
                "underwear",
                "lingerie",
                "revealing",
                "suggestive",
                "intimate",
                "shirtless",
                "topless",
            ),
            "violenza": ("violence", "violent", "blood", "injury", "wound"),
            "armi": ("weapon", "gun", "rifle", "knife", "firearm"),
            "droghe": ("drug", "cocaine", "heroin", "marijuana", "syringe"),
            "linguaggio offensivo visibile": ("offensive language", "slur", "insult"),
        }
        categories = [
            category
            for category, keywords in category_keywords.items()
            if any(keyword in description for keyword in keywords)
        ]
        return {
            "flagged": bool(categories),
            "categories": categories,
            "summary": content,
        }

    result = extract_json(content)
    allowed_categories = {
        "nudita",
        "contenuto sessuale esplicito",
        "contenuto intimo o non professionale",
        "violenza",
        "armi",
        "droghe",
        "linguaggio offensivo visibile",
    }
    categories = [
        str(category)
        for category in result.get("categories") or []
        if str(category).lower() in allowed_categories
    ]
    return {
        "flagged": bool(result.get("flagged")),
        "categories": categories,
        "summary": str(result.get("summary", "")).strip(),
    }


CV_IMAGE_BLOCKED_CATEGORIES = {
    "animale",
    "nudita",
    "contenuto sessuale",
    "violenza",
    "sangue o ferite",
    "armi",
    "droghe",
    "immagine non pertinente",
    "contenuto ambiguo",
}


def normalize_cv_image_result(result: Dict[str, Any]) -> Dict[str, Any]:
    categories = [
        str(category).strip().lower()
        for category in result.get("categories") or []
        if str(category).strip().lower() in CV_IMAGE_BLOCKED_CATEGORIES
    ]
    return {
        "blocked": bool(result.get("blocked")) or bool(categories),
        "categories": sorted(set(categories)),
        "summary": str(result.get("summary") or "").strip(),
    }


def classify_cv_image_description(description: str) -> Dict[str, Any]:
    normalized = (description or "").lower()
    category_keywords = {
        "animale": (
            "dog", "cat", "puppy", "kitten", "pet", "animal", "cane", "gatto",
            "cucciolo", "animale", "bird", "horse",
        ),
        "nudita": ("nude", "nudity", "naked", "topless", "senza vestiti", "nud"),
        "contenuto sessuale": ("sexual", "sexually", "lingerie", "underwear", "intimate"),
        "violenza": ("violence", "violent", "fight", "assault", "violenza"),
        "sangue o ferite": ("blood", "injury", "wound", "gore", "sangue", "ferita"),
        "armi": ("weapon", "gun", "rifle", "knife", "firearm", "arma", "pistola"),
        "droghe": ("drug", "cocaine", "heroin", "marijuana", "syringe", "droga"),
        "immagine non pertinente": (
            "landscape", "food", "meal", "car", "motorcycle", "vacation", "scenery",
            "paesaggio", "cibo", "automobile", "vacanza",
        ),
    }
    categories = [
        category
        for category, keywords in category_keywords.items()
        if any(keyword in normalized for keyword in keywords)
    ]
    allowed_markers = (
        "portrait", "headshot", "professional photo", "person", "man", "woman",
        "logo", "icon", "graphic", "chart", "diagram", "ritratto", "persona",
        "uomo", "donna", "grafico", "diagramma",
    )
    if not categories and not any(marker in normalized for marker in allowed_markers):
        categories.append("contenuto ambiguo")
    return normalize_cv_image_result({
        "blocked": bool(categories),
        "categories": categories,
        "summary": description,
    })


def analyze_cv_image_with_ollama(image_input: Dict[str, Any]) -> Dict[str, Any]:
    encoded_image = extract_image_base64(image_input)
    if not encoded_image:
        raise ValueError("Immagine Base64 non disponibile.")

    uses_lightweight_description = OLLAMA_VISION_MODEL.split(":", 1)[0] == "moondream"
    prompt = (
        "Describe only what is visibly present in this image in one short factual sentence."
        if uses_lightweight_description
        else (
            "Analyze this image embedded in a resume. Return JSON only: "
            '{"blocked": boolean, "categories": string[], "summary": string}. '
            "Allowed categories: animale, nudita, contenuto sessuale, violenza, "
            "sangue o ferite, armi, droghe, immagine non pertinente, contenuto ambiguo. "
            "Block animals including dogs and cats; nudity or sexual content; violence, gore, "
            "weapons or drugs; unrelated photos such as food, landscapes, vehicles or holidays; "
            "and anything ambiguous. Allow only a normal professional portrait of the candidate "
            "or harmless resume graphics such as logos, icons, charts and diagrams."
        )
    )
    payload = {
        "model": OLLAMA_VISION_MODEL,
        "stream": False,
        "messages": [{
            "role": "user",
            "content": prompt,
            "images": [encoded_image],
        }],
    }
    if not uses_lightweight_description:
        payload["format"] = "json"
    response = requests.post(f"{OLLAMA_URL}/api/chat", json=payload, timeout=240)
    if not response.ok:
        raise RuntimeError(f"Ollama: {response.text}")
    content = response.json().get("message", {}).get("content", "").strip()
    if uses_lightweight_description:
        return classify_cv_image_description(content)
    return normalize_cv_image_result(extract_json(content))


def analyze_cv_image_with_openai(image_input: Dict[str, Any]) -> Dict[str, Any]:
    if not openai_moderation_client:
        raise RuntimeError("OpenAI non configurato.")
    moderation = openai_moderation_client.moderations.create(
        model="omni-moderation-latest",
        input=[{
            "type": "image_url",
            "image_url": image_input.get("image_url", {}),
        }],
    )
    flagged_categories = sorted({
        category
        for item in moderation.results
        for category, flagged in item.categories.model_dump().items()
        if flagged
    })
    if flagged_categories:
        return {
            "blocked": True,
            "categories": ["contenuto ambiguo"],
            "summary": f"Moderazione sensibile: {', '.join(flagged_categories)}",
        }

    response = openai_moderation_client.chat.completions.create(
        model=OPENAI_VISION_MODEL,
        temperature=0,
        response_format={"type": "json_object"},
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "Analizza questa immagine incorporata in un CV. Rispondi solo con JSON: "
                        '{"blocked": boolean, "categories": string[], "summary": string}. '
                        "Categorie consentite: animale, nudita, contenuto sessuale, violenza, "
                        "sangue o ferite, armi, droghe, immagine non pertinente, contenuto ambiguo. "
                        "Blocca cani, gatti e altri animali, immagini sensibili, immagini estranee "
                        "al CV e casi ambigui. Consenti soltanto un normale ritratto professionale "
                        "del candidato oppure loghi, icone, grafici e diagrammi innocui."
                    ),
                },
                {"type": "image_url", "image_url": image_input.get("image_url", {})},
            ],
        }],
    )
    content = response.choices[0].message.content or "{}"
    return normalize_cv_image_result(extract_json(content))


def analyze_embedded_cv_image(image_input: Dict[str, Any]) -> Dict[str, Any]:
    if VISION_PROVIDER == "ollama":
        return analyze_cv_image_with_ollama(image_input)
    if VISION_PROVIDER == "openai":
        return analyze_cv_image_with_openai(image_input)
    raise RuntimeError("Provider visuale non configurato.")


def moderate_visual_inputs(image_inputs: List[Dict], source: str, discovered_count: int) -> Dict:
    global openai_visual_rate_limited_until

    if not image_inputs:
        return {
            "status": "no_media_found",
            "provider": VISION_PROVIDER,
            "source": source,
            "discovered_count": discovered_count,
            "analyzed_count": 0,
            "flagged_count": 0,
            "sensitive_flagged_count": 0,
            "message": "Non sono stati recuperati media analizzabili automaticamente.",
        }

    if VISION_PROVIDER == "ollama":
        analyzed = []
        failed_count = 0
        last_error = ""
        for image_input in image_inputs[:8]:
            try:
                analyzed.append(analyze_image_with_ollama(image_input))
            except Exception as exc:
                failed_count += 1
                last_error = str(exc)
                print(f"Analisi visuale Ollama non riuscita per un media: {exc}")
                if (
                    "connection" in last_error.lower()
                    or "connessione" in last_error.lower()
                    or last_error.startswith("Ollama:")
                ):
                    break
        if analyzed:
            return build_visual_analysis_result(source, discovered_count, analyzed, failed_count)
        return {
            "status": "provider_unavailable",
            "provider": "ollama",
            "source": source,
            "discovered_count": discovered_count,
            "analyzed_count": 0,
            "flagged_count": 0,
            "sensitive_flagged_count": 0,
            "failed_count": failed_count,
            "message": (
                "Ollama non risponde oppure il modello visuale non e disponibile. "
                f"Avvia Ollama e scarica {OLLAMA_VISION_MODEL}. "
                f"Dettaglio: {last_error}"
            ),
        }

    if VISION_PROVIDER != "openai" or not openai_moderation_client:
        return {
            "status": "provider_not_configured",
            "provider": VISION_PROVIDER,
            "source": source,
            "discovered_count": discovered_count,
            "analyzed_count": 0,
            "flagged_count": 0,
            "sensitive_flagged_count": 0,
            "message": "Provider visuale non configurato. Usa VISION_PROVIDER=ollama oppure configura OpenAI.",
        }
    if openai_visual_rate_limited_until and datetime.utcnow() < openai_visual_rate_limited_until:
        return {
            "status": "rate_limited",
            "provider": "openai",
            "source": source,
            "discovered_count": discovered_count,
            "analyzed_count": 0,
            "flagged_count": 0,
            "sensitive_flagged_count": 0,
            "message": "Il servizio OpenAI ha raggiunto il limite temporaneo. Attendi qualche minuto.",
        }

    analyzed = []
    failed_count = 0
    for image_input in image_inputs[:8]:
        try:
            response = openai_moderation_client.moderations.create(
                model="omni-moderation-latest",
                input=[image_input],
            )
            for result in response.results:
                analyzed.append({
                    "flagged": bool(result.flagged),
                    "categories": [
                        category
                        for category, flagged in result.categories.model_dump().items()
                        if flagged
                    ],
                })
        except Exception as exc:
            failed_count += 1
            error_text = str(exc)
            print(f"Moderazione visuale OpenAI non riuscita per un media: {exc}")
            if "429" in error_text or "too many requests" in error_text.lower():
                openai_visual_rate_limited_until = datetime.utcnow() + timedelta(minutes=5)
                break
    if analyzed:
        return build_visual_analysis_result(source, discovered_count, analyzed, failed_count)
    return {
        "status": "analysis_failed",
        "provider": "openai",
        "source": source,
        "discovered_count": discovered_count,
        "analyzed_count": 0,
        "flagged_count": 0,
        "sensitive_flagged_count": 0,
        "failed_count": failed_count,
        "message": "I media sono stati trovati, ma l'analisi visuale OpenAI non e riuscita.",
    }


def clean_social_ocr_text(value: Any) -> str:
    text = str(value or "").strip()
    if not text or text.lower() in {
        "no text",
        "no text visible",
        "no legible text",
        "nessun testo",
        "nessun testo leggibile",
    }:
        return ""
    lines = []
    seen = set()
    for raw_line in text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip(" \t-•")
        key = line.casefold()
        if len(line) < 2 or key in seen:
            continue
        seen.add(key)
        lines.append(line)
    return "\n".join(lines)[:8000]


@lru_cache(maxsize=1)
def get_rapid_ocr_engine():
    from rapidocr_onnxruntime import RapidOCR

    return RapidOCR()


def extract_social_text_with_rapidocr(image_input: Dict) -> str:
    encoded_image = extract_image_base64(image_input)
    if not encoded_image:
        raise ValueError("Immagine Base64 non disponibile.")
    image_bytes = base64.b64decode(encoded_image)
    result, _elapsed = get_rapid_ocr_engine()(image_bytes)
    lines = [
        str(item[1]).strip()
        for item in (result or [])
        if len(item) >= 3 and float(item[2] or 0) >= 0.45 and str(item[1]).strip()
    ]
    return clean_social_ocr_text("\n".join(lines))


def extract_social_text_with_ollama(image_input: Dict) -> str:
    encoded_image = extract_image_base64(image_input)
    if not encoded_image:
        raise ValueError("Immagine Base64 non disponibile.")
    try:
        from PIL import Image, ImageEnhance, ImageFilter, ImageOps

        image = Image.open(io.BytesIO(base64.b64decode(encoded_image))).convert("RGB")
        scale = min(3.0, max(1.0, 2400 / max(image.width, 1)))
        if scale > 1:
            image = image.resize(
                (round(image.width * scale), round(image.height * scale)),
                Image.Resampling.LANCZOS,
            )
        image = ImageOps.autocontrast(image)
        image = ImageEnhance.Contrast(image).enhance(1.35)
        image = ImageEnhance.Sharpness(image).enhance(1.6)
        image = image.filter(ImageFilter.SHARPEN)
        output = io.BytesIO()
        image.save(output, format="PNG", optimize=True)
        encoded_image = base64.b64encode(output.getvalue()).decode("ascii")
    except Exception as exc:
        print(f"Pre-elaborazione OCR non disponibile, uso immagine originale: {exc}")

    response = requests.post(
        f"{OLLAMA_URL}/api/chat",
        json={
            "model": OLLAMA_OCR_MODEL,
            "stream": False,
            "messages": [{
                "role": "user",
                "content": (
                    "You are an OCR engine. Transcribe every legible word visible in this social "
                    "profile screenshot, including username, display name, bio, buttons, labels "
                    "and links. Preserve reading order and line breaks. Return transcription only. "
                    "Do not describe, translate, summarize, correct spelling, infer hidden text, "
                    "or add commentary. If no text is legible, answer exactly: NO TEXT."
                ),
                "images": [encoded_image],
            }],
            "options": {
                "temperature": 0,
                "num_ctx": 4096,
                "num_predict": 900,
            },
        },
        timeout=180,
    )
    if not response.ok:
        raise RuntimeError(f"Ollama OCR: {response.text}")
    return clean_social_ocr_text(response.json().get("message", {}).get("content", ""))


def extract_social_text_with_openai(image_input: Dict) -> str:
    if not openai_moderation_client:
        raise RuntimeError("OpenAI non configurato.")
    response = openai_moderation_client.chat.completions.create(
        model=OPENAI_VISION_MODEL,
        temperature=0,
        max_completion_tokens=700,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "Trascrivi tutto il testo leggibile nello screenshot del profilo social. "
                        "Mantieni le interruzioni di riga. Non descrivere, tradurre o dedurre testo "
                        "nascosto. Se non è leggibile alcun testo, rispondi NO TEXT."
                    ),
                },
                {"type": "image_url", "image_url": image_input.get("image_url", {})},
            ],
        }],
    )
    return clean_social_ocr_text(response.choices[0].message.content or "")


def extract_social_screenshot_texts(image_inputs: List[Dict]) -> Dict:
    extracted = []
    failed_count = 0
    last_error = ""
    for image_input in image_inputs[:8]:
        try:
            text = extract_social_text_with_rapidocr(image_input)
            if text:
                extracted.append(text)
        except Exception as exc:
            print(f"RapidOCR non riuscito, provo fallback visuale: {exc}")
            try:
                if VISION_PROVIDER == "ollama":
                    text = extract_social_text_with_ollama(image_input)
                elif VISION_PROVIDER == "openai":
                    text = extract_social_text_with_openai(image_input)
                else:
                    raise RuntimeError("Provider OCR di fallback non configurato.")
                if text:
                    extracted.append(text)
            except Exception as fallback_exc:
                failed_count += 1
                last_error = str(fallback_exc)
                print(f"OCR screenshot social non riuscito: {fallback_exc}")

    combined = clean_social_ocr_text("\n".join(extracted))
    status = "completed" if combined else ("failed" if failed_count else "no_text_found")
    return {
        "status": status,
        "provider": "rapidocr",
        "fallback_provider": VISION_PROVIDER,
        "screenshots_checked": min(len(image_inputs), 8),
        "screenshots_with_text": len(extracted),
        "failed_count": failed_count,
        "extracted_text": combined,
        "message": (
            f"OCR completato: testo leggibile trovato in {len(extracted)} screenshot."
            if combined
            else (
                f"OCR non completato. Dettaglio: {last_error}"
                if failed_count
                else "Nessun testo leggibile trovato negli screenshot."
            )
        ),
    }


def social_text_tokens(value: Any) -> List[str]:
    normalized = strip_accents(str(value or "").lower())
    ignored = {
        "sono", "della", "delle", "degli", "dello", "alla", "alle", "agli",
        "con", "per", "nel", "nella", "nelle", "the", "and", "for", "with",
        "profilo", "profile", "instagram", "facebook", "follower", "following",
    }
    return [
        token
        for token in re.findall(r"[a-z0-9+#.]{3,}", normalized)
        if token not in ignored
    ]


def select_social_bio_candidate(extracted_text: str) -> str:
    ignored = {
        "home", "search", "reels", "messages", "following", "followers", "posts",
        "seguiti", "follower", "post", "modifica profilo", "edit profile",
    }
    candidates = []
    for line in (extracted_text or "").splitlines():
        cleaned = re.sub(r"\s+", " ", line).strip()
        plain = strip_accents(cleaned.lower())
        if (
            len(cleaned) < 4
            or plain in ignored
            or re.fullmatch(r"[@#]?[a-z0-9_.]+", plain)
            or re.fullmatch(r"[\d\s.,km+]+", plain)
        ):
            continue
        candidates.append(cleaned)
        if len(candidates) >= 5:
            break
    return " | ".join(candidates)[:500]


def evaluate_social_profile_text(extracted_text: str, profile_type: str, user: Dict) -> Dict:
    text = clean_social_ocr_text(extracted_text)
    if not text:
        return {
            "status": "not_available",
            "profile_type": profile_type,
            "bio_candidate": "",
            "matched_role_terms": [],
            "suggestions": [],
            "message": "La bio o il testo del profilo non risultano leggibili nello screenshot.",
        }

    target_role = str(user.get("target_role") or "").strip()
    role_tokens = list(dict.fromkeys(social_text_tokens(target_role)))
    text_tokens = set(social_text_tokens(text))
    matched_role_terms = [token for token in role_tokens if token in text_tokens]
    professional_markers = {
        "student", "studentessa", "studente", "engineer", "ingegnere", "developer",
        "analyst", "analista", "designer", "manager", "researcher", "ricercatore",
        "data", "software", "marketing", "finance", "cybersecurity", "machine",
        "linkedin", "university", "universita",
    }
    has_professional_marker = bool(text_tokens.intersection(professional_markers))
    suggestions = []
    if target_role and not matched_role_terms:
        suggestions.append(f"Rendi esplicito nella bio il ruolo o l'ambito target: {target_role}.")
    if not has_professional_marker:
        suggestions.append("Aggiungi una breve identità professionale o accademica verificabile.")
    if not text_tokens.intersection({"linkedin"}):
        suggestions.append(
            "Se pertinente, inserisci un collegamento professionale verificabile come LinkedIn."
        )

    aligned = bool(matched_role_terms or has_professional_marker)
    return {
        "status": "aligned" if aligned else "review",
        "profile_type": profile_type,
        "bio_candidate": select_social_bio_candidate(text),
        "matched_role_terms": matched_role_terms[:8],
        "suggestions": suggestions[:3],
        "message": (
            f"Testo del profilo coerente almeno in parte con il ruolo {target_role}."
            if aligned and target_role
            else "Il testo è leggibile, ma non comunica chiaramente un posizionamento professionale."
        ),
    }


def extract_profile_name_candidates(
    sources: List[Dict[str, str]],
    kinds: set[str],
    fallback_values: Optional[List[str]] = None,
) -> Dict[str, Any]:
    candidates: List[str] = []
    for source in sources or []:
        if source.get("kind") not in kinds:
            continue
        for raw_value in [source.get("title", ""), source.get("content", ""), source.get("url", "")]:
            value = str(raw_value or "").strip()
            if not value:
                continue
            matches = re.findall(
                r"\b[A-ZÀ-ÖØ-Ý][A-Za-zÀ-ÖØ-öø-ÿ'`-]{1,}\s+[A-ZÀ-ÖØ-Ý][A-Za-zÀ-ÖØ-öø-ÿ'`-]{1,}\b",
                value,
            )
            candidates.extend(matches[:3])
    for value in fallback_values or []:
        item = str(value or "").strip()
        if item:
            candidates.append(item)
    unique_candidates = _unique_preserve_order(candidates)
    return {
        "display_name_candidate": unique_candidates[0] if unique_candidates else "",
        "all_candidates": unique_candidates[:5],
    }


def public_image_url_to_input(image_url: str) -> Optional[Dict]:
    if not is_safe_public_url(image_url):
        return None

    try:
        response = requests.get(
            image_url,
            headers={"User-Agent": "Mozilla/5.0 (compatible; CareerCoach/1.0; public-media-check)"},
            timeout=8,
            stream=True,
        )
        response.raise_for_status()
        content_type = response.headers.get("content-type", "").split(";", 1)[0].lower()
        if content_type not in {"image/jpeg", "image/png", "image/webp"}:
            return None

        content = bytearray()
        for chunk in response.iter_content(chunk_size=64 * 1024):
            content.extend(chunk)
            if len(content) > 5 * 1024 * 1024:
                return None
        if not content:
            return None

        encoded = base64.b64encode(bytes(content)).decode("ascii")
        return {
            "type": "image_url",
            "image_url": {"url": f"data:{content_type};base64,{encoded}"},
        }
    except Exception as exc:
        print(f"Download media pubblico non riuscito per '{image_url}': {exc}")
        return None


def analyze_public_social_media(user: Dict) -> Dict:
    discovered = collect_public_image_urls(user)
    instagram_images = [
        image
        for image in discovered["image_urls"]
        if "instagram.com" in (urllib.parse.urlparse(image["profile_url"]).hostname or "").lower()
    ]
    image_inputs = [
        image_input
        for image in instagram_images
        if (image_input := public_image_url_to_input(image["url"]))
    ]
    analysis = moderate_visual_inputs(image_inputs, "public_links", len(instagram_images))
    analysis["analyzed_content_count"] = 0
    analysis["analyzed_preview_count"] = analysis.get("analyzed_count", 0)
    analysis["pages_checked"] = discovered["pages_checked"]
    return analysis


def describe_visual_media_analysis(evidence: Dict, has_instagram: bool) -> str:
    media_analysis = evidence.get("visual_media_analysis") or {}
    if media_analysis.get("status") == "completed":
        return media_analysis["message"]
    if media_analysis.get("status") == "provider_not_configured":
        return media_analysis["message"]
    if has_instagram and evidence.get("instagram_metadata_found"):
        return (
            "Il profilo Instagram risulta rintracciabile sul web, ma non sono stati recuperati media "
            "analizzabili automaticamente. Puoi caricare uno o piu screenshot per completare il controllo."
        )
    if has_instagram:
        return (
            "Instagram e stato collegato, ma foto e post non risultano accessibili automaticamente. "
            "Puoi caricare uno o piu screenshot per completare il controllo."
        )
    return "Non sono stati trovati media pubblici da analizzare automaticamente."


def visual_media_finding_status(evidence: Dict) -> str:
    media_analysis = evidence.get("visual_media_analysis") or {}
    if media_analysis.get("status") == "completed" and not media_analysis.get("flagged_count"):
        return "success"
    return "warning"


VISUAL_PROFILE_LABELS = {
    "instagram": "Instagram",
    "facebook": "Facebook",
    "other": "Altro profilo",
}

SOCIAL_SCREENSHOT_REJECTION_MESSAGE = (
    "Questo sembra uno screenshot di un CV o di un documento. Per l'analisi digitale carica "
    "screenshot di profili social o piattaforme professionali, ad esempio LinkedIn o Instagram."
)

SOCIAL_SCREENSHOT_MIN_LENGTH = 18


def classify_social_screenshot_text(extracted_text: str) -> Dict[str, Any]:
    text = clean_social_ocr_text(extracted_text)
    normalized = normalize_plain_text(text)
    if not text:
        return {
            "valid": False,
            "reason": "Lo screenshot non contiene testo leggibile sufficiente.",
            "kind": "unreadable",
        }

    social_markers = [
        "linkedin",
        "github",
        "behance",
        "dribbble",
        "portfolio",
        "projects",
        "followers",
        "following",
        "posts",
        "repository",
        "commit",
        "profile",
        "profilo",
        "bio",
        "website",
        "www.",
        "http://",
        "https://",
        "@",
    ]
    cv_markers = [
        "curriculum vitae",
        "cv",
        "resume",
        "esperienze professionali",
        "esperienza professionale",
        "formazione",
        "istruzione",
        "competenze",
        "hard skills",
        "soft skills",
        "lingue",
        "certificazioni",
        "progetti",
        "obiettivo professionale",
        "profilo professionale",
        "email",
        "telefono",
        "indirizzo",
    ]

    cv_hits = sum(1 for marker in cv_markers if marker in normalized)
    social_hits = sum(1 for marker in social_markers if marker in normalized)
    has_typical_layout = bool(re.search(r"\b(?:ruolo|azienda|date|periodo|esperienze|formazione)\b", normalized))
    has_table_like_layout = bool(re.search(r"\b(?:tabella|colonna|column|row)\b", normalized))

    if cv_hits >= 2 and social_hits == 0:
        return {
            "valid": False,
            "reason": SOCIAL_SCREENSHOT_REJECTION_MESSAGE,
            "kind": "cv_or_document",
        }

    if has_table_like_layout and cv_hits >= 1 and social_hits == 0:
        return {
            "valid": False,
            "reason": SOCIAL_SCREENSHOT_REJECTION_MESSAGE,
            "kind": "document_layout",
        }

    if social_hits >= 1 and cv_hits == 0:
        return {
            "valid": True,
            "reason": "Screenshot di profilo social o piattaforma professionale riconosciuto.",
            "kind": "social_profile",
        }

    if social_hits >= 1 and cv_hits <= 1:
        return {
            "valid": True,
            "reason": "Screenshot compatibile con una presenza digitale professionale.",
            "kind": "social_profile",
        }

    if cv_hits >= 1 and len(normalized) >= SOCIAL_SCREENSHOT_MIN_LENGTH:
        return {
            "valid": False,
            "reason": SOCIAL_SCREENSHOT_REJECTION_MESSAGE,
            "kind": "cv_or_document",
        }

    return {
        "valid": False,
        "reason": SOCIAL_SCREENSHOT_REJECTION_MESSAGE,
        "kind": "unknown",
    }


def calculate_visual_score_adjustment(profile_analyses: Dict[str, Dict]) -> int:
    adjustment = 0
    for analysis in profile_analyses.values():
        if analysis.get("status") != "completed" or analysis.get("analyzed_content_count", 0) <= 0:
            continue
        flagged_count = int(analysis.get("flagged_count", 0) or 0)
        if flagged_count <= 0:
            continue
        sensitive_count = int(analysis.get("sensitive_flagged_count", 0) or 0)
        generic_count = max(0, flagged_count - sensitive_count)
        adjustment -= min(24, (sensitive_count * 12) + (generic_count * 5))
    return max(-35, min(0, adjustment))


def calculate_social_screenshot_score_adjustment(batches: List[Dict[str, Any]]) -> int:
    positive_adjustment = 0
    negative_adjustment = 0
    for batch in batches or []:
        if not isinstance(batch, dict):
            continue
        if not batch.get("valid"):
            continue
        positive_adjustment += 2
        flagged_count = int(batch.get("flagged_count", 0) or 0)
        sensitive_count = int(batch.get("sensitive_flagged_count", 0) or 0)
        generic_count = max(0, flagged_count - sensitive_count)
        negative_adjustment += min(18, (sensitive_count * 10) + (generic_count * 4))
    adjustment = min(8, positive_adjustment) - min(28, negative_adjustment)
    return max(-28, min(8, adjustment))


def describe_profile_screenshot_analyses(profile_analyses: Dict[str, Dict]) -> str:
    descriptions = []
    for profile_type, analysis in profile_analyses.items():
        label = VISUAL_PROFILE_LABELS.get(profile_type, "Profilo")
        descriptions.append(f"{label}: {analysis.get('message', '')}")
    return " ".join(descriptions)


def classify_additional_link(url: str, sources: List[Dict[str, str]], identity: Dict) -> Dict:
    if not url:
        return {
            "status": "not_provided",
            "type": "none",
            "platform": "none",
            "message": "Non hai inserito un link aggiuntivo.",
        }

    hostname = (urllib.parse.urlparse(normalize_public_profile_url(url)).hostname or "").lower().removeprefix("www.")
    is_github_link = hostname == "github.com"

    if not has_public_other_profile_signals(sources):
        blocked_note = (
            " Facebook spesso richiede il login e impedisce il recupero automatico dei contenuti."
            if hostname.endswith("facebook.com")
            else ""
        )
        return {
            "status": "unverified",
            "type": "unknown",
            "platform": "github" if is_github_link else "generic",
            "message": (
                "Non e stato possibile analizzare direttamente il profilo GitHub dal link fornito. "
                "Carica uno screenshot del profilo GitHub per rendere l'analisi digitale piu completa."
                if is_github_link
                else "Il link aggiuntivo e stato registrato, ma non risultano contenuti pubblici accessibili."
                f"{blocked_note}"
            ),
        }

    social_hosts = {
        "linkedin.com", "instagram.com", "x.com", "twitter.com", "github.com", "behance.net", "dribbble.com",
        "facebook.com", "tiktok.com", "youtube.com", "medium.com",
    }
    link_type = "profilo personale" if hostname in social_hosts else "sito o pagina pubblica"
    platform = "github" if is_github_link else "generic"

    if identity["status"] == "matched":
        return {
            "status": "matched",
            "type": link_type,
            "platform": platform,
            "message": (
                "Il profilo GitHub e stato analizzato tramite il link fornito."
                if is_github_link
                else f"Il link aggiuntivo risulta essere un {link_type} compatibile con il nome del candidato. "
            ) + (
                " "
                "L'analisi usa solo testo pubblico indicizzato: eventuali foto, video e post non sono stati analizzati."
                if hostname in social_hosts
                else ""
            ),
        }

    return {
        "status": identity["status"],
        "type": link_type,
        "platform": platform,
        "message": (
            "Non e stato possibile analizzare direttamente il profilo GitHub dal link fornito. "
            "Carica uno screenshot del profilo GitHub per rendere l'analisi digitale piu completa."
            if is_github_link and identity["status"] == "unverified"
            else f"Il link risulta essere un {link_type}, ma non posso attribuirlo con certezza al candidato."
            if identity["status"] == "unverified"
            else f"Il link risulta essere un {link_type}, ma le evidenze sembrano appartenere a un'altra persona."
        ),
    }


def evaluate_profile_identity(user: Dict, sources: List[Dict[str, str]], kinds: set[str], label: str) -> Dict:
    candidate_tokens = normalize_identity_tokens(user.get("name"))
    evidence = [
        (source.get("title", label), source.get("content", ""))
        for source in sources
        if source.get("kind") in kinds
    ]

    if not evidence or not candidate_tokens:
        return {
            "status": "unverified",
            "message": f"Non ci sono abbastanza dati per verificare che il {label} appartenga al candidato.",
        }

    matching_sources = []
    conflicting_sources = []
    for title, text in evidence:
        source_tokens = normalize_identity_tokens(f"{title} {text}")
        common_tokens = candidate_tokens.intersection(source_tokens)
        if len(common_tokens) >= min(2, len(candidate_tokens)):
            matching_sources.append(title)
        elif source_tokens:
            conflicting_sources.append(title)

    if matching_sources:
        return {
            "status": "matched",
            "message": f"Le evidenze disponibili per il {label} sono compatibili con il nome del candidato.",
        }

    if conflicting_sources:
        return {
            "status": "mismatch",
            "message": f"Le evidenze disponibili per il {label} non sembrano appartenere al candidato. Verifica il dato inserito.",
        }

    return {
        "status": "unverified",
        "message": f"Non ci sono abbastanza dati per verificare che il {label} appartenga al candidato.",
    }


def evaluate_official_profile_identity(user: Dict, official_sources: List[Dict], platform: str) -> Dict:
    matching_sources = [source for source in official_sources if source.get("platform") == platform]
    if not matching_sources:
        return {
            "status": "not_connected",
            "message": f"Nessun collegamento OAuth ufficiale disponibile per {platform}.",
        }

    candidate_tokens = normalize_identity_tokens(user.get("name"))
    provider_name = matching_sources[0].get("identity", {}).get("display_name", "")
    provider_tokens = normalize_identity_tokens(provider_name)

    if not candidate_tokens or not provider_tokens:
        return {
            "status": "unverified",
            "message": f"{platform} e collegato tramite OAuth, ma non ci sono abbastanza dati per confrontare il nome.",
        }

    common_tokens = candidate_tokens.intersection(provider_tokens)
    if len(common_tokens) >= min(2, len(candidate_tokens)):
        return {
            "status": "matched",
            "message": f"Identita {platform} verificata tramite OAuth ufficiale: il nome e compatibile con il candidato.",
        }

    return {
        "status": "mismatch",
        "message": f"Il profilo {platform} collegato via OAuth non sembra compatibile con il nome del candidato.",
    }


def evaluate_cv_profile_name_match(
    cv_text: str,
    sources: List[Dict[str, str]],
    kinds: set[str],
    label: str,
    fallback_values: Optional[List[str]] = None,
) -> Dict[str, Any]:
    detected = extract_candidate_name_from_cv(cv_text or "")
    detected_name = str(detected.get("name") or "").strip()
    cv_tokens = normalize_identity_tokens(detected_name)
    ordered_cv_tokens = [token for token in normalize_name(detected_name).split() if token]
    profile_names = extract_profile_name_candidates(sources, kinds, fallback_values)
    if len(cv_tokens) < 2:
        return {
            "status": "unverified",
            "detected_name": detected_name,
            "profile_name_candidate": profile_names.get("display_name_candidate", ""),
            "profile_name_candidates": profile_names.get("all_candidates", []),
            "message": f"Non sono riuscito a leggere con affidabilita nome e cognome dal CV per confrontarli con {label}.",
        }

    evidence_tokens = set()
    evidence_text_fragments = []
    for source in sources or []:
        if source.get("kind") not in kinds:
            continue
        title = str(source.get("title", "") or "")
        content = str(source.get("content", "") or "")
        url = str(source.get("url", "") or "")
        evidence_text_fragments.extend([title, content, url])
        evidence_tokens.update(normalize_identity_tokens(title))
        evidence_tokens.update(normalize_identity_tokens(content))
        evidence_tokens.update(normalize_identity_tokens(url))

    for value in fallback_values or []:
        evidence_text_fragments.append(str(value or ""))
        evidence_tokens.update(normalize_identity_tokens(value))

    if not evidence_tokens:
        return {
            "status": "unverified",
            "detected_name": detected_name,
            "profile_name_candidate": profile_names.get("display_name_candidate", ""),
            "profile_name_candidates": profile_names.get("all_candidates", []),
            "message": f"Non ci sono abbastanza elementi pubblici per confrontare il nome del CV con {label}.",
        }

    matches = sum(1 for token in cv_tokens if token in evidence_tokens)
    compact_evidence = strip_accents(" ".join(evidence_text_fragments).lower()).replace(" ", "")
    if matches == 0 and len(ordered_cv_tokens) >= 2:
        first_token = ordered_cv_tokens[0]
        last_token = ordered_cv_tokens[-1]
        if first_token[:4] in compact_evidence and last_token[:4] in compact_evidence:
            matches = 1
    if matches >= 2:
        return {
            "status": "matched",
            "detected_name": detected_name,
            "profile_name_candidate": profile_names.get("display_name_candidate", ""),
            "profile_name_candidates": profile_names.get("all_candidates", []),
            "message": f"Il nome rilevato nel CV risulta coerente con {label}.",
        }
    if matches == 1:
        return {
            "status": "similar",
            "detected_name": detected_name,
            "profile_name_candidate": profile_names.get("display_name_candidate", ""),
            "profile_name_candidates": profile_names.get("all_candidates", []),
            "message": f"Il nome del CV e solo parzialmente coerente con {label}, ma ci sono alcuni elementi compatibili.",
        }
    return {
        "status": "mismatch",
        "detected_name": detected_name,
        "profile_name_candidate": profile_names.get("display_name_candidate", ""),
        "profile_name_candidates": profile_names.get("all_candidates", []),
        "message": (
            f"Attenzione: il nome rilevato nel CV non sembra corrispondere pienamente a {label}. "
            "Verifica che i profili appartengano alla stessa persona. L'analisi verra comunque eseguita."
        ),
    }


def build_github_profile_evidence(user: Dict, sources: List[Dict[str, str]]) -> Dict[str, Any]:
    url = normalize_public_profile_url(user.get("portfolio_url", ""))
    hostname = (urllib.parse.urlparse(url).hostname or "").lower().removeprefix("www.")
    is_github = hostname == "github.com"
    path_parts = [part for part in urllib.parse.urlparse(url).path.strip("/").split("/") if part]
    username = path_parts[0] if is_github and path_parts else ""
    github_sources = [
        source
        for source in sources or []
        if source.get("kind") in {"other_profile_reference", "other_profile_public_snippet"}
    ]
    text_blob = " ".join(
        f"{source.get('title', '')} {source.get('content', '')} {source.get('url', '')}"
        for source in github_sources
    )
    normalized_blob = normalize_plain_text(text_blob)
    repositories_visible = any(
        marker in normalized_blob
        for marker in [
            "repository", "repositories", "repo", "repos", "pinned", "stars", "followers",
            "readme", "python", "javascript", "typescript", "java", "docker", "machine learning",
        ]
    )
    bio_coherent = any(
        token in normalized_blob
        for token in social_text_tokens(user.get("target_role"))
    ) or any(
        marker in normalized_blob
        for marker in ["developer", "engineer", "analyst", "data", "software", "ai", "ml"]
    )
    profile_curated = repositories_visible or any(
        marker in normalized_blob for marker in ["readme", "pinned", "contributions", "followers", "following"]
    )
    public_name_match = evaluate_cv_profile_name_match(
        user.get("cv_text") or "",
        sources,
        {"other_profile_public_snippet", "other_profile_reference"},
        "il profilo GitHub",
        fallback_values=[username],
    )
    username_match = (
        match_personal_brand_username(user, username)
        if is_github and username
        else {
            "matched": False,
            "confidence": 0.0,
            "method": "not_github",
            "message": "Nessun username GitHub disponibile.",
        }
    )
    if public_name_match.get("status") == "matched":
        name_match = public_name_match
    elif username_match.get("matched"):
        name_match = {
            "status": "matched",
            "detected_name": resolve_candidate_name(user),
            "profile_name_candidate": username,
            "message": f"Username GitHub coerente con il candidato. {username_match['message']}",
            **username_match,
        }
    else:
        name_match = public_name_match
    return {
        "is_github_link": is_github,
        "username": username,
        "public_accessible": bool(is_github and github_sources),
        "analyzed_via_link": bool(is_github and github_sources),
        "requires_screenshot_fallback": bool(is_github and not github_sources),
        "repositories_visible": repositories_visible,
        "bio_coherent": bio_coherent,
        "profile_curated": profile_curated,
        "cv_name_match": name_match,
        "username_match": username_match,
        "snippet_count": len(github_sources),
    }


def infer_instagram_visibility(user: Dict, sources: List[Dict[str, str]], evidence: Dict[str, Any]) -> Dict[str, Any]:
    has_instagram = bool(normalize_instagram_handle(user.get("instagram_handle")))
    if not has_instagram:
        return {
            "status": "not_provided",
            "message": "Non hai collegato un profilo Instagram.",
        }
    has_metadata = has_public_instagram_metadata(sources)
    has_screenshots = any(
        item.get("profile_type") == "instagram"
        for item in (evidence.get("social_screenshot_batches") or [])
    )
    if has_metadata:
        return {
            "status": "public",
            "message": (
                "Il profilo Instagram risulta pubblico. Questo puo migliorare leggermente la valutazione "
                "della presenza digitale, se i contenuti sono coerenti e professionali."
            ),
        }
    if has_screenshots:
        return {
            "status": "private_or_limited",
            "message": (
                "Il profilo Instagram non risulta pienamente accessibile dal web pubblico. "
                "Gli screenshot aiutano comunque a completare il controllo."
            ),
        }
    return {
        "status": "private",
        "message": (
            "Il profilo Instagram risulta privato. Questo non e necessariamente un problema, "
            "ma potrebbe ridurre la visibilita della tua presenza digitale per un recruiter."
        ),
    }


def summarize_screenshot_evidence(evidence: Dict[str, Any]) -> Dict[str, Any]:
    batches = [
        batch
        for batch in (evidence.get("social_screenshot_batches") or [])
        if isinstance(batch, dict)
    ]
    valid_batches = [batch for batch in batches if batch.get("valid")]
    screenshots_count = sum(int(batch.get("analyzed_count", 0) or 0) for batch in valid_batches)
    flagged_count = sum(int(batch.get("flagged_count", 0) or 0) for batch in valid_batches)
    sensitive_flagged_count = sum(
        int(batch.get("sensitive_flagged_count", 0) or 0)
        for batch in valid_batches
    )
    return {
        "uploaded": bool(batches),
        "valid_uploaded": bool(valid_batches),
        "count": screenshots_count,
        "flagged_count": flagged_count,
        "sensitive_flagged_count": sensitive_flagged_count,
        "safe_content": bool(valid_batches) and flagged_count == 0,
        "profile_types": sorted({str(batch.get("profile_type") or "") for batch in valid_batches if batch.get("profile_type")}),
        "message": (
            f"Sono stati caricati {screenshots_count} screenshot validi di profili digitali."
            if valid_batches and screenshots_count
            else "Non sono stati caricati screenshot dei profili digitali. L'analisi e stata eseguita con le informazioni disponibili, ma il punteggio digitale potrebbe essere leggermente inferiore perche mancano elementi visivi di verifica."
        ),
    }


SOCIAL_BRANDING_AFFIXES = (
    "official", "officiale", "real", "thereal", "iam", "im", "its", "the",
    "mr", "mrs", "miss", "dr", "prof", "hello", "hey", "thisis",
)


def extract_social_username(value: str, expected_host: str = "") -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    normalized_url = raw if "://" in raw else f"https://{raw.lstrip('@')}"
    parsed = urllib.parse.urlparse(normalized_url)
    hostname = (parsed.hostname or "").lower().removeprefix("www.")
    if hostname and (not expected_host or hostname == expected_host):
        path_parts = [urllib.parse.unquote(part) for part in parsed.path.split("/") if part]
        if path_parts:
            return path_parts[0].strip().lower()
    return raw.lstrip("@").split("?")[0].strip("/").lower()


def remove_branding_affixes(value: str) -> str:
    cleaned = value
    changed = True
    while changed and cleaned:
        changed = False
        for affix in SOCIAL_BRANDING_AFFIXES:
            if cleaned.startswith(affix) and len(cleaned) - len(affix) >= 4:
                cleaned = cleaned[len(affix):]
                changed = True
            if cleaned.endswith(affix) and len(cleaned) - len(affix) >= 4:
                cleaned = cleaned[:-len(affix)]
                changed = True
    return cleaned


def resolve_candidate_name(user: Dict) -> str:
    cv_detected = extract_candidate_name_from_cv(user.get("cv_text") or "")
    detected_name = str(cv_detected.get("name") or "").strip()
    return detected_name or str(user.get("name") or "").strip()


def match_personal_brand_username(user: Dict, username: str) -> Dict[str, Any]:
    detected_name = resolve_candidate_name(user)
    if not detected_name:
        return {
            "matched": False,
            "confidence": 0.0,
            "method": "missing_candidate_name",
            "message": "Impossibile rilevare il nome dell'utente per la verifica di corrispondenza.",
        }

    ordered_tokens = [
        token
        for token in re.findall(r"[a-z]+", strip_accents(detected_name).lower())
        if len(token) >= 2
    ]
    if len(ordered_tokens) < 2:
        return {
            "matched": False,
            "confidence": 0.0,
            "method": "incomplete_candidate_name",
            "message": "Servono nome e cognome per verificare lo username.",
        }

    first_name = ordered_tokens[0]
    surname = ordered_tokens[-1]
    alphabetic_username = "".join(re.findall(r"[a-z]+", strip_accents(username).lower()))
    branding_core = remove_branding_affixes(alphabetic_username)
    if not branding_core:
        return {
            "matched": False,
            "confidence": 0.0,
            "method": "empty_username",
            "message": "Lo username non contiene caratteri utili per il confronto.",
        }

    direct_variants = (first_name + surname, surname + first_name)
    if any(variant in branding_core for variant in direct_variants):
        return {
            "matched": True,
            "confidence": 1.0,
            "method": "full_name",
            "message": "Lo username contiene nome e cognome, anche in ordine invertito.",
        }

    best_structural_match = None
    for surname_length in range(len(surname), 2, -1):
        surname_root = surname[:surname_length]
        position = branding_core.find(surname_root)
        if position < 0:
            continue
        remainder = branding_core[:position] + branding_core[position + surname_length:]
        name_prefix_length = max(
            (
                length
                for length in range(len(first_name), 0, -1)
                if first_name[:length] in remainder
            ),
            default=0,
        )
        if name_prefix_length:
            confidence = min(
                0.98,
                0.45
                + 0.35 * (surname_length / len(surname))
                + 0.20 * (name_prefix_length / len(first_name)),
            )
            best_structural_match = {
                "matched": True,
                "confidence": round(confidence, 3),
                "method": "surname_root_and_name_prefix",
                "surname_root": surname_root,
                "name_prefix": first_name[:name_prefix_length],
                "message": (
                    "Lo username usa una variante riconoscibile di nome e cognome "
                    f"({first_name[:name_prefix_length]} + {surname_root})."
                ),
            }
            break
    if best_structural_match:
        return best_structural_match

    similarity = max(
        SequenceMatcher(None, branding_core, variant).ratio()
        for variant in direct_variants
    )
    if similarity >= 0.78:
        return {
            "matched": True,
            "confidence": round(similarity, 3),
            "method": "fuzzy_similarity",
            "message": "Lo username e una variante sufficientemente simile al nome del candidato.",
        }

    return {
        "matched": False,
        "confidence": round(similarity, 3),
        "method": "no_identity_signal",
        "message": "Lo username non contiene elementi sufficienti riconducibili a nome e cognome.",
    }


def verify_instagram_slug(user: Dict, instagram_handle: str) -> Dict[str, Any]:
    if not instagram_handle:
        return {"matched": False, "message": "Link o handle Instagram non presente."}

    slug = extract_social_username(instagram_handle, "instagram.com")
    result = match_personal_brand_username(user, slug)
    result["slug"] = slug
    if result["matched"]:
        result["message"] = f"Slug Instagram coerente. {result['message']}"
        return result

    detected_name = resolve_candidate_name(user)
    ordered_tokens = [
        token
        for token in re.findall(r"[a-z]+", strip_accents(detected_name).lower())
        if len(token) >= 2
    ]
    suggested_slug = "".join(ordered_tokens[:1] + ordered_tokens[-1:])
    result["message"] = (
        f"Il nome e cognome non sono riconoscibili nello slug Instagram '{slug}'. "
        f"{result['message']}"
    )
    result["coach_tip"] = (
        f"Valuta uno username riconoscibile, ad esempio instagram.com/{suggested_slug}, "
        "anche con iniziali, separatori, numeri o prefissi di branding."
    )
    return result


PROFESSIONAL_SKILL_TERMS = {
    "python", "sql", "java", "javascript", "typescript", "c++", "c#", "react",
    "angular", "vue", "node.js", "fastapi", "django", "flask", "spring", "docker",
    "kubernetes", "aws", "azure", "gcp", "git", "excel", "power bi", "tableau",
    "machine learning", "deep learning", "pytorch", "tensorflow", "opencv",
    "data analysis", "data analytics", "project management", "agile", "scrum",
    "salesforce", "hubspot", "seo", "sem", "figma", "photoshop",
}

PROFESSIONAL_ROLE_TERMS = {
    "analyst", "developer", "engineer", "manager", "consultant", "specialist",
    "designer", "researcher", "recruiter", "marketer", "intern", "tirocinio",
    "stage", "freelance", "founder", "coordinator", "responsabile",
}


def extract_professional_profile_signals(text: str) -> Dict[str, set[str]]:
    normalized = normalize_plain_text(strip_accents(text or ""))
    return {
        "skills": {term for term in PROFESSIONAL_SKILL_TERMS if term in normalized},
        "roles": {term for term in PROFESSIONAL_ROLE_TERMS if term in normalized},
        "years": set(re.findall(r"\b(?:19|20)\d{2}\b", normalized)),
    }


def linkedin_professional_text(sources: List[Dict[str, str]], export_text: str = "") -> str:
    parts = [export_text] if export_text else []
    parts.extend(
        str(source.get("content") or "")
        for source in sources or []
        if source.get("kind") in {"linkedin_export", "linkedin_public_snippet"}
    )
    return clean_social_ocr_text("\n".join(parts))


def evaluate_linkedin_cv_coherence(cv_text: str, linkedin_text: str) -> Dict[str, Any]:
    if not cv_text or not linkedin_text:
        return {
            "score_adjustment": 0,
            "status": "unverified",
            "message": "Confronto non eseguibile per mancanza di testi del CV o di LinkedIn.",
            "details": []
        }
    
    cv_signals = extract_professional_profile_signals(cv_text)
    linkedin_signals = extract_professional_profile_signals(linkedin_text)
    details = []
    score_adj = 0

    shared_skills = cv_signals["skills"] & linkedin_signals["skills"]
    missing_skills = cv_signals["skills"] - linkedin_signals["skills"]
    if shared_skills:
        score_adj += min(4, len(shared_skills))
        details.append(f"Competenze allineate: {', '.join(sorted(shared_skills))}")
    if cv_signals["skills"] and not shared_skills:
        score_adj -= 3
        details.append("Non risultano competenze in comune tra CV e dati LinkedIn accessibili")
    elif missing_skills:
        score_adj -= min(2, len(missing_skills))
        details.append(f"Competenze presenti solo nel CV: {', '.join(sorted(missing_skills)[:6])}")

    shared_roles = cv_signals["roles"] & linkedin_signals["roles"]
    if shared_roles:
        score_adj += min(3, len(shared_roles))
        details.append(f"Esperienze o ruoli coerenti: {', '.join(sorted(shared_roles))}")
    elif cv_signals["roles"] and linkedin_signals["roles"]:
        score_adj -= 3
        details.append("I ruoli professionali rilevati nel CV e su LinkedIn non coincidono")

    shared_years = cv_signals["years"] & linkedin_signals["years"]
    linkedin_only_years = linkedin_signals["years"] - cv_signals["years"]
    if shared_years:
        score_adj += min(3, len(shared_years))
        details.append(f"Date in comune: {', '.join(sorted(shared_years))}")
    if linkedin_only_years:
        score_adj -= min(2, len(linkedin_only_years))
        details.append(f"Date LinkedIn da verificare nel CV: {', '.join(sorted(linkedin_only_years))}")

    score_adj = max(-10, min(10, score_adj))
    status = "success" if score_adj >= 3 else "warning" if score_adj < 0 else "review"
    message = "Analisi coerenza LinkedIn vs CV completata. "
    if score_adj >= 3:
        message += "Le informazioni accessibili risultano complessivamente allineate. "
    elif score_adj < 0:
        message += "Sono presenti elementi da verificare. "
    else:
        message += "I dati accessibili non sono sufficienti per confermare una piena coerenza. "
    message += "; ".join(details)
        
    return {
        "score_adjustment": score_adj,
        "status": status,
        "message": message,
        "details": details,
        "matched_skills": sorted(shared_skills),
        "matched_roles": sorted(shared_roles),
        "matched_years": sorted(shared_years),
    }


def build_analysis_evidence(user: Dict, sources: List[Dict[str, str]]) -> Dict:
    official_profile_sources = build_official_profile_sources(user)
    linkedin_export_identity = evaluate_profile_identity(user, sources, {"linkedin_export"}, "PDF LinkedIn")
    linkedin_public_identity = evaluate_profile_identity(user, sources, {"linkedin_public_snippet"}, "profilo LinkedIn pubblico")
    linkedin_official_identity = evaluate_official_profile_identity(user, official_profile_sources, "linkedin")
    instagram_identity = evaluate_profile_identity(user, sources, {"instagram_public_metadata"}, "profilo Instagram")
    linkedin_export_verified = linkedin_export_identity["status"] == "matched"
    linkedin_public_verified = linkedin_public_identity["status"] == "matched"
    linkedin_official_verified = linkedin_official_identity["status"] == "matched"
    linkedin_public_link_present = bool(user.get("linkedin_url"))
    if linkedin_public_link_present:
        linkedin_identity = linkedin_public_identity
        linkedin_verified = linkedin_public_verified
    elif linkedin_export_verified or user.get("linkedin_profile_text"):
        linkedin_identity = linkedin_export_identity
        linkedin_verified = linkedin_export_verified
    else:
        linkedin_identity = linkedin_official_identity
        linkedin_verified = False
    visual_media_analysis = user.get("visual_media_analysis") or {
        "status": "not_requested",
        "discovered_count": 0,
        "analyzed_count": 0,
        "flagged_count": 0,
        "sensitive_flagged_count": 0,
    }
    social_screenshot_batches = list((user.get("digital_analysis") or {}).get("analysis_evidence", {}).get("social_screenshot_batches", []))
    cv_linkedin_match = evaluate_cv_profile_name_match(
        user.get("cv_text") or "",
        sources,
        {"linkedin_export", "linkedin_public_snippet", "linkedin_reference"},
        "il profilo LinkedIn",
        fallback_values=[user.get("linkedin_url", "")],
    )
    cv_instagram_match = evaluate_cv_profile_name_match(
        user.get("cv_text") or "",
        sources,
        {"instagram_public_metadata", "instagram_reference"},
        "il profilo Instagram",
        fallback_values=[user.get("instagram_handle", "")],
    )
    screenshot_summary = summarize_screenshot_evidence({
        "social_screenshot_batches": social_screenshot_batches,
    })
    github_profile = build_github_profile_evidence(user, sources)
    verified_profiles = [
        profile
        for profile, verified in [
            ("linkedin", linkedin_verified),
            ("instagram", instagram_identity["status"] == "matched"),
            ("github", github_profile.get("cv_name_match", {}).get("status") == "matched"),
        ]
        if verified
    ]
    base_evidence = {
        "social_screenshot_batches": social_screenshot_batches,
    }
    instagram_visibility = infer_instagram_visibility(user, sources, base_evidence)
    instagram_slug_verification = verify_instagram_slug(user, user.get("instagram_handle", ""))
    linkedin_cv_text = linkedin_professional_text(
        sources,
        user.get("linkedin_profile_text") or "",
    )
    linkedin_cv_coherence = evaluate_linkedin_cv_coherence(
        user.get("cv_text") or "",
        linkedin_cv_text,
    )
    linkedin_link_provided = bool(user.get("linkedin_url"))
    instagram_link_provided = bool(normalize_instagram_handle(user.get("instagram_handle")))
    github_link_provided = bool(github_profile.get("is_github_link"))
    return {
        "cv_profile_loaded": bool(user.get("cv_text")),
        "cv_filename": user.get("cv_filename") or "",
        "target_role": user.get("target_role") or "",
        "cv_detected_name": extract_candidate_name_from_cv(user.get("cv_text") or ""),
        "instagram_slug_verification": instagram_slug_verification,
        "github_profile": github_profile,
        "cv_github_name_match": github_profile.get("cv_name_match", {}),
        "linkedin_cv_coherence": linkedin_cv_coherence,
        "linkedin_identity": linkedin_identity,
        "linkedin_export_identity": linkedin_export_identity,
        "linkedin_public_identity": linkedin_public_identity,
        "linkedin_official_identity": linkedin_official_identity,
        "cv_linkedin_name_match": cv_linkedin_match,
        "instagram_identity": instagram_identity,
        "cv_instagram_name_match": cv_instagram_match,
        "instagram_visibility": instagram_visibility,
        "instagram_metadata_found": has_public_instagram_metadata(sources),
        "instagram_media_analyzed": visual_media_analysis.get("analyzed_content_count", 0) > 0,
        "public_preview_analyzed": visual_media_analysis.get("analyzed_preview_count", 0) > 0,
        "visual_media_analysis": visual_media_analysis,
        "official_profile_sources": official_profile_sources,
        "official_profile_source_count": len(official_profile_sources),
        "official_profile_capabilities": OFFICIAL_PROFILE_CAPABILITIES,
        "linkedin_export_compared": bool(user.get("linkedin_profile_text")),
        "linkedin_export_filename": user.get("linkedin_profile_filename") or "",
        "linkedin_export_verified": linkedin_export_verified,
        "linkedin_official_verified": linkedin_official_verified,
        "linkedin_public_link_present": linkedin_public_link_present,
        "linkedin_link_provided": linkedin_link_provided,
        "instagram_link_provided": instagram_link_provided,
        "github_link_provided": github_link_provided,
        "provided_link_count": (
            int(linkedin_link_provided)
            + int(instagram_link_provided)
            + int(github_link_provided)
        ),
        "linkedin_public_verified": linkedin_public_verified,
        "linkedin_public_snippet_found": any(source.get("kind") == "linkedin_public_snippet" for source in sources),
        "screenshots_summary": screenshot_summary,
        "verified_profiles": verified_profiles,
        "verified_profile_count": len(verified_profiles),
        "can_compare_with_cv": bool(
            user.get("cv_text")
            and (
                linkedin_link_provided
                or instagram_link_provided
                or github_link_provided
            )
        ),
        "zero_score_reason": (
            ""
            if (
                linkedin_link_provided
                or instagram_link_provided
                or github_link_provided
            )
            else "Non hai inserito alcun link LinkedIn, Instagram o GitHub: il punteggio di coerenza digitale e 0."
        ),
    }


def describe_linkedin_evidence(evidence: Dict) -> str:
    messages = []

    if evidence["linkedin_official_identity"]["status"] != "not_connected":
        messages.append(f"OAuth LinkedIn ufficiale: {evidence['linkedin_official_identity']['message']}")

    if evidence["linkedin_export_compared"]:
        messages.append(f"PDF LinkedIn: {evidence['linkedin_export_identity']['message']}")

    if evidence["linkedin_public_link_present"]:
        messages.append(f"Link pubblico LinkedIn: {evidence['linkedin_public_identity']['message']}")

    return " ".join(messages) or "LinkedIn non e stato collegato tramite link pubblico, PDF o OAuth."


def compute_digital_presence_score(evidence: Dict[str, Any]) -> int:
    linkedin_provided = bool(
        evidence.get("linkedin_link_provided")
        or evidence.get("linkedin_public_link_present")
    )
    instagram_provided = bool(evidence.get("instagram_link_provided"))
    github_provided = bool(evidence.get("github_link_provided"))
    if not (linkedin_provided or instagram_provided or github_provided):
        return 0

    score = 10 if evidence.get("cv_profile_loaded") else 0
    score += 15 if linkedin_provided else 0
    score += 10 if instagram_provided else 0
    score += 10 if github_provided else 0

    if evidence.get("linkedin_export_verified"):
        score += 8
    elif evidence.get("linkedin_public_verified"):
        score += 8
    elif evidence.get("linkedin_official_verified"):
        score += 4

    linkedin_match_status = str((evidence.get("cv_linkedin_name_match") or {}).get("status") or "")
    if linkedin_match_status == "matched":
        score += 8
    elif linkedin_match_status == "similar":
        score += 3
    elif linkedin_match_status == "mismatch":
        score -= 8

    instagram_slug_match = evidence.get("instagram_slug_verification") or {}
    if instagram_provided and instagram_slug_match.get("matched"):
        score += 6
    elif instagram_provided:
        score -= 5

    github_match_status = str((evidence.get("cv_github_name_match") or {}).get("status") or "")
    if github_provided and github_match_status == "matched":
        score += 6
    elif github_provided and github_match_status == "similar":
        score += 2
    elif github_provided and github_match_status == "mismatch":
        score -= 5

    instagram_visibility = str((evidence.get("instagram_visibility") or {}).get("status") or "")
    if instagram_visibility == "public":
        score += 3
    elif instagram_visibility == "private":
        score -= 4
    elif instagram_visibility == "private_or_limited":
        score -= 2

    verified_count = int(evidence.get("verified_profile_count", 0) or 0)
    if verified_count > 1:
        score += min((verified_count - 1) * 3, 9)

    social_text_analyses = evidence.get("social_text_analyses") or {}
    for analysis in social_text_analyses.values():
        status = str((analysis.get("evaluation") or {}).get("status") or "").lower()
        if status == "aligned":
            score += 3
        elif status in {"misaligned", "warning", "review"}:
            score -= 2

    screenshots_summary = evidence.get("screenshots_summary") or {}
    if not screenshots_summary.get("valid_uploaded"):
        score -= 4
    else:
        score += min(4, int(screenshots_summary.get("count", 0) or 0))

    if screenshots_summary.get("safe_content"):
        score += 3

    coherence = evidence.get("linkedin_cv_coherence") or {}
    score += int(coherence.get("score_adjustment", 0) or 0)

    score += int(evidence.get("visual_score_adjustment", 0) or 0)
    return clamp_score(score)


def describe_cv_profile_name_matches(evidence: Dict[str, Any]) -> str:
    messages = []
    for platform, label in [
        ("cv_linkedin_name_match", "LinkedIn"),
        ("cv_github_name_match", "GitHub"),
        ("instagram_slug_verification", "Instagram"),
    ]:
        provided = False
        if label == "LinkedIn":
            provided = bool(
                evidence.get("linkedin_public_link_present")
                or evidence.get("linkedin_export_compared")
                or evidence.get("linkedin_official_verified")
                or (platform in evidence and (evidence.get(platform) or {}).get("status") != "unverified")
            )
        elif label == "GitHub":
            provided = bool(evidence.get("github_link_provided"))
        elif label == "Instagram":
            provided = bool(
                evidence.get("instagram_link_provided")
                or evidence.get("instagram_metadata_found")
                or evidence.get("instagram_media_analyzed")
                or evidence.get("public_preview_analyzed")
                or (evidence.get("instagram_identity") or {}).get("status") not in {"not_connected", None}
            )

        if not provided:
            continue

        if label in {"LinkedIn", "GitHub"}:
            match = evidence.get(platform) or {}
            status = str(match.get("status") or "unverified")
            detected_name = str(match.get("detected_name") or "").strip()
            profile_name = str(match.get("profile_name_candidate") or "").strip()
            if status == "matched":
                messages.append(
                    f"Nome CV ↔ {label}: coerente."
                    + (f" CV: {detected_name}." if detected_name else "")
                    + (f" Profilo: {profile_name}." if profile_name else "")
                )
            elif status == "similar":
                messages.append(
                    f"Nome CV ↔ {label}: parzialmente coerente."
                    + (f" CV: {detected_name}." if detected_name else "")
                    + (f" Profilo: {profile_name}." if profile_name else "")
                )
            elif status == "mismatch":
                messages.append(
                    f"Nome CV ↔ {label}: non corrispondente. Il nome sul CV ({detected_name}) non corrisponde a quello del profilo {label} ({profile_name})."
                )
            else:
                messages.append(
                    f"Nome CV ↔ {label}: non verificabile con i dati pubblici disponibili."
                )
        elif label == "Instagram":
            slug_verification = evidence.get(platform) or {}
            matched = slug_verification.get("matched", False)
            message = slug_verification.get("message", "")
            if matched:
                messages.append(f"Nome CV ↔ Instagram: coerente. {message}")
            else:
                messages.append(f"Nome CV ↔ Instagram: non corrispondente. {message}")
    return " ".join(messages)


def describe_screenshot_impact(evidence: Dict[str, Any]) -> str:
    screenshots_summary = evidence.get("screenshots_summary") or {}
    if screenshots_summary.get("valid_uploaded"):
        return (
            f"Sono stati caricati screenshot validi di {', '.join(screenshots_summary.get('profile_types') or ['profili digitali'])}. "
            f"Questo ha contribuito in modo leggero al punteggio digitale ({int(evidence.get('visual_score_adjustment', 0) or 0):+d})."
        )
    return screenshots_summary.get("message") or (
        "Non sono stati caricati screenshot dei profili digitali. "
        "L'analisi resta parziale e il punteggio puo risultare leggermente piu basso."
    )


def digital_name_checks_are_coherent(evidence: Dict[str, Any]) -> bool:
    summary = describe_cv_profile_name_matches(evidence).lower()
    return bool(summary) and "non corrispondente" not in summary


def build_fallback_digital_analysis(user: Dict, sources: List[Dict[str, str]]) -> Dict:
    has_linkedin = bool(user.get("linkedin_url"))
    has_linkedin_export = bool(user.get("linkedin_profile_text"))
    evidence = build_analysis_evidence(user, sources)
    has_linkedin_official = evidence["linkedin_official_identity"]["status"] != "not_connected"
    has_linkedin_input = has_linkedin or has_linkedin_export or has_linkedin_official
    has_instagram = bool(user.get("instagram_handle"))
    linkedin_identity = evidence["linkedin_identity"]
    linkedin_basic_info = build_linkedin_basic_info(user.get("linkedin_url", ""))
    score = compute_digital_presence_score(evidence)

    return {
        "score": score,
        "headline": "Analisi digitale completata" if evidence["can_compare_with_cv"] else "Analisi digitale parziale",
        "summary": (
            "Ho confrontato il CV con LinkedIn, Instagram e gli screenshot disponibili, usando solo dati pubblici o caricati dall'utente."
            if evidence["can_compare_with_cv"]
            else evidence["zero_score_reason"]
        ),
        "findings": [
            {
                "title": "Coerenza LinkedIn",
                "status": "success" if linkedin_identity["status"] == "matched" else "warning",
                "description": (
                    describe_linkedin_evidence(evidence)
                    if has_linkedin_input
                    else "Manca un link LinkedIn: per un recruiter e spesso il primo punto di verifica."
                ),
                "coach_tip": (
                    "Allinea headline, esperienze, competenze e date con il CV prima di candidarti."
                ),
            },
            {
                "title": "Coerenza CV e profili",
                "status": "success" if digital_name_checks_are_coherent(evidence) else "warning",
                "description": (
                    describe_cv_profile_name_matches(evidence)
                    or "Non ci sono abbastanza dati pubblici per confrontare in modo affidabile il nome del CV con i profili digitali."
                ),
                "coach_tip": "Controlla che nome, cognome, username e bio dei profili appartengano alla stessa persona.",
            },
            {
                "title": "Coerenza LinkedIn vs CV",
                "status": (evidence.get("linkedin_cv_coherence") or {}).get("status", "warning"),
                "description": (evidence.get("linkedin_cv_coherence") or {}).get(
                    "message",
                    "Confronto non disponibile.",
                ),
                "coach_tip": "Allinea esperienze, date e competenze tra CV e LinkedIn.",
            },
            {
                "title": "GitHub",
                "status": (
                    "success"
                    if (evidence.get("cv_github_name_match") or {}).get("status") == "matched"
                    else "warning"
                ),
                "description": (
                    (evidence.get("cv_github_name_match") or {}).get("message")
                    if evidence.get("github_link_provided")
                    else "Non hai inserito un profilo GitHub."
                ),
                "coach_tip": "Usa nome, cognome o una loro variante riconoscibile nello username GitHub.",
            },
            {
                "title": "Instagram",
                "status": "success" if str((evidence.get("instagram_visibility") or {}).get("status")) == "public" else "warning",
                "description": (evidence.get("instagram_visibility") or {}).get("message") or "Instagram non è stato collegato.",
                "coach_tip": "Se il profilo e visibile pubblicamente, mantieni bio e contenuti coerenti con il ruolo target.",
            },
            {
                "title": "Foto e contenuti pubblici",
                "status": visual_media_finding_status(evidence),
                "description": describe_visual_media_analysis(evidence, has_instagram),
                "coach_tip": "Mantieni foto profilo, bio e contenuti recenti coerenti con il ruolo per cui ti candidi.",
            },
            {
                "title": "Screenshot caricati",
                "status": "success" if (evidence.get("screenshots_summary") or {}).get("valid_uploaded") else "warning",
                "description": describe_screenshot_impact(evidence),
                "coach_tip": "Carica screenshot validi di LinkedIn o Instagram per rendere l'analisi piu verificabile.",
            },
        ],
        "sources": sources,
        "linkedin_basic_info": linkedin_basic_info,
        "analysis_evidence": evidence,
    }


def build_clean_digital_analysis(user: Dict, sources: List[Dict[str, str]], score: int) -> Dict:
    has_linkedin = bool(user.get("linkedin_url"))
    has_linkedin_export = bool(user.get("linkedin_profile_text"))
    evidence = build_analysis_evidence(user, sources)
    has_linkedin_official = evidence["linkedin_official_identity"]["status"] != "not_connected"
    has_linkedin_input = has_linkedin or has_linkedin_export or has_linkedin_official
    has_instagram = bool(user.get("instagram_handle"))
    linkedin_identity = evidence["linkedin_identity"]
    can_compare_with_cv = evidence["can_compare_with_cv"]
    linkedin_basic_info = build_linkedin_basic_info(user.get("linkedin_url", ""))

    findings = [
        {
            "title": "LinkedIn",
            "status": "success" if linkedin_identity["status"] == "matched" else "warning",
            "description": (
                describe_linkedin_evidence(evidence)
                if has_linkedin_input
                else "Non hai inserito un profilo LinkedIn: aggiungerlo rende piu forte la candidatura."
            ),
            "coach_tip": "Allinea headline, ruolo target, esperienze, date e competenze con il CV.",
        },
        {
            "title": "Coerenza CV/profili",
            "status": "success" if can_compare_with_cv and digital_name_checks_are_coherent(evidence) else "warning",
            "description": (
                describe_cv_profile_name_matches(evidence)
                or "L'analisi usa solo CV e profili pubblici verificabili, evitando confronti con omonimi o risultati non verificati."
                if can_compare_with_cv
                else evidence["zero_score_reason"]
            ),
            "coach_tip": "Controlla che ruolo target, formazione e competenze principali dicano la stessa cosa su CV e LinkedIn.",
        },
        {
            "title": "Coerenza LinkedIn vs CV",
            "status": (evidence.get("linkedin_cv_coherence") or {}).get("status", "warning"),
            "description": (evidence.get("linkedin_cv_coherence") or {}).get(
                "message",
                "Confronto non disponibile.",
            ),
            "coach_tip": "Allinea esperienze, date e competenze tra CV e LinkedIn.",
        },
        {
            "title": "GitHub",
            "status": (
                "success"
                if (evidence.get("cv_github_name_match") or {}).get("status") == "matched"
                else "warning"
            ),
            "description": (
                (evidence.get("cv_github_name_match") or {}).get("message")
                if evidence.get("github_link_provided")
                else "Non hai inserito un profilo GitHub."
            ),
            "coach_tip": "Usa nome, cognome o una loro variante riconoscibile nello username GitHub.",
        },
        {
            "title": "Instagram",
            "status": "success" if str((evidence.get("instagram_visibility") or {}).get("status")) == "public" else "warning",
            "description": (evidence.get("instagram_visibility") or {}).get("message") or "Instagram non e stato collegato.",
            "coach_tip": "Se il profilo e pubblico, mantieni visibili bio e contenuti coerenti con la tua immagine professionale.",
        },
        {
            "title": "Foto e contenuti pubblici",
            "status": visual_media_finding_status(evidence),
            "description": describe_visual_media_analysis(evidence, has_instagram),
            "coach_tip": "Evita contenuti pubblici che possano confondere il posizionamento professionale.",
        },
        {
            "title": "Screenshot caricati",
            "status": "success" if (evidence.get("screenshots_summary") or {}).get("valid_uploaded") else "warning",
            "description": describe_screenshot_impact(evidence),
            "coach_tip": "Gli screenshot validi aiutano a verificare nome, bio e coerenza del profilo digitale.",
        },
    ]

    return {
        "score": score,
        "headline": "Presenza digitale da allineare con cura" if can_compare_with_cv else "Analisi non disponibile",
        "summary": (
            "La valutazione considera solo i profili pubblici verificabili che hai inserito."
            if can_compare_with_cv
            else evidence["zero_score_reason"]
        ),
        "findings": findings,
        "sources": sources,
        "linkedin_basic_info": linkedin_basic_info,
        "analysis_evidence": evidence,
    }


def analyze_digital_profile(user: Dict, sources: List[Dict[str, str]]) -> Dict:
    fallback = build_fallback_digital_analysis(user, sources)
    has_linkedin = bool(user.get("linkedin_url"))
    has_instagram = bool(user.get("instagram_handle"))
    evidence = fallback["analysis_evidence"]
    linkedin_identity = evidence["linkedin_identity"]
    instagram_identity = evidence["instagram_identity"]
    prompt = f"""
Sei un consulente di personal branding e recruiting.

Valuta la coerenza professionale della presenza digitale del candidato usando solo:
- dati del profilo candidato;
- testo CV disponibile;
- link inseriti;
- dati API ufficiali gia standardizzati dal backend;
- estratti web pubblici forniti.

I dati in official_profile_sources arrivano da API/OAuth ufficiali e sono gia normalizzati dal backend.
Usali per verificare identita, provenienza e qualita della fonte. Non usarli per dedurre esperienze, competenze,
formazione, post o media se i relativi campi standardizzati sono vuoti o se la capability matrix indica che non sono leggibili.
Non affermare di aver analizzato foto o post Instagram se visual_media_analysis.analyzed_content_count e 0. Le immagini recuperate automaticamente sono soltanto anteprime pubbliche della pagina: non provano che i post siano stati analizzati. Riporta soltanto l'esito verificato dal backend senza inventare dettagli sulle immagini.
LinkedIn protegge molte sezioni del profilo. Non affermare di aver letto headline, esperienze, date, competenze o formazione da LinkedIn se queste informazioni non compaiono esplicitamente negli estratti pubblici forniti.
Se Instagram e indicato ma non sono presenti metadati pubblici, segnala che il profilo potrebbe essere privato o non accessibile. La presenza di metadati o anteprime non prova che foto o post siano stati analizzati: usa sempre visual_media_analysis.analyzed_content_count.

Profilo candidato:
- Nome: {user.get("name", "")}
- Email: {user.get("email", "")}
- Percorso di studi: {user.get("education", "")}
- Ruolo target: {user.get("target_role", "")}
- Settore: {user.get("sector", "")}
- Livello esperienza: {user.get("experience_level", "")}
- LinkedIn: {user.get("linkedin_url", "")}
- Instagram: {user.get("instagram_handle", "")}

Estratto CV:
{(user.get("cv_text") or "")[:5000]}

Esportazione LinkedIn caricata dal candidato:
{(user.get("linkedin_profile_text") or "Non disponibile")[:6000]}

Dati API/OAuth ufficiali standardizzati:
{json.dumps(evidence.get("official_profile_sources", []), ensure_ascii=False)}

Fonti pubbliche trovate:
{sources_to_prompt(sources)}

Evidenze verificate dal backend:
{json.dumps(evidence, ensure_ascii=False)}

Restituisci SOLO JSON valido con questa struttura:
{{
  "score": 0,
  "headline": "titolo breve del risultato",
  "summary": "sintesi concreta per il candidato",
  "findings": [
    {{
      "title": "area analizzata",
      "status": "success oppure warning",
      "description": "cosa e stato trovato o non trovato",
      "coach_tip": "azione consigliata"
    }}
  ],
  "sources": []
}}

Regole:
- score intero 0-100.
- Se can_compare_with_cv e false, score deve essere 0: non ci sono profili pubblici verificabili da confrontare con il CV.
- Se il web restituisce poche fonti, esplicita che l'analisi e limitata e non premiare la sola presenza di un link.
- Se linkedin_identity.status e mismatch, segnala chiaramente che il link LinkedIn potrebbe appartenere a un'altra persona. Se non restano altri profili verificati, score deve essere 0.
- findings deve includere almeno LinkedIn, coerenza CV/profili, foto o contenuti pubblici, impatto recruiter.
- Per LinkedIn considera sempre il link pubblico e l'identificativo del profilo. Considera headline, esperienze/date, competenze e formazione solo se compaiono esplicitamente negli estratti pubblici.
- Se linkedin_export_compared e true, il PDF LinkedIn caricato e una fonte dettagliata: non dire che mancano informazioni dettagliate su LinkedIn. Puoi invece segnalare che il link pubblico LinkedIn espone solo informazioni limitate.
- Se LinkedIn OAuth ufficiale e collegato, usalo come verifica identita di base, ma non come prova di coerenza professionale se non contiene sezioni professionali.
- Analizza solo i profili social inseriti dal candidato e gli snippet che corrispondono esattamente a quei link.
- L'handle Instagram inserito dal candidato e autoritativo: non segnalare profili Instagram multipli o omonimi se non compaiono tra le fonti esatte.
- Se le fonti non contengono altri profili, non parlare di "diverse fonti", "profili multipli", omonimi o incongruenze con persone diverse.
- Per ogni confronto descrivi elementi concreti realmente presenti nelle fonti, ad esempio ruolo, formazione, esperienze, date o competenze.
- Se non sono disponibili contenuti sufficienti per un confronto concreto, dichiaralo esplicitamente invece di formulare una valutazione generica.
- Non premiare la semplice presenza di URL: un link collegato ma non leggibile non dimostra coerenza con il CV.
- Non inventare dati non presenti.
- sources deve restare una lista vuota: le fonti reali vengono aggiunte dal backend.
"""

    try:
        result = extract_json(call_groq(prompt, temperature=0.25, max_tokens=1400))
        result["score"] = compute_digital_presence_score(evidence)
        result["headline"] = result.get("headline") or fallback["headline"]
        result["summary"] = result.get("summary") or fallback["summary"]
        result["findings"] = result.get("findings") or fallback["findings"]
        result["sources"] = sources
        result["linkedin_basic_info"] = fallback["linkedin_basic_info"]
        result["analysis_evidence"] = evidence
        if (
            has_instagram
            and not evidence["instagram_media_analyzed"]
            and "instagram" in str(result.get("summary", "")).lower()
        ):
            result["summary"] = (
                "Il confronto usa il CV e i profili pubblici verificabili disponibili."
            )
        if not evidence["can_compare_with_cv"]:
            result["headline"] = "Analisi non disponibile"
            result["summary"] = evidence["zero_score_reason"]
        
        if has_linkedin or evidence["linkedin_export_compared"] or evidence["linkedin_official_verified"]:
            for finding in result["findings"]:
                title = str(finding.get("title", "")).lower()
                if "linkedin" in title and "coerenza" not in title and "cv" not in title:
                    finding["status"] = "success" if linkedin_identity["status"] == "matched" else "warning"
                    finding["description"] = describe_linkedin_evidence(evidence)
                    finding["coach_tip"] = (
                        "Mantieni allineati date, titoli e descrizioni tra CV, PDF LinkedIn e profilo pubblico."
                        if linkedin_identity["status"] == "matched"
                        else "Controlla separatamente che il PDF LinkedIn caricato e il link pubblico appartengano al candidato."
                    )

        coherence = evidence.get("linkedin_cv_coherence") or {}
        if coherence and coherence.get("status") != "unverified":
            coh_finding = None
            for finding in result["findings"]:
                title = str(finding.get("title", "")).lower()
                if "coerenza cv" in title or "coerenza linkedin" in title or "confronto" in title:
                    coh_finding = finding
                    break
            if not coh_finding:
                coh_finding = {
                    "title": "Coerenza LinkedIn vs CV",
                    "status": coherence.get("status", "warning"),
                    "description": coherence.get("message", ""),
                    "coach_tip": "Risolvi le discrepanze di date e allinea le competenze tra il CV e LinkedIn."
                }
                result["findings"].append(coh_finding)
            else:
                coh_finding["status"] = coherence.get("status", "warning")
                coh_finding["description"] = coherence.get("message", "")
                if coherence.get("details"):
                    coh_finding["coach_tip"] = "Risolvi le discrepanze di date e allinea le competenze tra il CV e LinkedIn."
                else:
                    coh_finding["coach_tip"] = "I profili sono coerenti. Continua a mantenerli aggiornati."

        for finding in result["findings"]:
            title = str(finding.get("title", "")).lower()
            if "foto" in title or "contenuti pubblici" in title:
                finding["status"] = visual_media_finding_status(evidence)
                finding["description"] = describe_visual_media_analysis(evidence, has_instagram)
                finding["coach_tip"] = (
                    "Controlla manualmente cosa risulta visibile a chi non segue il profilo. "
                    "Per un controllo automatico dei contenuti servono media realmente accessibili."
                )

        if has_instagram:
            instagram_visibility = str((evidence.get("instagram_visibility") or {}).get("status") or "")
            vis_finding = None
            for finding in result["findings"]:
                title = str(finding.get("title", "")).lower()
                if "visibilità instagram" in title or title == "profilo instagram" or title == "instagram":
                    vis_finding = finding
                    break
            if not vis_finding:
                vis_finding = {
                    "title": "Profilo Instagram",
                    "status": "warning" if instagram_visibility != "public" else "success",
                    "description": "",
                    "coach_tip": ""
                }
                result["findings"].append(vis_finding)
            
            if instagram_visibility == "public":
                vis_finding["status"] = "success"
                vis_finding["description"] = "Il profilo Instagram risulta pubblico. Questo facilita la trasparenza e la verifica dell'identità professionale."
                vis_finding["coach_tip"] = "Ottima scelta. Mantenere il profilo pubblico (con contenuti appropriati) favorisce il personal branding."
            else:
                vis_finding["status"] = "warning"
                vis_finding["description"] = "Il profilo Instagram risulta privato (chiuso) o non accessibile."
                vis_finding["coach_tip"] = "Si consiglia di impostare il profilo come pubblico per migliorare la trasparenza professionale."

        name_coh_finding = None
        for finding in result["findings"]:
            title = str(finding.get("title", "")).lower()
            if "nome" in title or "nominativ" in title or "coerenza cv/profil" in title or "coerenza cv e profil" in title:
                name_coh_finding = finding
                break
        if name_coh_finding:
            name_matches_summary = describe_cv_profile_name_matches(evidence)
            is_matched = "non corrispondente" not in name_matches_summary.lower()
            name_coh_finding["status"] = "success" if is_matched else "warning"
            name_coh_finding["description"] = name_matches_summary or "Non ci sono abbastanza dati pubblici per confrontare in modo affidabile il nome del CV con i profili digitali."
            
            slug_verification = evidence.get("instagram_slug_verification") or {}
            if has_instagram and not slug_verification.get("matched", False):
                coach_tip_msg = slug_verification.get("coach_tip") or "Si consiglia di allineare lo slug del link Instagram con il tuo nome e cognome per fini di personal branding."
                name_coh_finding["coach_tip"] = coach_tip_msg
            else:
                name_coh_finding["coach_tip"] = "Controlla che nome, cognome e bio dei profili siano allineati al tuo CV."

        screenshot_summary = evidence.get("screenshots_summary") or {}
        sc_finding = None
        for finding in result["findings"]:
            title = str(finding.get("title", "")).lower()
            if "screenshot" in title or "verifica visiva" in title:
                sc_finding = finding
                break
        if not screenshot_summary.get("valid_uploaded"):
            if not sc_finding:
                sc_finding = {
                    "title": "Verifica Screenshot",
                    "status": "warning",
                    "description": "Non sono stati caricati screenshot dei profili digitali. L'analisi è stata eseguita con le informazioni disponibili, ma il punteggio digitale è inferiore.",
                    "coach_tip": "Carica gli screenshot dei tuoi profili professionali per sbloccare la verifica visiva completa."
                }
                result["findings"].append(sc_finding)
            else:
                sc_finding["status"] = "warning"
                sc_finding["description"] = "Nessun screenshot inserito per il controllo visivo."
                sc_finding["coach_tip"] = "Si consiglia di caricare gli screenshot dei profili per sbloccare la verifica visiva completa."
        else:
            if not sc_finding:
                sc_finding = {
                    "title": "Verifica Screenshot",
                    "status": "success",
                    "description": f"Sono stati caricati {screenshot_summary.get('count', 0)} screenshot validi.",
                    "coach_tip": "Screenshot caricati con successo. Il controllo visuale è stato eseguito."
                }
                result["findings"].append(sc_finding)
            else:
                sc_finding["status"] = "success"
                sc_finding["description"] = f"Sono stati caricati {screenshot_summary.get('count', 0)} screenshot validi per la verifica visiva."
                sc_finding["coach_tip"] = "Verifica visiva completata correttamente."

        coherence = evidence.get("linkedin_cv_coherence") or {}
        if coherence and coherence.get("status") != "unverified":
            coh_finding = None
            for finding in result["findings"]:
                title = str(finding.get("title", "")).lower()
                if "coerenza cv" in title or "coerenza linkedin" in title or "confronto" in title:
                    coh_finding = finding
                    break
            if not coh_finding:
                coh_finding = {
                    "title": "Coerenza LinkedIn vs CV",
                    "status": coherence.get("status", "warning"),
                    "description": coherence.get("message", ""),
                    "coach_tip": "Risolvi le discrepanze di date e allinea le competenze tra il CV e LinkedIn."
                }
                result["findings"].append(coh_finding)
            else:
                coh_finding["status"] = coherence.get("status", "warning")
                coh_finding["description"] = coherence.get("message", "")
                if coherence.get("details"):
                    coh_finding["coach_tip"] = "Risolvi le discrepanze di date e allinea le competenze tra il CV e LinkedIn."
                else:
                    coh_finding["coach_tip"] = "I profili sono coerenti. Continua a mantenerli aggiornati."

        unsafe_text = json.dumps(result, ensure_ascii=False).lower()
        unsafe_patterns = [
            "profili multipli",
            "diverse fonti",
            "omonim",
            "persona diversa",
            "nomi diversi",
        ]
        if any(pattern in unsafe_text for pattern in unsafe_patterns):
            return build_clean_digital_analysis(user, sources, result["score"])
        return result
    except Exception as exc:
        print(f"Analisi digitale AI non riuscita, uso fallback: {exc}")
        return fallback


def build_fallback_cv_strategy(
    user: Dict,
    company: str,
    role: str,
    goal: str,
    job_link: str,
    sources: List[Dict[str, str]]
) -> Dict:
    cv_text = (user.get("cv_text") or "").strip()
    has_cv_text = bool(cv_text)
    scorecard = build_deterministic_cv_scorecard(
        cv_text,
        company,
        role,
        goal,
    )

    return {
        "score": scorecard["overall_score"],
        "overall_score": scorecard["overall_score"],
        "ats_score": scorecard["ats_score"],
        "keyword_score": scorecard["keyword_score"],
        "format_score": scorecard["format_score"],
        "role_match_score": scorecard["role_match_score"],
        "company_fit_score": scorecard["company_fit_score"],
        "company_provided": scorecard["company_provided"],
        "completeness_score": scorecard["completeness_score"],
        "clarity_score": scorecard["clarity_score"],
        "professionalism_score": scorecard["professionalism_score"],
        "ats_analysis": scorecard["ats_analysis"],
        "score_explanation": scorecard["score_explanation"],
        "scoring_context": scorecard["scoring_context"],
        "headline": "Analisi strategica pronta",
        "summary": (
            "Ho confrontato il CV disponibile con i dati della candidatura. "
            "Per rendere il report piu preciso, assicurati che il CV caricato contenga testo estraibile."
        ),
        "strengths": [
            {
                "title": "Contesto candidatura definito",
                "description": f"La candidatura e impostata su {role or 'un ruolo da definire'} presso {company or 'Azienda Generica'}.",
                "coach_tip": "Mantieni nel CV le parole chiave piu vicine al ruolo scelto."
            },
            {
                "title": "Base CV disponibile" if has_cv_text else "CV caricato ma testo limitato",
                "description": (
                    "Il testo del CV e disponibile per confrontare competenze, formazione ed esperienze."
                    if has_cv_text
                    else "Il file risulta caricato, ma il testo estraibile e scarso o assente. L'analisi resta preliminare."
                ),
                "coach_tip": "Se hai caricato PDF, verifica che il testo sia selezionabile o carica una versione DOCX per un'analisi piu accurata."
            }
        ],
        "improvements": [
            {
                "title": "Quantifica i risultati",
                "description": "Inserisci metriche concrete nelle esperienze: percentuali, volumi, tempi, impatto o miglioramenti ottenuti.",
                "coach_tip": "Trasforma responsabilita generiche in risultati misurabili."
            },
            {
                "title": "Allinea le competenze",
                "description": "Evidenzia nel riepilogo e nelle esperienze le competenze richieste dal ruolo e dall'annuncio.",
                "coach_tip": "Usa le stesse parole chiave dell'annuncio quando sono vere per il tuo profilo."
            }
        ],
        "sources": sources,
        "target": {
            "company": company,
            "role": role,
            "goal": goal,
            "job_link": job_link
        }
    }


def normalize_cv_strategy_result(result: Dict, fallback: Dict, sources: List[Dict[str, str]], target: Dict) -> Dict:
    strengths = result.get("strengths") or fallback["strengths"]
    improvements = result.get("improvements") or fallback["improvements"]

    return {
        "score": fallback["overall_score"],
        "overall_score": fallback["overall_score"],
        "ats_score": fallback["ats_score"],
        "keyword_score": fallback["keyword_score"],
        "format_score": fallback["format_score"],
        "role_match_score": fallback["role_match_score"],
        "company_fit_score": fallback["company_fit_score"],
        "company_provided": fallback["company_provided"],
        "completeness_score": fallback["completeness_score"],
        "clarity_score": fallback["clarity_score"],
        "professionalism_score": fallback["professionalism_score"],
        "ats_analysis": fallback["ats_analysis"],
        "score_explanation": fallback["score_explanation"],
        "scoring_context": fallback["scoring_context"],
        "headline": result.get("headline") or fallback["headline"],
        "summary": result.get("summary") or fallback["summary"],
        "strengths": strengths[:6],
        "improvements": improvements[:6],
        "sources": sources,
        "target": target,
    }


def analyze_cv_strategy(user: Dict, company: str, role: str, goal: str, job_link: str, sources: List[Dict[str, str]]) -> Dict:
    target = {
        "company": company,
        "role": role,
        "goal": goal,
        "job_link": job_link,
    }
    fallback = build_fallback_cv_strategy(user, company, role, goal, job_link, sources)
    if not CV_LLM_ENABLED:
        print("Analisi strategica CV: uso motore locale deterministico, LLM disabilitato.")
        return fallback

    prompt = f"""
Sei un career coach e recruiter esperto di ottimizzazione CV.

Analizza il CV del candidato rispetto a una candidatura specifica. Usa solo:
- dati del profilo;
- testo CV disponibile;
- azienda, ruolo, obiettivo e link annuncio forniti dall'utente;
- fonti web pubbliche fornite.

Profilo candidato:
- Nome: {user.get("name", "")}
- Percorso di studi: {user.get("education", "")}
- Settore: {user.get("sector", "")}
- Livello esperienza: {user.get("experience_level", "")}

Candidatura:
- Azienda: {company}
- Ruolo: {role}
- Obiettivo dichiarato: {goal or "Non specificato"}
- Link annuncio o azienda: {job_link or "Non specificato"}

Estratto CV:
{(user.get("cv_text") or "")[:8000]}

Contesto web/annuncio:
{sources_to_prompt(sources)}

Restituisci SOLO JSON valido con questa struttura:
{{
  "score": 0,
  "headline": "titolo breve",
  "summary": "sintesi concreta",
  "strengths": [
    {{
      "title": "punto di forza",
      "description": "perche e rilevante rispetto al ruolo",
      "coach_tip": "come valorizzarlo nel CV"
    }}
  ],
  "improvements": [
    {{
      "title": "area da migliorare",
      "description": "cosa manca o cosa non e abbastanza forte",
      "coach_tip": "azione pratica per sistemarla"
    }}
  ]
}}

Regole:
- score intero 0-100.
- strengths deve contenere 2-5 elementi.
- improvements deve contenere 2-5 elementi.
- Sii specifico su competenze, formazione, esperienze, parole chiave e risultati quantificabili.
- Distingui cio che e presente nel CV da cio che manca o va reso piu evidente.
- Non inventare esperienze, aziende, titoli o competenze non presenti.
- Se il testo CV e scarso o assente, dichiaralo e dai consigli su come rendere il CV analizzabile.
- Scrivi in italiano, tono professionale e diretto.
"""

    try:
        result = extract_json(call_groq(prompt, temperature=0.25, max_tokens=1000))
        return normalize_cv_strategy_result(result, fallback, sources, target)
    except Exception as exc:
        print(f"Analisi strategica CV AI non riuscita, uso fallback: {exc}")
        return fallback


def strategy_item_to_text(item, fallback_title: str) -> str:
    if isinstance(item, str):
        return item

    if not isinstance(item, dict):
        return str(item or "")

    title = item.get("title") or fallback_title
    description = item.get("description") or ""
    coach_tip = item.get("coach_tip") or ""
    parts = [part for part in [title, description, coach_tip] if part]
    return " - ".join(parts)


def build_fallback_optimized_cv_text(
    cv_text: str,
    analysis: Dict,
    company: str,
    role: str,
    goal: str,
    job_link: str,
    accepted_suggestions: Optional[List[Dict]] = None,
    user_additional_data: Optional[Dict[str, Any]] = None,
) -> str:
    section_titles = [
        "CONTATTI", "LINGUE", "HARD SKILLS", "SOFT SKILLS", "CHI SONO",
        "PROFILO", "PROFILO PROFESSIONALE", "FORMAZIONE", "ISTRUZIONE",
        "ESPERIENZE PROFESSIONALI", "ESPERIENZA PROFESSIONALE", "ESPERIENZE",
        "PROGETTI", "CERTIFICAZIONI", "COMPETENZE",
    ]
    optimized_text = clean_extracted_text(cv_text or "")

    for title in sorted(section_titles, key=len, reverse=True):
        optimized_text = re.sub(
            rf"(?<!\n)\b{re.escape(title)}\b",
            f"\n\n{title}",
            optimized_text,
            flags=re.IGNORECASE,
        )

    optimized_text = re.sub(r"\n{3,}", "\n\n", optimized_text).strip()
    if not optimized_text:
        raise HTTPException(
            status_code=422,
            detail="Non riesco a generare un CV ottimizzato perche il testo del CV originale non e leggibile.",
        )

    supported_text = normalize_plain_text("\n".join([
        optimized_text,
        flatten_cv_support_data(user_additional_data),
    ]))
    target_role = (role or "").strip()
    target_company = (company or "").strip()
    profile_line = ""
    if target_role and any(token in supported_text for token in tokenize_meaningful(target_role)):
        profile_line = f"Obiettivo professionale: valorizzare competenze gia presenti nel percorso per il ruolo di {target_role}"
        if target_company and target_company.lower() != " Azienda Generica":
            profile_line += f" presso {target_company}"
        profile_line += "."

    additional_lines = []
    for key, value in (user_additional_data or {}).items():
        if key == "adaptation_answers" or not isinstance(value, str) or not value.strip():
            continue
        additional_lines.append(value.strip())
    for item in (user_additional_data or {}).get("adaptation_answers", []):
        if isinstance(item, dict) and str(item.get("answer", "")).strip():
            additional_lines.append(str(item["answer"]).strip())

    if profile_line:
        if re.search(r"\b(CHI SONO|PROFILO|PROFILO PROFESSIONALE)\b", optimized_text, flags=re.IGNORECASE):
            optimized_text = re.sub(
                r"(\b(?:CHI SONO|PROFILO|PROFILO PROFESSIONALE)\b\s*)",
                rf"\1\n{profile_line}\n",
                optimized_text,
                count=1,
                flags=re.IGNORECASE,
            )
        else:
            optimized_text = f"PROFILO PROFESSIONALE\n{profile_line}\n\n{optimized_text}"

    if additional_lines:
        optimized_text = f"{optimized_text}\n\nNOTE PROFESSIONALI\n" + "\n".join(
            f"- {line}" for line in additional_lines[:5]
        )

    return optimized_text.strip()


def optimize_cv_text_for_job(
    cv_text: str,
    analysis: Dict,
    company: str,
    role: str,
    goal: str,
    job_link: str,
    sources: List[Dict[str, str]],
    cv_evaluation: Optional[Any] = None,
    strategic_analysis: Optional[Any] = None,
    recommended_adaptations: Optional[Any] = None,
    accepted_suggestions: Optional[Any] = None,
    user_additional_data: Optional[Dict[str, Any]] = None,
) -> str:
    result = build_resume_rewrite_result(
        cv_text=cv_text,
        company=company,
        role=role,
        goal=goal,
        accepted_suggestions=accepted_suggestions,
        user_additional_data=user_additional_data,
    )
    return result["optimized_text"]


def build_resume_rewrite_result(
    cv_text: str,
    company: str,
    role: str,
    goal: str,
    accepted_suggestions: Optional[Any] = None,
    user_additional_data: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    from services.cv_optimizer.rewrite import build_resume_rewrite_result as build_result

    return build_result(
        cv_text=cv_text,
        company=company,
        role=role,
        goal=goal,
        accepted_suggestions=accepted_suggestions,
        user_additional_data=user_additional_data,
    )


def extract_docx_text_bytes(file_bytes: bytes) -> str:
    try:
        from docx import Document

        document = Document(io.BytesIO(file_bytes))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        seen_cells = set()
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if id(cell._tc) not in seen_cells:
                        seen_cells.add(id(cell._tc))
                        parts.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text)
        return "\n".join(parts)
    except Exception:
        return ""


def validate_optimized_docx_structure(final_text: str) -> List[str]:
    normalized_lines = [line.strip() for line in (final_text or "").splitlines() if line.strip()]
    normalized_text = "\n".join(normalized_lines)
    warnings = []
    section_titles = [
        "CONTATTI", "LINGUE", "HARD SKILLS", "SOFT SKILLS", "FORMAZIONE",
        "ISTRUZIONE", "ESPERIENZE PROFESSIONALI", "ESPERIENZA PROFESSIONALE",
        "PROGETTI", "CERTIFICAZIONI",
    ]
    profile_aliases = {"CHI SONO", "PROFILO", "PROFILO PROFESSIONALE"}

    upper_lines = [line.upper().strip(":") for line in normalized_lines]
    for profile_title in profile_aliases:
        if profile_title not in upper_lines:
            continue
        start = upper_lines.index(profile_title)
        segment = []
        for line in upper_lines[start + 1:]:
            if line in profile_aliases or line in section_titles:
                break
            segment.append(line)
        leaked = [line for line in segment if line in section_titles]
        if leaked:
            warnings.append(
                f"La sezione {profile_title} contiene titoli di altre sezioni: {', '.join(leaked)}."
            )

    for title in section_titles + list(profile_aliases):
        count = upper_lines.count(title)
        if count > 1:
            warnings.append(f"La sezione {title} compare {count} volte nel CV finale.")

    if normalize_plain_text("CV ottimizzato - bozza guidata") in normalize_plain_text(normalized_text):
        warnings.append("Il CV finale contiene testo di bozza/report.")

    return warnings


def review_generated_cv_quality(
    final_text: str,
    original_cv_text: str,
    role: str,
    company: str,
    accepted_instructions: List[RewriteInstruction],
) -> Dict[str, Any]:
    if not CV_QUALITY_LLM_ENABLED:
        return review_generated_cv_quality_locally(
            final_text=final_text,
            accepted_instructions=accepted_instructions,
            llm_error="Revisione LLM disabilitata: controllo locale eseguito.",
        )

    accepted_payload = [
        {
            "id": instruction.source_id,
            "section": instruction.section,
            "replacement": instruction.replacement,
        }
        for instruction in accepted_instructions
    ]
    prompt = f"""
Agisci come revisore finale senior di curriculum. Restituisci SOLO JSON valido.

Devi decidere se il CV seguente è pulito, professionale e pronto per essere inviato a una candidatura reale.

Ruolo target: {role or "Non specificato"}
Azienda target: {company or "Non specificata"}

CV originale, usato solo per verificare che non siano stati inventati fatti:
{original_cv_text[:3500]}

Modifiche accettate che devono essere tutte conservate:
{json.dumps(accepted_payload, ensure_ascii=False)}

CV finale da revisionare:
{final_text[:4000]}

Schema:
{{
  "ready_to_send": true,
  "score": 0,
  "issues": [
    {{
      "severity": "critical",
      "section": "nome sezione",
      "description": "problema concreto"
    }}
  ],
  "revisions": [
    {{
      "section": "titolo della sezione",
      "original_text": "testo esatto del CV finale da sostituire",
      "replacement": "testo corretto pronto per il CV",
      "reason": "motivo"
    }}
  ]
}}

Criteri obbligatori:
- Il CV deve essere pronto da inviare, non una bozza o un report.
- Controlla duplicazioni, frasi spezzate, elenchi illeggibili, tono debole, incoerenze e ripetizioni.
- Controlla che ogni sezione sia chiara, sintetica e coerente con il ruolo target.
- Controlla che tutte le modifiche accettate siano ancora rappresentate.
- Non proporre fatti, skill, risultati, strumenti, ruoli, date o aziende non supportati.
- Le revisioni devono correggere solo problemi reali e conservare tutti i fatti validi.
- original_text deve essere copiato esattamente dal CV finale.
- severity deve essere sempre uno tra: critical, major, minor.
- Se non ci sono problemi critical o major, ready_to_send deve essere true.
- score deve essere un intero da 0 a 100.
"""
    try:
        result = call_analysis_llm(
            prompt,
            context="final_cv_quality_review",
            temperature=0.05,
            max_tokens=1500,
            timeout=60,
        )
    except Exception as exc:
        print(f"Revisione finale CV non disponibile: {exc}")
        return review_generated_cv_quality_locally(
            final_text=final_text,
            accepted_instructions=accepted_instructions,
            llm_error=str(exc),
        )

    issues = result.get("issues") if isinstance(result.get("issues"), list) else []
    revisions = result.get("revisions") if isinstance(result.get("revisions"), list) else []

    local_review = review_generated_cv_quality_locally(
        final_text=final_text,
        accepted_instructions=accepted_instructions,
        llm_error="Controllo locale eseguito come validazione del review LLM.",
    )

    normalized_issues = []
    malformed_issue_payload = False
    for item in issues:
        if not isinstance(item, dict):
            continue
        severity_raw = str(item.get("severity") or "").strip().lower()
        if severity_raw in {"critical", "major", "minor"}:
            severity = severity_raw
        else:
            malformed_issue_payload = True
            severity = "minor"
        normalized_issues.append({
            **item,
            "severity": severity,
            "section": str(item.get("section") or "").strip(),
            "description": str(item.get("description") or "").strip(),
        })

    normalized_revisions = [item for item in revisions if isinstance(item, dict)][:8]
    llm_score = clamp_score(result.get("score", 0))
    local_blocking = any(
        isinstance(item, dict) and str(item.get("severity") or "").lower() in {"critical", "major"}
        for item in (local_review.get("issues") or [])
    )
    llm_blocking = any(item["severity"] in {"critical", "major"} for item in normalized_issues)

    if malformed_issue_payload or (llm_score < 50 and not llm_blocking and not local_blocking):
        return {
            **local_review,
            "issues": normalized_issues or local_review.get("issues", []),
            "revisions": normalized_revisions,
            "review_provider": "llm+local",
            "llm_review_suspect": True,
        }

    if llm_blocking and not local_blocking:
        return {
            **local_review,
            "issues": [
                {
                    **item,
                    "severity": "minor",
                    "description": (
                        f"Avviso del revisore AI non bloccante: {item['description']}"
                    ),
                }
                for item in normalized_issues
            ],
            "revisions": [],
            "review_provider": "llm+local",
            "llm_review_suspect": True,
            "llm_score": llm_score,
        }

    return {
        "ready_to_send": bool(result.get("ready_to_send")),
        "score": llm_score,
        "issues": normalized_issues[:12],
        "revisions": normalized_revisions,
        "review_unavailable": False,
        "review_provider": "llm",
        "local_checks_completed": True,
    }


def review_generated_cv_quality_locally(
    final_text: str,
    accepted_instructions: List[RewriteInstruction],
    llm_error: str = "",
) -> Dict[str, Any]:
    issues = [
        {
            "severity": "major",
            "section": "struttura",
            "description": warning,
        }
        for warning in validate_optimized_docx_structure(final_text)
    ]
    normalized_final = normalize_plain_text(final_text)
    def _instruction_present(replacement: str) -> bool:
        normalized_replacement = normalize_plain_text(replacement)
        if not normalized_replacement:
            return True
        if normalized_replacement in normalized_final:
            return True

        replacement_tokens = [token for token in normalized_replacement.split() if len(token) > 2]
        if len(replacement_tokens) < 3:
            return False

        final_tokens = set(normalized_final.split())
        overlap = sum(1 for token in replacement_tokens if token in final_tokens)
        return (overlap / len(replacement_tokens)) >= 0.7

    missing_changes = [
        instruction.source_id or instruction.section or "modifica"
        for instruction in accepted_instructions
        if instruction.replacement
        and (
            not _instruction_present(instruction.replacement)
            or not rewrite_preserves_instruction_content(final_text, instruction)
        )
    ]
    if missing_changes:
        issues.append({
            "severity": "major",
            "section": "modifiche",
            "description": (
                "Non risultano presenti tutte le modifiche richieste: "
                + ", ".join(missing_changes[:8])
            ),
        })
    forbidden_markers = [
        marker
        for marker in ResumeRewriter.FORBIDDEN_OUTPUT_MARKERS
        if normalize_plain_text(marker) in normalized_final
    ]
    if forbidden_markers:
        issues.append({
            "severity": "major",
            "section": "contenuto",
            "description": "Il documento contiene testo operativo o di report non destinato al CV.",
        })
    ready = not any(issue["severity"] in {"critical", "major"} for issue in issues)
    return {
        "ready_to_send": ready,
        "score": 85 if ready else 55,
        "issues": issues,
        "revisions": [],
        "review_unavailable": False,
        "review_provider": "local",
        "local_checks_completed": True,
        "llm_review_unavailable": True,
        "technical_detail": llm_error[:300],
    }


def quality_rewrite_instructions(review: Dict[str, Any], final_text: str) -> List[RewriteInstruction]:
    instructions: List[RewriteInstruction] = []
    for index, item in enumerate(review.get("revisions") or []):
        section = str(item.get("section") or "").strip()
        original = str(item.get("original_text") or "").strip()
        replacement = str(item.get("replacement") or "").strip()
        suggestion = {
            "type": "actionableEdit",
            "section": section,
            "original_text": original,
            "proposed_text": replacement,
        }
        if (
            not section
            or not original
            or not replacement
            or not is_valid_actionable_suggestion(suggestion)
            or not suggestion_targets_current_cv(suggestion, final_text)
        ):
            continue
        instructions.append(RewriteInstruction(
            section=section,
            original=original,
            replacement=replacement,
            reason=str(item.get("reason") or "Correzione della revisione finale.").strip(),
            category="quality_review",
            source_id=f"quality_review_{index}",
        ))
    return instructions


def get_optimized_cv_filename(filename: Optional[str] = None, extension: str = "pdf") -> str:
    base = "cv-ottimizzato"
    if filename:
        sanitized = re.sub(r"[^0-9A-Za-z_.-]", "-", os.path.splitext(filename)[0]).strip("-_")
        if sanitized:
            base = f"{sanitized}-ottimizzato"
    return f"{base}.{extension}"


def get_target_optimized_cv_filename(user_name: str, role: str, company: str, extension: str = "docx") -> str:
    raw = "_".join(part for part in ["CV", user_name, role, company, "ottimizzato"] if part)
    sanitized = re.sub(r"[^0-9A-Za-z]+", "_", strip_accents(raw)).strip("_")
    return f"{sanitized or 'CV_ottimizzato'}.{extension}"


def build_professional_extra_text(user_additional_data: Dict[str, Any], role: str) -> str:
    support = flatten_cv_support_data(user_additional_data)
    if not support:
        return ""

    prompt = "\n".join(
        [
            "Sei un resume editor. Restituisci SOLO JSON valido.",
            "",
            f"Ruolo target:\n{role or 'Non specificato'}",
            "",
            f"Informazioni vere inserite dall utente:\n{support[:3500]}",
            "",
            "Schema:",
            "{",
            '  "text": "testo breve direttamente utilizzabile in un CV"',
            "}",
            "",
            "Regole:",
            "- Usa esclusivamente i fatti presenti nelle informazioni dell'utente.",
            "- Non dedurre che si tratti di universita, lavoro, tirocinio o progetto se non e scritto.",
            "- Non inventare strumenti, risultati, contesti, aziende, date o responsabilita.",
            "- Migliora solo grammatica, chiarezza e tono professionale.",
            "- Mantieni la prima persona implicita e non aggiungere etichette o spiegazioni.",
        ]
    )
    try:
        result = call_rewrite_llm(prompt, context="professional_extra_text", temperature=0.05, max_tokens=250, timeout=20)
        text = clean_extracted_text(str(result.get("text") or ""))
        if text:
            return text[:700]
    except Exception as exc:
        print(f"Riformulazione informazioni extra non riuscita, uso testo conservativo: {exc}")

    fallback = re.split(r"(?<=[.!?])\s+", support.strip())[0][:500].strip()
    replacements = [
        (r"^ho\s+applicato\s+", "Applicazione di "),
        (r"^ho\s+(usato|utilizzato)\s+", "Utilizzo di "),
        (r"^(l'?ho|lo|la|li|le)\s+(usato|utilizzato|applicato)\s*", "Utilizzo in "),
        (r"\b(visto|vista)\s+all'?esame\b", "approfondito in ambito universitario"),
        (r"\bin esame di basi di dati\b", "nell'ambito delle basi di dati"),
        (r"\bin esame di basi dati\b", "nell'ambito delle basi di dati"),
    ]
    for pattern, replacement in replacements:
        fallback = re.sub(pattern, replacement, fallback, flags=re.IGNORECASE)
    fallback = re.sub(r"\s{2,}", " ", fallback).strip(" ,;:-")
    return fallback[:1].upper() + fallback[1:] if fallback else ""


def infer_extra_content_section(value: str) -> tuple[str, str]:
    plain = normalize_plain_text(value)
    def score_terms(terms: Iterable[str]) -> int:
        return sum(1 for term in terms if term in plain)

    certification_terms = ("certificazione", "certificato", "attestato", "licenza")
    language_terms = (
        "lingua", "lingue", "inglese", "italiano", "francese", "spagnolo",
        "tedesco", "portoghese", "madrelingua",
    )
    education_terms = (
        "laurea", "universita", "università", "corso", "formazione", "studio", "esame",
        "master", "diploma", "triennale", "magistrale",
    )
    experience_terms = ("azienda", "cliente", "lavoro", "tirocinio", "stage", "impiego")
    soft_terms = (
        "pensiero analitico", "attenzione ai dettagli", "comunicazione dei risultati",
        "problem solving", "team working", "collaborazione", "comunicazione tecnica",
        "pensiero logico", "precisione",
    )
    hard_terms = (
        "excel", "excel avanzato", "power bi", "tableau", "bigquery", "google analytics",
        "sql", "database", "kpi", "reporting", "data visualization", "analisi dei dati",
        "statistica", "python", "c++", "java", "javascript", "docker", "git", "github",
        "gitlab", "aws", "software engineering", "debugging", "unit testing", "version control",
        "design architetturale",
    )
    project_terms = (
        "progetto", "progetti", "project", "portfolio", "prototipo", "realizzato",
        "sviluppato", "implementazione", "pipeline", "dataset", "analisi predittiva",
        "feature engineering", "cross-validation", "deploy", "deployment",
    )

    # Una "frase descrittiva" e' un periodo con almeno 4 parole e un verbo in
    # prima persona o un marker temporale: l'utente sta raccontando qualcosa,
    # non sta elencando una competenza secca. Va in ATTIVITA RILEVANTI come
    # paragrafo riformulato, NON in COMPETENZE TECNICHE come parola chiave.
    word_count = len(plain.split())
    descriptive_markers = (
        "ho usato", "ho utilizzato", "ho applicato", "ho fatto", "ho lavorato",
        "ho realizzato", "ho sviluppato", "ho implementato", "ho creato",
        "ho collaborato", "ho coordinato", "ho gestito", "ho analizzato",
        "ho costruito", "ho condotto", "sto usando", "ho preparato",
        "mi sono occupat", "mi occupo", "per fare", "per creare",
        "per analizzare", "durante", "nell'ambito", "nell ambito",
        # marker piu' generici: l'utente sta descrivendo un USO/applicazione
        " per i ", " per il ", " per la ", " per le ", " per gli ",
        " per analisi ", " per analizzare ", " per creare ",
        " usato ", " usata ", " utilizzato ", " utilizzata ",
        " applicat", " gestit", " realizzat", " sviluppat",
    )
    has_verb_marker = any(marker in f" {plain} " for marker in descriptive_markers)
    is_descriptive_sentence = word_count >= 4 and has_verb_marker

    # Priorita' di routing per il testo "info extra" dell'utente:
    # 1. CERTIFICAZIONI se ne parla esplicitamente
    # 2. PROGETTI se l'utente menziona un progetto / dataset / dashboard
    #    (anche se cita strumenti tecnici - l'intento e' descrivere un progetto reale)
    # 3. ESPERIENZE se cita azienda/tirocinio/stage
    # 4. CERTIFICAZIONI/FORMAZIONE specifiche
    # 5. Frase descrittiva (verbo + >=4 parole) senza marker espliciti -> ATTIVITA RILEVANTI
    # 6. solo come ultima risorsa: COMPETENZE TECNICHE
    has_language_level = bool(re.search(r"\b[abc][12]\b", plain))
    if score_terms(language_terms) and (has_language_level or "madrelingua" in plain):
        return "LINGUE", "languages"
    if score_terms(certification_terms):
        return "CERTIFICAZIONI", "certification"
    if score_terms(project_terms):
        return "PROGETTI", "project"
    if score_terms(experience_terms):
        return "ESPERIENZE PROFESSIONALI", "experience"
    # se l'utente cita laurea/università/esame senza altro contesto
    # ma menziona uno strumento tecnico, e' tipicamente un'attivita di studio
    # -> meglio PROGETTI o ATTIVITA RILEVANTI di una nuova competenza generica
    if any(term in plain for term in ["universita", "università", "esame", "corso"]):
        return "ATTIVITA RILEVANTI", "extra_page"
    # Frase narrativa con verbo: l'utente racconta un'esperienza/uso reale,
    # non sta elencando una skill. Mai trattarla come parola chiave secca.
    if is_descriptive_sentence:
        return "ATTIVITA RILEVANTI", "extra_page"

    scores = {
        "education": score_terms(education_terms),
        "experience": score_terms(experience_terms),
        "soft_skill": score_terms(soft_terms),
        "skill": score_terms(hard_terms),
    }
    best_section = max(scores, key=scores.get)
    if scores[best_section] > 0:
        if best_section == "education":
            return "FORMAZIONE", "education"
        if best_section == "experience":
            return "ESPERIENZE PROFESSIONALI", "experience"
        if best_section == "soft_skill":
            return "SOFT SKILLS", "soft_skill"
        if best_section == "skill":
            return "COMPETENZE TECNICHE", "skill"

    if any(term in plain for term in ["laurea", "universita", "università", "corso", "formazione", "studio", "esame"]):
        return "FORMAZIONE", "education"
    if any(term in plain for term in ["azienda", "cliente", "lavoro", "tirocinio", "stage", "impiego"]):
        return "ESPERIENZE PROFESSIONALI", "experience"
    if any(term in plain for term in soft_terms):
        return "SOFT SKILLS", "soft_skill"
    if any(term in plain for term in hard_terms):
        return "COMPETENZE TECNICHE", "skill"
    return "ATTIVITA RILEVANTI", "extra_page"


def is_role_like_confirmation(name: str, role: str) -> bool:
    normalized_name = normalize_plain_text(name)
    normalized_role = normalize_plain_text(role)
    if not normalized_name:
        return False
    if normalized_name == normalized_role:
        return True
    if normalized_role and normalized_role in normalized_name:
        return True
    generic_role_terms = {
        "project manager", "project manager stage", "data scientist", "data analyst",
        "software engineer", "manager", "developper", "developer", "engineer", "analyst",
        "specialist", "consultant", "researcher", "assistant", "intern", "stage"
    }
    if normalized_name in generic_role_terms:
        return True
    return False


def build_additional_rewrite_instructions(
    user_additional_data: Dict[str, Any],
    role: str,
    cv_text: str = "",
) -> List[RewriteInstruction]:
    instructions: List[RewriteInstruction] = []
    seen_fragments: set[tuple[str, str]] = set()

    def is_grounded_rewrite(source: str, replacement: str) -> bool:
        ignored = {
            "che", "con", "dalla", "dalle", "dati", "degli", "dei", "del", "della",
            "delle", "di", "e", "gli", "ha", "ho", "il", "in", "la", "le", "lo",
            "nel", "nella", "nelle", "per", "su", "un", "una",
        }
        source_terms = {
            term for term in re.findall(r"[a-z0-9]+", normalize_plain_text(source))
            if len(term) >= 3 and term not in ignored
        }
        replacement_terms = {
            term for term in re.findall(r"[a-z0-9]+", normalize_plain_text(replacement))
            if len(term) >= 3 and term not in ignored
        }
        if not source_terms:
            return False
        preserved = len(source_terms & replacement_terms) / len(source_terms)
        return preserved >= 0.55

    def professionalize_user_fact(fragment: str) -> str:
        text = re.sub(r"\s+", " ", fragment or "").strip(" .")
        if CV_REWRITE_LLM_ENABLED and text:
            prompt = f"""
Sei un resume editor. Riscrivi il fatto seguente in italiano professionale e conciso.
Non aggiungere informazioni, tecnologie, risultati, aziende, date o responsabilita.
Mantieni tutti i fatti presenti e restituisci SOLO JSON valido.

Fatto dell'utente:
{text}

Schema:
{{"replacement": "frase professionale"}}
"""
            try:
                result = call_rewrite_llm(
                    prompt,
                    context="professionalize_user_fact",
                    temperature=0.05,
                    max_tokens=220,
                    timeout=25,
                )
                replacement = str(result.get("replacement") or "").strip()
                if replacement and is_grounded_rewrite(text, replacement):
                    return replacement.rstrip(".") + "."
            except Exception as exc:
                print(f"Riscrittura Gemini del fatto utente non riuscita: {exc}")
        company_experience = re.match(
            r"^ho lavorato (?:in|presso)\s+(.+?)\s+dove ho "
            r"(realizzato|sviluppato|creato)\s+(.+)$",
            text,
            flags=re.IGNORECASE,
        )
        if company_experience:
            company, action, activity = company_experience.groups()
            action_label = {
                "realizzato": "realizzazione",
                "sviluppato": "sviluppo",
                "creato": "creazione",
            }.get(action.lower(), "realizzazione")
            text = f"Presso {company.strip()}, {action_label} di {activity.strip()}"
        replacements = [
            (r"^ho lavorato in\s+", "Attivita svolta presso "),
            (r"^ho lavorato presso\s+", "Attivita svolta presso "),
            (r"^ho realizzato\s+", "Realizzazione di "),
            (r"^ho sviluppato\s+", "Sviluppo di "),
            (r"^ho creato\s+", "Creazione di "),
            (r"^ho coordinato\s+", "Coordinamento di "),
            (r"^ho utilizzato\s+", "Utilizzo di "),
            (r"^ho usato\s+", "Utilizzo di "),
            (r"^l[' ]?ho utilizzato\s+", "Utilizzo di "),
            (r"^l[' ]?ho usato\s+", "Utilizzo di "),
            (r"^ho preso\s+", "Conseguimento di "),
        ]
        for pattern, replacement in replacements:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        text = re.sub(r"\bdove ho realizzato\b", "con realizzazione di", text, flags=re.IGNORECASE)
        text = re.sub(r"\bdove ho sviluppato\b", "con sviluppo di", text, flags=re.IGNORECASE)
        text = re.sub(r"\bper la spedizione di pacchi\b", "per la gestione delle spedizioni di pacchi", text, flags=re.IGNORECASE)
        return text[:1].upper() + text[1:] + "." if text else ""

    def project_block(fragment: str, force_project: bool = False) -> str:
        plain = normalize_plain_text(fragment)
        explicit_project = force_project or any(term in plain for term in [
            "ho sviluppato un progetto", "ho realizzato un progetto",
            "ho implementato un progetto", "ho coordinato un progetto",
            "ho lavorato a un progetto", "ho collaborato a un progetto",
            "progetto universitario", "progetto personale", "in progetto",
            "in progetti", "durante i progetti", "nei progetti",
        ])
        if not explicit_project:
            return ""
        if "machine learning" in plain or re.search(r"\bml\b", plain):
            title = "Progetto di Machine Learning"
        elif "ingegneria dei dati" in plain or "data engineering" in plain:
            title = "Progetto di Data Engineering"
        elif "sviluppo software" in plain:
            title = "Progetto di sviluppo software"
        elif "analisi dati" in plain or "data visualization" in plain:
            title = "Progetto di analisi dati"
        elif "dashboard" in plain or "power bi" in plain or "tableau" in plain:
            title = "Progetto di data visualization"
        elif "universitari" in plain or "esame" in plain:
            title = "Progetto universitario"
        else:
            title = "Progetto personale"

        short_project_phrases = [
            (("provenienza", "dati"), "Analisi della provenienza dei dati utilizzati nel progetto"),
            (("qualita", "dati"), "Verifica della qualita dei dati utilizzati nel progetto"),
            (("raccolta", "dati"), "Raccolta e organizzazione dei dati utilizzati nel progetto"),
            (("pulizia", "dati"), "Pulizia e preparazione dei dati utilizzati nel progetto"),
            (("dashboard",), "Sviluppo di dashboard per la lettura e la sintesi dei dati"),
            (("visualizzazione", "dati"), "Realizzazione di visualizzazioni per l'analisi dei dati"),
        ]
        description = professionalize_user_fact(fragment).rstrip(".")
        description = re.sub(r"^(progetto di [^\n]+[:\-]\s*)", "", description, flags=re.IGNORECASE)
        description = re.sub(r"^(data analyst|data scientist|project manager|software engineer)\s*:\s*", "", description, flags=re.IGNORECASE)
        replacements = [
            (r"^ho coordinato\b", "Coordinamento di"),
            (r"^ho sviluppato\b", "Sviluppo di"),
            (r"^ho realizzato\b", "Realizzazione di"),
            (r"^ho implementato\b", "Implementazione di"),
            (r"^ho lavorato a\b", "Attività svolte in"),
            (r"^ho collaborato a\b", "Collaborazione a"),
        ]
        for pattern, replacement in replacements:
            description = re.sub(pattern, replacement, description, flags=re.IGNORECASE)
        description = re.sub(r"^(usata?|utilizzata?)\s+", "Applicazione della competenza ", description, flags=re.IGNORECASE)
        if "allineare requisiti e aspettative" in plain:
            description = "Collaborazione con il team per l'analisi dei requisiti e l'allineamento delle attivita progettuali"
        description = re.sub(r"\s{2,}", " ", description).strip(" ,;:-")
        if not description:
            return ""
        if len(description.split()) <= 5:
            for markers, rewritten in short_project_phrases:
                if all(marker in plain for marker in markers):
                    description = rewritten
                    break
        if len(description.split()) <= 4:
            description = f"Attivita progettuale incentrata su {description.lower()}"
        description = description[:1].upper() + description[1:]
        return f"{title}\n{description.rstrip('.')}."

    def skill_text(fragment: str, category: str) -> str:
        professional = build_professional_extra_text({"additional_notes": fragment}, role)
        if not professional:
            return ""
        if category == "soft_skill":
            known = [
                "Pensiero analitico", "Attenzione ai dettagli", "Comunicazione dei risultati",
                "Problem solving", "Team working", "Collaborazione", "Comunicazione",
                "Organizzazione", "Precisione", "Gestione requisiti",
            ]
        else:
            known = [
                "Excel", "Power BI", "Tableau", "SQL", "Database relazionali",
                "Data visualization", "Analisi dati", "Python", "Java", "C++",
                "Docker", "Git", "Machine Learning", "Deep Learning", "Apache Spark",
                "Apache Hive", "MapReduce",
            ]
        hits = [
            skill for skill in known
            if normalize_plain_text(skill) in normalize_plain_text(fragment)
        ]
        if hits:
            return " Â· ".join(dict.fromkeys(hits))
        return professional.rstrip(".")

    def add_fragment(fragment: str, source_id: str, reason: str, category_hint: str = "") -> None:
        cleaned_fragment = clean_extracted_text(fragment).strip()
        if not cleaned_fragment:
            return
        normalized_hint = normalize_plain_text(category_hint)
        fragment_plain = normalize_plain_text(cleaned_fragment)
        language_values = []
        for language in ("inglese", "italiano", "francese", "spagnolo", "tedesco", "portoghese"):
            match = re.search(
                rf"(?:\b([abc][12])\b(?:\s+(?:di|in))?\s+{language}\b|\b{language}\b(?:\s+livello)?\s+\b([abc][12])\b)",
                fragment_plain,
                flags=re.IGNORECASE,
            )
            if match:
                language_values.append(f"{language.capitalize()} {(match.group(1) or match.group(2)).upper()}")
        if language_values and (
            normalized_hint in {"certifications", "certificazioni", "languages", "lingue"}
            or any(language.lower() in fragment_plain for language in language_values)
        ):
            instructions.append(RewriteInstruction(
                section="LINGUE",
                original="",
                replacement="\n".join(language_values),
                reason="Livello linguistico confermato dall'utente integrato nella sezione lingue.",
                category="languages",
                source_id=source_id,
            ))
            seen_fragments.add(("LINGUE", normalize_plain_text(cleaned_fragment)))
            return
        inferred_section, inferred_category = infer_extra_content_section(cleaned_fragment)
        section, category = inferred_section, inferred_category
        if normalized_hint in {"technical skills", "technical_skills", "tools"}:
            section, category = "COMPETENZE TECNICHE", "skill"
        elif normalized_hint in {"soft skills", "soft_skills"}:
            section, category = "SOFT SKILLS", "soft_skill"
        elif normalized_hint in {"certifications", "certificazioni"} and category != "languages":
            section, category = "CERTIFICAZIONI", "certification"
        elif normalized_hint in {"languages", "lingue"}:
            section, category = "LINGUE", "languages"
        elif normalized_hint in {"company role notes", "company_role_notes"}:
            section, category = "PROFILO", "profile"
        else:
            mapped_section_key = additional_field_section_key(category_hint)
            mapped_sections = {
                "experience": ("ESPERIENZE PROFESSIONALI", "experience"),
                "projects": ("PROGETTI", "project"),
                "certifications": ("CERTIFICAZIONI", "certification"),
                "languages": ("LINGUE", "languages"),
                "education": ("FORMAZIONE", "education"),
                "hard_skills": ("COMPETENZE TECNICHE", "skill"),
                "soft_skills": ("SOFT SKILLS", "soft_skill"),
                "profile": ("PROFILO", "profile"),
            }
            if mapped_section_key in mapped_sections:
                section, category = mapped_sections[mapped_section_key]
        if normalized_hint in {"measurable results", "measurable_results"}:
            has_project_context = bool(str(
                (user_additional_data or {}).get("projects") or ""
            ).strip()) or any(term in fragment_plain for term in [
                "progetto", "progetti", "analisi", "metriche", "kpi",
                "indicatori", "dataset", "modello", "report",
            ])
            section, category = (
                ("PROGETTI", "project")
                if has_project_context
                else ("ATTIVITA RILEVANTI", "extra_page")
            )
        if category == "project" or normalized_hint in {"project", "projects", "progetto", "progetti"}:
            professional_text = project_block(
                cleaned_fragment,
                force_project=(
                    category == "project"
                    or normalized_hint in {"project", "projects", "progetto", "progetti"}
                ),
            )
            if professional_text:
                project_lines = [line.strip() for line in professional_text.splitlines() if line.strip()]
                if len(project_lines) == 1:
                    project_lines = [project_lines[0], "Attivita progettuale descritta dall'utente."]
                if len(project_lines) >= 2:
                    project_lines[1] = project_lines[1].rstrip(".") + "."
                    professional_text = "\n".join(project_lines[:2])
                section, category = "PROGETTI", "project"
            else:
                professional_text = skill_text(cleaned_fragment, "skill")
                section, category = "COMPETENZE TECNICHE", "skill"
        elif category in {"skill", "soft_skill"}:
            if normalized_hint not in {
                "technical skills", "technical_skills", "tools",
                "soft skills", "soft_skills",
            }:
                # L'utente ha scritto qualcosa che contiene parole chiave
                # tecniche ma non e' stato indirizzato esplicitamente come
                # skill (es. additional_notes o adaptation answer senza
                # category): non scartare silenziosamente. Riformula come
                # paragrafo in ATTIVITA RILEVANTI.
                professional_text = build_professional_extra_text(
                    {"additional_notes": cleaned_fragment}, role
                )
                section, category = "ATTIVITA RILEVANTI", "extra_page"
                if not professional_text:
                    return
            else:
                professional_text = skill_text(cleaned_fragment, category)
                section = "SOFT SKILLS" if category == "soft_skill" else "COMPETENZE TECNICHE"
                if not professional_text:
                    return
        elif category == "extra_page":
            professional_text = build_professional_extra_text({"additional_notes": cleaned_fragment}, role)
            section, category = "ATTIVITA RILEVANTI", "extra_page"
            if not professional_text:
                return
        elif category == "profile":
            professional_text = build_professional_extra_text({"additional_notes": cleaned_fragment}, role)
            if not professional_text:
                return
        elif category == "experience":
            if not any(
                term in fragment_plain
                for term in ["azienda", "cliente", "lavoro", "lavorato", "presso", "tirocinio", "stage", "impiego", "ruolo"]
            ):
                # Se l'utente ha descritto un'esperienza in forma sintetica,
                # la manteniamo comunque ma la teniamo nella sezione corretta.
                section = "ESPERIENZE PROFESSIONALI"
            professional_text = professionalize_user_fact(cleaned_fragment)
            if not professional_text:
                return
        elif category in {"education", "certification", "languages"}:
            professional_text = professionalize_user_fact(cleaned_fragment)
            if category == "certification":
                professional_text = re.sub(r"\bFrancece\b", "Francese", professional_text, flags=re.IGNORECASE)
                cert_plain = normalize_plain_text(professional_text)
                level_match = re.search(r"\b([abc][12])\b", cert_plain)
                if level_match and "cambridge" in cert_plain:
                    language = "Francese" if "francese" in cert_plain else "linguistica"
                    professional_text = f"Certificazione {language} {level_match.group(1).upper()} Cambridge."
            if not professional_text:
                return
        else:
            professional_text = build_professional_extra_text({"additional_notes": cleaned_fragment}, role)
            if not professional_text:
                return
        print(
            f"[EXTRA-INFO] source_field={category_hint or '-'} "
            f"inferred_section={inferred_section!r}, final_section={section!r}, "
            f"category={category!r}, fragment='{cleaned_fragment[:120]}'"
        )
        key = (section, normalize_plain_text(cleaned_fragment))
        if key in seen_fragments:
            return
        seen_fragments.add(key)
        instructions.append(RewriteInstruction(
            section=section,
            original="",
            replacement=professional_text[:500],
            reason=reason,
            category=category,
            source_id=source_id,
        ))

    answers = (user_additional_data or {}).get("adaptation_answers", [])
    if isinstance(answers, list):
        for index, item in enumerate(answers):
            if not isinstance(item, dict):
                continue
            answer = str(item.get("answer") or "").strip()
            if not answer:
                continue
            category_hint = str(item.get("category") or "").strip()
            fragments = [frag.strip() for frag in re.split(r"\n+|(?<=[.!?])\s+", answer) if frag.strip()]
            if not fragments:
                fragments = [answer]
            for fragment_index, fragment in enumerate(fragments):
                safe_category_hint = re.sub(r"[^a-z0-9_]+", "_", normalize_plain_text(category_hint)).strip("_")
                add_fragment(
                    fragment,
                    f"user_additional_answer_{index}_{fragment_index}_{safe_category_hint or 'additional_notes'}",
                    "Risposta aggiuntiva confermata dall'utente trasformata in contenuto CV.",
                    category_hint,
                )

    for field_name, raw_value in (user_additional_data or {}).items():
        if field_name in {"adaptation_answers", "confirmed_skills"}:
            continue
        if not isinstance(raw_value, str):
            continue
        category_hint = field_name if additional_field_section_key(field_name) else ""
        value = str((user_additional_data or {}).get(field_name) or "").strip()
        if not value:
            continue
        split_pattern = r"\n+"
        if field_name == "projects":
            split_pattern += r"|\s+(?:e|ed)\s+(?=progett[oi]\s+(?:di|su|per)\b)"
        elif field_name not in {"experiences", "company_role_notes", "additional_notes"}:
            split_pattern += r"|(?<=[.!?])\s+"
        fragments = [frag.strip() for frag in re.split(split_pattern, value, flags=re.IGNORECASE) if frag.strip()] or [value]
        for fragment_index, fragment in enumerate(fragments):
            add_fragment(
                fragment,
                f"user_box_{field_name}_{fragment_index}",
                "Informazione inserita dall'utente e riformulata per il CV.",
                category_hint,
            )
    return instructions


def _extract_sections_for_structured_suggestions(cv_text: str) -> Dict[str, str]:
    return extract_resume_sections(cv_text)


def consolidate_rewrite_instructions(
    cv_text: str,
    instructions: List[RewriteInstruction],
    company: str,
    role: str,
    goal: str,
    use_llm: bool = True,
) -> List[RewriteInstruction]:
    if len(instructions) <= 1:
        return instructions

    grouped: Dict[str, List[RewriteInstruction]] = {}
    for instruction in instructions:
        section = canonical_edit_section_name(instruction.section) or instruction.section.strip().upper()
        grouped.setdefault(section, []).append(instruction)

    section_map = extract_resume_sections(cv_text)
    consolidated: List[RewriteInstruction] = []
    for section, section_instructions in grouped.items():
        if len(section_instructions) == 1:
            consolidated.append(section_instructions[0])
            continue

        section_key = _resume_section_key(section)
        original_section = section_map.get(section_key, "").strip()
        source_ids = [
            instruction.source_id or f"{section_key}_{index}"
            for index, instruction in enumerate(section_instructions)
        ]
        changes = [
            {
                "id": source_ids[index],
                "original": instruction.original,
                "replacement": instruction.replacement,
                "reason": instruction.reason,
                "category": instruction.category,
            }
            for index, instruction in enumerate(section_instructions)
        ]
        prompt = f"""
Sei un senior resume editor. Restituisci SOLO JSON valido.

Devi creare la versione finale di UNA SOLA sezione del CV incorporando TUTTE le modifiche accettate.

Candidatura corrente:
- Azienda: {company or "Non specificata"}
- Ruolo: {role or "Non specificato"}
- Obiettivo/annuncio: {goal or "Non specificato"}

Sezione:
{section}

Testo originale completo della sezione:
{original_section or "Sezione non presente"}

Modifiche accettate da integrare tutte:
{json.dumps(changes, ensure_ascii=False)}

Schema:
{{
  "replacement": "testo finale completo della sezione",
  "applied_ids": {json.dumps(source_ids, ensure_ascii=False)}
}}

Regole obbligatorie:
- Integra tutte le modifiche elencate in un unico testo finale coerente.
- Non perdere nessuna skill, informazione o riscrittura accettata.
- Non duplicare concetti sovrapposti: fondili mantenendo tutto il contenuto utile.
- Usa esclusivamente fatti presenti nel testo originale o nelle modifiche accettate.
- Non inventare aziende, date, strumenti, risultati, esperienze o titoli.
- Mantieni lingua, tono, sintesi, punteggiatura e forma del CV originale.
- Se la sezione contiene skill, conserva tutte le skill originali e tutte quelle confermate, organizzandole in modo leggibile.
- Non inserire il titolo della sezione nel replacement.
- applied_ids deve contenere esattamente tutti gli ID ricevuti.
"""
        replacement = ""
        applied_ids: List[str] = []
        contains_user_facts = any(
            (instruction.source_id or "").startswith((
                "user_box_",
                "user_additional_",
                "confirmed_skill_detail_",
            ))
            for instruction in section_instructions
        )
        if use_llm and not contains_user_facts:
            try:
                result = call_rewrite_llm(
                    prompt,
                    context="consolidate_rewrite_instructions",
                    temperature=0.05,
                    max_tokens=1200,
                    timeout=60,
                )
                replacement = str(result.get("replacement") or "").strip()
                applied_ids = [
                    str(value).strip()
                    for value in result.get("applied_ids", [])
                    if str(value).strip()
                ] if isinstance(result.get("applied_ids"), list) else []
            except Exception as exc:
                print(f"Consolidamento sezione {section} non riuscito: {exc}")

        if (
            not replacement
            or set(applied_ids) != set(source_ids)
            or not ResumeRewriter().is_safe_replacement(section, replacement)
            or not all(
                rewrite_preserves_instruction_content(replacement, instruction)
                for instruction in section_instructions
            )
        ):
            replacement = deterministic_section_consolidation(
                original_section,
                section_instructions,
            )

        if not replacement:
            consolidated.extend(section_instructions)
            continue

        consolidated.append(RewriteInstruction(
            section=section,
            original=original_section or section_instructions[0].original,
            replacement=replacement,
            reason="Tutte le modifiche accettate per la sezione sono state consolidate.",
            category=section_instructions[0].category,
            source_id="consolidated:" + "|".join(source_ids),
        ))

    return consolidated


def deterministic_section_consolidation(
    original_section: str,
    instructions: List[RewriteInstruction],
) -> str:
    section = canonical_edit_section_name(instructions[0].section) if instructions else ""
    if section in {"HARD SKILLS", "SOFT SKILLS", "COMPETENZE TECNICHE"}:
        is_soft = section == "SOFT SKILLS"
        values: List[str] = extract_clean_skill_items(
            original_section,
            is_soft=is_soft,
        )
        seen_skill_keys = set()
        clean_values: List[str] = []
        for value in values:
            key = canonical_skill_identity(value)
            if key and key not in seen_skill_keys:
                seen_skill_keys.add(key)
                clean_values.append(value)
        for source in [item.replacement for item in instructions]:
            for value in re.split(r"[,;|•·\n]+", source or ""):
                clean = value.strip(" -\t")
                key = canonical_skill_identity(clean)
                if (
                    clean
                    and 1 <= len(clean.split()) <= 4
                    and key
                    and key not in seen_skill_keys
                ):
                    seen_skill_keys.add(key)
                    clean_values.append(clean)
        limit = 6 if is_soft else 14
        return format_skill_list_like_original(original_section, clean_values[:limit])

    replacement_instructions = [item for item in instructions if item.original.strip()]
    additions = [item.replacement.strip() for item in instructions if not item.original.strip()]
    base = (
        replacement_instructions[-1].replacement.strip()
        if replacement_instructions
        else original_section.strip()
    )
    parts = [base, *additions]
    unique_parts: List[str] = []
    for part in parts:
        if not part:
            continue
        normalized = normalize_plain_text(part)
        if normalized and normalized not in {
            normalize_plain_text(existing) for existing in unique_parts
        }:
            unique_parts.append(part)
    return "\n".join(unique_parts).strip()


def rewrite_preserves_instruction_content(
    consolidated_text: str,
    instruction: RewriteInstruction,
) -> bool:
    consolidated_tokens = set(re.findall(
        r"[a-z0-9+#.]{2,}",
        normalize_plain_text(consolidated_text),
    ))
    replacement_tokens = set(re.findall(
        r"[a-z0-9+#.]{2,}",
        normalize_plain_text(instruction.replacement),
    ))
    original_tokens = set(re.findall(
        r"[a-z0-9+#.]{2,}",
        normalize_plain_text(instruction.original),
    ))
    required_tokens = replacement_tokens - original_tokens
    if not required_tokens:
        required_tokens = replacement_tokens
    if not required_tokens:
        return True
    overlap = len(required_tokens & consolidated_tokens) / len(required_tokens)
    return overlap >= 0.6


def build_confirmed_skill_rewrite_instructions(
    cv_text: str,
    user_additional_data: Dict[str, Any],
    role: str = "",
) -> List[RewriteInstruction]:
    confirmed = user_additional_data.get("confirmed_skills", [])
    if not isinstance(confirmed, list) or not confirmed:
        return []
    skill_names_by_section: Dict[str, List[str]] = {}
    detailed_skills: List[Dict[str, str]] = []
    for item in confirmed:
        if isinstance(item, dict):
            name = str(item.get("name") or "").strip()
            detail = str(item.get("user_example") or item.get("detail") or "").strip()
            category = normalize_plain_text(str(item.get("category") or "hard_skill"))
        else:
            name = str(item or "").strip()
            detail = ""
            category = "hard_skill"
        if category == "keyword":
            category = "hard skill"
        if category not in {"hard skill", "soft skill", "tool", "language"}:
            continue
        if (
            not name
            or is_role_like_confirmation(name, role)
            or skill_semantically_present(cv_text, name)
        ):
            continue
        section = "SOFT SKILLS" if category == "soft skill" else "HARD SKILLS"
        existing_normalized = {
            canonical_skill_identity(existing)
            for existing in skill_names_by_section.get(section, [])
        }
        if canonical_skill_identity(name) not in existing_normalized:
            skill_names_by_section.setdefault(section, []).append(name)
        if detail:
            detailed_skills.append({
                "name": name,
                "detail": detail,
                "category": category,
            })
    if not skill_names_by_section and not detailed_skills:
        return []

    sections = extract_resume_sections(cv_text)
    instructions: List[RewriteInstruction] = []
    for section, skill_names in skill_names_by_section.items():
        is_soft = normalize_plain_text(section) == "soft skills"
        raw_original = sections.get("soft_skills" if is_soft else "hard_skills", "")
        original = clean_skill_section_source(raw_original)
        existing_parts = extract_clean_skill_items(original, is_soft=is_soft)
        merged = list(dict.fromkeys([*existing_parts, *skill_names]))
        instructions.append(RewriteInstruction(
            section=section if original else ("SOFT SKILLS" if is_soft else "COMPETENZE TECNICHE"),
            original=original,
            replacement=format_skill_list_like_original(original, merged),
            reason="Skill confermate dall'utente integrate nella sezione corretta.",
            category="soft_skills" if is_soft else "skills",
            source_id=f"confirmed_{normalize_plain_text(section).replace(' ', '_')}",
        ))
    for index, item in enumerate(detailed_skills[:6]):
        detail_instruction = fallback_skill_detail_instruction(item, index, cv_text)
        if detail_instruction is not None:
            instructions.append(detail_instruction)
    return instructions


def clean_skill_section_source(value: str) -> str:
    """Keep only skill-list lines when DOCX extraction crosses column boundaries."""
    value = re.split(
        r"\b[A-ZÀ-ÖØ-Þ]{2,}(?:\s+[A-ZÀ-ÖØ-Þ]{2,})+\b",
        str(value or ""),
        maxsplit=1,
    )[0]
    kept: List[str] = []
    blocked_terms = {
        "universita", "università", "laurea", "diploma", "liceo",
        "formazione", "esperienza", "tirocinio", "curriculare",
    }
    for raw_line in str(value or "").splitlines():
        line = raw_line.strip()
        plain = normalize_plain_text(line)
        if not line:
            continue
        if kept and (
            len(line.split()) > 8
            or any(term in plain for term in blocked_terms)
            or bool(re.search(r"\b(?:19|20)\d{2}\b", line))
            or line.endswith((".", "!", "?"))
        ):
            break
        kept.append(line)
        if len(kept) >= 3:
            break
    return "\n".join(kept).strip()


def canonical_skill_identity(value: str) -> str:
    normalized = normalize_plain_text(value)
    normalized = re.sub(
        r"\b(programming|programmazione|language|linguaggio|framework|tool|strumento|skills?|competenze?)\b",
        "",
        normalized,
    )
    normalized = re.sub(r"\s+", " ", normalized).strip()
    aliases = {
        "ml ai": "machine learning",
        "ml & ai": "machine learning",
        "ai ml": "machine learning",
        "ai & ml": "machine learning",
        "machine learning ai": "machine learning",
        "team working": "collaborazione",
        "teamwork": "collaborazione",
        "collaborazione in team": "collaborazione",
        "lavoro di squadra": "collaborazione",
        "lavoro in team": "collaborazione",
        "team collaboration": "collaborazione",
        "collaboration": "collaborazione",
        "collaborative working": "collaborazione",
        "risoluzione dei problemi": "problem solving",
        "risoluzione problemi": "problem solving",
        "gestione delle priorita": "gestione priorita",
        "priority management": "gestione priorita",
        "communication": "comunicazione",
        "communication skills": "comunicazione",
        "organization": "organizzazione",
        "organisational skills": "organizzazione",
        "organizational skills": "organizzazione",
        "leadership skills": "leadership",
        "time management": "gestione tempo",
        "gestione del tempo": "gestione tempo",
        "critical thinking": "pensiero critico",
        "analytical thinking": "pensiero analitico",
        "attention to detail": "attenzione ai dettagli",
        "adaptability": "adattabilita",
        "flexibility": "flessibilita",
        "data visualisation": "data visualization",
        "visualizzazione dati": "data visualization",
        "visualizzazione dei dati": "data visualization",
        "analisi dei dati": "analisi dati",
        "data analysis": "analisi dati",
        "data analytics": "analisi dati",
        "powerbi": "power bi",
        "google bigquery": "bigquery",
        "version control": "controllo versione",
        "controllo di versione": "controllo versione",
        "source control": "controllo versione",
        "unit tests": "unit testing",
        "test unitari": "unit testing",
        "rest api": "api rest",
        "restful api": "api rest",
        "api development": "sviluppo api",
        "sviluppo di api": "sviluppo api",
        "database design": "progettazione database",
        "progettazione di database": "progettazione database",
        "risk management": "gestione rischi",
        "budget management": "gestione budget",
        "stakeholder communication": "comunicazione stakeholder",
        "comunicazione con stakeholder": "comunicazione stakeholder",
    }
    return aliases.get(normalized, normalized)


def skill_semantically_present(cv_text: str, skill: str) -> bool:
    identity = canonical_skill_identity(skill)
    if not identity:
        return True
    equivalent_groups = {
        "collaborazione": {
            "collaborazione", "team working", "teamwork", "lavoro di squadra",
            "lavoro in team", "collaborazione in team", "team collaboration",
        },
        "problem solving": {
            "problem solving", "problem-solving", "risoluzione dei problemi",
        },
        "gestione priorita": {
            "gestione priorita", "gestione delle priorita", "priority management",
        },
        "comunicazione": {"comunicazione", "communication"},
        "organizzazione": {"organizzazione", "organization"},
        "gestione tempo": {"gestione tempo", "gestione del tempo", "time management"},
        "pensiero critico": {"pensiero critico", "critical thinking"},
        "pensiero analitico": {"pensiero analitico", "analytical thinking"},
        "attenzione ai dettagli": {"attenzione ai dettagli", "attention to detail"},
        "adattabilita": {"adattabilita", "adaptability"},
        "flessibilita": {"flessibilita", "flexibility"},
        "data visualization": {
            "data visualization", "data visualisation", "visualizzazione dati",
            "visualizzazione dei dati",
        },
        "analisi dati": {"analisi dati", "analisi dei dati", "data analysis", "data analytics"},
        "controllo versione": {
            "controllo versione", "controllo di versione", "version control", "source control",
        },
        "unit testing": {"unit testing", "unit tests", "test unitari"},
        "api rest": {"api rest", "rest api", "restful api"},
        "sviluppo api": {"sviluppo api", "sviluppo di api", "api development"},
        "progettazione database": {
            "progettazione database", "progettazione di database", "database design",
        },
        "gestione rischi": {"gestione rischi", "risk management"},
        "gestione budget": {"gestione budget", "budget management"},
        "comunicazione stakeholder": {
            "comunicazione stakeholder", "comunicazione con stakeholder",
            "stakeholder communication",
        },
    }
    variants = {
        identity,
        normalize_plain_text(skill),
        *equivalent_groups.get(identity, set()),
    }
    cv_plain = normalize_plain_text(cv_text)
    if keyword_group_present(cv_plain, list(variants)):
        return True

    # Catch close wording variants without equating unrelated one-word skills.
    identity_tokens = {
        token for token in re.findall(r"[a-z0-9+#.]+", identity)
        if len(token) > 2
    }
    if len(identity_tokens) < 2:
        return False
    cv_tokens = set(re.findall(r"[a-z0-9+#.]+", cv_plain))
    return len(identity_tokens & cv_tokens) / len(identity_tokens) >= 0.8


def extract_clean_skill_items(value: str, is_soft: bool = False) -> List[str]:
    source = clean_skill_section_source(value)
    if not source:
        return []
    if is_soft:
        candidates = [
            "Pensiero analitico", "Attenzione ai dettagli",
            "Comunicazione dei risultati", "Problem solving",
            "Collaborazione", "Creatività", "Flessibilità",
            "Capacità di adattamento", "Apprendimento continuo",
            "Team working", "Comunicazione", "Organizzazione",
            "Precisione", "Leadership", "Negoziazione",
        ]
        source_plain = normalize_plain_text(source)
        found = [
            (source_plain.find(normalize_plain_text(skill)), skill)
            for skill in candidates
            if normalize_plain_text(skill) in source_plain
        ]
        return [
            skill for _, skill in sorted(found, key=lambda item: item[0])
        ]

    try:
        from services.cv_optimizer.structured_cv_engine import extract_skill_terms

        extracted = extract_skill_terms(source)
    except Exception:
        extracted = re.split(r"[,;|•·\n]+", source)
    return [
        str(skill).strip()
        for skill in extracted
        if str(skill).strip()
        and not normalize_plain_text(str(skill)).startswith(("in ", "con ", "e "))
    ]


def format_skill_list_like_original(original: str, skills: List[str]) -> str:
    clean_skills: List[str] = []
    seen = set()
    for skill in skills:
        clean = re.sub(r"\s+", " ", str(skill or "")).strip(" -·•|,;")
        key = canonical_skill_identity(clean)
        if not clean or key in seen:
            continue
        seen.add(key)
        clean_skills.append(clean)
    if not clean_skills:
        return ""

    rows: List[str] = []
    current = ""
    for skill in clean_skills:
        candidate = f"{current} · {skill}" if current else skill
        if current and len(candidate) > 44:
            rows.append(current)
            current = skill
        else:
            current = candidate
    if current:
        rows.append(current)
    return "\n".join(rows)


def build_skill_detail_rewrite_instructions(
    cv_text: str,
    detailed_skills: List[Dict[str, str]],
    role: str,
) -> List[RewriteInstruction]:
    if not detailed_skills:
        return []

    sections = ResumeParser().parse_text(cv_text)
    section_payload = [
        {
            "name": section.name,
            "heading": section.heading,
            "text": section.text[:1800],
        }
        for section in sections
    ]
    prompt = f"""
Sei un resume editor. Restituisci SOLO JSON valido.

Devi trasformare esempi reali forniti dall'utente in brevi contenuti professionali da aggiungere al CV.

Ruolo target:
{role or "Non specificato"}

Sezioni e stile testuale del CV originale:
{json.dumps(section_payload, ensure_ascii=False)}

Skill confermate con dettagli reali:
{json.dumps(detailed_skills, ensure_ascii=False)}

Schema:
{{
  "items": [
    {{
      "skill": "nome skill",
      "section": "PROGETTI | ESPERIENZE PROFESSIONALI | FORMAZIONE | CERTIFICAZIONI | ATTIVITA RILEVANTI",
      "text": "testo professionale breve da inserire"
    }}
  ]
}}

Regole:
- Usa esclusivamente i fatti scritti dall'utente; non inventare aziende, ruoli, date, risultati, strumenti o responsabilita.
- Scegli per ogni dettaglio la sezione semanticamente piu utile. Non mettere tutto automaticamente in PROGETTI.
- Se il dettaglio riguarda lavoro o tirocinio usa ESPERIENZE PROFESSIONALI; studi o corsi usa FORMAZIONE; progetti usa PROGETTI; attestati usa CERTIFICAZIONI.
- Mantieni nel testo un riferimento naturale alla skill indicata, soprattutto per le soft skill.
- Scrivi una frase discorsiva: non usare il formato "Nome skill: descrizione".
- Integra la skill nel periodo, ad esempio "Approccio analitico applicato alla..." o "Collaborazione con il team durante...".
- Se il dettaglio non contiene un contesto sufficiente, usa ATTIVITA RILEVANTI invece delle sezioni HARD SKILLS o SOFT SKILLS.
- Mantieni lingua, tono, lunghezza, forma dei bullet e livello di sintesi osservati nel CV originale.
- Puoi migliorare grammatica e chiarezza senza cambiare il significato.
- Ogni testo deve essere breve e direttamente utilizzabile nel CV, senza note, spiegazioni o titoli interni.
"""
    try:
        result = call_rewrite_llm(prompt, context="skill_detail_rewrite", temperature=0.1, max_tokens=700, timeout=30)
        raw_items = result.get("items") if isinstance(result, dict) else []
    except Exception as exc:
        print(f"Collocazione dettagli skill non riuscita, uso fallback: {exc}")
        raw_items = []

    allowed_sections = {
        "PROGETTI": "project",
        "ESPERIENZE PROFESSIONALI": "experience",
        "FORMAZIONE": "education",
        "CERTIFICAZIONI": "certification",
        "ATTIVITA RILEVANTI": "extra_page",
    }
    instructions = []
    for index, item in enumerate(raw_items if isinstance(raw_items, list) else []):
        if not isinstance(item, dict):
            continue
        section = str(item.get("section") or "").strip().upper()
        replacement = clean_extracted_text(str(item.get("text") or ""))
        if section not in allowed_sections or not replacement:
            continue
        skill_name = str(item.get("skill") or "").strip()
        if skill_name and normalize_plain_text(skill_name) not in normalize_plain_text(replacement):
            replacement = build_discursive_skill_evidence(skill_name, replacement)
        instructions.append(RewriteInstruction(
            section=section,
            original="",
            replacement=replacement[:500],
            reason="Dettaglio reale fornito dall'utente, riformulato e collocato nella sezione più coerente.",
            category=allowed_sections[section],
            source_id=f"confirmed_skill_detail_{index}",
        ))

    if instructions:
        return instructions[:6]

    return [
        fallback_skill_detail_instruction(item, index)
        for index, item in enumerate(detailed_skills[:6])
    ]


def fallback_skill_detail_instruction(
    item: Dict[str, str],
    index: int,
    cv_text: str = "",
) -> Optional[RewriteInstruction]:
    skill_name = str(item.get("name") or "").strip()
    detail = str(item.get("detail") or "").strip()
    plain = normalize_plain_text(detail)
    professional = build_deterministic_skill_evidence(skill_name, detail)
    original_sections = extract_resume_sections(cv_text)
    category_name = normalize_plain_text(str(item.get("category") or "hard_skill"))
    soft_skill_names = {
        normalize_plain_text(value)
        for value in [
            "Pensiero analitico", "Attenzione ai dettagli", "Comunicazione dei risultati",
            "Problem solving", "Team working", "Collaborazione", "Comunicazione",
            "Organizzazione", "Precisione", "Creativita", "Flessibilita",
            "Capacita di adattamento", "Apprendimento continuo", "Leadership",
            "Negoziazione", "Gestione requisiti",
        ]
    }
    is_soft_skill = category_name == "soft_skill" or normalize_plain_text(skill_name) in soft_skill_names

    if is_soft_skill:
        evidence = build_discursive_skill_evidence(skill_name, detail)
        original_profile = (original_sections.get("profile") or "").strip()
        if original_profile:
            normalized_profile = normalize_plain_text(original_profile)
            replacement = (
                original_profile
                if normalize_plain_text(evidence) in normalized_profile
                else f"{original_profile.rstrip('.')}.\n{evidence}"
            )
        else:
            replacement = evidence
        return RewriteInstruction(
            section="PROFILO",
            original=original_profile,
            replacement=replacement,
            reason="Esempio reale di soft skill integrato nel profilo professionale.",
            category="profile",
            source_id=f"confirmed_skill_detail_profile_{index}",
        )

    if any(term in plain for term in ["lavoro", "azienda", "tirocinio", "stage", "cliente"]):
        section, category = "ESPERIENZE PROFESSIONALI", "experience"
    elif any(term in plain for term in ["certificazione", "attestato", "certificato"]):
        section, category = "CERTIFICAZIONI", "certification"
    elif any(term in plain for term in ["progetto", "progetti", "dashboard", "prototipo", "dataset", "portfolio"]):
        section, category = "PROGETTI", "project"
    elif any(term in plain for term in ["laurea", "universita", "corso", "esame", "formazione"]):
        section, category = "FORMAZIONE", "education"
    else:
        # Senza un contesto reale la hard skill resta nell'elenco competenze:
        # non creiamo una sezione narrativa generica.
        return None

    if section == "PROGETTI":
        if "sviluppo software" in plain:
            title = "Progetto di sviluppo software"
        elif "analisi dati" in plain or "dashboard" in plain:
            title = "Progetto di analisi dati"
        else:
            title = "Progetto universitario" if any(term in plain for term in ["universita", "universitario", "esame"]) else "Progetto personale"
        if "allineare requisiti e aspettative" in plain:
            replacement = (
                f"{title}\n"
                "Collaborazione con il team per l'analisi dei requisiti e "
                "l'allineamento delle attivita progettuali."
            )
        else:
            replacement = f"{title}\n{professional.rstrip('.')}."
    elif section in {"COMPETENZE TECNICHE", "SOFT SKILLS"}:
        context = professional.rstrip(".")
        context = re.sub(
            rf"^(utilizzo|applicazione)\s+(di|della|del)?\s*{re.escape(skill_name)}\s*",
            "",
            context,
            flags=re.IGNORECASE,
        ).strip(" :-")
        replacement = f"{skill_name}: {context}" if context else skill_name
    else:
        context = professional.rstrip(".")
        if skill_name and normalize_plain_text(skill_name) not in normalize_plain_text(context):
            replacement = build_discursive_skill_evidence(skill_name, context)
        else:
            replacement = context + "."

    return RewriteInstruction(
        section=section,
        original="",
        replacement=replacement,
        reason="Dettaglio reale fornito dall'utente collocato con fallback conservativo.",
        category=category,
        source_id=f"confirmed_skill_detail_fallback_{index}",
    )


def build_discursive_skill_evidence(skill_name: str, detail: str) -> str:
    skill_plain = normalize_plain_text(skill_name)
    clean_detail = re.sub(r"\s+", " ", detail or "").strip(" .,:;-")
    if not clean_detail:
        return skill_name
    clean_detail = re.sub(
        r"^(applicazione|utilizzo)\s+(in|durante)\s+|^usat[oaie]\s+in\s+",
        "",
        clean_detail,
        flags=re.IGNORECASE,
    ).strip()

    starters = {
        "pensiero analitico": "Approccio analitico applicato",
        "problem solving": "Capacita di problem solving applicata",
        "collaborazione": "Collaborazione con il team dimostrata",
        "team working": "Lavoro in team svolto",
        "comunicazione": "Comunicazione efficace utilizzata",
        "comunicazione dei risultati": "Comunicazione dei risultati curata",
        "attenzione ai dettagli": "Attenzione ai dettagli mantenuta",
        "organizzazione": "Capacita organizzativa applicata",
        "precisione": "Precisione applicata",
        "leadership": "Leadership esercitata",
        "gestione requisiti": "Gestione dei requisiti svolta",
    }
    starter = starters.get(skill_plain, f"Competenza di {skill_name.lower()} applicata")
    detail_plain = normalize_plain_text(clean_detail)
    if detail_plain.startswith("progetto "):
        clean_detail = f"un {clean_detail}"
        detail_plain = normalize_plain_text(clean_detail)
    if detail_plain.startswith(("a ", "al ", "alla ", "alle ", "ai ", "allo ", "per ", "durante ", "nella ", "nel ")):
        sentence = f"{starter} {clean_detail}"
    else:
        sentence = f"{starter} durante {clean_detail[:1].lower() + clean_detail[1:]}"
    return sentence.rstrip(".") + "."


def build_deterministic_skill_evidence(skill_name: str, detail: str) -> str:
    clean_detail = re.sub(r"\s+", " ", detail or "").strip(" .,:;-")
    if not clean_detail:
        return skill_name
    if normalize_plain_text(skill_name) in normalize_plain_text(clean_detail):
        return clean_detail[:1].upper() + clean_detail[1:]
    without_lead = re.sub(
        r"^(?:l[' ]?ho\s+)?usat[oaie]?\s+(?:in|durante)\s+",
        "",
        clean_detail,
        flags=re.IGNORECASE,
    ).strip()
    if without_lead != clean_detail:
        return f"Utilizzo di {skill_name} in {without_lead}"
    without_personal_lead = re.sub(
        r"^(?:l[' ]?ho|ho)\s+(?:usato|utilizzato|applicato)\s+",
        "",
        clean_detail,
        flags=re.IGNORECASE,
    ).strip()
    if without_personal_lead != clean_detail:
        return f"Utilizzo di {skill_name} durante {without_personal_lead}"
    return f"Utilizzo di {skill_name}: {clean_detail[:1].lower() + clean_detail[1:]}"


def format_confirmed_skill_example(skill_name: str, detail: str) -> str:
    clean_detail = re.sub(r"\s+", " ", detail or "").strip().rstrip(".")
    if not clean_detail:
        return skill_name
    if len(clean_detail) > 180:
        clean_detail = clean_detail[:180].rsplit(" ", 1)[0].strip()
    return f"Utilizzo di {skill_name}: {clean_detail}."


def group_rewrite_instructions(instructions: List[RewriteInstruction]) -> Dict[str, List[Dict[str, str]]]:
    grouped = {
        "profile_updates": [],
        "hard_skills_updates": [],
        "soft_skills_updates": [],
        "experience_updates": [],
        "education_updates": [],
        "projects_updates": [],
        "extra_sections": [],
    }
    for instruction in instructions:
        section = normalize_plain_text(instruction.section)
        category = normalize_plain_text(instruction.category)
        payload = {
            "id": instruction.source_id,
            "section": instruction.section,
            "replacement": instruction.replacement,
        }
        if any(term in section for term in ["chi sono", "profilo"]) or category == "profile":
            grouped["profile_updates"].append(payload)
        elif "soft" in section or category == "soft_skills":
            grouped["soft_skills_updates"].append(payload)
        elif any(term in section for term in ["hard", "competenze", "skills"]) or category == "skills":
            grouped["hard_skills_updates"].append(payload)
        elif any(term in section for term in ["esperienze", "esperienza", "tirocinio"]) or category in {"experience", "experiences"}:
            grouped["experience_updates"].append(payload)
        elif any(term in section for term in ["formazione", "istruzione"]) or category == "education":
            grouped["education_updates"].append(payload)
        elif any(term in section for term in ["progetti", "attivita"]) or category == "project":
            grouped["projects_updates"].append(payload)
        else:
            grouped["extra_sections"].append(payload)
    return grouped


def sort_rewrite_instructions(instructions: List[RewriteInstruction]) -> List[RewriteInstruction]:
    order = {
        "profile_updates": 0,
        "hard_skills_updates": 1,
        "soft_skills_updates": 2,
        "experience_updates": 3,
        "education_updates": 4,
        "projects_updates": 5,
        "extra_sections": 6,
    }
    grouped = group_rewrite_instructions(instructions)
    rank_by_id = {
        item["id"]: rank
        for group_name, rank in order.items()
        for item in grouped[group_name]
    }
    return sorted(instructions, key=lambda item: rank_by_id.get(item.source_id, 99))


def create_plain_optimized_pdf(optimized_text: str) -> tuple[bytes, str, str]:
    try:
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        margin_x = 48
        y = 50
        line_height = 13.2
        paragraph_gap = 4
        page_bottom = 800
        heading_color = (0.15, 0.22, 0.30)
        body_color = (0.18, 0.24, 0.31)
        accent_color = (0.11, 0.42, 0.67)

        lines = [line.rstrip() for line in (optimized_text or "").splitlines()]
        title_line = next((line.strip() for line in lines if line.strip() and len(line.strip()) <= 70 and not line.strip().startswith("-")), "")
        if title_line:
            page.insert_text((margin_x, y), title_line, fontsize=15.5, fontname="helv", color=accent_color)
            y += 22
            lines = lines[1:] if lines and lines[0].strip() == title_line else lines

        def is_heading(value: str) -> bool:
            stripped = value.strip()
            return bool(stripped) and len(stripped) <= 42 and stripped.upper() == stripped and any(char.isalpha() for char in stripped)

        for paragraph in lines:
            stripped = paragraph.strip()
            if not stripped:
                y += paragraph_gap
                continue
            heading = is_heading(stripped)
            font_size = 12 if heading else 10.2
            font_name = "helv" if not heading else "helv"
            wrap_width = 76 if heading else 92
            wrapped_lines = textwrap.wrap(stripped, width=wrap_width) or [stripped]
            if heading:
                y += 6
            for line in wrapped_lines:
                if y > page_bottom:
                    page = doc.new_page()
                    y = 50
                page.insert_text(
                    (margin_x, y),
                    line,
                    fontsize=font_size,
                    fontname=font_name,
                    color=heading_color if heading else body_color,
                )
                y += line_height if not heading else 15
            y += paragraph_gap + (2 if heading else 0)

        pdf_bytes = doc.write()
        doc.close()
        return pdf_bytes, "application/pdf", "pdf"
    except Exception as exc:
        print(f"Errore generazione PDF ottimizzato: {exc}")
        return optimized_text.encode("utf-8"), "text/plain; charset=utf-8", "txt"


def create_optimized_pdf_from_template(optimized_text: str, original_file_bytes: bytes) -> Optional[tuple[bytes, str, str]]:
    try:
        import fitz

        original = fitz.open(stream=original_file_bytes, filetype="pdf")
        if original.page_count <= 0:
            original.close()
            return None

        first_page = original[0]
        page_rect = first_page.rect
        page_width = page_rect.width
        page_height = page_rect.height
        text_blocks = []
        for page_index, page in enumerate(original):
            blocks = [
                block
                for block in page.get_text("blocks")
                if len(block) >= 5 and str(block[4]).strip()
            ]
            blocks.sort(key=lambda block: (round(block[1], 1), round(block[0], 1)))
            for block in blocks:
                rect = fitz.Rect(block[:4])
                if rect.width < 24 or rect.height < 8:
                    continue
                text_blocks.append((page_index, rect, str(block[4])))

        if not text_blocks:
            original.close()
            return None

        lines = [line.strip() for line in optimized_text.splitlines()]
        lines = [line for line in lines if line]
        if not lines:
            original.close()
            return None

        left_blocks = [
            rect
            for _, rect, _ in text_blocks
            if rect.x0 < page_width * 0.22 and rect.x1 < page_width * 0.52
        ]
        sidebar_width = 0
        if len(left_blocks) >= 3:
            sidebar_width = min(max(max(rect.x1 for rect in left_blocks) + 18, 145), page_width * 0.38)

        def is_heading(value: str) -> bool:
            stripped = value.strip()
            return (
                bool(stripped)
                and len(stripped) <= 52
                and stripped.upper() == stripped
                and any(char.isalpha() for char in stripped)
            )

        sections = []
        current = {"heading": "", "lines": []}
        for line in lines:
            if is_heading(line):
                if current["heading"] or current["lines"]:
                    sections.append(current)
                current = {"heading": line, "lines": []}
            else:
                current["lines"].append(line)
        if current["heading"] or current["lines"]:
            sections.append(current)

        sidebar_headings = {
            "CONTATTI", "LINGUE", "HARD SKILLS", "SOFT SKILLS", "COMPETENZE",
            "CERTIFICAZIONI", "SKILLS", "COMPETENZE TECNICHE",
        }

        doc = fitz.open()
        body_color = (0.14, 0.18, 0.22)
        heading_color = (0.10, 0.22, 0.34)
        sidebar_fill = (0.93, 0.95, 0.96)
        accent_fill = (0.16, 0.30, 0.42)
        margin = 38
        bottom_margin = 42
        gutter = 24

        def new_page():
            page = doc.new_page(width=page_width, height=page_height)
            if sidebar_width:
                page.draw_rect(fitz.Rect(0, 0, sidebar_width, page_height), color=None, fill=sidebar_fill)
                page.draw_rect(fitz.Rect(sidebar_width, 0, sidebar_width + 2, page_height), color=None, fill=accent_fill)
            else:
                page.draw_rect(fitz.Rect(0, 0, page_width, 16), color=None, fill=accent_fill)
            return page

        page = new_page()
        main_x = sidebar_width + gutter if sidebar_width else margin
        main_width = page_width - main_x - margin
        side_x = 22
        side_width = max(80, sidebar_width - 40) if sidebar_width else 0
        y_main = margin
        y_side = margin
        page_bottom = page_height - bottom_margin

        header_lines = []
        while sections and not sections[0]["heading"]:
            header_lines.extend(sections.pop(0)["lines"])
        if header_lines:
            name = header_lines[0]
            page.insert_text((main_x, y_main), name, fontsize=18, fontname="helv", color=heading_color)
            y_main += 24
            for line in header_lines[1:4]:
                wrapped = textwrap.wrap(line, width=max(42, int(main_width / 5.6))) or [line]
                for wrapped_line in wrapped:
                    page.insert_text((main_x, y_main), wrapped_line, fontsize=9.8, fontname="helv", color=body_color)
                    y_main += 12.5
            y_main += 10

        def draw_section(current_page, x, y, width, section, allow_page_break: bool = True):
            def ensure_space(target_page, current_y, target_x, target_width):
                if current_y <= page_bottom:
                    return target_page, current_y, target_x, target_width
                if not allow_page_break:
                    return target_page, current_y, target_x, target_width
                next_page = new_page()
                return next_page, margin, main_x, main_width

            if section["heading"]:
                y += 8
                current_page, y, x, width = ensure_space(current_page, y, x, width)
                current_page.insert_text((x, y), section["heading"], fontsize=11.5, fontname="helv", color=heading_color)
                y += 14
            for line in section["lines"]:
                normalized_line = line if line.startswith(("-", "•")) else line
                wrap_width = max(22, int(width / 5.2))
                for wrapped_line in textwrap.wrap(normalized_line, width=wrap_width) or [normalized_line]:
                    current_page, y, x, width = ensure_space(current_page, y, x, width)
                    current_page.insert_text((x, y), wrapped_line, fontsize=9.7, fontname="helv", color=body_color)
                    y += 12.5
                y += 2
            return current_page, y + 6

        for section in sections:
            use_sidebar = bool(sidebar_width and section["heading"].upper() in sidebar_headings)
            x = side_x if use_sidebar else main_x
            width = side_width if use_sidebar else main_width
            section_line_count = len(section["lines"]) + (1 if section["heading"] else 0)
            needed_height = 24 + section_line_count * 16
            if use_sidebar and y_side + needed_height <= page_bottom:
                page, y_side = draw_section(page, x, y_side, width, section, allow_page_break=False)
                continue

            if y_main + needed_height > page_bottom:
                page = new_page()
                y_main = margin
                y_side = margin
            page, y_main = draw_section(page, main_x, y_main, main_width, section)

        pdf_bytes = doc.write()
        doc.close()
        original.close()
        return pdf_bytes, "application/pdf", "pdf"
    except Exception as exc:
        print(f"Errore riuso template PDF originale: {exc}")
        return None


def create_optimized_cv_file(
    optimized_text: str,
    original_file_bytes: Optional[bytes] = None,
    original_filename: Optional[str] = None,
) -> tuple[bytes, str, str]:
    if original_file_bytes and (original_filename or "").lower().endswith(".pdf"):
        templated_pdf = create_optimized_pdf_from_template(optimized_text, original_file_bytes)
        if templated_pdf:
            return templated_pdf

    return create_plain_optimized_pdf(optimized_text)


def create_optimized_docx_file(optimized_text: str, original_file_bytes: Optional[bytes] = None) -> tuple[bytes, str, str]:
    try:
        from docx import Document

        document = Document(io.BytesIO(original_file_bytes)) if original_file_bytes else Document()
        optimized_lines = [line.strip() for line in optimized_text.splitlines() if line.strip()]

        def replace_paragraph_text_preserving_runs(paragraph, value: str) -> None:
            # Gestione hyperlink: se il nuovo testo contiene il display text
            # gia' presente dentro un <w:hyperlink> (es. "linkedin.com/in/..."),
            # lo lasciamo dentro l'hyperlink per non perdere la cliccabilita',
            # e lo rimuoviamo dal testo dei run normali (evitando il duplicato).
            ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            remaining = value
            try:
                for hyperlink in paragraph._p.iter(f"{ns_w}hyperlink"):
                    t_elements = list(hyperlink.iter(f"{ns_w}t"))
                    link_text = "".join((t.text or "") for t in t_elements)
                    stripped = link_text.strip()
                    if stripped and stripped in remaining:
                        remaining = remaining.replace(stripped, "", 1)
                    else:
                        for t in t_elements:
                            t.text = ""
                import re as _re
                remaining = _re.sub(r"[ \t]{2,}", " ", remaining).strip()
            except Exception:
                remaining = value
            runs = list(paragraph.runs)
            if not runs:
                paragraph.add_run(remaining)
                return
            runs[0].text = remaining
            for run in runs[1:]:
                run.text = ""

        def clone_paragraph_format(source, target) -> None:
            try:
                if source.style is not None:
                    target.style = source.style
                target.alignment = source.alignment
                target.paragraph_format.left_indent = source.paragraph_format.left_indent
                target.paragraph_format.right_indent = source.paragraph_format.right_indent
                target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
                target.paragraph_format.space_before = source.paragraph_format.space_before
                target.paragraph_format.space_after = source.paragraph_format.space_after
                target.paragraph_format.line_spacing = source.paragraph_format.line_spacing
                target.paragraph_format.keep_together = source.paragraph_format.keep_together
                target.paragraph_format.keep_with_next = source.paragraph_format.keep_with_next
                target.paragraph_format.page_break_before = source.paragraph_format.page_break_before
                target.paragraph_format.widow_control = source.paragraph_format.widow_control
                for src_run, dst_run in zip(source.runs, target.runs):
                    if src_run.bold is not None:
                        dst_run.bold = src_run.bold
                    if src_run.italic is not None:
                        dst_run.italic = src_run.italic
                    if src_run.underline is not None:
                        dst_run.underline = src_run.underline
                    if src_run.font.name:
                        dst_run.font.name = src_run.font.name
                    if src_run.font.size:
                        dst_run.font.size = src_run.font.size
            except Exception:
                pass

        paragraphs = list(document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)

        if original_file_bytes and paragraphs:
            last_paragraph = None
            for index, paragraph in enumerate(paragraphs):
                if index < len(optimized_lines):
                    replace_paragraph_text_preserving_runs(paragraph, optimized_lines[index])
                    last_paragraph = paragraph
                else:
                    replace_paragraph_text_preserving_runs(paragraph, "")
            for line in optimized_lines[len(paragraphs):]:
                new_paragraph = document.add_paragraph(line)
                if last_paragraph:
                    clone_paragraph_format(last_paragraph, new_paragraph)
        else:
            for line in optimized_lines:
                paragraph = document.add_paragraph(line)
                if len(line) <= 42 and line.upper() == line and any(char.isalpha() for char in line):
                    paragraph.style = document.styles["Heading 2"]

        output = io.BytesIO()
        document.save(output)
        return output.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
    except Exception as exc:
        print(f"Errore generazione DOCX ottimizzato: {exc}")
        return b"", "", ""


def sanitize_cv_additional_data(data: Optional[Dict[str, Any]]) -> tuple[Dict[str, Any], List[str]]:
    sanitized: Dict[str, Any] = {}
    rejected_candidates = []

    def clean_answer(value: Any, question: Any) -> str:
        answer = str(value or "").strip()
        prompt = str(question or "").strip()
        if prompt:
            answer = re.sub(
                rf"^\s*(?:\*\*)?\s*{re.escape(prompt)}\s*(?:\*\*)?\s*[:-]?\s*",
                "",
                answer,
                flags=re.IGNORECASE,
            )
        return answer.strip().strip("*").strip()

    for key, value in (data or {}).items():
        if key in {"adaptation_answers", "confirmed_skills"}:
            continue
        if not isinstance(value, str) or not value.strip():
            continue
        cleaned = value.strip()
        is_language_level = (
            key in {"certifications", "languages"}
            and bool(re.fullmatch(r"(?i)(?:inglese\s+)?[abc][12]", cleaned))
        )
        if (
            not is_language_level
            and not is_meaningful_cv_detail(cleaned)
            and is_low_quality_text(cleaned, min_chars=8, min_words=2)
        ):
            rejected_candidates.append(key.replace("_", " "))
            continue
        sanitized[key] = cleaned

    answers = []
    for index, item in enumerate((data or {}).get("adaptation_answers", [])):
        if not isinstance(item, dict):
            continue
        question = str(item.get("question", "")).strip()
        answer = clean_answer(item.get("answer", ""), question)
        if not answer:
            continue
        if not is_meaningful_cv_detail(answer) and is_low_quality_text(answer, min_chars=8, min_words=2):
            rejected_candidates.append(f"risposta domanda {index + 1}")
            continue
        answers.append({
            "question": question,
            "reason": str(item.get("reason", "")).strip(),
            "category": str(item.get("category", "")).strip(),
            "answer": answer,
        })

    if answers:
        sanitized["adaptation_answers"] = answers

    confirmed_skills = []
    seen_confirmed_skills = set()
    for item in (data or {}).get("confirmed_skills", []):
        if isinstance(item, str):
            continue
        elif isinstance(item, dict):
            name = str(item.get("name") or item.get("skill") or "").strip()
            detail = str(item.get("user_example") or item.get("detail") or item.get("example") or "").strip()
            category = str(item.get("category") or "hard_skill").strip()
            item_type = str(item.get("type") or "skillConfirmation").strip()
        else:
            continue
        if not name:
            continue
        if category == "keyword":
            category = "hard_skill"
        if item_type != "skillConfirmation" or category not in {
            "hard_skill", "soft_skill", "tool", "language",
        }:
            continue
        target_section = "SOFT SKILLS" if category == "soft_skill" else "HARD SKILLS"
        if detail and not is_meaningful_cv_detail(detail) and is_low_quality_text(detail, min_chars=6, min_words=2):
            rejected_candidates.append(f"conferma skill {name}")
            continue
        item_id = str(item.get("id") or name).strip() if isinstance(item, dict) else name
        dedupe_key = (
            normalize_plain_text(target_section),
            canonical_skill_identity(name),
        )
        if item_id in seen_confirmed_skills or dedupe_key in seen_confirmed_skills:
            continue
        seen_confirmed_skills.add(item_id)
        seen_confirmed_skills.add(dedupe_key)
        confirmed_skills.append({
            "id": item_id,
            "type": item_type,
            "name": name,
            "category": category,
            "detail": detail,
            "user_example": detail,
            "target_section": target_section,
            "status": "confirmed",
        })
    if confirmed_skills:
        sanitized["confirmed_skills"] = confirmed_skills

    rejected = rejected_candidates
    return sanitized, rejected


def flatten_cv_support_data(data: Optional[Dict[str, Any]]) -> str:
    parts = []
    for key, value in (data or {}).items():
        if isinstance(value, str) and value.strip():
            parts.append(value.strip())
        elif key == "adaptation_answers" and isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    parts.append(str(item.get("answer", "")).strip())
        elif key == "confirmed_skills" and isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    parts.append(" ".join([str(item.get("name", "")).strip(), str(item.get("detail", "")).strip()]).strip())
                elif isinstance(item, str):
                    parts.append(item.strip())
    return "\n".join(part for part in parts if part)


def detect_unsupported_optimized_claims(
    optimized_text: str,
    cv_text: str,
    user_additional_data: Optional[Dict[str, Any]],
    company: str,
    role: str,
    goal: str,
) -> List[Dict[str, str]]:
    support_text = normalize_plain_text(
        "\n".join([
            cv_text or "",
            flatten_cv_support_data(user_additional_data),
            company or "",
            role or "",
            goal or "",
        ])
    )
    if not support_text:
        return []

    risky_patterns = [
        (r"\b(certificazione|certificazioni|certificato|corso|master|laurea|dottorato)\b", "titolo, corso o certificazione"),
        (r"\b\d{1,3}%\b|\b\d+\s*(anni|mesi|utenti|clienti|progetti|record|righe)\b", "risultato numerico"),
        (r"\b(manager|lead|responsabile|coordinatore|senior|director)\b", "ruolo o seniority"),
        (r"\b(power bi|tableau|aws|azure|google cloud|docker|kubernetes|sap)\b", "strumento o tecnologia"),
    ]
    warnings = []
    for raw_line in optimized_text.splitlines():
        line = raw_line.strip()
        if len(line) < 35:
            continue
        line_plain = normalize_plain_text(line)
        matched_reason = next((reason for pattern, reason in risky_patterns if re.search(pattern, line_plain)), "")
        if not matched_reason:
            continue
        meaningful_tokens = [
            token for token in tokenize_meaningful(line_plain)
            if len(token) >= 4 and token not in {"curriculum", "profilo", "esperienza", "competenze"}
        ]
        if not meaningful_tokens:
            continue
        supported_tokens = sum(1 for token in meaningful_tokens if token in support_text)
        if supported_tokens / max(len(meaningful_tokens), 1) < 0.35:
            warnings.append({
                "claim": line[:220],
                "reason": f"Possibile {matched_reason} non supportato dal CV originale o dai dati aggiunti.",
            })
        if len(warnings) >= 8:
            break
    return warnings


def get_question_type_instructions(interview_type: str, role: str, company: str) -> str:
    """
    Restituisce istruzioni specifiche per generare domande diverse
    in base alla tipologia scelta dall'utente.
    """
    if interview_type == "conoscitive_motivazionali":
        return f"""
TIPOLOGIA: DOMANDE CONOSCITIVE E MOTIVAZIONALI

Obiettivo:
Capire chi è il candidato, cosa cerca, come ragiona, cosa si aspetta dall'azienda,
come si presenta e quanto è motivato per il ruolo e per {company}.

Le domande devono riguardare:
- presentazione personale;
- percorso di studi o esperienze;
- obiettivi professionali;
- aspettative rispetto all'azienda;
- motivazione per il ruolo;
- conoscenza dell'azienda;
- punti di forza e aree di miglioramento;
- hobby o interessi, se utili a capire la persona;
- lavoro di gruppo;
- gestione di difficoltà, conflitti o responsabilità;
- come il candidato si vede tra alcuni anni.

Le domande NON devono essere tecniche.
Le domande devono sembrare domande da recruiter HR.
"""

    if interview_type == "tecniche":
        return f"""
TIPOLOGIA: DOMANDE TECNICHE

Obiettivo:
Valutare le competenze tecniche del candidato rispetto al ruolo target: {role}.

Le domande devono essere specifiche per:
- ruolo selezionato;
- settore indicato dal candidato;
- strumenti, metodi, tecnologie o competenze richieste;
- progetti svolti;
- capacità di spiegare concetti tecnici;
- capacità di risolvere problemi operativi;
- capacità di applicare conoscenze teoriche a casi pratici.

Le domande NON devono essere generiche o motivazionali.
Devono servire a capire se il candidato sa ragionare tecnicamente.
"""

    if interview_type == "logica":
        return f"""
TIPOLOGIA: DOMANDE DI LOGICA, RAGIONAMENTO E PROBLEM SOLVING

Obiettivo:
Valutare il modo in cui il candidato ragiona, formula ipotesi, affronta problemi ambigui,
fa stime, riconosce trabocchetti e spiega il proprio processo mentale.

Le 10 domande devono alternare queste sottocategorie:

1. Domande di logica legate al ruolo o alle abilità della persona:
   - problemi di ragionamento applicati al ruolo {role};
   - casi in cui bisogna ordinare informazioni, prendere decisioni o analizzare vincoli;
   - situazioni realistiche ma mentalmente sfidanti.

2. Domande di stima/Fermi questions:
   - "Quanti Big Mac vengono venduti da McDonald’s ogni anno negli Stati Uniti?"
   - "Quante palline da ping pong servirebbero per riempire questa stanza?"
   - "Quante birre vengono consumate in media in un anno in Italia?"
   - domande simili, ma possibilmente adattate al settore o all'azienda {company}.

3. Domande a trabocchetto:
   - domande dove bisogna stare attenti alle assunzioni;
   - domande apparentemente semplici ma con un inganno logico;
   - domande che valutano attenzione e lucidità.

4. Serie numeriche o alfabetiche:
   - completare una sequenza di numeri;
   - completare una sequenza di lettere;
   - individuare la regola di una sequenza.

5. Mini business case/logica aziendale:
   - problemi di priorità, trade-off, risorse limitate;
   - decisioni con dati incompleti;
   - ragionamenti collegati almeno in parte all'azienda {company} o al ruolo {role}.

Regole specifiche:
- Le domande devono richiedere uno sforzo mentale, non una risposta motivazionale.
- Almeno 2 domande devono essere di stima.
- Almeno 2 domande devono essere a trabocchetto.
- Almeno 2 domande devono contenere una serie numerica o alfabetica.
- Almeno 2 domande devono essere collegate al ruolo {role} o all'azienda {company}.
- Ogni domanda deve chiedere al candidato di spiegare il ragionamento, non solo dare una risposta secca.
- Le domande devono essere realistiche per un colloquio.
"""

    return """
Genera domande di colloquio realistiche, coerenti con il profilo del candidato.
"""


def get_difficulty_instructions(difficulty: str) -> str:
    """
    Restituisce istruzioni specifiche in base al livello scelto.
    """

    if difficulty == "base":
        return """
LIVELLO: BASE

Le domande devono essere realistiche ma semplici.
Devono essere adatte a un candidato junior o a una persona alle prime esperienze.
Evita domande troppo tecniche, troppo lunghe o troppo astratte.
Le domande devono permettere al candidato di rispondere anche con esperienze universitarie, personali o di progetto.
Per la logica, usa problemi semplici, guidati e spiegabili.
"""

    if difficulty == "intermedio":
        return """
LIVELLO: INTERMEDIO

Le domande devono essere realistiche e leggermente sfidanti.
Devono richiedere esempi concreti, ragionamento e collegamento con il ruolo.
Le domande possono contenere casi pratici, situazioni lavorative o problemi da analizzare.
Per la logica, usa problemi di stima, ragionamento e piccoli trabocchetti di difficoltà media.
"""

    if difficulty == "avanzato":
        return """
LIVELLO: AVANZATO

Le domande devono essere più specifiche, complesse e selettive.
Devono simulare un colloquio più esigente.
Devono richiedere ragionamento strutturato, esempi solidi, capacità critica e conoscenza del ruolo.
Per le domande tecniche, usa casi più dettagliati e realistici.
Per le domande di logica, usa stime complesse, vincoli multipli, serie meno immediate, business case e problemi a trabocchetto.
"""

    return """
LIVELLO: INTERMEDIO

Le domande devono essere realistiche, coerenti con il ruolo e abbastanza sfidanti.
"""
def save_sources_for_question(cursor, question_id: int, sources: List[Dict[str, str]]):
    for source in sources:
        title = source.get("title", "").strip()
        url = source.get("url", "").strip()
        content = source.get("content", "").strip()

        if not url:
            continue

        cursor.execute("""
        INSERT OR IGNORE INTO web_sources (
            title,
            url,
            content
        )
        VALUES (?, ?, ?)
        """, (
            title,
            url,
            content
        ))

        cursor.execute("""
        SELECT id
        FROM web_sources
        WHERE url = ?
        """, (url,))

        row = cursor.fetchone()

        if not row:
            continue

        source_id = row[0]

        cursor.execute("""
        SELECT id
        FROM question_web_sources
        WHERE question_id = ? AND source_id = ?
        """, (
            question_id,
            source_id
        ))

        already_linked = cursor.fetchone()

        if already_linked:
            continue

        cursor.execute("""
        INSERT INTO question_web_sources (
            question_id,
            source_id
        )
        VALUES (?, ?)
        """, (
            question_id,
            source_id
        ))


# =========================
# ENDPOINT BASE / DEBUG
# =========================

@app.get("/")
def home():
    return {
        "message": "CareerCoach backend attivo",
        "rewrite_model": GEMINI_MODEL if GEMINI_API_KEY else OLLAMA_TEXT_MODEL,
    }


@app.get("/debug")
def debug():
    return {
        "status": "ok",
        "message": "Backend FastAPI attivo",
        "groq_model": GROQ_MODEL,
        "has_groq_key": bool(GROQ_API_KEY),
        "gemini_model": GEMINI_MODEL,
        "has_gemini_key": bool(GEMINI_API_KEY),
        "rewrite_fallback_model": OLLAMA_TEXT_MODEL,
        "has_tavily_key": bool(TAVILY_API_KEY)
    }


# =========================
# ENDPOINT UTENTI
# =========================

@app.post("/auth/register")
def register(data: RegisterRequest):
    email = validate_email_address(data.email)
    phone = validate_phone(data.phone)
    validate_password(data.password)

    name = data.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="Inserisci il nome.")

    verification_token = make_token()
    expires_at = utc_now() + timedelta(hours=24)

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT id FROM users WHERE lower(email) = lower(?)", (email,))
    existing_email = cursor.fetchone()
    if existing_email:
        conn.close()
        raise HTTPException(status_code=409, detail="Esiste già un account con questa email.")

    if phone:
        cursor.execute("SELECT id FROM users WHERE phone = ?", (phone,))
        existing_phone = cursor.fetchone()
        if existing_phone:
            conn.close()
            raise HTTPException(status_code=409, detail="Questo numero è già associato a un account.")

    cursor.execute("""
    INSERT INTO users (
        name,
        email,
        phone,
        password_hash,
        email_verified,
        email_verification_token,
        email_verification_expires_at,
        education,
        target_role,
        sector,
        experience_level,
        interview_language
    )
    VALUES (?, ?, ?, ?, 0, ?, ?, '', '', '', 'Junior', 'Italiano')
    """, (
        name,
        email,
        phone,
        hash_password(data.password),
        verification_token,
        expires_at.isoformat(),
    ))

    conn.commit()
    conn.close()

    verification_link = make_frontend_link("verify", verification_token)
    email_sent = send_email(
        email,
        "Verifica il tuo account CareerCoach",
        (
            "Ciao,\n\n"
            "per attivare il tuo account CareerCoach apri questo link:\n"
            f"{verification_link}\n\n"
            "Il link scade tra 24 ore."
        ),
    )

    return {
        "message": "Account creato. Controlla la tua email per verificare l'accesso.",
        "email_sent": email_sent,
        "preview_link": None if email_sent else verification_link,
    }


@app.post("/auth/verify-email")
def verify_email(data: TokenRequest):
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT id, email_verification_expires_at
    FROM users
    WHERE email_verification_token = ?
    """, (data.token,))
    row = cursor.fetchone()

    if not row:
        conn.close()
        raise HTTPException(status_code=400, detail="Link di verifica non valido.")

    user_id, expires_at = row
    if expires_at and datetime.fromisoformat(expires_at) < utc_now():
        conn.close()
        raise HTTPException(status_code=400, detail="Link di verifica scaduto. Registrati di nuovo o richiedi un nuovo link.")

    cursor.execute("""
    UPDATE users
    SET email_verified = 1,
        email_verification_token = NULL,
        email_verification_expires_at = NULL
    WHERE id = ?
    """, (user_id,))

    session_token = create_session(cursor, user_id)
    user = fetch_user_by_id(cursor, user_id)
    conn.commit()
    conn.close()

    return {
        "token": session_token,
        "user": user_to_response(user),
        "message": "Email verificata correttamente.",
    }


@app.post("/auth/login")
def login(data: LoginRequest):
    identifier = data.identifier.strip()
    password = data.password

    conn = get_connection()
    cursor = conn.cursor()

    if "@" in identifier:
        identifier = validate_email_address(identifier)
        cursor.execute("""
        SELECT id, password_hash, email_verified
        FROM users
        WHERE lower(email) = lower(?)
        """, (identifier,))
    else:
        phone = validate_phone(identifier)
        cursor.execute("""
        SELECT id, password_hash, email_verified
        FROM users
        WHERE phone = ?
        """, (phone,))

    row = cursor.fetchone()
    if not row or not verify_password(password, row[1]):
        conn.close()
        raise HTTPException(status_code=401, detail="Email/telefono o password non corretti.")

    if not row[2]:
        conn.close()
        raise HTTPException(status_code=403, detail="Verifica prima la tua email tramite il link ricevuto.")

    session_token = create_session(cursor, row[0])
    user = fetch_user_by_id(cursor, row[0])
    conn.commit()
    conn.close()

    return {
        "token": session_token,
        "user": user_to_response(user),
    }


@app.get("/auth/me")
def get_current_user(authorization: Optional[str] = Header(default=None)):
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Sessione mancante.")

    token = authorization.replace("Bearer ", "", 1).strip()
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT user_id, expires_at
    FROM user_sessions
    WHERE token = ?
    """, (token,))
    session = cursor.fetchone()

    if not session:
        conn.close()
        raise HTTPException(status_code=401, detail="Sessione non valida.")

    if datetime.fromisoformat(session[1]) < utc_now():
        cursor.execute("DELETE FROM user_sessions WHERE token = ?", (token,))
        conn.commit()
        conn.close()
        raise HTTPException(status_code=401, detail="Sessione scaduta.")

    user = fetch_user_by_id(cursor, session[0])
    conn.close()

    if not user:
        raise HTTPException(status_code=404, detail="Utente non trovato.")

    return {"user": user_to_response(user)}


@app.post("/auth/logout")
def logout(data: TokenRequest):
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM user_sessions WHERE token = ?", (data.token,))
    conn.commit()
    conn.close()
    return {"message": "Logout effettuato."}


@app.get("/auth/oauth/redirect-uris")
def oauth_redirect_uris():
    return {
        "google": make_oauth_callback_url("google"),
        "linkedin": make_oauth_callback_url("linkedin"),
    }


@app.get("/auth/oauth/{provider}/start")
def start_oauth(provider: str):
    provider = normalize_oauth_provider(provider)
    conn = get_connection()
    cursor = conn.cursor()
    state = create_oauth_state(cursor, provider)
    conn.commit()
    conn.close()

    return RedirectResponse(build_oauth_authorization_url(provider, state))


@app.get("/auth/oauth/{provider}/url")
def get_oauth_url(provider: str):
    provider = normalize_oauth_provider(provider)
    conn = get_connection()
    cursor = conn.cursor()
    state = create_oauth_state(cursor, provider)
    auth_url = build_oauth_authorization_url(provider, state)
    conn.commit()
    conn.close()

    return {"auth_url": auth_url}


@app.get("/auth/oauth/{provider}/callback")
def oauth_callback(provider: str, code: Optional[str] = None, state: Optional[str] = None, error: Optional[str] = None):
    provider = normalize_oauth_provider(provider)

    if error:
        raise HTTPException(status_code=400, detail=f"Accesso social annullato: {error}")

    if not code or not state:
        raise HTTPException(status_code=400, detail="Risposta OAuth incompleta.")

    conn = get_connection()
    cursor = conn.cursor()

    try:
        consume_oauth_state(cursor, provider, state)
        profile = fetch_oauth_profile(provider, code)
        user_id = find_or_create_oauth_user(cursor, provider, profile)
        session_token = create_session(cursor, user_id)
        conn.commit()
    finally:
        conn.close()

    return RedirectResponse(make_frontend_oauth_link(session_token))


@app.get("/auth/linkedin/callback")
def linkedin_callback(code: Optional[str] = None, state: Optional[str] = None, error: Optional[str] = None):
    return oauth_callback("linkedin", code=code, state=state, error=error)


@app.post("/auth/forgot-password")
def forgot_password(data: ForgotPasswordRequest):
    identifier = data.identifier.strip()
    conn = get_connection()
    cursor = conn.cursor()

    if "@" in identifier:
        identifier = validate_email_address(identifier)
        cursor.execute("SELECT id, email FROM users WHERE lower(email) = lower(?)", (identifier,))
    else:
        phone = validate_phone(identifier)
        cursor.execute("SELECT id, email FROM users WHERE phone = ?", (phone,))

    row = cursor.fetchone()
    if not row:
        conn.close()
        return {"message": "Se l'account esiste, riceverai un link per reimpostare la password."}

    token = make_token()
    expires_at = utc_now() + timedelta(hours=1)
    cursor.execute("""
    UPDATE users
    SET password_reset_token = ?,
        password_reset_expires_at = ?
    WHERE id = ?
    """, (token, expires_at.isoformat(), row[0]))
    conn.commit()
    conn.close()

    reset_link = make_frontend_link("reset", token)
    email_sent = send_email(
        row[1],
        "Recupero password CareerCoach",
        (
            "Ciao,\n\n"
            "per scegliere una nuova password apri questo link:\n"
            f"{reset_link}\n\n"
            "Il link scade tra 1 ora."
        ),
    )

    return {
        "message": "Se l'account esiste, riceverai un link per reimpostare la password.",
        "email_sent": email_sent,
        "preview_link": None if email_sent else reset_link,
    }


@app.post("/auth/reset-password")
def reset_password(data: ResetPasswordRequest):
    validate_password(data.password)

    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("""
    SELECT id, password_reset_expires_at
    FROM users
    WHERE password_reset_token = ?
    """, (data.token,))
    row = cursor.fetchone()

    if not row:
        conn.close()
        raise HTTPException(status_code=400, detail="Link di recupero non valido.")

    if row[1] and datetime.fromisoformat(row[1]) < utc_now():
        conn.close()
        raise HTTPException(status_code=400, detail="Link di recupero scaduto.")

    cursor.execute("""
    UPDATE users
    SET password_hash = ?,
        password_reset_token = NULL,
        password_reset_expires_at = NULL
    WHERE id = ?
    """, (hash_password(data.password), row[0]))
    cursor.execute("DELETE FROM user_sessions WHERE user_id = ?", (row[0],))
    conn.commit()
    conn.close()

    return {"message": "Password aggiornata. Ora puoi accedere."}


@app.post("/users")
def create_user(data: UserCreate):
    email = validate_email_address(data.email) if data.email else None
    phone = validate_phone(data.phone)

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    INSERT INTO users (
        name,
        email,
        phone,
        education,
        target_role,
        sector,
        experience_level,
        interview_language
    )
    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        data.name,
        email,
        phone,
        data.education,
        data.target_role,
        data.sector,
        data.experience_level,
        data.interview_language
    ))

    conn.commit()
    user_id = cursor.lastrowid
    conn.close()

    return {
        "user_id": user_id,
        "message": "Profilo creato correttamente"
    }


@app.get("/users/{user_id}")
def get_user(user_id: int, authorization: Optional[str] = Header(default=None)):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT 
        id,
        name,
        email,
        education,
        target_role,
        sector,
        experience_level,
        interview_language,
        phone,
        email_verified,
        cv_filename,
        cv_content_type,
        cv_size,
        cv_text,
        cv_uploaded_at,
        linkedin_url,
        portfolio_url,
        instagram_handle,
        auth_provider,
        profile_image_data_url
    FROM users
    WHERE id = ?
    """, (user_id,))

    row = cursor.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="Utente non trovato")

    return {
        "id": row[0],
        "name": row[1],
        "email": row[2],
        "education": row[3],
        "target_role": row[4],
        "sector": row[5],
        "experience_level": row[6],
        "interview_language": row[7],
        "phone": row[8],
        "email_verified": bool(row[9]),
        "cv_filename": row[10],
        "cv_uploaded": bool(row[10]),
        "cv_uploaded_at": row[14],
        "linkedin_url": row[15],
        "portfolio_url": row[16],
        "instagram_handle": row[17],
        "auth_provider": row[18],
        "profile_image_data_url": row[19],
    }


def require_user_session(user_id: int, authorization: Optional[str]) -> None:
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Sessione mancante.")

    token = authorization.replace("Bearer ", "", 1).strip()
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("""
    SELECT user_id, expires_at
    FROM user_sessions
    WHERE token = ?
    """, (token,))
    session = cursor.fetchone()
    conn.close()

    if not session or session[0] != user_id:
        raise HTTPException(status_code=403, detail="Sessione non autorizzata.")
    if datetime.fromisoformat(session[1]) < utc_now():
        raise HTTPException(status_code=401, detail="Sessione scaduta.")


@app.put("/users/{user_id}/profile-image")
def update_profile_image(
    user_id: int,
    data: ProfileImageUpdate,
    authorization: Optional[str] = Header(default=None),
):
    require_user_session(user_id, authorization)
    match = re.fullmatch(
        r"data:(image/(?:jpeg|png|webp));base64,([A-Za-z0-9+/=\s]+)",
        data.image_data_url.strip(),
    )
    if not match:
        raise HTTPException(status_code=400, detail="Formato immagine non valido.")

    try:
        image_bytes = base64.b64decode(match.group(2), validate=True)
    except (ValueError, binascii.Error):
        raise HTTPException(status_code=400, detail="Immagine non valida.")

    if not image_bytes or len(image_bytes) > 2 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="L'immagine deve essere inferiore a 2 MB.")

    content_type = match.group(1)
    valid_signature = (
        (content_type == "image/jpeg" and image_bytes.startswith(b"\xff\xd8\xff"))
        or (content_type == "image/png" and image_bytes.startswith(b"\x89PNG\r\n\x1a\n"))
        or (content_type == "image/webp" and len(image_bytes) >= 12 and image_bytes[:4] == b"RIFF" and image_bytes[8:12] == b"WEBP")
    )
    if not valid_signature:
        raise HTTPException(status_code=400, detail="Immagine non supportata.")

    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE users SET profile_image_data_url = ? WHERE id = ?",
        (data.image_data_url.strip(), user_id),
    )
    conn.commit()
    row = fetch_user_by_id(cursor, user_id)
    conn.close()
    return {"message": "Foto profilo aggiornata.", "user": user_to_response(row), "profile_image_data_url": data.image_data_url.strip()}


@app.delete("/users/{user_id}/profile-image")
def delete_profile_image(user_id: int, authorization: Optional[str] = Header(default=None)):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("UPDATE users SET profile_image_data_url = NULL WHERE id = ?", (user_id,))
    conn.commit()
    row = fetch_user_by_id(cursor, user_id)
    conn.close()
    return {"message": "Foto profilo rimossa.", "user": user_to_response(row)}


@app.delete("/users/{user_id}")
def delete_user(user_id: int, authorization: Optional[str] = Header(default=None)):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT id FROM users WHERE id = ?", (user_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="Utente non trovato.")

    cursor.execute("""
    DELETE FROM question_web_sources
    WHERE question_id IN (
        SELECT q.id
        FROM questions q
        JOIN interview_sessions s ON q.session_id = s.id
        WHERE s.user_id = ?
    )
    """, (user_id,))

    cursor.execute("""
    DELETE FROM answers
    WHERE question_id IN (
        SELECT q.id
        FROM questions q
        JOIN interview_sessions s ON q.session_id = s.id
        WHERE s.user_id = ?
    )
    """, (user_id,))

    cursor.execute("""
    DELETE FROM questions
    WHERE session_id IN (
        SELECT id
        FROM interview_sessions
        WHERE user_id = ?
    )
    """, (user_id,))

    cursor.execute("DELETE FROM interview_sessions WHERE user_id = ?", (user_id,))
    cursor.execute("DELETE FROM user_sessions WHERE user_id = ?", (user_id,))
    cursor.execute("DELETE FROM users WHERE id = ?", (user_id,))

    conn.commit()
    conn.close()

    return {"message": "Profilo eliminato correttamente."}


@app.put("/users/{user_id}")
def update_user(user_id: int, data: UserUpdate, authorization: Optional[str] = Header(default=None)):
    require_user_session(user_id, authorization)
    email = validate_email_address(data.email) if data.email else None
    phone = validate_phone(data.phone)

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT id, email FROM users WHERE id = ?", (user_id,))
    existing_user = cursor.fetchone()
    if not existing_user:
        conn.close()
        raise HTTPException(status_code=404, detail="Utente non trovato.")

    if email:
        cursor.execute("SELECT id FROM users WHERE lower(email) = lower(?) AND id != ?", (email, user_id))
        if cursor.fetchone():
            conn.close()
            raise HTTPException(status_code=409, detail="Email già associata a un altro account.")

    if phone:
        cursor.execute("SELECT id FROM users WHERE phone = ? AND id != ?", (phone, user_id))
        if cursor.fetchone():
            conn.close()
            raise HTTPException(status_code=409, detail="Numero già associato a un altro account.")

    email_changed = email and existing_user[1] and email.lower() != existing_user[1].lower()
    verification_token = make_token() if email_changed else None
    verification_expires = (utc_now() + timedelta(hours=24)).isoformat() if email_changed else None

    cursor.execute("""
    UPDATE users
    SET name = ?,
        email = ?,
        phone = ?,
        education = ?,
        target_role = ?,
        sector = ?,
        experience_level = ?,
        interview_language = ?,
        email_verified = CASE WHEN ? THEN 0 ELSE email_verified END,
        email_verification_token = CASE WHEN ? THEN ? ELSE email_verification_token END,
        email_verification_expires_at = CASE WHEN ? THEN ? ELSE email_verification_expires_at END
    WHERE id = ?
    """, (
        data.name,
        email,
        phone,
        data.education,
        data.target_role,
        data.sector,
        data.experience_level,
        data.interview_language,
        1 if email_changed else 0,
        1 if email_changed else 0,
        verification_token,
        1 if email_changed else 0,
        verification_expires,
        user_id,
    ))

    conn.commit()
    user = fetch_user_by_id(cursor, user_id)
    conn.close()

    preview_link = None
    email_sent = False
    if email_changed:
        preview_link = make_frontend_link("verify", verification_token)
        email_sent = send_email(
            email,
            "Verifica il nuovo indirizzo email CareerCoach",
            (
                "Ciao,\n\n"
                "per confermare il nuovo indirizzo email apri questo link:\n"
                f"{preview_link}\n\n"
                "Il link scade tra 24 ore."
            ),
        )

    return {
        "user": user_to_response(user),
        "message": "Profilo aggiornato correttamente.",
        "email_sent": email_sent,
        "preview_link": None if email_sent else preview_link,
    }


@app.delete("/users/{user_id}/cv")
def delete_user_cv(user_id: int, authorization: Optional[str] = Header(default=None)):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT id FROM users WHERE id = ?", (user_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="Utente non trovato.")

    cursor.execute("""
    UPDATE users
    SET cv_filename = NULL,
        cv_content_type = NULL,
        cv_size = NULL,
        cv_text = NULL,
        cv_file_base64 = NULL,
        cv_uploaded_at = NULL,
        digital_analysis_json = NULL
    WHERE id = ?
    """, (user_id,))

    conn.commit()
    user = fetch_user_by_id(cursor, user_id)
    conn.close()

    return {
        "user": user_to_response(user),
        "message": "CV eliminato correttamente.",
    }


GENERIC_BAD_INPUTS = {
    "aaaaaa", "aaa", "asdf", "asdfgh", "asdfghjkl", "qwerty", "boh",
    "non lo so", "non saprei", "chi lo sa", "azienda a caso",
    "lavoro qualunque", "qualunque", "a caso", "test", "prova",
    "nessuna idea", "non ho idea", "n/a", "na"
}

ROLE_KEYWORDS = {
    "analyst", "analista", "engineer", "developer", "sviluppatore",
    "designer", "manager", "specialist", "consultant", "consulente",
    "marketing", "sales", "account", "hr", "human resources",
    "recruiter", "project", "product", "data", "software", "frontend",
    "backend", "full stack", "cybersecurity", "security", "ux", "ui",
    "finance", "controller", "assistant", "coordinator", "operations",
    "researcher", "scientist", "copywriter", "content", "seo", "devops",
    "cloud", "qa", "tester", "business", "amministrativo", "contabile"
}

KNOWN_COMPANIES = {
    "google", "amazon", "microsoft", "apple", "meta", "facebook",
    "netflix", "tesla", "ibm", "oracle", "accenture", "deloitte",
    "pwc", "kpmg", "ey", "linkedin", "spotify", "salesforce",
    "adobe", "sap", "siemens", "enel", "eni", "intesa sanpaolo",
    "unicredit", "poste italiane", "telecom", "tim", "ferrari",
    "lamborghini", "barilla", "luxottica", "gucci", "prada"
}

KNOWN_COMPANY_DISPLAY_NAMES = {
    "google": "Google",
    "amazon": "Amazon",
    "microsoft": "Microsoft",
    "apple": "Apple",
    "meta": "Meta",
    "facebook": "Facebook",
    "netflix": "Netflix",
    "tesla": "Tesla",
    "ibm": "IBM",
    "oracle": "Oracle",
    "accenture": "Accenture",
    "deloitte": "Deloitte",
    "pwc": "PwC",
    "kpmg": "KPMG",
    "ey": "EY",
    "linkedin": "LinkedIn",
    "spotify": "Spotify",
    "salesforce": "Salesforce",
    "adobe": "Adobe",
    "sap": "SAP",
    "siemens": "Siemens",
    "enel": "Enel",
    "eni": "Eni",
    "intesa sanpaolo": "Intesa Sanpaolo",
    "unicredit": "UniCredit",
    "poste italiane": "Poste Italiane",
    "telecom": "Telecom",
    "tim": "TIM",
    "ferrari": "Ferrari",
    "lamborghini": "Lamborghini",
    "barilla": "Barilla",
    "luxottica": "Luxottica",
    "gucci": "Gucci",
    "prada": "Prada",
}

COMPANY_SECTOR_HINTS = {
    "technology": {
        "google", "amazon", "microsoft", "apple", "meta", "facebook",
        "netflix", "ibm", "oracle", "linkedin", "spotify", "salesforce",
        "adobe", "sap",
    },
    "consulting": {"accenture", "deloitte", "pwc", "kpmg", "ey"},
    "finance": {"intesa sanpaolo", "unicredit"},
    "automotive": {"tesla", "ferrari", "lamborghini"},
    "energy": {"enel", "eni"},
    "telecommunications": {"telecom", "tim"},
    "food": {"barilla"},
    "fashion": {"luxottica", "gucci", "prada"},
    "industrial": {"siemens"},
    "logistics": {"poste italiane"},
}

ROLE_SECTOR_HINTS = {
    "beauty": {
        "estetista", "parrucchiere", "parrucchiera", "barbiere",
        "onicotecnica", "make up artist", "truccatore", "truccatrice",
    },
    "healthcare": {
        "infermiere", "infermiera", "medico", "fisioterapista",
        "odontoiatra", "farmacista", "ostetrica",
    },
    "hospitality": {
        "cameriere", "cameriera", "cuoco", "cuoca", "chef",
        "barista", "pizzaiolo", "receptionist albergo",
    },
    "construction": {
        "muratore", "imbianchino", "idraulico", "elettricista",
        "carpentiere", "piastrellista",
    },
    "automotive_service": {
        "meccanico", "meccanica", "gommista", "carrozziere",
    },
    "education": {
        "maestra", "maestro", "insegnante scuola", "educatrice asilo",
    },
}

SECTOR_COMPATIBILITY = {
    "beauty": {"fashion"},
    "healthcare": set(),
    "hospitality": {"food"},
    "construction": {"industrial"},
    "automotive_service": {"automotive", "industrial"},
    "education": set(),
}


def strip_accents(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value or "")
    return "".join(char for char in normalized if not unicodedata.combining(char))


def normalize_plain_text(value: Optional[str]) -> str:
    value = strip_accents(value or "").lower()
    value = re.sub(r"https?://\S+", " ", value)
    value = re.sub(r"[^a-z0-9\s&.+#/-]", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def canonical_known_company_name(company: Optional[str]) -> str:
    cleaned = normalize_plain_text(company)
    if not cleaned:
        return ""
    if cleaned in KNOWN_COMPANY_DISPLAY_NAMES:
        return KNOWN_COMPANY_DISPLAY_NAMES[cleaned]

    compact = re.sub(r"[^a-z0-9]", "", cleaned)
    matches = [
        canonical
        for canonical in KNOWN_COMPANY_DISPLAY_NAMES
        if len(compact) >= 4
        and re.sub(r"[^a-z0-9]", "", canonical).startswith(compact)
    ]
    if len(matches) == 1:
        return KNOWN_COMPANY_DISPLAY_NAMES[matches[0]]

    return str(company or "").strip()


def clean_job_role_title(value: Optional[str]) -> str:
    cleaned = str(value or "").strip()
    if any(marker in cleaned for marker in ("Ã", "Â")):
        try:
            cleaned = cleaned.encode("latin-1").decode("utf-8")
        except (UnicodeEncodeError, UnicodeDecodeError):
            pass
    cleaned = re.sub(
        r"(?i)\b(analyst|developer|engineer|manager|specialist|designer|scientist)[àáèéìíòóùú]\b",
        r"\1",
        cleaned,
    )
    return re.sub(r"\s+", " ", cleaned).strip()


def is_low_quality_text(value: Optional[str], min_chars: int = 3, min_words: int = 1) -> bool:
    cleaned = normalize_plain_text(value)
    compact = re.sub(r"[^a-z0-9]", "", cleaned)

    if len(compact) < min_chars:
        return True

    if cleaned in GENERIC_BAD_INPUTS or compact in GENERIC_BAD_INPUTS:
        return True

    if len(re.findall(r"[a-z0-9]+", cleaned)) < min_words:
        return True

    if re.fullmatch(r"(.)\1{3,}", compact or ""):
        return True

    # Reject inputs made almost entirely from a repeated short token (for
    # example "abcabcabcabc"), without flagging valid dates or words that
    # happen to contain a repeated sequence.
    if re.fullmatch(r"([a-z0-9]{1,3})\1{3,}", compact or ""):
        return True

    vowels = len(re.findall(r"[aeiou]", compact))
    letters = len(re.findall(r"[a-z]", compact))
    if letters >= 8 and vowels == 0:
        return True

    unique_ratio = len(set(compact)) / max(len(compact), 1)
    # Character diversity is meaningful only for short inputs. In normal
    # sentences the ratio naturally decreases as the text gets longer.
    if 8 <= len(compact) <= 24 and unique_ratio < 0.28:
        return True

    return False


def is_meaningful_cv_detail(value: Optional[str]) -> bool:
    """Return True when a short CV note still contains concrete evidence."""
    cleaned = normalize_plain_text(value)
    if not cleaned:
        return False

    concrete_patterns = (
        r"\b(?:19|20)\d{2}\b",
        r"\b[abc][12]\b",
        r"\b(?:presso|azienda|cliente|tirocinio|stage|impiego|lavorato)\b",
        r"\b(?:certificazione|certificato|attestato|corso|laurea|diploma)\b",
        r"\b(?:progetto|dashboard|dataset|portfolio|tesi)\b",
        r"\b(?:sql|python|java|javascript|typescript|excel|power bi|tableau|kpi|docker|git|github|aws|azure)\b",
        r"\b(?:problem solving|team working|leadership|comunicazione|organizzazione|precisione|collaborazione)\b",
    )
    return any(re.search(pattern, cleaned) for pattern in concrete_patterns)


def validate_role_plausibility(role: Optional[str]) -> Dict:
    cleaned = normalize_plain_text(role)

    if is_low_quality_text(cleaned, min_chars=4, min_words=1):
        return {
            "is_valid": False,
            "message": "Il ruolo inserito non sembra coerente. Inserisci un ruolo reale, ad esempio Data Analyst, Software Engineer o Marketing Specialist.",
        }

    if cleaned in {"lavoro", "ruolo", "impiego", "posto", "qualsiasi lavoro", "lavoro qualunque"}:
        return {
            "is_valid": False,
            "message": "Il ruolo inserito e troppo generico. Inserisci un ruolo lavorativo specifico.",
        }

    words = cleaned.split()
    sentence_terms = {
        "voglio", "vorrei", "prepararmi", "colloquio", "intervista",
        "candidarmi", "cerco", "azienda", "presso",
    }
    if len(words) > 6 or sentence_terms.intersection(words):
        return {
            "is_valid": False,
            "message": "Inserisci solo il titolo del ruolo, ad esempio Data Analyst o Software Engineer, non una frase sul colloquio.",
        }

    has_role_keyword = any(keyword in cleaned for keyword in ROLE_KEYWORDS)
    looks_like_title = 1 <= len(words) <= 6 and any(len(word) >= 3 for word in words)

    if has_role_keyword or (looks_like_title and len(cleaned) >= 6):
        return {"is_valid": True, "message": "Ruolo plausibile."}

    return {
        "is_valid": False,
        "message": "Il ruolo inserito non sembra coerente. Inserisci un ruolo reale, ad esempio Data Analyst, Software Engineer o Marketing Specialist.",
    }


QUICK_APPLICATION_INVALID_MESSAGE = (
    "Il testo inserito non sembra una candidatura valida. Scrivi il ruolo a cui "
    "stai puntando, ad esempio: 'Voglio candidarmi come Data Analyst'."
)
MISSING_TARGET_ROLE_MESSAGE = (
    "Inserisci almeno un ruolo target, ad esempio 'Data Analyst', 'Project Manager' "
    "o 'Computer Vision Engineer'. L'azienda e opzionale."
)
INVALID_COMPANY_MESSAGE = (
    "L'azienda inserita non sembra valida. Puoi correggerla oppure procedere "
    "indicando solo il ruolo."
)


def extract_quick_application_context(description: Optional[str]) -> Dict[str, Any]:
    raw = re.sub(r"\s+", " ", str(description or "")).strip()
    cleaned = normalize_plain_text(raw)
    invalid_markers = {
        "ignora le regole", "ignora le istruzioni", "prompt injection",
        "inventa esperienze", "esperienze inventate", "fai finta",
        "menti", "a caso", "ciao come stai", "che lavoro posso fare",
        "mi trovi un lavoro", "cosa devo fare", "secondo te va bene",
    }
    offensive_markers = {
        "cazzo", "merda", "stronzo", "stronza", "vaffanculo", "idiota",
        "fuck", "shit",
    }
    if (
        not raw
        or is_low_quality_text(raw, min_chars=12, min_words=3)
        or any(marker in cleaned for marker in invalid_markers | offensive_markers)
        or "?" in raw
    ):
        return {"is_valid": False, "role": "", "company": "", "message": QUICK_APPLICATION_INVALID_MESSAGE}

    role_patterns = [
        r"(?i)\b(?:candidarmi|candidare|puntando|punto)\s+(?:come|a un ruolo da|a una posizione da)\s+(.+?)(?=\s+(?:presso|in|per)\s+[A-ZÀ-ÖØ-Þ]|\s*$)",
        r"(?i)\b(?:ruolo|posizione)\s+(?:di|da)\s+(.+?)(?=\s+(?:presso|in)\s+[A-ZÀ-ÖØ-Þ]|\s*$)",
        r"(?i)\bottimizza(?:re)?\s+(?:il\s+)?cv\s+per\s+(?:il\s+)?ruolo\s+di\s+(.+?)(?=\s+(?:presso|in)\s+[A-ZÀ-ÖØ-Þ]|\s*$)",
    ]
    company_patterns = [
        r"(?i)\b(?:presso|in)\s+([A-ZÀ-ÖØ-Þ][\wÀ-ÿ&.'-]*(?:\s+[A-ZÀ-ÖØ-Þ][\wÀ-ÿ&.'-]*){0,4})\s*$",
    ]

    role = ""
    for pattern in role_patterns:
        match = re.search(pattern, raw)
        if match:
            role = clean_job_role_title(match.group(1).strip(" .,:;-"))
            break
    company = ""
    for pattern in company_patterns:
        match = re.search(pattern, raw)
        if match:
            company = match.group(1).strip(" .,:;-")
            break

    if company and role.lower().endswith(f" {company.lower()}"):
        role = role[:-(len(company) + 1)].strip()
    role_validation = validate_role_plausibility(role)
    if not role or not role_validation["is_valid"]:
        return {"is_valid": False, "role": "", "company": company, "message": QUICK_APPLICATION_INVALID_MESSAGE}
    return {
        "is_valid": True,
        "role": role,
        "company": company,
        "message": "Metodo rapido valido.",
    }


def classify_company_sector(company: Optional[str], sector: Optional[str] = "") -> Optional[str]:
    cleaned_company = normalize_plain_text(company)
    cleaned_sector = normalize_plain_text(sector)
    sector_aliases = {
        "tech": "technology",
        "tecnologia": "technology",
        "informatica": "technology",
        "consulenza": "consulting",
        "finanza": "finance",
        "banca": "finance",
        "automotive": "automotive",
        "energia": "energy",
        "telecomunicazioni": "telecommunications",
        "alimentare": "food",
        "moda": "fashion",
        "industria": "industrial",
        "logistica": "logistics",
    }
    for hint, family in sector_aliases.items():
        if hint in cleaned_sector:
            return family
    for family, companies in COMPANY_SECTOR_HINTS.items():
        if cleaned_company in companies:
            return family
    return None


def classify_role_sector(role: Optional[str]) -> Optional[str]:
    cleaned = normalize_plain_text(role)
    for family, hints in ROLE_SECTOR_HINTS.items():
        if any(hint in cleaned for hint in hints):
            return family
    return None


def validate_company_role_coherence(company: str, role: str, sector: str = "") -> Dict:
    company_family = classify_company_sector(company, sector)
    role_family = classify_role_sector(role)
    if not company_family or not role_family:
        return {"is_valid": True, "message": "Coerenza non determinabile con certezza."}

    compatible_sectors = SECTOR_COMPATIBILITY.get(role_family, set())
    if company_family in compatible_sectors:
        return {"is_valid": True, "message": "Azienda e ruolo risultano compatibili."}

    return {
        "is_valid": False,
        "message": (
            f"Il ruolo '{role.strip()}' non sembra coerente con l'attivita di {company.strip()}. "
            "Controlla azienda e ruolo oppure aggiungi una descrizione dell'offerta che chiarisca il collegamento."
        ),
    }


def tokenize_meaningful(value: Optional[str]) -> set:
    stopwords = {
        "per", "un", "una", "uno", "il", "lo", "la", "i", "gli", "le",
        "di", "da", "in", "a", "con", "su", "del", "della", "dello",
        "voglio", "prepararmi", "colloquio", "lavoro", "annuncio",
        "azienda", "ruolo", "presso", "come", "for", "the", "and"
    }
    return {
        token
        for token in re.findall(r"[a-z0-9][a-z0-9+#.-]*", normalize_plain_text(value))
        if len(token) > 2 and token not in stopwords
    }


def fields_are_coherent(description: str, company: str, role: str) -> tuple[bool, List[str]]:
    warnings = []
    desc_tokens = tokenize_meaningful(description)
    company_tokens = tokenize_meaningful(company)
    role_tokens = tokenize_meaningful(role)

    if description and len(normalize_plain_text(description)) >= 20:
        if company_tokens and not company_tokens.intersection(desc_tokens):
            warnings.append("La descrizione non cita l'azienda indicata.")

        if role_tokens and not role_tokens.intersection(desc_tokens):
            warnings.append("La descrizione non cita il ruolo indicato.")

    coherent = len(warnings) == 0
    if len(warnings) == 1 and len(desc_tokens) >= 10:
        coherent = True

    return coherent, warnings


def validate_job_link(link: Optional[str], company: str, role: str) -> Dict:
    link = (link or "").strip()
    if not link:
        return {"is_valid": True, "message": "Link non inserito.", "normalized_link": ""}

    if " " in link or is_low_quality_text(link, min_chars=8, min_words=1):
        return {
            "is_valid": False,
            "message": "Il link inserito non sembra valido. Inserisci un URL completo, ad esempio https://careers.google.com.",
            "normalized_link": "",
        }

    normalized_link = link if re.match(r"^https?://", link, re.I) else f"https://{link}"

    try:
        parsed = urllib.parse.urlparse(normalized_link)
    except Exception:
        parsed = None

    if not parsed or parsed.scheme not in {"http", "https"} or "." not in parsed.netloc:
        return {
            "is_valid": False,
            "message": "Il link inserito non e un URL valido.",
            "normalized_link": "",
        }

    return {
        "is_valid": True,
        "message": "Link valido.",
        "normalized_link": normalized_link,
    }


def verify_company_exists(company: str) -> Dict:
    cleaned = normalize_plain_text(company)
    canonical_name = canonical_known_company_name(company)
    if is_low_quality_text(cleaned, min_chars=3, min_words=1):
        return {
            "exists": False,
            "confidence": 0,
            "sources": [],
            "normalized_name": "",
            "message": "Non sono riuscito a verificare l'esistenza dell'azienda inserita. Controlla che il nome sia corretto.",
        }

    if not TAVILY_API_KEY:
        words = cleaned.split()
        compact = re.sub(r"[^a-z]", "", cleaned)
        plausible = (
            cleaned in KNOWN_COMPANIES
            or (
                1 <= len(words) <= 6
                and len(compact) >= 4
                and len(re.findall(r"[aeiou]", compact)) >= 2
                and not re.search(r"[bcdfghjklmnpqrstvwxyz]{7,}", compact)
                and not is_low_quality_text(cleaned, min_chars=4, min_words=1)
            )
        )
        return {
            "exists": plausible,
            "confidence": 55 if plausible else 20,
            "sources": [],
            "normalized_name": canonical_name if plausible else "",
            "message": (
                "Azienda plausibile in base ai controlli locali. Configura TAVILY_API_KEY per la verifica web."
                if plausible
                else "Non sono riuscito a verificare l'esistenza dell'azienda inserita. Controlla che il nome sia corretto."
            ),
        }

    query = f'{company} official website LinkedIn careers company'
    try:
        response = requests.post(
            "https://api.tavily.com/search",
            json={
                "api_key": TAVILY_API_KEY,
                "query": query,
                "search_depth": "basic",
                "max_results": 5,
                "include_answer": False,
                "include_raw_content": False,
            },
            timeout=10,
        )
        response.raise_for_status()
        results = response.json().get("results", [])
    except Exception as exc:
        print(f"Verifica azienda Tavily non riuscita, uso fallback locale: {exc}")
        return {
            "exists": cleaned in KNOWN_COMPANIES or len(cleaned) >= 3,
            "confidence": 45,
            "sources": [],
            "normalized_name": canonical_name,
            "message": "Verifica web non disponibile al momento; ho applicato un controllo locale di plausibilita.",
        }

    company_tokens = tokenize_meaningful(company)
    sources = []
    score = 0

    for item in results:
        title = item.get("title", "")
        url = item.get("url", "")
        content = item.get("content", "")
        haystack = normalize_plain_text(f"{title} {url} {content}")
        token_hits = sum(1 for token in company_tokens if token in haystack)
        trusted_hit = any(marker in haystack for marker in ["official", "linkedin", "careers", "azienda", "company"])

        if token_hits:
            score += 25 + (10 if trusted_hit else 0)
            sources.append({"title": title, "url": url, "content": content})

    exists = score >= 25
    return {
        "exists": exists,
        "confidence": clamp_score(score),
        "sources": sources[:4],
        "normalized_name": canonical_name if exists else "",
        "message": (
            "Azienda verificata con fonti web coerenti."
            if exists
            else "Non sono riuscito a verificare l'esistenza dell'azienda inserita. Controlla che il nome sia corretto."
        ),
    }


def validate_required_skills(required_skills: Optional[str], role: str, description: str) -> tuple[bool, List[str], str]:
    cleaned = normalize_plain_text(required_skills or "")
    if not cleaned:
        return True, [], "Competenze richieste non inserite."

    if is_low_quality_text(cleaned, min_chars=4, min_words=1):
        return False, [], "Le competenze richieste non sembrano coerenti. Inserisci strumenti, tecnologie o capacità reali."

    skills = [
        item.strip()
        for item in re.split(r"[,;\n]+", required_skills or "")
        if item.strip()
    ][:16]
    if not skills:
        return False, [], "Inserisci almeno una competenza richiesta reale."

    context_tokens = tokenize_meaningful(f"{role} {description}")
    skill_tokens = tokenize_meaningful(" ".join(skills))
    if context_tokens and skill_tokens and not context_tokens.intersection(skill_tokens) and len(skill_tokens) <= 2:
        return False, skills, "Le competenze richieste non sembrano collegate al ruolo o alla descrizione inserita."

    return True, skills, "Competenze richieste plausibili."


def validate_job_input(
    description: str,
    company: str,
    role: str,
    link: str,
    sector: str = "",
    required_skills: str = "",
) -> Dict:
    description = (description or "").strip()
    company = (company or "").strip()
    role = clean_job_role_title(role)
    errors = {}
    warnings = []
    quick_context = extract_quick_application_context(description) if description else {
        "is_valid": False,
        "role": "",
        "company": "",
    }
    if description and not quick_context["is_valid"]:
        errors["description"] = QUICK_APPLICATION_INVALID_MESSAGE

    effective_role = role or str(quick_context.get("role") or "").strip()
    effective_company = company or str(quick_context.get("company") or "").strip()
    role_validation = validate_role_plausibility(effective_role)
    link_validation = validate_job_link(link, effective_company, effective_role)
    has_valid_link = link_validation["is_valid"] and bool((link or "").strip())

    if not role_validation["is_valid"] and not has_valid_link:
        errors["role"] = MISSING_TARGET_ROLE_MESSAGE
    elif effective_role and not role_validation["is_valid"]:
        errors["role"] = role_validation["message"]

    company_check = {
        "exists": False,
        "confidence": 0,
        "sources": [],
        "normalized_name": "",
    }
    if effective_company:
        if is_low_quality_text(effective_company, min_chars=3, min_words=1):
            errors["company"] = INVALID_COMPANY_MESSAGE
        else:
            company_check = verify_company_exists(effective_company)
            if not company_check["exists"]:
                errors["company"] = INVALID_COMPANY_MESSAGE

    company_role_validation = validate_company_role_coherence(
        effective_company,
        effective_role,
        sector,
    )
    if effective_company and effective_role and not company_role_validation["is_valid"]:
        errors["coherence"] = company_role_validation["message"]
    if not link_validation["is_valid"]:
        errors["link"] = link_validation["message"]
    if sector and is_low_quality_text(sector, min_chars=3, min_words=1):
        errors["sector"] = "Il settore inserito non sembra coerente."

    skills_valid, normalized_skills, skills_message = validate_required_skills(
        required_skills,
        effective_role,
        description,
    )
    if not skills_valid:
        errors["required_skills"] = skills_message

    coherent, coherence_warnings = fields_are_coherent(
        description,
        effective_company,
        effective_role,
    )
    if effective_company and effective_role and description and not coherent and "coherence" not in errors:
        errors["coherence"] = "La descrizione non sembra coerente con azienda e ruolo indicati."
    else:
        warnings.extend(coherence_warnings)

    normalized_company = company_check.get("normalized_name") or effective_company

    is_valid = not errors
    return {
        "is_valid": is_valid,
        "company_exists": bool(company_check["exists"]),
        "role_is_plausible": role_validation["is_valid"],
        "fields_are_coherent": coherent,
        "link_is_valid": link_validation["is_valid"],
        "errors": errors,
        "warnings": warnings,
        "sources": company_check.get("sources", []),
        "normalized_link": link_validation.get("normalized_link", ""),
        "required_skills": normalized_skills,
        "normalized_role": effective_role,
        "normalized_company": normalized_company,
        "quick_method_used": bool(description and quick_context["is_valid"]),
        "message": "I dati inseriti sono validi." if is_valid else "Correggi i campi evidenziati prima di continuare.",
    }

    has_description = not is_low_quality_text(description, min_chars=20, min_words=4)
    has_specific_details = (
        not is_low_quality_text(company, min_chars=3, min_words=1)
        and validate_role_plausibility(role)["is_valid"]
    )

    link_validation = validate_job_link(link, company, role)
    has_valid_link = link_validation["is_valid"] and bool(link.strip())

    if not has_description and not has_specific_details and not has_valid_link:
        errors["description"] = "Compila il metodo rapido, inserisci l'azienda e il ruolo, oppure fornisci un link valido."

    if company and is_low_quality_text(company, min_chars=3, min_words=1):
        errors["company"] = "Il nome azienda non sembra valido."

    role_validation = validate_role_plausibility(role)
    if role and not role_validation["is_valid"]:
        errors["role"] = role_validation["message"]
    elif role and not infer_role_family(role, description, required_skills) and not classify_role_sector(role):
        errors["role"] = (
            "Il ruolo non è stato riconosciuto. Inserisci un titolo professionale specifico e completo, "
            "ad esempio Data Analyst, Software Engineer, Project Manager, Estetista o Infermiere."
        )

    company_role_validation = validate_company_role_coherence(company, role, sector)
    if company and role and not company_role_validation["is_valid"]:
        errors["coherence"] = company_role_validation["message"]

    if not link_validation["is_valid"]:
        errors["link"] = link_validation["message"]

    if sector and is_low_quality_text(sector, min_chars=3, min_words=1):
        errors["sector"] = "Il settore inserito non sembra coerente."

    skills_valid, normalized_skills, skills_message = validate_required_skills(required_skills, role, description)
    if not skills_valid:
        errors["required_skills"] = skills_message

    coherent, coherence_warnings = fields_are_coherent(description, company, role)
    if company and role and description and not coherent and "coherence" not in errors:
        errors["coherence"] = "La descrizione non sembra coerente con azienda e ruolo indicati."
    else:
        warnings.extend(coherence_warnings)

    company_check = verify_company_exists(company) if company and "company" not in errors else {
        "exists": False,
        "confidence": 0,
        "sources": [],
        "message": "Azienda non inserita o non verificata.",
    }
    if company and not company_check["exists"]:
        warnings.append(company_check["message"])

    is_valid = not errors
    return {
        "is_valid": is_valid,
        "company_exists": bool(company_check["exists"]),
        "role_is_plausible": role_validation["is_valid"],
        "fields_are_coherent": coherent,
        "link_is_valid": link_validation["is_valid"],
        "errors": errors,
        "warnings": warnings,
        "sources": company_check.get("sources", []),
        "normalized_link": link_validation.get("normalized_link", ""),
        "required_skills": normalized_skills,
        "message": "I dati inseriti sono validi." if is_valid else "Correggi i campi evidenziati prima di continuare.",
    }


def extract_requested_keywords(role: str, description: str, required_skills: str = "") -> List[str]:
    return JobAnalyzer(normalize_plain_text).extract_keywords(role, description, required_skills)


def filter_cv_keyword_list(values: List[Any]) -> List[str]:
    # Block generic single-word role nouns and weak tokens from appearing as ATS keywords
    blocked = {"data", "analyst", "analysis", "business", "project", "manager", "team", "office",
               "scientist", "engineer", "developer", "specialist", "consultant", "designer", "researcher"}
    cleaned = []
    seen = set()
    for value in values or []:
        text = str(value or "").strip()
        normalized = normalize_plain_text(text)
        if not normalized or normalized in blocked:
            continue
        if normalized == "data analyst":
            continue
        if normalized in seen:
            continue
        seen.add(normalized)
        cleaned.append(text)
    return cleaned


ATS_HARD_SKILL_TERMS = [
    "python", "sql", "postgresql", "mysql", "excel", "power bi", "tableau",
    "database", "etl", "data cleaning", "data analysis", "analisi dati",
    "machine learning", "ai", "ml", "llm", "nlp", "clustering testuale",
    "accuratezza", "tempi di risposta", "preparazione dati", "confronto modelli",
    "pandas", "numpy", "cplex", "ottimizzazione",
    "analisi finanziaria", "modellazione finanziaria", "valutazione d'azienda",
    "bilancio", "contabilita generale", "ifrs", "gaap", "corporate finance",
    "mergers and acquisitions", "due diligence", "risk management", "credit risk",
    "market risk", "controllo di gestione", "budgeting", "forecasting",
    "scienza attuariale", "econometria", "audit", "fiscalita", "asset allocation",
    "portfolio management", "cad", "progettazione meccanica", "analisi fem",
    "termodinamica", "bim", "progettazione strutturale", "calcolo strutturale",
    "impianti elettrici", "circuit design", "lean manufacturing", "six sigma",
    "ingegneria di processo", "simulazione di processo", "valutazione ambientale",
    "life cycle assessment", "dispositivi medici", "segnali biomedicali",
    "inferenza statistica", "disegno sperimentale", "ricerca operativa",
    "programmazione lineare", "analisi quantitativa", "stochastic calculus",
    "modellazione matematica", "strategia di marketing", "marketing automation",
    "seo", "sem", "content strategy", "brand management", "media relations",
    "talent acquisition", "selezione del personale", "diritto del lavoro",
    "contrattualistica", "compliance normativa", "gdpr", "user research",
    "interaction design", "information architecture", "graphic design",
    "tipografia", "industrial design", "clinical data management", "gcp",
    "bioinformatica", "genomica", "metodologia scientifica", "laboratory research",
    "supply chain planning", "demand planning", "gestione inventario",
    "logistica", "procurement", "strategic sourcing", "vendor management",
]

ATS_SOFT_SKILL_TERMS = [
    "comunicazione", "problem solving", "team", "collaborazione",
    "pensiero analitico", "ragionamento analitico", "decisioni",
    "organizzazione", "precisione", "autonomia",
    "attenzione ai dettagli", "leadership", "negoziazione",
    "gestione stakeholder", "orientamento al risultato", "integrita professionale",
    "riservatezza", "creativita", "pensiero sistemico", "gestione priorita",
    "comunicazione scientifica", "comunicazione visiva", "empatia",
    "capacita decisionale", "gestione del cambiamento", "orientamento al cliente",
]


def split_requested_skill_terms(role: str, description: str, required_skills: str = "") -> tuple[List[str], List[str]]:
    requirement_text = normalize_plain_text(f"{role} {description} {required_skills}")
    hard_skills = [term for term in ATS_HARD_SKILL_TERMS if normalize_plain_text(term) in requirement_text]
    soft_skills = [term for term in ATS_SOFT_SKILL_TERMS if normalize_plain_text(term) in requirement_text]
    return hard_skills[:12], soft_skills[:12]


ROLE_KEYWORD_GROUPS = {
    "data analyst": [
        ("Python", ["python"]),
        ("SQL", ["sql"]),
        ("Machine Learning / AI", ["machine learning", "intelligenza artificiale", " ai ", " ml "]),
        ("Analisi dei dati", ["analisi dei dati", "analisi dati", "data analysis"]),
        ("NLP / LLM", ["nlp", "natural language processing", "llm", "large language"]),
        ("clustering testuale", ["clustering testuale", "clustering"]),
        ("accuratezza", ["accuratezza"]),
        ("tempi di risposta", ["tempi di risposta"]),
        ("preparazione dati", ["preparazione dati", "preparazione dei dati"]),
        ("confronto modelli", ["confronto modelli", "confronto delle prestazioni", "valutazione delle prestazioni"]),
        ("problem solving", ["problem solving", "risoluzione di problemi"]),
    ],
    "game design": [
        ("Game design", ["game design", "game designer"]),
        ("Level design", ["level design"]),
        ("Storytelling", ["storytelling", "narrazione"]),
        ("UX", ["ux", "user experience", "esperienza utente"]),
        ("Unity", ["unity"]),
        ("Unreal Engine", ["unreal engine", "unreal"]),
        ("C#", ["c#"]),
        ("prototipazione", ["prototipazione", "prototype", "prototipi"]),
        ("playtesting", ["playtesting", "play test"]),
        ("portfolio progetti", ["portfolio", "progetti"]),
    ],
    "financial analyst": [
        ("Analisi finanziaria", ["analisi finanziaria", "financial analysis"]),
        ("Modellazione finanziaria", ["modellazione finanziaria", "financial modeling"]),
        ("Valutazione d'azienda", ["valutazione d'azienda", "business valuation", "company valuation"]),
        ("Budgeting e forecasting", ["budgeting", "forecasting"]),
        ("Excel avanzato", ["excel avanzato", "advanced excel"]),
        ("Bloomberg Terminal", ["bloomberg terminal", "bloomberg"]),
    ],
    "accountant": [
        ("Contabilita generale", ["contabilita generale", "general ledger"]),
        ("Bilancio", ["bilancio", "financial statements"]),
        ("IFRS / GAAP", ["ifrs", "gaap"]),
        ("Riconciliazioni", ["riconciliazioni", "reconciliation"]),
        ("SAP", ["sap"]),
        ("Zucchetti", ["zucchetti"]),
    ],
    "risk manager": [
        ("Risk management", ["risk management", "gestione dei rischi"]),
        ("Credit risk", ["credit risk", "rischio di credito"]),
        ("Market risk", ["market risk", "rischio di mercato"]),
        ("Stress testing", ["stress testing", "stress test"]),
        ("VaR", ["value at risk", "var"]),
        ("SAS", ["sas"]),
    ],
    "mechanical engineer": [
        ("Progettazione meccanica", ["progettazione meccanica", "mechanical design"]),
        ("CAD 3D", ["cad 3d", "3d cad"]),
        ("Analisi FEM", ["analisi fem", "finite element analysis", "fea"]),
        ("GD&T", ["gd&t", "geometric dimensioning"]),
        ("SolidWorks", ["solidworks"]),
        ("ANSYS", ["ansys"]),
    ],
    "civil engineer": [
        ("Progettazione civile", ["progettazione civile", "civil design"]),
        ("BIM", ["bim", "building information modeling"]),
        ("Calcolo strutturale", ["calcolo strutturale", "structural analysis"]),
        ("Direzione lavori", ["direzione lavori", "construction supervision"]),
        ("AutoCAD", ["autocad"]),
        ("Revit", ["revit"]),
    ],
    "electrical engineer": [
        ("Progettazione elettrica", ["progettazione elettrica", "electrical design"]),
        ("Circuit design", ["circuit design", "progettazione circuiti"]),
        ("PLC", ["plc", "programmable logic controller"]),
        ("Sistemi di potenza", ["sistemi di potenza", "power systems"]),
        ("MATLAB", ["matlab"]),
        ("ETAP", ["etap"]),
    ],
    "marketing manager": [
        ("Strategia di marketing", ["strategia di marketing", "marketing strategy"]),
        ("Market research", ["market research", "ricerche di mercato"]),
        ("Campaign management", ["campaign management", "gestione campagne"]),
        ("Brand positioning", ["brand positioning", "posizionamento del brand"]),
        ("Google Analytics", ["google analytics", "ga4"]),
        ("HubSpot", ["hubspot"]),
    ],
    "digital marketer": [
        ("Digital advertising", ["digital advertising", "online advertising"]),
        ("SEO / SEM", ["seo", "sem", "search engine marketing"]),
        ("Marketing automation", ["marketing automation"]),
        ("Conversion rate optimization", ["conversion rate optimization", "cro"]),
        ("Google Ads", ["google ads"]),
        ("Meta Ads Manager", ["meta ads manager", "facebook ads manager"]),
    ],
    "recruiter": [
        ("Talent acquisition", ["talent acquisition", "acquisizione talenti"]),
        ("Sourcing", ["candidate sourcing", "sourcing candidati"]),
        ("Selezione del personale", ["selezione del personale", "recruitment"]),
        ("Interviste strutturate", ["interviste strutturate", "structured interviews"]),
        ("LinkedIn Recruiter", ["linkedin recruiter"]),
        ("ATS recruiting", ["applicant tracking system", "ats recruiting"]),
    ],
    "legal counsel": [
        ("Contrattualistica", ["contrattualistica", "contract drafting"]),
        ("Diritto societario", ["diritto societario", "corporate law"]),
        ("Legal research", ["legal research", "ricerca giuridica"]),
        ("Compliance normativa", ["compliance normativa", "regulatory compliance"]),
        ("Westlaw", ["westlaw"]),
        ("LexisNexis", ["lexisnexis"]),
    ],
    "ux designer": [
        ("User research", ["user research", "ricerca utente"]),
        ("Interaction design", ["interaction design"]),
        ("Wireframing", ["wireframing", "wireframe"]),
        ("Usability testing", ["usability testing", "test di usabilita"]),
        ("Figma", ["figma"]),
        ("Miro", ["miro"]),
    ],
    "supply chain manager": [
        ("Supply chain planning", ["supply chain planning", "pianificazione supply chain"]),
        ("Demand planning", ["demand planning", "pianificazione della domanda"]),
        ("Gestione inventario", ["gestione inventario", "inventory management"]),
        ("S&OP", ["s&op", "sales and operations planning"]),
        ("SAP", ["sap"]),
        ("Oracle SCM Cloud", ["oracle scm cloud", "oracle scm"]),
    ],
}

DATA_ANALYST_KEYWORDS_TO_CONFIRM = ["Excel avanzato", "Power BI", "Tableau", "dashboard", "KPI", "reporting"]
GAME_DESIGN_KEYWORDS_TO_CONFIRM = ["Unity", "Unreal Engine", "C#", "prototipazione", "playtesting", "portfolio progetti"]

ROLE_SKILL_LIBRARY = {
    "data analyst": {
        "hard_skills": ["Analisi dei dati", "SQL", "Python", "Machine Learning", "Statistica", "Data visualization", "Reporting", "KPI", "Database"],
        "soft_skills": ["Problem solving", "Attenzione ai dettagli", "Pensiero analitico", "Comunicazione dei risultati", "Team working"],
        "programming_languages": ["Python", "SQL"],
        "tools": ["Excel avanzato", "Power BI", "Tableau", "Google Analytics", "BigQuery"],
    },
    "game design": {
        "hard_skills": ["Game design", "Level design", "Prototipazione", "Game mechanics", "User experience", "Playtesting", "Storytelling"],
        "soft_skills": ["Creativita", "Problem solving", "Collaborazione", "Comunicazione", "Iterazione su feedback"],
        "programming_languages": ["C#", "C++", "Python"],
        "tools": ["Unity", "Unreal Engine", "Blender", "Figma", "Miro"],
    },
    "backend developer": {
        "hard_skills": ["API development", "Database design", "Autenticazione", "Testing backend", "Architetture REST", "Deploy"],
        "soft_skills": ["Problem solving", "Precisione", "Collaborazione", "Pensiero logico", "Documentazione tecnica"],
        "programming_languages": ["Python", "Java", "JavaScript", "Node.js"],
        "tools": ["FastAPI", "Django", "Spring", "PostgreSQL", "MongoDB", "Docker", "Cloud"],
    },
    "data scientist": {
        "hard_skills": ["Python", "Machine Learning", "SQL", "Analisi predittiva", "Modelli statistici", "Data preprocessing", "Feature engineering", "Data visualization"],
        "soft_skills": ["Problem solving", "Pensiero analitico", "Comunicazione scientifica", "Collaborazione", "Attenzione ai dettagli"],
        "programming_languages": ["Python", "SQL", "R"],
        "tools": ["pandas", "scikit-learn", "Jupyter Notebook", "TensorFlow", "Tableau", "Excel avanzato"],
    },
    "project manager": {
        "hard_skills": ["Pianificazione attività", "Gestione scadenze", "Monitoraggio avanzamento progetto", "Gestione budget", "Risk management"],
        "soft_skills": ["Comunicazione con stakeholder", "Coordinamento team", "Organizzazione", "Problem solving", "Gestione priorità", "Leadership", "Negoziazione"],
        "programming_languages": [],
        "tools": ["Excel avanzato", "Trello", "Jira", "Notion", "Microsoft Project", "Asana", "Monday.com"],
    },
    "software engineer": {
        "hard_skills": ["Sviluppo software", "Design architetturale", "Debugging", "Version control", "Unit testing", "Code review"],
        "soft_skills": ["Problem solving", "Collaborazione", "Comunicazione tecnica", "Precisione", "Pensiero logico"],
        "programming_languages": ["Python", "JavaScript", "Java", "C++"],
        "tools": ["Git", "GitHub", "VS Code", "Docker", "AWS", "GitLab"],
    },
    "frontend developer": {
        "hard_skills": ["HTML/CSS", "JavaScript", "User Interface Design", "User Experience", "Responsive Design", "Accessibilita"],
        "soft_skills": ["Creativita", "Problem solving", "Comunicazione", "Attenzione ai dettagli", "Collaborazione"],
        "programming_languages": ["JavaScript", "TypeScript", "React"],
        "tools": ["Figma", "Adobe XD", "VS Code", "Git", "Webpack"],
    },
    "financial analyst": {
        "hard_skills": ["Analisi finanziaria", "Modellazione finanziaria", "Valutazione d'azienda", "Analisi di bilancio", "Budgeting", "Forecasting", "Variance analysis", "KPI finanziari"],
        "soft_skills": ["Pensiero analitico", "Attenzione ai dettagli", "Comunicazione con stakeholder", "Gestione delle scadenze", "Integrita professionale"],
        "programming_languages": ["SQL", "Python", "VBA"],
        "tools": ["Excel avanzato", "Bloomberg Terminal", "Power BI", "Tableau", "SAP", "Oracle Hyperion"],
    },
    "accountant": {
        "hard_skills": ["Contabilita generale", "Scritture contabili", "Chiusure mensili", "Redazione del bilancio", "Riconciliazioni contabili", "Principi IFRS", "Principi GAAP", "Contabilita fornitori e clienti"],
        "soft_skills": ["Precisione", "Affidabilita", "Gestione delle scadenze", "Riservatezza", "Organizzazione"],
        "programming_languages": [],
        "tools": ["SAP", "Oracle NetSuite", "Zucchetti", "TeamSystem", "Excel avanzato", "Microsoft Dynamics 365"],
    },
    "investment banker": {
        "hard_skills": ["Mergers and Acquisitions", "Corporate finance", "Financial modeling", "Valutazione d'azienda", "Due diligence finanziaria", "Leveraged Buyout modeling", "Analisi dei capital markets", "Preparazione di pitch book"],
        "soft_skills": ["Resistenza allo stress", "Negoziazione", "Comunicazione executive", "Attenzione ai dettagli", "Gestione delle priorita"],
        "programming_languages": ["VBA", "Python"],
        "tools": ["Bloomberg Terminal", "Refinitiv Workspace", "S&P Capital IQ", "PitchBook", "FactSet", "Excel avanzato", "PowerPoint"],
    },
    "risk manager": {
        "hard_skills": ["Enterprise Risk Management", "Credit risk", "Market risk", "Operational risk", "Value at Risk", "Stress testing", "Risk assessment", "Regulatory reporting"],
        "soft_skills": ["Pensiero critico", "Capacita decisionale", "Comunicazione con stakeholder", "Integrita professionale", "Gestione delle priorita"],
        "programming_languages": ["Python", "R", "SQL", "SAS"],
        "tools": ["SAS Risk Management", "Moody's Analytics", "MATLAB", "Bloomberg Terminal", "Power BI", "Excel avanzato"],
    },
    "controller": {
        "hard_skills": ["Controllo di gestione", "Budgeting", "Forecasting", "Variance analysis", "Cost accounting", "Management reporting", "Analisi della marginalita", "Definizione KPI"],
        "soft_skills": ["Pensiero analitico", "Orientamento al risultato", "Comunicazione manageriale", "Precisione", "Pianificazione"],
        "programming_languages": ["SQL", "VBA"],
        "tools": ["SAP CO", "Oracle Hyperion", "Tagetik", "Power BI", "Excel avanzato", "Qlik Sense"],
    },
    "actuary": {
        "hard_skills": ["Matematica attuariale", "Pricing assicurativo", "Reserving", "Solvency II", "Modelli di mortalita", "Analisi di sopravvivenza", "Risk modeling", "Valutazione delle passivita"],
        "soft_skills": ["Pensiero quantitativo", "Precisione", "Comunicazione tecnica", "Pensiero critico", "Apprendimento continuo"],
        "programming_languages": ["R", "Python", "SQL", "SAS"],
        "tools": ["Prophet", "MoSes", "SAS", "RStudio", "Excel avanzato", "Power BI"],
    },
    "economist": {
        "hard_skills": ["Econometria", "Analisi macroeconomica", "Analisi microeconomica", "Modelli economici", "Forecasting economico", "Valutazione delle politiche pubbliche", "Analisi delle serie storiche", "Causal inference"],
        "soft_skills": ["Pensiero critico", "Comunicazione dei risultati", "Rigore metodologico", "Curiosita intellettuale", "Sintesi"],
        "programming_languages": ["R", "Python", "Stata", "SQL"],
        "tools": ["Stata", "EViews", "RStudio", "MATLAB", "Excel avanzato", "World Bank DataBank"],
    },
    "auditor": {
        "hard_skills": ["Revisione contabile", "Audit planning", "Internal controls", "Risk assessment", "Test di conformita", "Analisi di bilancio", "Campionamento di audit", "Principi ISA"],
        "soft_skills": ["Integrita professionale", "Scetticismo professionale", "Precisione", "Comunicazione con il cliente", "Gestione delle scadenze"],
        "programming_languages": ["SQL"],
        "tools": ["CaseWare", "TeamMate+", "IDEA", "ACL Analytics", "SAP", "Excel avanzato"],
    },
    "tax consultant": {
        "hard_skills": ["Fiscalita d'impresa", "Imposte dirette e indirette", "IVA", "Transfer pricing", "Tax compliance", "Pianificazione fiscale", "Contenzioso tributario", "Dichiarazioni fiscali"],
        "soft_skills": ["Precisione", "Riservatezza", "Aggiornamento continuo", "Comunicazione con il cliente", "Gestione delle scadenze"],
        "programming_languages": [],
        "tools": ["TeamSystem", "Zucchetti", "SAP", "Bloomberg Tax", "OneSource", "Excel avanzato"],
    },
    "portfolio manager": {
        "hard_skills": ["Asset allocation", "Portfolio construction", "Security analysis", "Performance attribution", "Risk-adjusted return", "Fixed income analysis", "Equity valuation", "Investment strategy"],
        "soft_skills": ["Capacita decisionale", "Gestione del rischio", "Comunicazione con investitori", "Disciplina", "Pensiero strategico"],
        "programming_languages": ["Python", "R", "SQL"],
        "tools": ["Bloomberg Terminal", "FactSet", "Morningstar Direct", "BlackRock Aladdin", "Refinitiv Workspace", "Excel avanzato"],
    },
    "mechanical engineer": {
        "hard_skills": ["Progettazione meccanica", "CAD 3D", "Analisi FEM", "GD&T", "Termodinamica", "Meccanica dei materiali", "Design for Manufacturing", "Tolleranze dimensionali"],
        "soft_skills": ["Problem solving tecnico", "Precisione", "Collaborazione multidisciplinare", "Pensiero sistemico", "Gestione delle priorita"],
        "programming_languages": ["MATLAB", "Python"],
        "tools": ["SolidWorks", "CATIA", "Siemens NX", "AutoCAD", "ANSYS", "Abaqus", "PTC Creo"],
    },
    "civil engineer": {
        "hard_skills": ["Progettazione civile", "Calcolo strutturale", "BIM", "Geotecnica", "Idraulica", "Direzione lavori", "Computo metrico", "Normativa delle costruzioni"],
        "soft_skills": ["Gestione del cantiere", "Comunicazione con stakeholder", "Problem solving sul campo", "Precisione", "Pianificazione"],
        "programming_languages": ["Python", "MATLAB"],
        "tools": ["AutoCAD Civil 3D", "Revit", "SAP2000", "ETABS", "Primus", "ArcGIS", "Microsoft Project"],
    },
    "electrical engineer": {
        "hard_skills": ["Progettazione elettrica", "Circuit design", "Sistemi di potenza", "Controlli automatici", "Elettronica analogica e digitale", "PLC", "Dimensionamento impianti", "Normative IEC"],
        "soft_skills": ["Pensiero analitico", "Problem solving tecnico", "Precisione", "Collaborazione multidisciplinare", "Orientamento alla sicurezza"],
        "programming_languages": ["MATLAB", "C", "C++", "VHDL"],
        "tools": ["MATLAB Simulink", "ETAP", "EPLAN", "AutoCAD Electrical", "LTspice", "Altium Designer", "Siemens TIA Portal"],
    },
    "structural engineer": {
        "hard_skills": ["Analisi strutturale", "Progettazione in cemento armato", "Progettazione di strutture in acciaio", "Analisi sismica", "Analisi FEM", "Verifiche agli stati limite", "Eurocodici", "Dettagli costruttivi"],
        "soft_skills": ["Precisione", "Pensiero critico", "Responsabilita professionale", "Collaborazione con progettisti", "Gestione delle scadenze"],
        "programming_languages": ["Python", "MATLAB"],
        "tools": ["SAP2000", "ETABS", "Tekla Structures", "Revit", "MIDAS Gen", "Abaqus", "AutoCAD"],
    },
    "industrial engineer": {
        "hard_skills": ["Lean manufacturing", "Six Sigma", "Ottimizzazione dei processi", "Tempi e metodi", "Operations management", "Quality management", "Capacity planning", "Supply chain analysis"],
        "soft_skills": ["Pensiero sistemico", "Orientamento al miglioramento continuo", "Leadership operativa", "Comunicazione interfunzionale", "Capacita decisionale"],
        "programming_languages": ["Python", "R", "SQL", "MATLAB"],
        "tools": ["Minitab", "Arena Simulation", "SAP", "Power BI", "Excel avanzato", "Microsoft Visio", "IBM CPLEX"],
    },
    "marketing manager": {
        "hard_skills": ["Strategia di marketing", "Market research", "Campaign management", "Brand positioning", "Go-to-market strategy", "Marketing analytics", "Budget marketing", "Customer segmentation"],
        "soft_skills": ["Creativita strategica", "Leadership", "Comunicazione persuasiva", "Orientamento al cliente", "Pensiero analitico"],
        "programming_languages": [],
        "tools": ["Google Analytics", "HubSpot", "Salesforce", "Meta Ads Manager", "Google Ads", "Power BI", "Excel avanzato"],
    },
    "digital marketer": {
        "hard_skills": ["Digital advertising", "SEO", "SEM", "Marketing automation", "Conversion rate optimization", "Email marketing", "Content performance analysis", "A/B testing"],
        "soft_skills": ["Curiosita digitale", "Pensiero analitico", "Creativita", "Adattabilita", "Orientamento ai risultati"],
        "programming_languages": [],
        "tools": ["Google Ads", "Meta Ads Manager", "Google Analytics", "Search Console", "Mailchimp", "HubSpot", "SEMrush"],
    },
    "recruiter": {
        "hard_skills": ["Talent acquisition", "Sourcing candidati", "Selezione del personale", "Interviste strutturate", "Screening CV", "Employer branding", "Candidate pipeline management", "ATS recruiting"],
        "soft_skills": ["Ascolto attivo", "Comunicazione", "Empatia professionale", "Negoziazione", "Valutazione critica"],
        "programming_languages": [],
        "tools": ["LinkedIn Recruiter", "ATS", "Greenhouse", "Lever", "Workday", "Excel", "Google Workspace"],
    },
    "legal counsel": {
        "hard_skills": ["Contrattualistica", "Diritto societario", "Legal research", "Compliance normativa", "Corporate governance", "Risk assessment legale", "Privacy e GDPR", "Negoziazione contrattuale"],
        "soft_skills": ["Pensiero critico", "Precisione", "Riservatezza", "Comunicazione consulenziale", "Gestione delle priorita"],
        "programming_languages": [],
        "tools": ["Westlaw", "LexisNexis", "OneTrust", "DocuSign", "Microsoft Word", "SharePoint"],
    },
    "ux designer": {
        "hard_skills": ["User research", "Interaction design", "Wireframing", "Usability testing", "Information architecture", "Prototipazione", "Design system", "User journey mapping"],
        "soft_skills": ["Empatia utente", "Creativita", "Collaborazione", "Comunicazione visiva", "Pensiero critico"],
        "programming_languages": [],
        "tools": ["Figma", "Miro", "FigJam", "Adobe XD", "Maze", "Optimal Workshop", "Notion"],
    },
    "supply chain manager": {
        "hard_skills": ["Supply chain planning", "Demand planning", "Gestione inventario", "S&OP", "Logistics management", "Procurement planning", "Supplier performance analysis", "Forecasting operativo"],
        "soft_skills": ["Pensiero sistemico", "Negoziazione", "Problem solving operativo", "Comunicazione interfunzionale", "Gestione delle priorita"],
        "programming_languages": ["SQL", "Python"],
        "tools": ["SAP", "Oracle SCM Cloud", "Excel avanzato", "Power BI", "Microsoft Project", "Tableau"],
    },
}


# --- SEZIONE 1 ---
ROLE_SKILL_LIBRARY.update({
    "cloud_engineer": {
        "hard_skills": ["Cloud infrastructure", "Infrastructure as Code", "Kubernetes", "Terraform", "Networking cloud", "Monitoring", "High availability", "Security hardening", "Cost optimization"],
        "soft_skills": ["Problem solving", "Pensiero sistemico", "Precisione", "Collaborazione interfunzionale", "Gestione delle priorita"],
        "programming_languages": ["Python", "Bash", "YAML"],
        "tools": ["AWS", "Azure", "Google Cloud", "Terraform", "Kubernetes", "Docker", "Helm", "Prometheus"],
    },
    "cloud_architect": {
        "hard_skills": ["Cloud architecture", "Landing zone design", "Identity and access management", "Networking ibrido", "Disaster recovery", "Governance cloud", "FinOps", "Solution design", "Migration strategy"],
        "soft_skills": ["Pensiero strategico", "Leadership tecnica", "Comunicazione con stakeholder", "Decision making", "Visione d'insieme"],
        "programming_languages": ["Python", "Bash", "YAML"],
        "tools": ["AWS", "Azure", "Google Cloud", "Terraform", "CloudFormation", "Kubernetes", "Well-Architected Framework", "Lucidchart"],
    },
    "devops_engineer": {
        "hard_skills": ["CI/CD", "Infrastructure as Code", "Containerization", "Release automation", "Configuration management", "Observability", "Version control workflow", "Scripting", "Infrastructure monitoring"],
        "soft_skills": ["Problem solving", "Precisione", "Collaborazione", "Gestione incident", "Pensiero sistematico"],
        "programming_languages": ["Bash", "Python", "YAML"],
        "tools": ["Jenkins", "GitHub Actions", "GitLab CI", "Docker", "Kubernetes", "Terraform", "Ansible", "Prometheus"],
    },
    "site_reliability_engineer": {
        "hard_skills": ["SRE practices", "Incident management", "Monitoring and alerting", "Capacity planning", "Service level objectives", "Postmortem analysis", "Reliability engineering", "Performance tuning"],
        "soft_skills": ["Calma sotto pressione", "Problem solving", "Pensiero analitico", "Collaborazione", "Responsabilita operativa"],
        "programming_languages": ["Python", "Go", "Bash"],
        "tools": ["Kubernetes", "Prometheus", "Grafana", "PagerDuty", "ELK Stack", "Jaeger", "Terraform"],
    },
    "platform_engineer": {
        "hard_skills": ["Platform engineering", "Developer experience", "Internal developer platforms", "CI/CD pipelines", "Kubernetes administration", "Service provisioning", "Automation", "Observability", "Security controls"],
        "soft_skills": ["Pensiero sistemico", "Collaborazione", "Problem solving", "Leadership tecnica", "Orientamento al servizio"],
        "programming_languages": ["Python", "Go", "Bash"],
        "tools": ["Kubernetes", "Argo CD", "Terraform", "Backstage", "Helm", "Docker", "GitHub Actions", "Prometheus"],
    },
    "kubernetes_specialist": {
        "hard_skills": ["Kubernetes administration", "Cluster hardening", "Helm charts", "Networking Kubernetes", "Pod security", "Workload orchestration", "Backup and recovery", "Autoscaling"],
        "soft_skills": ["Precisione", "Problem solving", "Pensiero sistemico", "Autonomia", "Collaborazione"],
        "programming_languages": ["YAML", "Bash", "Python"],
        "tools": ["Kubernetes", "Helm", "kubectl", "Rancher", "Argo CD", "Prometheus", "Grafana", "Istio"],
    },
    "devsecops_engineer": {
        "hard_skills": ["Secure CI/CD", "Threat modeling", "Container security", "Secrets management", "Vulnerability management", "Policy as code", "Security automation", "IAM", "Compliance by design"],
        "soft_skills": ["Rigore metodologico", "Problem solving", "Pensiero critico", "Collaborazione", "Gestione del rischio"],
        "programming_languages": ["Python", "Bash", "YAML"],
        "tools": ["Snyk", "Trivy", "HashiCorp Vault", "OWASP ZAP", "GitHub Actions", "Terraform", "Kubernetes", "SonarQube"],
    },
    "ml_engineer": {
        "hard_skills": ["Machine learning pipelines", "Model deployment", "Feature engineering", "Model monitoring", "Experiment tracking", "Data preprocessing", "Model optimization", "MLOps basics"],
        "soft_skills": ["Problem solving", "Pensiero analitico", "Collaborazione", "Precisione", "Curiosita tecnica"],
        "programming_languages": ["Python", "SQL"],
        "tools": ["scikit-learn", "PyTorch", "TensorFlow", "MLflow", "Docker", "Kubernetes", "Airflow", "Jupyter Notebook"],
    },
    "deep_learning_engineer": {
        "hard_skills": ["Neural networks", "Computer vision models", "Sequence modeling", "Model fine-tuning", "GPU training", "Hyperparameter tuning", "Transfer learning", "Model evaluation"],
        "soft_skills": ["Rigore metodologico", "Curiosita intellettuale", "Problem solving", "Precisione", "Resilienza"],
        "programming_languages": ["Python", "CUDA", "SQL"],
        "tools": ["PyTorch", "TensorFlow", "Keras", "Jupyter Notebook", "Weights & Biases", "MLflow", "NVIDIA CUDA", "Docker"],
    },
    "ai_researcher": {
        "hard_skills": ["Research design", "Experimental methodology", "Model evaluation", "Scientific writing", "Statistical analysis", "Benchmarking", "Ablation studies", "Literature review"],
        "soft_skills": ["Curiosita intellettuale", "Rigore metodologico", "Pensiero critico", "Autonomia", "Comunicazione scientifica"],
        "programming_languages": ["Python", "R", "SQL"],
        "tools": ["PyTorch", "Jupyter Notebook", "LaTeX", "Weights & Biases", "Hugging Face", "Git", "Zotero"],
    },
    "mlops_engineer": {
        "hard_skills": ["Model lifecycle management", "CI/CD for ML", "Model serving", "Data versioning", "Monitoring dei modelli", "Pipeline orchestration", "Feature store", "Automation"],
        "soft_skills": ["Problem solving", "Pensiero sistemico", "Collaborazione", "Precisione", "Gestione delle priorita"],
        "programming_languages": ["Python", "Bash", "SQL"],
        "tools": ["MLflow", "Kubeflow", "Airflow", "Docker", "Kubernetes", "Terraform", "DVC", "GitHub Actions"],
    },
    "ai_product_manager": {
        "hard_skills": ["Product strategy", "AI roadmap", "Requirements analysis", "User research", "Prioritization framework", "Go-to-market planning", "Experiment design", "Metrics definition"],
        "soft_skills": ["Leadership", "Comunicazione", "Gestione stakeholder", "Pensiero strategico", "Gestione priorita"],
        "programming_languages": [],
        "tools": ["Jira", "Confluence", "Figma", "Aha!", "Productboard", "Amplitude", "Miro", "Notion"],
    },
    "responsible_ai_specialist": {
        "hard_skills": ["AI governance", "Bias assessment", "Model explainability", "Risk assessment", "Privacy by design", "Policy compliance", "Human oversight", "Model documentation"],
        "soft_skills": ["Integrita professionale", "Pensiero critico", "Rigore metodologico", "Comunicazione", "Sensibilita etica"],
        "programming_languages": ["Python", "SQL"],
        "tools": ["Hugging Face", "SHAP", "LIME", "TensorFlow", "PyTorch", "MLflow", "OneTrust", "Jupyter Notebook"],
    },
    "computer_vision_engineer": {
        "hard_skills": ["Image classification", "Object detection", "Segmentation", "Feature extraction", "Model training", "Annotation pipelines", "Edge deployment", "Performance evaluation"],
        "soft_skills": ["Precisione", "Problem solving", "Curiosita tecnica", "Collaborazione", "Pensiero analitico"],
        "programming_languages": ["Python", "C++"],
        "tools": ["OpenCV", "PyTorch", "TensorFlow", "Detectron2", "MMDetection", "Label Studio", "Jupyter Notebook", "Docker"],
    },
    "image_processing_specialist": {
        "hard_skills": ["Image enhancement", "Filtering", "Segmentation", "Feature detection", "Noise reduction", "Color correction", "Image analysis", "Computer vision preprocessing"],
        "soft_skills": ["Precisione", "Problem solving", "Pensiero analitico", "Collaborazione", "Rigore"],
        "programming_languages": ["Python", "MATLAB", "C++"],
        "tools": ["OpenCV", "MATLAB", "ImageJ", "Python", "NumPy", "SciPy", "Jupyter Notebook"],
    },
    "robotics_perception_engineer": {
        "hard_skills": ["Sensor fusion", "SLAM", "Object tracking", "3D perception", "Point cloud processing", "Localization", "Kalman filtering", "Perception stack integration"],
        "soft_skills": ["Problem solving", "Pensiero sistemico", "Precisione", "Collaborazione multidisciplinare", "Adattabilita"],
        "programming_languages": ["Python", "C++"],
        "tools": ["ROS", "OpenCV", "PCL", "Gazebo", "RViz", "TensorRT", "Docker"],
    },
    "autonomous_systems_engineer": {
        "hard_skills": ["Autonomous navigation", "Path planning", "Control systems", "Sensor integration", "Perception algorithms", "Simulation testing", "Real-time constraints", "Functional safety"],
        "soft_skills": ["Pensiero sistemico", "Problem solving", "Precisione", "Responsabilita", "Collaborazione"],
        "programming_languages": ["Python", "C++", "MATLAB"],
        "tools": ["ROS", "Gazebo", "CARLA", "MATLAB Simulink", "Docker", "NVIDIA Isaac", "Git"],
    },
    "nlp_engineer": {
        "hard_skills": ["Text preprocessing", "Information extraction", "Named entity recognition", "Text classification", "Model fine-tuning", "Evaluation metrics", "Prompt engineering", "Linguistic annotation"],
        "soft_skills": ["Curiosita intellettuale", "Problem solving", "Precisione", "Collaborazione", "Pensiero analitico"],
        "programming_languages": ["Python", "SQL"],
        "tools": ["spaCy", "Hugging Face", "PyTorch", "TensorFlow", "NLTK", "Jupyter Notebook", "MLflow"],
    },
    "llm_engineer": {
        "hard_skills": ["Prompt engineering", "RAG pipelines", "LLM fine-tuning", "Vector search", "Evaluation of LLM outputs", "Context window optimization", "Guardrails", "Tool calling"],
        "soft_skills": ["Problem solving", "Pensiero sistemico", "Curiosita tecnica", "Precisione", "Collaborazione"],
        "programming_languages": ["Python", "SQL"],
        "tools": ["LangChain", "LlamaIndex", "Hugging Face", "OpenAI API", "FAISS", "Pinecone", "Docker", "MLflow"],
    },
    "conversational_ai_engineer": {
        "hard_skills": ["Dialog management", "Intent classification", "Conversation design", "RAG integration", "Multimodal assistants", "NLU", "Conversation testing", "Fallback handling"],
        "soft_skills": ["Empatia", "Comunicazione", "Problem solving", "Collaborazione", "Orientamento all utente"],
        "programming_languages": ["Python", "JavaScript"],
        "tools": ["Dialogflow", "Rasa", "Botpress", "OpenAI API", "LangChain", "Figma", "Jira"],
    },
    "speech_recognition_engineer": {
        "hard_skills": ["Acoustic modeling", "Speech-to-text", "Language modeling", "Audio preprocessing", "Forced alignment", "WER evaluation", "Streaming inference", "Noise robustness"],
        "soft_skills": ["Precisione", "Curiosita tecnica", "Problem solving", "Pensiero analitico", "Persistenza"],
        "programming_languages": ["Python", "C++"],
        "tools": ["PyTorch", "Kaldi", "Whisper", "Hugging Face", "Librosa", "Jupyter Notebook", "Docker"],
    },
    "data_engineer": {
        "hard_skills": ["Data pipeline design", "ETL/ELT", "Data modeling", "Data warehousing", "Orchestration", "Data quality", "Batch processing", "Streaming ingestion"],
        "soft_skills": ["Pensiero sistemico", "Precisione", "Problem solving", "Collaborazione", "Gestione delle priorita"],
        "programming_languages": ["Python", "SQL", "Scala"],
        "tools": ["Airflow", "Spark", "dbt", "Snowflake", "BigQuery", "Kafka", "Databricks", "Docker"],
    },
    "big_data_engineer": {
        "hard_skills": ["Distributed computing", "Spark tuning", "Data lake architecture", "Stream processing", "Large-scale ETL", "Partitioning strategy", "Performance optimization", "Data governance"],
        "soft_skills": ["Pensiero sistemico", "Problem solving", "Precisione", "Collaborazione", "Capacita decisionale"],
        "programming_languages": ["Python", "SQL", "Scala"],
        "tools": ["Apache Spark", "Hadoop", "Kafka", "Databricks", "Hive", "Airflow", "S3", "Delta Lake"],
    },
    "streaming_data_engineer": {
        "hard_skills": ["Event-driven architecture", "Real-time processing", "Kafka topics", "Stream transformation", "Low-latency pipelines", "Schema evolution", "Data observability", "Fault tolerance"],
        "soft_skills": ["Problem solving", "Precisione", "Collaborazione", "Pensiero sistemico", "Gestione del tempo"],
        "programming_languages": ["Python", "Scala", "SQL"],
        "tools": ["Kafka", "Flink", "Spark Structured Streaming", "Kinesis", "Databricks", "Airflow", "Docker"],
    },
    "analytics_engineer": {
        "hard_skills": ["Semantic layer modeling", "dbt modeling", "Data transformation", "Metric definition", "BI layer optimization", "SQL optimization", "Data documentation", "Analytics engineering"],
        "soft_skills": ["Precisione", "Pensiero analitico", "Collaborazione", "Problem solving", "Orientamento al dato"],
        "programming_languages": ["SQL", "Python"],
        "tools": ["dbt", "BigQuery", "Snowflake", "Looker", "Power BI", "Tableau", "Git", "Airflow"],
    },
    "cybersecurity_analyst": {
        "hard_skills": ["Threat analysis", "Security monitoring", "Incident response", "Vulnerability assessment", "SIEM analysis", "Log review", "Malware triage", "Access control"],
        "soft_skills": ["Pensiero critico", "Precisione", "Calma sotto pressione", "Collaborazione", "Integrita professionale"],
        "programming_languages": ["Python", "Bash", "SQL"],
        "tools": ["Splunk", "Microsoft Sentinel", "QRadar", "Wireshark", "Nmap", "Burp Suite", "CrowdStrike"],
    },
    "penetration_tester": {
        "hard_skills": ["Penetration testing", "Web application testing", "Network exploitation", "Privilege escalation", "Reporting tecnico", "Social engineering assessment", "Exploit validation", "Remediation guidance"],
        "soft_skills": ["Curiosita tecnica", "Precisione", "Pensiero critico", "Riservatezza", "Problem solving"],
        "programming_languages": ["Python", "Bash", "JavaScript"],
        "tools": ["Kali Linux", "Burp Suite", "Metasploit", "Nmap", "OWASP ZAP", "Wireshark", "Nessus"],
    },
    "security_architect": {
        "hard_skills": ["Security architecture", "Zero Trust", "IAM design", "Network segmentation", "Threat modeling", "Security controls", "Cloud security", "Compliance mapping"],
        "soft_skills": ["Pensiero strategico", "Leadership tecnica", "Comunicazione con stakeholder", "Gestione del rischio", "Decision making"],
        "programming_languages": ["Python", "Bash", "YAML"],
        "tools": ["Microsoft Defender", "Palo Alto", "Okta", "Terraform", "Kubernetes", "OneTrust", "Splunk"],
    },
    "soc_analyst": {
        "hard_skills": ["Security event monitoring", "Alert triage", "Incident escalation", "Threat intelligence", "Log correlation", "Use case tuning", "IOC analysis", "Forensic basics"],
        "soft_skills": ["Calma sotto pressione", "Precisione", "Problem solving", "Collaborazione", "Prontezza operativa"],
        "programming_languages": ["Python", "Bash"],
        "tools": ["Splunk", "Microsoft Sentinel", "QRadar", "TheHive", "Wireshark", "VirusTotal", "MISP"],
    },
    "embedded_systems_engineer": {
        "hard_skills": ["Embedded C", "Real-time programming", "Microcontroller programming", "Peripheral integration", "RTOS", "Debugging su hardware", "Protocol integration", "Low-power design"],
        "soft_skills": ["Precisione", "Problem solving", "Pensiero sistemico", "Pazienza", "Collaborazione"],
        "programming_languages": ["C", "C++", "Assembly"],
        "tools": ["STM32CubeIDE", "Keil", "IAR Embedded Workbench", "Oscilloscope", "Logic analyzer", "FreeRTOS", "Git"],
    },
    "iot_engineer": {
        "hard_skills": ["IoT architecture", "Sensor integration", "Edge computing", "MQTT", "Device provisioning", "Connectivity protocols", "Remote monitoring", "Data ingestion"],
        "soft_skills": ["Pensiero sistemico", "Problem solving", "Precisione", "Collaborazione", "Gestione delle priorita"],
        "programming_languages": ["C", "C++", "Python"],
        "tools": ["AWS IoT", "Azure IoT Hub", "MQTT", "Node-RED", "Raspberry Pi", "Arduino IDE", "Docker"],
    },
    "firmware_engineer": {
        "hard_skills": ["Firmware development", "Bare-metal programming", "Hardware debugging", "Device drivers", "Communication protocols", "Memory optimization", "Interrupt handling", "Bootloader integration"],
        "soft_skills": ["Precisione", "Problem solving", "Persistenza", "Pensiero analitico", "Collaborazione"],
        "programming_languages": ["C", "C++", "Assembly"],
        "tools": ["STM32CubeIDE", "Keil", "JTAG debugger", "Oscilloscope", "IAR Embedded Workbench", "Git", "FreeRTOS"],
    },
    "real_time_systems_engineer": {
        "hard_skills": ["Real-time systems", "Deterministic scheduling", "RTOS", "Latency analysis", "Concurrency control", "Safety-critical design", "Performance profiling", "Deadline management"],
        "soft_skills": ["Precisione", "Pensiero sistemico", "Problem solving", "Responsabilita", "Collaborazione"],
        "programming_languages": ["C", "C++", "Ada"],
        "tools": ["FreeRTOS", "QNX", "VxWorks", "Tracealyzer", "MATLAB Simulink", "Git", "Jenkins"],
    },
    "network_engineer": {
        "hard_skills": ["Routing and switching", "Network troubleshooting", "LAN/WAN design", "Firewall configuration", "VPN setup", "Network monitoring", "Subnetting", "TCP/IP"],
        "soft_skills": ["Problem solving", "Precisione", "Pensiero analitico", "Collaborazione", "Calma sotto pressione"],
        "programming_languages": ["Python", "Bash"],
        "tools": ["Cisco IOS", "Juniper", "Wireshark", "SolarWinds", "NetBox", "pfSense", "Nmap"],
    },
    "network_architect": {
        "hard_skills": ["Network architecture", "WAN optimization", "SD-WAN", "Network segmentation", "High availability", "Capacity planning", "Security networking", "Hybrid cloud networking"],
        "soft_skills": ["Pensiero strategico", "Leadership tecnica", "Comunicazione con stakeholder", "Problem solving", "Decision making"],
        "programming_languages": ["Python", "Bash"],
        "tools": ["Cisco DNA Center", "Juniper", "SolarWinds", "NetBox", "Aruba Central", "Palo Alto", "Visio"],
    },
    "wireless_engineer": {
        "hard_skills": ["Wireless planning", "RF engineering", "Site survey", "Wi-Fi optimization", "Spectrum analysis", "Antenna design", "Coverage analysis", "Interference mitigation"],
        "soft_skills": ["Precisione", "Problem solving", "Pensiero analitico", "Collaborazione", "Capacita di diagnosi"],
        "programming_languages": ["Python", "MATLAB"],
        "tools": ["Ekahau", "AirMagnet", "Wireshark", "Cisco Wireless", "NetSpot", "Spectrum analyzer", "Visio"],
    },
    "operations_manager": {
        "hard_skills": ["Operations planning", "Process management", "KPI management", "Budget monitoring", "Resource allocation", "Continuous improvement", "Capacity planning", "Workflow optimization"],
        "soft_skills": ["Leadership", "Comunicazione", "Gestione priorita", "Decision making", "Negoziazione"],
        "programming_languages": [],
        "tools": ["Excel avanzato", "Power BI", "SAP", "Asana", "Monday.com", "Notion", "Trello"],
    },
    "lean_manager": {
        "hard_skills": ["Lean manufacturing", "Kaizen", "Value stream mapping", "Waste reduction", "Root cause analysis", "Standard work", "Continuous improvement", "A3 problem solving"],
        "soft_skills": ["Pensiero sistemico", "Leadership", "Coinvolgimento team", "Problem solving", "Orientamento al miglioramento"],
        "programming_languages": [],
        "tools": ["Minitab", "Visio", "Excel avanzato", "Power BI", "5S boards", "A3 templates", "SAP"],
    },
    "quality_manager": {
        "hard_skills": ["Quality management system", "ISO 9001", "Nonconformity management", "Audit di qualita", "CAPA", "Root cause analysis", "Process validation", "Supplier quality"],
        "soft_skills": ["Precisione", "Leadership", "Comunicazione", "Problem solving", "Gestione delle priorita"],
        "programming_languages": [],
        "tools": ["SAP QM", "Minitab", "QMS software", "Excel avanzato", "Power BI", "Jira", "Confluence"],
    },
    "process_engineer": {
        "hard_skills": ["Process design", "Process optimization", "Root cause analysis", "FMEA", "SPC", "Process simulation", "Yield improvement", "Control plans"],
        "soft_skills": ["Pensiero sistemico", "Precisione", "Problem solving", "Collaborazione", "Orientamento al miglioramento"],
        "programming_languages": ["Python", "MATLAB"],
        "tools": ["Aspen HYSYS", "Minitab", "Excel avanzato", "Power BI", "Visio", "AutoCAD", "SAP"],
    },
    "industrial engineer": {
        "hard_skills": ["Industrial engineering", "Time and motion study", "Production planning", "Capacity analysis", "Lean methods", "Process optimization", "Layout design", "Cost reduction"],
        "soft_skills": ["Pensiero sistemico", "Leadership operativa", "Problem solving", "Comunicazione interfunzionale", "Gestione priorita"],
        "programming_languages": ["Python", "SQL", "MATLAB"],
        "tools": ["Minitab", "SAP", "Power BI", "Excel avanzato", "Arena Simulation", "Visio", "AutoCAD"],
    },
    "supply_chain_analyst": {
        "hard_skills": ["Demand forecasting", "Inventory analysis", "Supply planning", "Logistics analytics", "Supplier performance", "S&OP support", "Data analysis", "Cost analysis"],
        "soft_skills": ["Pensiero analitico", "Comunicazione", "Collaborazione", "Gestione delle priorita", "Problem solving"],
        "programming_languages": ["SQL", "Python"],
        "tools": ["SAP", "Excel avanzato", "Power BI", "Tableau", "Oracle SCM Cloud", "Anaplan", "SQL Server"],
    },
    "logistics_manager": {
        "hard_skills": ["Logistics management", "Transport planning", "Warehouse operations", "Distribution planning", "Route optimization", "Customs coordination", "Inventory control", "Service level management"],
        "soft_skills": ["Leadership", "Negoziazione", "Gestione delle priorita", "Comunicazione", "Decision making"],
        "programming_languages": [],
        "tools": ["SAP TM", "WMS", "TMS", "Excel avanzato", "Power BI", "Oracle SCM Cloud", "Trello"],
    },
    "procurement_manager": {
        "hard_skills": ["Strategic sourcing", "Supplier negotiation", "Procurement planning", "Contract management", "Spend analysis", "Vendor management", "Category management", "Procurement compliance"],
        "soft_skills": ["Negoziazione", "Leadership", "Comunicazione", "Gestione delle priorita", "Orientamento al risultato"],
        "programming_languages": [],
        "tools": ["SAP Ariba", "Coupa", "Excel avanzato", "Power BI", "Oracle Procurement", "Jaggaer", "Concur"],
    },
    "erp_specialist": {
        "hard_skills": ["ERP implementation", "Business process mapping", "Module configuration", "User support", "Data migration", "Integration testing", "Change management", "Process optimization"],
        "soft_skills": ["Problem solving", "Comunicazione", "Supporto agli utenti", "Precisione", "Gestione del cambiamento"],
        "programming_languages": ["SQL", "ABAP"],
        "tools": ["SAP", "Oracle ERP Cloud", "Microsoft Dynamics 365", "NetSuite", "Jira", "Confluence", "Excel avanzato"],
    },
    "business_analyst": {
        "hard_skills": ["Requirements gathering", "Process analysis", "Gap analysis", "Business case development", "UML", "Data analysis", "User story writing", "Workshop facilitation"],
        "soft_skills": ["Comunicazione", "Pensiero analitico", "Gestione stakeholder", "Problem solving", "Facilitazione"],
        "programming_languages": ["SQL", "Python"],
        "tools": ["Jira", "Confluence", "Miro", "Excel avanzato", "Power BI", "Lucidchart", "Visio"],
    },
    "business_intelligence_analyst": {
        "hard_skills": ["Dashboard design", "Data modeling", "KPI definition", "Reporting", "SQL analysis", "ETL support", "Data storytelling", "Performance monitoring"],
        "soft_skills": ["Pensiero analitico", "Precisione", "Comunicazione", "Orientamento al dato", "Collaborazione"],
        "programming_languages": ["SQL", "Python"],
        "tools": ["Power BI", "Tableau", "Looker", "BigQuery", "Excel avanzato", "dbt", "Snowflake"],
    },
    "management_consultant": {
        "hard_skills": ["Business analysis", "Strategy development", "Operating model design", "Market analysis", "Financial analysis", "Project delivery", "Stakeholder mapping", "Process improvement"],
        "soft_skills": ["Comunicazione", "Leadership", "Pensiero strategico", "Gestione delle priorita", "Problem solving"],
        "programming_languages": [],
        "tools": ["Excel avanzato", "PowerPoint", "Power BI", "Miro", "Confluence", "Jira", "Tableau"],
    },
    "it_project_manager": {
        "hard_skills": ["IT project planning", "Scope management", "Risk management", "Agile delivery", "Budget tracking", "Stakeholder coordination", "Release planning", "Vendor coordination"],
        "soft_skills": ["Leadership", "Comunicazione", "Gestione priorita", "Negoziazione", "Decision making"],
        "programming_languages": [],
        "tools": ["Jira", "Confluence", "Microsoft Project", "Trello", "Asana", "Smartsheet", "Excel avanzato"],
    },
    "scrum_master": {
        "hard_skills": ["Scrum facilitation", "Sprint planning", "Backlog refinement", "Retrospectives", "Impediment removal", "Agile coaching", "Velocity tracking", "Team metrics"],
        "soft_skills": ["Facilitazione", "Leadership servente", "Comunicazione", "Gestione conflitti", "Coaching"],
        "programming_languages": [],
        "tools": ["Jira", "Confluence", "Miro", "Azure DevOps", "Trello", "Notion", "Slack"],
    },
    "product_owner": {
        "hard_skills": ["Product backlog management", "User story prioritization", "Roadmap planning", "Acceptance criteria", "Stakeholder alignment", "Release planning", "Value definition", "Agile requirements"],
        "soft_skills": ["Comunicazione", "Leadership", "Gestione priorita", "Negoziazione", "Pensiero strategico"],
        "programming_languages": [],
        "tools": ["Jira", "Confluence", "Productboard", "Aha!", "Miro", "Figma", "Notion"],
    },
    "change_manager": {
        "hard_skills": ["Change impact analysis", "Change planning", "Stakeholder engagement", "Adoption tracking", "Training coordination", "Communication planning", "Resistance management", "Benefits realization"],
        "soft_skills": ["Gestione del cambiamento", "Comunicazione", "Empatia", "Leadership", "Negoziazione"],
        "programming_languages": [],
        "tools": ["Prosci", "Miro", "Confluence", "PowerPoint", "Excel avanzato", "Jira", "SharePoint"],
    },
    "cost_controller": {
        "hard_skills": ["Cost control", "Variance analysis", "Budget monitoring", "Forecasting", "Cost reporting", "Cost allocation", "Financial planning", "KPI analysis"],
        "soft_skills": ["Precisione", "Pensiero analitico", "Gestione priorita", "Comunicazione", "Orientamento al risultato"],
        "programming_languages": ["SQL", "VBA"],
        "tools": ["SAP CO", "Excel avanzato", "Power BI", "Oracle Hyperion", "Tagetik", "Qlik Sense", "Tableau"],
    },
    "facility_manager": {
        "hard_skills": ["Facility operations", "Space planning", "Vendor management", "Maintenance planning", "Health and safety", "Energy monitoring", "Budget control", "Space utilization"],
        "soft_skills": ["Organizzazione", "Comunicazione", "Gestione priorita", "Problem solving", "Negoziazione"],
        "programming_languages": [],
        "tools": ["IWMS", "Excel avanzato", "SAP", "ServiceNow", "Planon", "AutoCAD", "Power BI"],
    },
    "maintenance_manager": {
        "hard_skills": ["Preventive maintenance", "Corrective maintenance", "Asset management", "Maintenance planning", "Reliability improvement", "Spare parts management", "CMMS", "Safety compliance"],
        "soft_skills": ["Leadership", "Problem solving", "Gestione priorita", "Comunicazione", "Decision making"],
        "programming_languages": [],
        "tools": ["SAP PM", "IBM Maximo", "CMMS", "Excel avanzato", "Power BI", "ServiceNow", "Minitab"],
    },
    "energy_manager": {
        "hard_skills": ["Energy management", "Energy auditing", "Consumption analysis", "ISO 50001", "Efficiency projects", "Utility monitoring", "Carbon reporting", "Energy saving planning"],
        "soft_skills": ["Pensiero analitico", "Orientamento alla sostenibilita", "Comunicazione", "Problem solving", "Gestione delle priorita"],
        "programming_languages": ["Python", "SQL"],
        "tools": ["SCADA", "Power BI", "Excel avanzato", "EnergyCAP", "ISO 50001 software", "SAP", "Tableau"],
    },
})

# --- SEZIONE 2 ---
ATS_HARD_SKILL_TERMS.extend([
    "kubernetes", "terraform", "ci/cd", "iac", "infrastructure as code", "docker",
    "helm", "prometheus", "grafana", "observability", "cloud architecture",
    "landing zone", "finops", "sre", "incident management", "service level objectives",
    "platform engineering", "devsecops", "threat modeling", "vulnerability management",
    "machine learning pipelines", "model monitoring", "experiment tracking", "mlops",
    "rag", "vector search", "prompt engineering", "llm fine-tuning", "computer vision",
    "object detection", "segmentation", "sensor fusion", "slam", "speech recognition",
    "data engineering", "etl/elt", "streaming ingestion", "kafka", "spark", "dbt",
    "siem", "incident response", "pen testing", "iso 9001", "lean manufacturing",
    "six sigma", "erp implementation", "s&op", "budget monitoring", "energy auditing",
])

# --- SEZIONE 3 ---
ATS_SOFT_SKILL_TERMS.extend([
    "leadership tecnica", "gestione del cambiamento", "pensiero sistemico", "rigore metodologico",
    "curiosita intellettuale", "comunicazione con stakeholder", "gestione incident", "calma sotto pressione",
    "responsabilita operativa", "orientamento al servizio", "adattabilita", "resilienza",
    "visione d'insieme", "decision making", "facilitazione", "coaching", "leadership servente",
    "gestione del rischio", "orientamento all utente", "empatia professionale", "pensiero strategico",
    "orientamento al miglioramento", "coinvolgimento team", "responsabilita professionale",
    "supporto agli utenti", "negoziazione", "comunicazione persuasiva",
])

# --- SEZIONE 4 ---
ROLE_KEYWORD_GROUPS.update({
    "cloud engineer": [
        ("Kubernetes", ["kubernetes", "k8s"]),
        ("Terraform", ["terraform", "infrastructure as code", "iac"]),
        ("Docker", ["docker", "containerization", "containerizzazione"]),
        ("AWS", ["aws", "amazon web services"]),
        ("Azure", ["azure", "microsoft azure"]),
        ("Google Cloud", ["google cloud", "gcp"]),
        ("Prometheus", ["prometheus"]),
        ("Grafana", ["grafana"]),
        ("Observability", ["observability", "monitoring", "osservabilita"]),
    ],
    "platform engineer": [
        ("Platform engineering", ["platform engineering"]),
        ("Internal developer platform", ["internal developer platform", "idp"]),
        ("Kubernetes", ["kubernetes", "k8s"]),
        ("Terraform", ["terraform", "iac"]),
        ("Argo CD", ["argo cd"]),
        ("Backstage", ["backstage"]),
        ("Developer experience", ["developer experience", "devex"]),
    ],
    "devops engineer": [
        ("CI/CD", ["ci/cd", "continuous integration", "continuous delivery"]),
        ("Jenkins", ["jenkins"]),
        ("GitHub Actions", ["github actions"]),
        ("GitLab CI", ["gitlab ci"]),
        ("Docker", ["docker"]),
        ("Kubernetes", ["kubernetes", "k8s"]),
        ("Terraform", ["terraform", "iac"]),
        ("Ansible", ["ansible"]),
    ],
    "site reliability engineer": [
        ("SRE", ["sre", "site reliability"]),
        ("Incident management", ["incident management", "incident response"]),
        ("Service level objectives", ["slo", "service level objectives"]),
        ("Prometheus", ["prometheus"]),
        ("Grafana", ["grafana"]),
        ("PagerDuty", ["pagerduty"]),
        ("Postmortem", ["postmortem", "blameless postmortem"]),
    ],
    "ml engineer": [
        ("Machine learning", ["machine learning", "ml"]),
        ("Feature engineering", ["feature engineering"]),
        ("Model deployment", ["model deployment", "serving"]),
        ("MLflow", ["mlflow"]),
        ("PyTorch", ["pytorch"]),
        ("TensorFlow", ["tensorflow"]),
        ("Kubernetes", ["kubernetes"]),
    ],
    "mlops engineer": [
        ("MLOps", ["mlops"]),
        ("MLflow", ["mlflow"]),
        ("Kubeflow", ["kubeflow"]),
        ("Airflow", ["airflow"]),
        ("DVC", ["dvc", "data versioning"]),
        ("Docker", ["docker"]),
        ("Kubernetes", ["kubernetes"]),
    ],
    "computer vision engineer": [
        ("OpenCV", ["opencv"]),
        ("Object detection", ["object detection"]),
        ("Segmentation", ["segmentation"]),
        ("PyTorch", ["pytorch"]),
        ("TensorFlow", ["tensorflow"]),
        ("Label Studio", ["label studio"]),
    ],
    "nlp engineer": [
        ("spaCy", ["spacy"]),
        ("Hugging Face", ["hugging face", "transformers"]),
        ("Text classification", ["text classification"]),
        ("NER", ["named entity recognition", "ner"]),
        ("Prompt engineering", ["prompt engineering"]),
        ("RAG", ["rag", "retrieval augmented generation"]),
    ],
    "data engineer": [
        ("Airflow", ["airflow"]),
        ("Spark", ["spark", "apache spark"]),
        ("dbt", ["dbt"]),
        ("Kafka", ["kafka"]),
        ("Snowflake", ["snowflake"]),
        ("BigQuery", ["bigquery"]),
    ],
    "cybersecurity analyst": [
        ("SIEM", ["siem"]),
        ("Splunk", ["splunk"]),
        ("Microsoft Sentinel", ["microsoft sentinel"]),
        ("Incident response", ["incident response"]),
        ("Vulnerability assessment", ["vulnerability assessment"]),
        ("Wireshark", ["wireshark"]),
    ],
    "embedded systems engineer": [
        ("Embedded C", ["embedded c", "c"]),
        ("RTOS", ["rtos"]),
        ("Microcontroller", ["microcontroller", "microcontrollore"]),
        ("FreeRTOS", ["freertos"]),
        ("JTAG", ["jtag"]),
        ("Oscilloscope", ["oscilloscope"]),
    ],
    "network engineer": [
        ("Routing", ["routing"]),
        ("Switching", ["switching"]),
        ("Firewall", ["firewall"]),
        ("VPN", ["vpn"]),
        ("Wireshark", ["wireshark"]),
        ("Subnetting", ["subnetting"]),
    ],
    "operations manager": [
        ("Operations management", ["operations management", "gestione operations"]),
        ("KPI", ["kpi"]),
        ("Capacity planning", ["capacity planning"]),
        ("Resource allocation", ["resource allocation"]),
        ("Process improvement", ["continuous improvement", "miglioramento continuo"]),
    ],
    "industrial engineer": [
        ("Lean manufacturing", ["lean manufacturing"]),
        ("Six Sigma", ["six sigma"]),
        ("Capacity planning", ["capacity planning"]),
        ("Process optimization", ["process optimization"]),
        ("Minitab", ["minitab"]),
        ("Value stream mapping", ["value stream mapping"]),
    ],
    "business intelligence analyst": [
        ("Power BI", ["power bi"]),
        ("Tableau", ["tableau"]),
        ("Looker", ["looker"]),
        ("Data modeling", ["data modeling"]),
        ("KPI", ["kpi"]),
        ("dbt", ["dbt"]),
    ],
})

for role_family in list(ROLE_SKILL_LIBRARY):
    spaced_role_family = role_family.replace("_", " ")
    if "_" in role_family and spaced_role_family in ROLE_KEYWORD_GROUPS:
        ROLE_KEYWORD_GROUPS[role_family] = ROLE_KEYWORD_GROUPS[spaced_role_family]
        ROLE_KEYWORD_GROUPS.pop(spaced_role_family, None)

# --- SEZIONE 5 ---
CLOUD_ENGINEER_KEYWORDS_TO_CONFIRM = ["Kubernetes", "Terraform", "Docker", "AWS", "Azure", "Google Cloud", "Prometheus", "Grafana"]
DEVOPS_ENGINEER_KEYWORDS_TO_CONFIRM = ["CI/CD", "Jenkins", "GitHub Actions", "GitLab CI", "Docker", "Kubernetes", "Terraform", "Ansible"]
SITE_RELIABILITY_ENGINEER_KEYWORDS_TO_CONFIRM = ["SRE", "Incident management", "Service level objectives", "Prometheus", "Grafana", "PagerDuty", "Postmortem"]
ML_ENGINEER_KEYWORDS_TO_CONFIRM = ["Machine learning", "Feature engineering", "Model deployment", "MLflow", "PyTorch", "TensorFlow", "Kubernetes"]
MLOPS_ENGINEER_KEYWORDS_TO_CONFIRM = ["MLOps", "MLflow", "Kubeflow", "Airflow", "DVC", "Docker", "Kubernetes"]
COMPUTER_VISION_ENGINEER_KEYWORDS_TO_CONFIRM = ["OpenCV", "Object detection", "Segmentation", "PyTorch", "TensorFlow", "Label Studio"]
NLP_ENGINEER_KEYWORDS_TO_CONFIRM = ["spaCy", "Hugging Face", "Text classification", "NER", "Prompt engineering", "RAG"]
DATA_ENGINEER_KEYWORDS_TO_CONFIRM = ["Airflow", "Spark", "dbt", "Kafka", "Snowflake", "BigQuery"]
CYBERSECURITY_ANALYST_KEYWORDS_TO_CONFIRM = ["SIEM", "Splunk", "Microsoft Sentinel", "Incident response", "Vulnerability assessment", "Wireshark"]
EMBEDDED_SYSTEMS_ENGINEER_KEYWORDS_TO_CONFIRM = ["Embedded C", "RTOS", "Microcontroller", "FreeRTOS", "JTAG", "Oscilloscope"]


def keyword_group_present(cv_plain: str, variants: List[str]) -> bool:
    padded = f" {cv_plain} "
    for variant in variants:
        normalized = normalize_plain_text(variant)
        if not normalized:
            continue
        if len(normalized) <= 3 and re.search(rf"(?<![a-z0-9]){re.escape(normalized)}(?![a-z0-9])", padded):
            return True
        if len(normalized) > 3 and normalized in padded:
            return True
    return False


def role_keyword_snapshot(cv_text: str, role: str, description: str = "", required_skills: str = "") -> Dict[str, List[str]]:
    cv_plain = normalize_plain_text(cv_text)
    target_plain = normalize_plain_text(f"{role} {description} {required_skills}")
    family = infer_role_family(role, description, required_skills)
    groups = ROLE_KEYWORD_GROUPS.get(family, [])
    present = [label for label, variants in groups if keyword_group_present(cv_plain, variants)]
    partially_present = []
    if family == "data analyst" and any(item in present for item in ["Machine Learning / AI", "NLP / LLM"]) and "Analisi dei dati" not in present:
        partially_present.append("Analisi dei dati")
    confirm_library = {
        "data analyst": DATA_ANALYST_KEYWORDS_TO_CONFIRM,
        "game design": GAME_DESIGN_KEYWORDS_TO_CONFIRM,
    }.get(family, [])
    to_confirm = [keyword for keyword in confirm_library if not keyword_group_present(cv_plain, [keyword])]
    return {
        "present": present,
        "partially_present": partially_present,
        "to_confirm": to_confirm,
    }


def is_cv_noise_keyword(value: str) -> bool:
    plain = normalize_plain_text(value)
    return not plain or plain in {
        "contatti",
        "lingue",
        "hard skills",
        "soft skills",
        "formazione",
        "esperienze",
        "esperienze professionali",
        "progetti",
        "certificazioni",
    } or len(plain) < 3


def infer_role_family(role: str, description: str = "", required_skills: str = "") -> str:
    target_plain = normalize_plain_text(f"{role} {description} {required_skills}")
    for family in ROLE_SKILL_LIBRARY:
        family_plain = normalize_plain_text(family)
        family_spaced_plain = normalize_plain_text(family.replace("_", " "))
        if target_plain in {family_plain, family_spaced_plain}:
            return family
    for family in ROLE_SKILL_LIBRARY:
        family_plain = normalize_plain_text(family)
        family_spaced_plain = normalize_plain_text(family.replace("_", " "))
        if family_plain and family_plain in target_plain:
            return family
        if family_spaced_plain and family_spaced_plain in target_plain:
            return family
    if any(term in target_plain for term in ["game design", "game designer", "level design", "unity", "unreal"]):
        return "game design"
    if any(term in target_plain for term in ["project manager", "project management", "gestione progetti", "pm"]):
        return "project manager"
    if any(term in target_plain for term in ["backend developer", "back end", "backend", "api", "fastapi", "django", "spring"]):
        return "backend developer"
    if any(term in target_plain for term in ["data scientist", "data science", "scientist"]):
        return "data scientist"
    if any(term in target_plain for term in ["data analyst", "analista dati", "analisi dati", "data analysis", "business intelligence"]):
        return "data analyst"
    if any(term in target_plain for term in ["software engineer", "software developer", "sviluppatore"]):
        return "software engineer"
    if any(term in target_plain for term in ["frontend", "ui", "ux", "designer"]):
        return "frontend developer"
    return ""


def infer_skill_library_from_role(role: str, description: str = "") -> Dict[str, List[str]]:
    """Genera una libreria di skill di fallback basata sul ruolo inserito."""
    role_plain = normalize_plain_text(role)
    description_plain = normalize_plain_text(description)
    
    if any(term in role_plain for term in ["project manager", "project management"]) or any(term in description_plain for term in ["pianificazione", "progetto", "scadenze"]):
        return {
            "hard_skills": ["Pianificazione attività", "Gestione scadenze", "Coordinamento team", "Monitoraggio avanzamento", "Gestione rischi", "Budget management"],
            "soft_skills": ["Comunicazione", "Organizzazione", "Problem solving", "Gestione priorità", "Leadership", "Negoziazione"],
            "programming_languages": [],
            "tools": ["Excel", "Trello", "Jira", "Notion", "Microsoft Project", "Asana"],
        }
    
    if any(term in role_plain for term in ["data scientist", "machine learning"]):
        return {
            "hard_skills": ["Python", "Machine Learning", "SQL", "Analisi predittiva", "Modelli statistici", "Data preprocessing", "Feature engineering"],
            "soft_skills": ["Problem solving", "Pensiero analitico", "Comunicazione", "Attenzione ai dettagli", "Collaborazione"],
            "programming_languages": ["Python", "SQL", "R"],
            "tools": ["pandas", "scikit-learn", "Jupyter", "TensorFlow", "Tableau", "Excel"],
        }
    
    if any(term in role_plain for term in ["data analyst", "analista", "analysis"]):
        return {
            "hard_skills": ["SQL", "Python", "Data visualization", "Analisi dati", "Reporting", "KPI", "Business intelligence"],
            "soft_skills": ["Attenzione ai dettagli", "Pensiero analitico", "Comunicazione", "Problem solving", "Organizzazione"],
            "programming_languages": ["SQL", "Python"],
            "tools": ["Excel avanzato", "Power BI", "Tableau", "Google Analytics", "Looker"],
        }
    
    if any(term in role_plain for term in ["software engineer", "developer", "sviluppatore", "programmatore"]):
        return {
            "hard_skills": ["Sviluppo software", "Design architetturale", "Debugging", "Version control", "Unit testing", "Code review"],
            "soft_skills": ["Problem solving", "Collaborazione", "Precisione", "Comunicazione tecnica", "Pensiero logico"],
            "programming_languages": ["Python", "JavaScript", "Java", "C++"],
            "tools": ["Git", "GitHub", "VS Code", "Docker", "AWS"],
        }
    
    if any(term in role_plain for term in ["backend", "api", "server"]):
        return {
            "hard_skills": ["API development", "Database design", "Autenticazione", "Testing backend", "Architetture REST", "Deploy", "Microservizi"],
            "soft_skills": ["Problem solving", "Precisione", "Collaborazione", "Documentazione", "Comunicazione tecnica"],
            "programming_languages": ["Python", "Java", "JavaScript", "Node.js"],
            "tools": ["FastAPI", "Django", "Spring", "PostgreSQL", "MongoDB", "Docker"],
        }
    
    if any(term in role_plain for term in ["frontend", "ui", "ux", "designer"]):
        return {
            "hard_skills": ["HTML/CSS", "JavaScript", "User Interface Design", "User Experience", "Responsive Design", "Accessibilita"],
            "soft_skills": ["Creativita", "Problem solving", "Comunicazione", "Attenzione ai dettagli", "Collaborazione"],
            "programming_languages": ["JavaScript", "TypeScript", "React"],
            "tools": ["Figma", "Adobe XD", "VS Code", "Git", "Webpack"],
        }
    
    return {}


def build_role_skill_suggestions(cv_text: str, role: str, description: str = "", required_skills: str = "") -> Dict[str, Any]:
    family = infer_role_family(role, description, required_skills)
    library = ROLE_SKILL_LIBRARY.get(family, {})
    cv_plain = normalize_plain_text(cv_text)
    role_plain = normalize_plain_text(f"{role} {description} {required_skills}")
    result = {
        "role_family": family,
        "hard_skills": [],
        "soft_skills": [],
        "programming_languages": [],
        "tools": [],
        "already_present": [],
        "to_highlight": [],
        "to_confirm": [],
        "confirmation_items": [],
    }
    
    # Log iniziale
    print(f"[build_role_skill_suggestions] family='{family}', role='{role}', library_found={bool(library)}")
    
    # Use only concrete libraries. Generic placeholder skills are never useful.
    if not library:
        library = infer_skill_library_from_role(role, description)
    if not library:
        print("[build_role_skill_suggestions] ruolo non riconosciuto: nessuna skill generica generata")
        return result

    def _score_skill_for_role(skill: str, group_name: str) -> int:
        skill_plain = normalize_plain_text(skill)
        score = 0
        if skill_plain and skill_plain in cv_plain:
            score += 4
        if skill_plain and skill_plain in role_plain:
            score += 3
        family_terms = {
            "software engineer": ["software", "api", "debug", "testing", "git", "docker", "ci/cd", "system design", "team", "collabor"],
            "backend developer": ["backend", "api", "database", "rest", "microserv", "docker", "redis", "testing", "collabor"],
            "frontend developer": ["frontend", "ui", "ux", "design", "react", "access", "responsive", "component", "state", "figma"],
            "data analyst": ["data", "sql", "excel", "power bi", "tableau", "report", "kpi", "analysis", "analit", "team"],
            "data scientist": ["data", "python", "machine", "ml", "model", "stat", "pandas", "scikit", "analysis", "team"],
            "project manager": ["project", "plan", "jira", "trello", "risk", "budget", "stakeholder", "team", "organ"],
        }.get(family, [])
        if any(term in skill_plain for term in family_terms):
            score += 2
        if group_name == "soft_skills":
            soft_role_terms = {
                "software engineer": ["problem solving", "collaborazione", "comunicazione tecnica", "precisione", "pensiero logico", "ownership"],
                "backend developer": ["problem solving", "precisione", "collaborazione", "documentazione tecnica", "comunicazione tecnica", "pensiero logico"],
                "frontend developer": ["creativita", "attenzione ai dettagli", "collaborazione", "comunicazione", "problem solving", "user empathy"],
                "data analyst": ["pensiero analitico", "attenzione ai dettagli", "comunicazione", "problem solving", "collaborazione"],
                "data scientist": ["pensiero analitico", "problem solving", "comunicazione scientifica", "collaborazione", "attenzione ai dettagli"],
                "project manager": ["organizzazione", "gestione priorita", "comunicazione", "coordinamento team", "leadership", "negoziazione"],
            }.get(family, [])
            if any(term in skill_plain for term in soft_role_terms):
                score += 3
        if group_name == "hard_skills":
            hard_role_terms = {
                "software engineer": ["software", "api", "debug", "testing", "version", "code review", "ci/cd", "system design"],
                "backend developer": ["api", "database", "rest", "backend", "microserv", "deploy", "scalabil", "cache"],
                "frontend developer": ["html", "css", "javascript", "typescript", "responsive", "access", "component", "state"],
                "data analyst": ["sql", "python", "excel", "power bi", "tableau", "report", "kpi", "dashboard"],
                "data scientist": ["python", "machine learning", "sql", "stat", "feature", "preprocessing", "visualization"],
                "project manager": ["pianificazione", "scadenze", "coordinamento", "risk", "budget", "stakeholder", "jira", "trello"],
            }.get(family, [])
            if any(term in skill_plain for term in hard_role_terms):
                score += 3
        return score

    def _ordered_unique_skills(skills: List[str], group_name: str) -> List[str]:
        scored = []
        seen = set()
        for skill in skills:
            normalized_skill = normalize_plain_text(skill)
            if not normalized_skill or normalized_skill in seen:
                continue
            seen.add(normalized_skill)
            scored.append((_score_skill_for_role(skill, group_name), skill))
        scored.sort(key=lambda item: (-item[0], normalize_plain_text(item[1])))
        return [skill for _score, skill in scored]
    
    group_limits = {
        "hard_skills": 6,
        "soft_skills": 5,
        "programming_languages": 3,
        "tools": 5,
    }
    seen_candidate_identities = set()
    for group_name in ["hard_skills", "programming_languages", "tools", "soft_skills"]:
        ordered_skills = _ordered_unique_skills(list(library.get(group_name, [])), group_name)
        ordered_skills = ordered_skills[:group_limits[group_name]]
        for skill in ordered_skills:
            identity = canonical_skill_identity(skill)
            if not identity or identity in seen_candidate_identities:
                continue
            seen_candidate_identities.add(identity)
            present = skill_semantically_present(cv_text, skill)
            item = {"name": skill, "status": "present" if present else "to_confirm"}
            result[group_name].append(item)
            category = {
                "soft_skills": "soft_skill",
                "programming_languages": "language",
                "tools": "tool",
            }.get(group_name, "hard_skill")
            if present:
                result["already_present"].append(skill)
                result["to_highlight"].append(skill)
                continue
            result["confirmation_items"].append({
                "id": f"{(family or normalize_plain_text(role)).replace(' ', '-')}-{group_name.replace('_', '-')}-{normalize_plain_text(skill).replace(' ', '-')}",
                "type": "skillConfirmation",
                "name": skill,
                "category": category,
                "reason": (
                    f"Competenza {group_name.replace('_', ' ')} pertinente al ruolo {role or family}; confermala solo se l'hai realmente utilizzata."
                    if not present
                    else f"Competenze già presenti nel CV e utili per il ruolo futuro {role or family}."
                ),
                "already_present": False,
                "requires_confirmation": True,
                "status": "pending",
                "user_example": "",
                "target_section": "SOFT SKILLS" if category == "soft_skill" else "HARD SKILLS",
            })
            result["to_confirm"].append(skill)

    deduped_confirmations = []
    seen_confirm_ids = set()
    for item in result["confirmation_items"]:
        item_id = str(item.get("id") or "").strip()
        if item_id and item_id in seen_confirm_ids:
            continue
        seen_confirm_ids.add(item_id)
        deduped_confirmations.append(item)
    result["confirmation_items"] = deduped_confirmations
    result["to_confirm"] = [
        skill for skill in result["to_confirm"]
        if not any(normalize_plain_text(skill) == normalize_plain_text(item.get("name") or "") for item in result["confirmation_items"])
    ]
    
    # Filtra falsi positivi generati da frasi dell'obiettivo/colloquio.
    # Esempi scartati: "voglio prepararmi", "per un colloquio", "design per google".
    generic_phrase_tokens = {
        "voglio", "prepararmi", "preparare", "colloquio", "intervista", "candidatura",
        "ruolo", "azienda", "presso", "per", "con", "una", "uno", "un", "di", "da",
    }
    role_plain_for_filter = normalize_plain_text(role)
    company_noise = {"google", "amazon", "microsoft", "apple", "meta", "poste", "italiane"}

    filtered_confirmation_items = []
    seen_confirmation_names = set()
    for item in result.get("confirmation_items", []):
        name = str(item.get("name") or "").strip()
        plain = normalize_plain_text(name)
        tokens = plain.split()
        if not plain:
            continue
        if item.get("type") != "skillConfirmation":
            continue
        if item.get("category") not in {"hard_skill", "soft_skill", "tool", "language"}:
            continue
        if skill_semantically_present(cv_text, name):
            continue
        if plain in seen_confirmation_names:
            continue
        if len(plain) <= 2:
            continue
        if plain in {"per un", "di game", "un colloquio", "voglio prepararmi", "prepararmi per"}:
            continue
        if any(noise in tokens for noise in ["voglio", "prepararmi", "colloquio", "intervista"]):
            continue
        if len(tokens) >= 3 and sum(1 for token in tokens if token in generic_phrase_tokens or token in company_noise) >= 2:
            continue
        if len(tokens) > 4 and not any(symbol in plain for symbol in ["+", "#", "."]) and plain not in role_plain_for_filter:
            continue
        seen_confirmation_names.add(plain)
        filtered_confirmation_items.append(item)

    technical_confirmations = [
        item for item in filtered_confirmation_items
        if item.get("category") != "soft_skill"
    ][:7]
    soft_confirmations = [
        item for item in filtered_confirmation_items
        if item.get("category") == "soft_skill"
    ][:5]
    result["confirmation_items"] = [*technical_confirmations, *soft_confirmations]
    result["to_confirm"] = [
        item.get("name")
        for item in result["confirmation_items"]
        if item.get("name") and not item.get("already_present")
    ]

    try:
        from services.cv_optimizer.structured_cv_engine import filter_confirmation_items, is_noise_keyword
        result["confirmation_items"] = filter_confirmation_items(result.get("confirmation_items", []))
        result["to_confirm"] = [
            item for item in result.get("to_confirm", [])
            if not is_noise_keyword(item)
        ]
    except Exception as exc:
        print(f"Filtro skill/keyword non disponibile: {exc}")
    print(f"[build_role_skill_suggestions] final_confirmation_items={len(result['confirmation_items'])}")
    return result


def filter_confirmed_skill_suggestions(
    cv_text: str,
    confirmed_skills: List[Dict[str, Any]],
    allowed_skill_suggestions: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    allowed_skills = {
        canonical_skill_identity(str(item.get("name") or "")): {
            "id": str(item.get("id") or "").strip(),
            "category": str(item.get("category") or "").strip(),
        }
        for item in allowed_skill_suggestions
        if isinstance(item, dict)
        and item.get("type") == "skillConfirmation"
        and item.get("category") in {"hard_skill", "soft_skill", "tool", "language", "keyword"}
    }
    filtered = []
    for item in confirmed_skills:
        if not isinstance(item, dict):
            continue
        name = str(item.get("name") or "").strip()
        identity = canonical_skill_identity(name)
        allowed = allowed_skills.get(identity)
        if (
            allowed
            and str(item.get("id") or "").strip() == allowed["id"]
            and (
                str(item.get("category") or "").strip() == allowed["category"]
                or {str(item.get("category") or "").strip(), allowed["category"]} == {"keyword", "hard_skill"}
            )
            and str(item.get("status") or "").strip().lower() in {"accepted", "confirmed"}
            and not skill_semantically_present(cv_text, name)
        ):
            filtered.append(item)
    return filtered


def _cv_content_metrics(cv_text: str) -> Dict[str, Any]:
    text = str(cv_text or "").strip()
    plain = normalize_plain_text(text)
    sections = extract_resume_sections(text)
    words = re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ0-9+#.%-]+", text)
    non_empty_lines = [line.strip() for line in text.splitlines() if line.strip()]
    bullet_lines = [
        line for line in non_empty_lines
        if re.match(r"^\s*(?:[-*•·▪◦]|\d+[.)])\s+", line)
    ]
    quantified_matches = re.findall(
        r"(?:\b\d+(?:[.,]\d+)?\s*(?:%|k|m|ore|giorni|mesi|anni|utenti|clienti|progetti|record|righe|dataset)\b|\b(?:aument|ridott|migliorat|ottimizzat)[a-z]*\b[^.\n]{0,45}\b\d+)",
        plain,
    )
    date_matches = re.findall(r"\b(?:19|20)\d{2}\b", plain)
    action_verbs = {
        "analizzato", "analisi", "sviluppato", "realizzato", "progettato",
        "gestito", "coordinato", "ottimizzato", "implementato", "creato",
        "monitorato", "automatizzato", "migliorato", "ridotto", "aumentato",
        "developed", "designed", "implemented", "managed", "improved",
    }
    action_count = sum(plain.count(verb) for verb in action_verbs)
    experience_text = sections.get("experience", "")
    project_text = sections.get("projects", "")
    profile_text = sections.get("profile", "")
    hard_skills_text = sections.get("hard_skills", "")
    soft_skills_text = sections.get("soft_skills", "")
    return {
        "text": text,
        "plain": plain,
        "sections": sections,
        "word_count": len(words),
        "line_count": len(non_empty_lines),
        "bullet_count": len(bullet_lines),
        "quantified_count": len(quantified_matches),
        "date_count": len(date_matches),
        "action_count": action_count,
        "has_contact": bool(
            re.search(r"[\w.+-]+@[\w.-]+\.\w+", text)
            or re.search(r"\+?\d[\d\s().-]{7,}", text)
        ),
        "has_profile": len(profile_text.split()) >= 8,
        "has_experience": len(experience_text.split()) >= 8,
        "has_education": len(sections.get("education", "").split()) >= 3,
        "has_hard_skills": len(hard_skills_text.split()) >= 1,
        "has_soft_skills": len(soft_skills_text.split()) >= 1,
        "has_languages": bool(sections.get("languages")) or any(
            term in plain for term in ["inglese", "english", "francese", "francais", "spagnolo"]
        ),
        "has_projects": len(project_text.split()) >= 5,
        "has_certifications": bool(sections.get("certifications")) or "certific" in plain,
        "experience_quality": min(
            1.0,
            (len(experience_text.split()) / 120)
            + (min(action_count, 5) * 0.08)
            + (min(len(quantified_matches), 3) * 0.10),
        ) if experience_text else 0.0,
        "profile_quality": min(1.0, len(profile_text.split()) / 55) if profile_text else 0.0,
    }


def _target_skill_candidates(role: str, description: str, required_skills: str) -> Dict[str, List[str]]:
    family = infer_role_family(role, description, required_skills)
    library = ROLE_SKILL_LIBRARY.get(family) or infer_skill_library_from_role(role, description)
    requested_hard, requested_soft = split_requested_skill_terms(role, description, required_skills)
    requested_keywords = extract_requested_keywords(role, description, required_skills)

    hard = [
        *(library.get("hard_skills", []) if library else []),
        *(library.get("programming_languages", []) if library else []),
        *(library.get("tools", []) if library else []),
        *requested_hard,
    ]
    soft = [
        *(library.get("soft_skills", []) if library else []),
        *requested_soft,
    ]
    generic = {
        "data", "analyst", "analysis", "business", "project", "manager",
        "engineer", "developer", "designer", "specialist", "team", "game",
        "design", "role", "company",
    }

    def unique(values: List[Any]) -> List[str]:
        result = []
        seen = set()
        for value in values:
            text = str(value or "").strip()
            identity = canonical_skill_identity(text)
            if not identity or identity in generic or identity in seen or len(identity) < 3:
                continue
            seen.add(identity)
            result.append(text)
        return result

    return {
        "hard": unique(hard)[:18],
        "soft": unique(soft)[:10],
        "keywords": unique(requested_keywords)[:20],
    }


def _semantic_coverage(cv_text: str, candidates: List[str]) -> tuple[List[str], List[str], float]:
    present = [skill for skill in candidates if skill_semantically_present(cv_text, skill)]
    missing = [skill for skill in candidates if not skill_semantically_present(cv_text, skill)]
    coverage = len(present) / len(candidates) if candidates else 0.0
    return present, missing, coverage


def compute_cv_completeness_score(cv_text: str, role: str = "", description: str = "", required_skills: str = "") -> int:
    metrics = _cv_content_metrics(cv_text)
    score = (
        (12 if metrics["has_contact"] else 0)
        + (10 * metrics["profile_quality"])
        + (22 * metrics["experience_quality"])
        + (12 if metrics["has_education"] else 0)
        + (12 if metrics["has_hard_skills"] else 0)
        + (6 if metrics["has_soft_skills"] else 0)
        + (6 if metrics["has_languages"] else 0)
        + (7 if metrics["has_projects"] else 0)
        + (4 if metrics["has_certifications"] else 0)
        + min(metrics["quantified_count"] * 3, 9)
    )
    length_factor = min(metrics["word_count"] / 450, 1)
    score *= 0.72 + (0.28 * length_factor)
    if metrics["word_count"] < 120:
        score = min(score, 48)
    if not metrics["has_experience"] and not metrics["has_projects"]:
        score = min(score, 58)
    return clamp_score(round(score))


def compute_role_match_score(cv_text: str, role: str, description: str = "", required_skills: str = "") -> int:
    if not str(role or "").strip():
        return 0
    metrics = _cv_content_metrics(cv_text)
    candidates = _target_skill_candidates(role, description, required_skills)
    hard_present, _hard_missing, hard_coverage = _semantic_coverage(cv_text, candidates["hard"])
    soft_present, _soft_missing, soft_coverage = _semantic_coverage(cv_text, candidates["soft"])
    target_tokens = tokenize_meaningful(f"{role} {description} {required_skills}")
    cv_tokens = tokenize_meaningful(cv_text)
    token_coverage = len(target_tokens & cv_tokens) / len(target_tokens) if target_tokens else 0
    role_family = infer_role_family(role, description, required_skills)
    profile_role_signal = (
        role_family and role_family in normalize_plain_text(metrics["sections"].get("profile", ""))
    )
    evidence_score = min(
        1.0,
        (metrics["experience_quality"] * 0.55)
        + (0.25 if metrics["has_projects"] else 0)
        + (min(metrics["quantified_count"], 3) * 0.07),
    )
    score = (
        (hard_coverage * 45)
        + (soft_coverage * 10)
        + (token_coverage * 20)
        + (evidence_score * 20)
        + (5 if profile_role_signal else 0)
    )
    if not hard_present:
        score = min(score, 45)
    if not metrics["has_experience"] and not metrics["has_projects"]:
        score = min(score, 52)
    return clamp_score(round(score))


def analyze_cv_ats(cv_text: str, role: str, description: str, required_skills: str = "") -> Dict:
    metrics = _cv_content_metrics(cv_text)
    heuristic = analyze_cv_heuristics(cv_text)
    candidates = _target_skill_candidates(role, description, required_skills)
    requested_keywords = candidates["keywords"]
    hard_present, missing_hard_skills, hard_coverage = _semantic_coverage(cv_text, candidates["hard"])
    soft_present, missing_soft_skills, soft_coverage = _semantic_coverage(cv_text, candidates["soft"])
    keyword_present, keyword_missing, keyword_coverage = _semantic_coverage(cv_text, requested_keywords)
    role_snapshot = role_keyword_snapshot(cv_text, role, description, required_skills)
    present_keywords = filter_cv_keyword_list([
        *keyword_present,
        *hard_present,
        *soft_present,
        *role_snapshot["present"],
    ])
    missing_keywords = filter_cv_keyword_list(keyword_missing)
    role_fragments = {
        token for token in tokenize_meaningful(role)
        if len(token) <= 3 or token in {"specialist", "manager", "engineer", "developer", "designer", "analyst", "scientist", "consultant", "researcher", "assistant"}
    }
    missing_keywords = [
        keyword for keyword in missing_keywords
        if normalize_plain_text(keyword) not in role_fragments
    ]

    required_sections = {
        "contatti": "Aggiungi o rendi piu visibili email, telefono o link professionali.",
        "formazione": "Mantieni una sezione formazione chiara e riconoscibile.",
        "esperienze professionali": "Rendi riconoscibile la sezione esperienze con ruolo, azienda e attività.",
        "competenze": "Inserisci una sezione competenze leggibile dagli ATS.",
    }
    detected_sections = set(heuristic.get("detected_sections", []))
    missing_sections = [
        {"section": section, "suggestion": suggestion}
        for section, suggestion in required_sections.items()
        if section not in detected_sections
    ]

    section_score = (len(required_sections) - len(missing_sections)) / len(required_sections)
    target_coverage_parts = [
        (hard_coverage, 0.60),
        (soft_coverage, 0.15),
    ]
    if requested_keywords:
        target_coverage_parts.append((keyword_coverage, 0.25))
    coverage_weight = sum(weight for _coverage, weight in target_coverage_parts)
    target_coverage = (
        sum(coverage * weight for coverage, weight in target_coverage_parts) / coverage_weight
        if coverage_weight else 0
    )
    keyword_score = clamp_score(round(target_coverage * 100))
    length_score = min(metrics["word_count"] / 450, 1)
    bullet_score = min(metrics["bullet_count"] / 8, 1)
    format_score = clamp_score(round(
        (section_score * 62)
        + (length_score * 23)
        + (bullet_score * 10)
        + (5 if metrics["line_count"] >= 8 else 0)
    ))
    ats_score = clamp_score(round(
        (target_coverage * 48)
        + (section_score * 32)
        + (length_score * 12)
        + (bullet_score * 8)
    ))
    if metrics["word_count"] < 120:
        ats_score = min(ats_score, 50)
        format_score = min(format_score, 58)
    missing_hard_skills = filter_cv_keyword_list(missing_hard_skills)
    missing_soft_skills = filter_cv_keyword_list(missing_soft_skills)

    issues = []
    if missing_keywords:
        issues.append("Alcune parole chiave rilevanti per ruolo e offerta non risultano abbastanza visibili.")
    if missing_sections:
        issues.append("Una o piu sezioni fondamentali del CV non sono riconoscibili in modo chiaro.")
    if len(cv_text.strip()) < 800:
        issues.append("Il testo estraibile dal CV e scarno: l'ottimizzazione richiede integrazioni reali dell'utente.")
    if not issues:
        issues.append("La base ATS e buona: conviene rifinire keyword, risultati e ordine delle sezioni.")

    suggestions = [
        "Usa titoli di sezione semplici e riconoscibili: Esperienze, Formazione, Competenze, Progetti, Certificazioni.",
        "Metti in alto 3-6 hard skills davvero rilevanti per il ruolo e raggruppale per tema: linguaggi, strumenti, metodi.",
        "Inserisci le soft skills solo se puoi sostenerle con esempi concreti in esperienza, progetto o studio.",
        "Evita testo in immagini, colonne troppo complesse e formattazioni difficili da leggere per ATS.",
    ]
    if missing_keywords:
        suggestions.insert(0, "Verifica se possiedi davvero le parole chiave mancanti prima di inserirle nel CV.")
    if missing_sections:
        suggestions.insert(0, "Completa le sezioni mancanti con informazioni vere e verificabili.")

    return {
        "ats_score": ats_score,
        "keyword_score": keyword_score,
        "format_score": format_score,
        "keyword_coverage": round(keyword_coverage, 2),
        "hard_skill_coverage": round(hard_coverage, 2),
        "soft_skill_coverage": round(soft_coverage, 2),
        "target_coverage": round(target_coverage, 2),
        "keywords_present": present_keywords[:12],
        "keywords_missing": missing_keywords[:12],
        "present_keywords": present_keywords[:12],
        "missing_keywords": missing_keywords[:12],
        "keywords_partially_present": role_snapshot["partially_present"][:8],
        "keywords_to_confirm": role_snapshot["to_confirm"][:8],
        "missing_hard_skills": missing_hard_skills[:10],
        "missing_soft_skills": missing_soft_skills[:10],
        "missing_sections": missing_sections,
        "sections_to_improve": missing_sections,
        "issues": issues,
        "suggestions": suggestions[:6],
    }


def build_deterministic_cv_scorecard(
    cv_text: str,
    company: str,
    role: str,
    description: str = "",
    required_skills: str = "",
) -> Dict[str, Any]:
    metrics = _cv_content_metrics(cv_text)
    ats_analysis = analyze_cv_ats(cv_text, role, description, required_skills)
    completeness = compute_cv_completeness_score(cv_text, role, description, required_skills)
    role_match = compute_role_match_score(cv_text, role, description, required_skills)

    average_line_length = (
        len(metrics["text"]) / metrics["line_count"]
        if metrics["line_count"] else 0
    )
    clarity = clamp_score(round(
        24
        + (min(metrics["bullet_count"], 8) * 4)
        + (min(metrics["action_count"], 6) * 3)
        + (12 if metrics["has_profile"] else 0)
        + (8 if 20 <= average_line_length <= 140 else 0)
        + (8 if metrics["word_count"] >= 180 else 0)
    ))
    professionalism = clamp_score(round(
        22
        + (12 if metrics["has_contact"] else 0)
        + (14 * metrics["experience_quality"])
        + (10 if metrics["has_education"] else 0)
        + (10 if metrics["has_hard_skills"] else 0)
        + min(metrics["quantified_count"] * 4, 12)
        + min(metrics["date_count"] * 2, 8)
    ))
    if metrics["word_count"] < 120:
        clarity = min(clarity, 48)
        professionalism = min(professionalism, 50)

    company_provided = bool(
        str(company or "").strip()
        and normalize_plain_text(company) not in {"generica", "azienda generica", "non specificata"}
    )
    if company_provided:
        description_tokens = tokenize_meaningful(description)
        cv_tokens = tokenize_meaningful(cv_text)
        description_coverage = (
            len(description_tokens & cv_tokens) / len(description_tokens)
            if description_tokens else 0
        )
        company_fit = clamp_score(round(
            (role_match * 0.55)
            + (description_coverage * 35)
            + (10 if normalize_plain_text(company) in metrics["plain"] else 0)
        ))
        if not description_tokens:
            company_fit = clamp_score(round(role_match * 0.70))
    else:
        company_fit = 0

    payload = {
        "role_match_score": role_match,
        "job_match_score": role_match,
        "company_fit_score": company_fit,
        "company_provided": company_provided,
        "completeness_score": completeness,
        "ats_score": ats_analysis["ats_score"],
        "keyword_score": ats_analysis["keyword_score"],
        "format_score": ats_analysis["format_score"],
        "clarity_score": clarity,
        "professionalism_score": professionalism,
    }
    payload["overall_score"] = compute_weighted_cv_job_score(payload)
    payload["ats_analysis"] = ats_analysis
    payload["score_explanation"] = build_cv_score_explanation(payload)
    payload["scoring_context"] = {
        "cv_fingerprint": cv_content_fingerprint(cv_text),
        "target_fingerprint": cv_analysis_target_fingerprint(
            role, company, description, required_skills
        ),
        "role": role,
        "company": company,
        "description": description,
        "required_skills": required_skills,
        "metrics": {
            key: value
            for key, value in metrics.items()
            if key not in {"text", "plain", "sections"}
        },
    }
    return payload


def compare_cv_scorecards(before: Dict[str, Any], after: Dict[str, Any]) -> Dict[str, Any]:
    score_fields = [
        "overall_score", "ats_score", "keyword_score", "format_score",
        "role_match_score", "company_fit_score", "completeness_score",
        "clarity_score", "professionalism_score",
    ]
    return {
        "before": {field: clamp_score(before.get(field, 0)) for field in score_fields},
        "after": {field: clamp_score(after.get(field, 0)) for field in score_fields},
        "delta": {
            field: clamp_score(after.get(field, 0)) - clamp_score(before.get(field, 0))
            for field in score_fields
        },
    }


def generate_cv_optimization_questions(cv_text: str, evaluation: Dict, ats_analysis: Dict) -> List[Dict[str, str]]:
    questions = []
    target = evaluation.get("target") if isinstance(evaluation.get("target"), dict) else {}
    role_family = infer_role_family(str(target.get("role") or ""), str(target.get("description") or ""))

    role_questions = {
        "data analyst": [
            ("data_projects", "Hai svolto progetti universitari o personali legati all'analisi dei dati?", "progetti"),
            ("data_tools", "Hai usato Python, SQL, Excel, Power BI, Tableau o strumenti simili?", "strumenti"),
            ("data_outputs", "Hai lavorato con dataset, report, dashboard o KPI?", "risultati"),
            ("data_internship", "Hai svolto stage, tirocini o esperienze coerenti con questo ruolo?", "esperienze"),
            ("data_metrics", "Puoi indicare risultati concreti, anche universitari, come accuratezza, tempi ridotti, modelli confrontati o quantita di dati analizzati?", "risultati"),
        ],
        "game design": [
            ("game_prototypes", "Hai creato prototipi di giochi, livelli, meccaniche o concept?", "progetti"),
            ("game_tools", "Hai usato Unity, Unreal Engine, C#, Blender o strumenti simili?", "strumenti"),
            ("game_projects", "Hai realizzato progetti universitari o personali legati al game design?", "progetti"),
            ("game_design_skills", "Hai esperienze in storytelling, level design, user experience o testing?", "competenze"),
        ],
        "backend developer": [
            ("backend_api", "Hai sviluppato API, database o servizi backend?", "progetti"),
            ("backend_stack", "Hai usato Python, Java, Node.js, Spring, FastAPI o Django?", "strumenti"),
            ("backend_data", "Hai lavorato con SQL, PostgreSQL, MongoDB o Docker?", "strumenti"),
            ("backend_projects", "Hai realizzato progetti universitari o personali backend?", "progetti"),
            ("backend_auth_cloud", "Hai esperienze con autenticazione, deploy o cloud?", "esperienze"),
        ],
    }
    for question_id, question, category in role_questions.get(role_family, []):
        questions.append({
            "id": question_id,
            "question": question,
            "reason": "Risposta facoltativa: verra usata solo se aggiungi informazioni reali.",
            "category": category,
        })

    if len(cv_text.strip()) < 800:
        questions.append({
            "id": "cv_too_sparse",
            "question": "Quali esperienze, attività, studi o progetti reali possiamo aggiungere per rendere il CV più completo?",
            "reason": "Il testo estraibile dal CV è scarno o poco dettagliato.",
            "category": "completezza",
        })

    for section in ats_analysis.get("missing_sections", [])[:3]:
        questions.append({
            "id": f"missing_section_{section.get('section', '').replace(' ', '_')}",
            "question": f"Quali informazioni reali puoi aggiungere per completare la sezione {section.get('section')}?",
            "reason": section.get("suggestion", "Sezione poco riconoscibile nel CV."),
            "category": "struttura",
        })

    missing_keywords = ats_analysis.get("keywords_missing", [])
    if missing_keywords:
        questions.append({
            "id": "missing_keywords",
            "question": f"Possiedi davvero alcune di queste competenze o parole chiave? Indica solo quelle reali con un esempio concreto: {', '.join(missing_keywords[:8])}.",
            "reason": "Keyword utili per ATS non abbastanza presenti nel CV.",
            "category": "ats",
        })

    if evaluation.get("missing_skills_for_role"):
        questions.append({
            "id": "missing_skills_for_role",
            "question": "Tra le competenze richieste dal ruolo, quali possiedi davvero e in quali esperienze o progetti le hai usate?",
            "reason": "Alcune competenze richieste dall'offerta non risultano evidenti.",
            "category": "competenze",
        })

    questions.append({
        "id": "measurable_results",
        "question": "Puoi indicare risultati misurabili, strumenti usati, volumi gestiti, tempi ridotti o obiettivi raggiunti?",
        "reason": "I risultati concreti aumentano chiarezza, impatto e credibilità del CV.",
        "category": "risultati",
    })
    questions.append({
        "id": "experience_details",
        "question": "Quali attività specifiche vuoi rendere più chiare nelle esperienze già presenti nel CV?",
        "reason": "Le descrizioni troppo generiche possono essere riscritte in modo più professionale.",
        "category": "ai_writer",
    })

    unique_questions = []
    seen = set()
    for question in questions:
        if question["id"] in seen:
            continue
        seen.add(question["id"])
        unique_questions.append(question)
    return unique_questions[:8]


def extract_text_from_file(file: UploadFile, file_bytes: bytes) -> str:
    text, _method = extract_text_from_file_bytes(file_bytes, file.filename or "")
    return clean_extracted_text(text)


def is_probably_cv(text: str) -> Dict:
    if not clean_extracted_text(text):
        return {"is_cv": False, "confidence": 0, "reason": "Il file e vuoto.", "detected_sections": []}

    heuristic = analyze_cv_heuristics(text)
    strong_sections = {"contatti", "formazione", "esperienze professionali", "competenze"}
    strong_count = len(strong_sections.intersection(set(heuristic["detected_sections"])))
    is_cv = heuristic["score"] >= 45 or (heuristic["score"] >= 35 and strong_count >= 2)
    return {
        "is_cv": is_cv,
        "confidence": clamp_score(heuristic["score"]),
        "reason": heuristic["reason"] if is_cv else "Il documento non contiene abbastanza elementi tipici di un CV.",
        "detected_sections": heuristic["detected_sections"],
    }


def extract_candidate_name_from_cv(text: str) -> Dict:
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    for line in lines[:12]:
        if "@" in line or re.search(r"\d", line):
            continue
        candidates = re.findall(r"\b[A-ZÀ-ÖØ-Ý][A-Za-zÀ-ÖØ-öø-ÿ'`-]{1,}\b", line)
        if 2 <= len(candidates) <= 4:
            return {"name": " ".join(candidates[:3]), "confidence": 0.78}

    normalized = re.sub(r"\s+", " ", text or "")
    match = re.search(r"(?:nome|name)\s*[:\-]\s*([A-ZÀ-ÖØ-Ý][A-Za-zÀ-ÖØ-öø-ÿ'`-]+(?:\s+[A-ZÀ-ÖØ-Ý][A-Za-zÀ-ÖØ-öø-ÿ'`-]+){1,3})", normalized)
    if match:
        return {"name": match.group(1).strip(), "confidence": 0.7}

    return {"name": "", "confidence": 0}


def normalize_name(name: str) -> str:
    return re.sub(r"[^a-z\s]", " ", strip_accents(name or "").lower()).strip()


def _name_fragment_in_text(text: str, name: str) -> bool:
    normalized_text = normalize_plain_text(text)
    normalized_name = normalize_name(name)
    if not normalized_text or not normalized_name:
        return False

    if normalized_name in normalized_text:
        return True

    compact_text = normalized_text.replace(" ", "")
    compact_name = normalized_name.replace(" ", "")
    return bool(compact_name) and compact_name in compact_text


def check_cv_identity(cv_text: str, user_first_name: str, user_last_name: str) -> Dict:
    expected = normalize_name(f"{user_first_name} {user_last_name}")
    expected_tokens = [token for token in expected.split() if token]
    reversed_expected = " ".join(reversed(expected_tokens))

    if len(expected_tokens) >= 2 and (
        _name_fragment_in_text(cv_text, expected)
        or _name_fragment_in_text(cv_text, reversed_expected)
    ):
        return {
            "matches_user": True,
            "confidence": 1.0,
            "detected_name": f"{user_first_name} {user_last_name}".strip(),
            "message": "Identità coerente: il nome presente nel CV corrisponde a quello dell'utente.",
        }

    detected = extract_candidate_name_from_cv(cv_text)
    detected_name = normalize_name(detected.get("name", ""))
    detected_tokens = [token for token in detected_name.split() if token]

    if len(expected_tokens) < 2:
        return {
            "matches_user": None,
            "confidence": 0,
            "detected_name": detected.get("name", ""),
            "message": "Non è stato possibile verificare la coerenza tra il nome indicato e il CV caricato, perché nome e cognome utente non sono disponibili.",
        }

    if not detected_tokens:
        return {
            "matches_user": None,
            "confidence": 0,
            "detected_name": "",
            "message": "Non sono riuscito a verificare con certezza il nome nel CV. Controlla che il documento sia corretto.",
        }

    token_matches = sum(
        1
        for expected_token in expected_tokens
        if any(
            expected_token == detected_token
            or SequenceMatcher(None, expected_token, detected_token).ratio() >= 0.86
            for detected_token in detected_tokens
        )
    )
    ordered_ratio = SequenceMatcher(None, expected, detected_name).ratio()
    reversed_ratio = SequenceMatcher(None, " ".join(reversed(expected_tokens)), detected_name).ratio()
    confidence = max(token_matches / max(len(expected_tokens), 1), ordered_ratio, reversed_ratio)

    if token_matches >= len(expected_tokens) or confidence >= 0.82:
        return {
            "matches_user": True,
            "confidence": round(confidence, 2),
            "detected_name": detected.get("name", ""),
            "message": "Identità coerente: il nome presente nel CV corrisponde a quello dell'utente.",
        }

    strong_detected_name = detected.get("confidence", 0) >= 0.70 and len(detected_tokens) >= 2
    very_low_match = confidence < 0.50 and token_matches == 0
    likely_other_person = strong_detected_name and very_low_match

    if likely_other_person:
        return {
            "matches_user": False,
            "confidence": round(confidence, 2),
            "detected_name": detected.get("name", ""),
            "message": (
                "Il nome rilevato nel CV sembra appartenere a un'altra persona. "
                "Controlla di aver caricato il file corretto."
            ),
        }

    if len(detected_tokens) >= 2 and confidence < 0.50:
        return {
            "matches_user": None,
            "confidence": round(confidence, 2),
            "detected_name": detected.get("name", ""),
            "message": (
                "Il nome nel CV non coincide in modo chiaro con quello del profilo. "
                "Verifica che il documento sia corretto prima di proseguire."
            ),
        }

    if detected.get("confidence", 0) >= 0.65 and confidence < 0.60:
        return {
            "matches_user": None,
            "confidence": round(confidence, 2),
            "detected_name": detected.get("name", ""),
            "message": (
                "Il nome rilevato nel CV e solo parzialmente coerente con il profilo. "
                "Controlla il file finale prima di usarlo."
            ),
        }

    return {
        "matches_user": None,
        "confidence": round(confidence, 2),
        "detected_name": detected.get("name", ""),
        "message": "Non sono riuscito a verificare con certezza il nome nel CV. Controlla che il documento sia corretto.",
    }


COACH_SUGGESTION_CATEGORY_LABELS = {
    "profile": "Profilo professionale",
    "experience": "Esperienze da riscrivere meglio",
    "phrases": "Frasi da migliorare",
    "skills": "Competenze da evidenziare",
    "soft_skills": "Soft skills",
    "education": "Formazione",
    "project": "Progetti",
    "extra_page": "Pagina aggiuntiva",
    "experiences": "Esperienze da riscrivere meglio",
    "missing_info": "Informazioni mancanti da confermare",
    "sections": "Sezioni poco chiare o poco efficaci",
    "company_alignment": "Allineamento azienda e ruolo",
}


def make_coach_suggestion(
    category: str,
    title: str,
    description: str,
    action: str = "",
    requires_confirmation: bool = False,
    section: str = "",
    original_text: str = "",
    proposed_text: str = "",
    supported_by_cv: bool = True,
    keywords_added: Optional[List[str]] = None,
    suggestion_type: Optional[str] = None,
) -> Dict:
    normalized_category = category if category in COACH_SUGGESTION_CATEGORY_LABELS else "phrases"
    computed_type = suggestion_type or ("actionableEdit" if section and original_text and proposed_text else "adviceOnly")
    suggestion_id = re.sub(
        r"[^a-z0-9]+",
        "-",
        normalize_plain_text(f"{normalized_category}-{title}-{description}")[:90],
    ).strip("-")
    return {
        "id": suggestion_id or secrets.token_hex(4),
        "type": computed_type,
        "category": normalized_category,
        "category_label": COACH_SUGGESTION_CATEGORY_LABELS[normalized_category],
        "title": title.strip() or COACH_SUGGESTION_CATEGORY_LABELS[normalized_category],
        "message": description.strip(),
        "description": description.strip(),
        "action": action.strip(),
        "section": section.strip(),
        "original_text": original_text.strip(),
        "proposed_text": proposed_text.strip(),
        "requires_confirmation": bool(requires_confirmation),
        "supported_by_cv": bool(supported_by_cv),
        "keywords_added": keywords_added or [],
    }


def text_from_strategy_item(item: Any, fallback: str = "") -> tuple[str, str]:
    if isinstance(item, dict):
        title = str(item.get("title") or item.get("section") or fallback or "Suggerimento").strip()
        description = str(item.get("description") or item.get("suggestion") or item.get("coach_tip") or "").strip()
        action = str(item.get("coach_tip") or item.get("action") or "").strip()
        return title, description or action, action
    text = str(item or "").strip()
    return fallback or "Suggerimento", text, ""


EDIT_SECTION_ALIASES = {
    "CHI SONO": ["chi sono", "profilo", "profilo professionale", "summary", "about me", "obiettivo", "obiettivo professionale"],
    "HARD SKILLS": ["hard skills", "competenze", "competenze tecniche", "technical skills", "skills"],
    "SOFT SKILLS": ["soft skills", "competenze trasversali"],
    "FORMAZIONE": ["formazione", "istruzione", "education", "studi"],
    "ESPERIENZE PROFESSIONALI": ["esperienze professionali", "esperienza professionale", "esperienze", "esperienza", "work experience"],
    "PROGETTI": ["progetti", "projects", "project work"],
    "PAGINA AGGIUNTIVA": ["pagina aggiuntiva", "esperienze aggiuntive", "attivita rilevanti", "attività rilevanti"],
    "LINGUE": ["lingue", "languages", "comunicazione"],
    "CONTATTI": ["contatti", "contact", "contacts"],
}

EDIT_SECTION_ALIASES.update({
    "CHI SONO": sorted(SHARED_CV_SECTION_ALIASES["profile"]),
    "HARD SKILLS": sorted(SHARED_CV_SECTION_ALIASES["hard_skills"]),
    "SOFT SKILLS": sorted(SHARED_CV_SECTION_ALIASES["soft_skills"]),
    "FORMAZIONE": sorted(SHARED_CV_SECTION_ALIASES["education"]),
    "ESPERIENZE PROFESSIONALI": sorted(SHARED_CV_SECTION_ALIASES["experience"]),
    "PROGETTI": sorted(SHARED_CV_SECTION_ALIASES["projects"]),
    "CERTIFICAZIONI": sorted(SHARED_CV_SECTION_ALIASES["certifications"]),
    "LINGUE": sorted(SHARED_CV_SECTION_ALIASES["languages"]),
    "CONTATTI": sorted(SHARED_CV_SECTION_ALIASES["contacts"]),
})

EDIT_SECTION_MARKERS = [
    "CONTATTI", "LINGUE", "COMUNICAZIONE", "HARD SKILLS", "SOFT SKILLS",
    "COMPETENZE TECNICHE", "COMPETENZE", "CHI SONO", "OBIETTIVO",
    "OBIETTIVO PROFESSIONALE", "PROFILO PROFESSIONALE", "PROFILO",
    "FORMAZIONE", "ISTRUZIONE", "ESPERIENZE PROFESSIONALI",
    "ESPERIENZA PROFESSIONALE", "ESPERIENZE", "ESPERIENZA",
    "PROGETTI", "PAGINA AGGIUNTIVA", "ATTIVITA RILEVANTI", "ATTIVITÀ RILEVANTI",
]

EDIT_SECTION_MARKERS = sorted({
    marker
    for canonical, aliases in EDIT_SECTION_ALIASES.items()
    for marker in [canonical, *aliases]
}, key=len, reverse=True)

CANONICAL_EDIT_SECTION_NAMES = {
    alias: canonical
    for canonical, aliases in EDIT_SECTION_ALIASES.items()
    for alias in [canonical.lower(), *aliases]
}

def canonical_edit_section_name(value: str) -> Optional[str]:
    cleaned = normalize_section_title(str(value or ""))
    if not cleaned:
        return None
    shared_key = canonical_section_key(cleaned)
    if shared_key:
        return {
            "profile": "CHI SONO",
            "experience": "ESPERIENZE PROFESSIONALI",
            "education": "FORMAZIONE",
            "hard_skills": "HARD SKILLS",
            "soft_skills": "SOFT SKILLS",
            "languages": "LINGUE",
            "projects": "PROGETTI",
            "certifications": "CERTIFICAZIONI",
            "contacts": "CONTATTI",
        }.get(shared_key, cleaned.upper())
    return CANONICAL_EDIT_SECTION_NAMES.get(cleaned, cleaned.upper())


def _resume_section_key(canonical_section: str) -> str:
    normalized = (canonical_section or "").strip().upper()
    return {
        "CHI SONO": "profile",
        "PROFILO": "profile",
        "PROFILO PROFESSIONALE": "profile",
        "OBIETTIVO": "profile",
        "OBIETTIVO PROFESSIONALE": "profile",
        "HARD SKILLS": "hard_skills",
        "COMPETENZE TECNICHE": "hard_skills",
        "COMPETENZE": "hard_skills",
        "SOFT SKILLS": "soft_skills",
        "COMUNICAZIONE": "languages",
        "FORMAZIONE": "education",
        "ISTRUZIONE": "education",
        "ESPERIENZE PROFESSIONALI": "experience",
        "ESPERIENZA PROFESSIONALE": "experience",
        "ESPERIENZE": "experience",
        "ESPERIENZA": "experience",
        "CONTATTI": "contacts",
        "LINGUE": "languages",
        "CERTIFICAZIONI": "certifications",
        "PROGETTI": "projects",
    }.get(normalized, normalized.lower().replace(" ", "_") if normalized else "")


def extract_resume_sections(cv_text: str) -> Dict[str, str]:
    section_map = split_cv_edit_sections(cv_text)
    extracted: Dict[str, str] = {}
    for raw_section, text in section_map.items():
        canonical = canonical_edit_section_name(raw_section)
        key = _resume_section_key(canonical or raw_section)
        if key:
            extracted.setdefault(key, text.strip())

    if "profile" not in extracted:
        extracted["profile"] = find_profile_fallback(cv_text)
    if "hard_skills" not in extracted:
        extracted["hard_skills"] = find_hard_skills_fallback(cv_text)
    if "experience" not in extracted:
        extracted["experience"] = find_experience_fallback(cv_text)
    if "education" not in extracted:
        extracted["education"] = find_education_fallback(cv_text)

    return {key: value.strip() for key, value in extracted.items() if value.strip()}


def find_profile_fallback(cv_text: str) -> str:
    lines = [line.strip() for line in (cv_text or "").splitlines() if line.strip()]
    for line in lines[:20]:
        plain = normalize_plain_text(line)
        if (
            45 <= len(line) <= 900
            and not re.search(r"@|https?://|\+?\d[\d\s().-]{7,}", line)
            and not canonical_edit_section_name(line) in EDIT_SECTION_ALIASES
            and len(plain.split()) >= 8
        ):
            return line
    return ""


def find_hard_skills_fallback(cv_text: str) -> str:
    lines = [line.strip() for line in (cv_text or "").splitlines() if line.strip()]
    for index, line in enumerate(lines):
        if canonical_edit_section_name(line) != "HARD SKILLS":
            continue
        return lines[index + 1] if index + 1 < len(lines) else ""
    return ""


def find_experience_fallback(cv_text: str) -> str:
    sections = split_cv_edit_sections(cv_text)
    for heading in ("ESPERIENZE PROFESSIONALI", "ESPERIENZA PROFESSIONALE", "ESPERIENZE"):
        text = sections.get(heading, "").strip()
        if text:
            return first_section_paragraph(text)
    return ""


def find_education_fallback(cv_text: str) -> str:
    sections = split_cv_edit_sections(cv_text)
    for heading in ("FORMAZIONE", "ISTRUZIONE"):
        text = sections.get(heading, "").strip()
        if text:
            return first_section_paragraph(text)
    return ""


def split_cv_edit_sections(cv_text: str) -> Dict[str, str]:
    prepared = re.sub(r"\r\n?", "\n", cv_text or "").strip()
    prepared = re.sub(r"[ \t]+", " ", prepared)
    marker_pattern = "|".join(
        re.escape(marker)
        for marker in sorted(EDIT_SECTION_MARKERS, key=len, reverse=True)
    )
    if marker_pattern:
        prepared = re.sub(
            rf"(?<![\wÀ-ÖØ-öø-ÿ])({marker_pattern})\s*:?(?=\s|$)",
            lambda match: f"\n{match.group(1).strip()}\n",
            prepared,
            flags=re.IGNORECASE,
        )
    lines = [line.strip() for line in prepared.splitlines() if line.strip()]
    sections: Dict[str, List[str]] = {}
    current = "INTESTAZIONE"
    for line in lines:
        normalized = normalize_plain_text(line).strip(":")
        matched = None
        for canonical, aliases in EDIT_SECTION_ALIASES.items():
            if normalized == normalize_plain_text(canonical) or normalized in [normalize_plain_text(alias) for alias in aliases]:
                matched = canonical
                break
        if matched:
            current = matched
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)
    return {
        key: "\n".join(value).strip()
        for key, value in sections.items()
        if value and "\n".join(value).strip()
    }


def first_section_paragraph(text: str, preferred_terms: Optional[List[str]] = None, max_chars: int = 900) -> str:
    chunks = [chunk.strip() for chunk in re.split(r"\n{2,}|(?<=[.!?])\s+(?=[A-ZÀ-ÖØ-Ý])", text or "") if chunk.strip()]
    preferred_terms = [normalize_plain_text(term) for term in (preferred_terms or [])]
    for chunk in chunks:
        normalized = normalize_plain_text(chunk)
        if preferred_terms and any(term in normalized for term in preferred_terms) and len(chunk) <= max_chars:
            return chunk
    for chunk in chunks:
        if len(chunk) <= max_chars:
            return chunk
    return ""


def count_section_markers(value: str) -> int:
    lines = [
        normalize_plain_text(line).strip().strip(":")
        for line in str(value or "").splitlines()
        if str(line).strip()
    ]
    normalized_markers = {
        normalize_plain_text(marker).strip().strip(":")
        for marker in EDIT_SECTION_MARKERS
    }
    return sum(1 for line in lines if line in normalized_markers)


def is_valid_actionable_suggestion(suggestion: Dict) -> bool:
    if not isinstance(suggestion, dict) or suggestion.get("type") != "actionableEdit":
        return False
    section = str(suggestion.get("section") or "").strip()
    original = str(suggestion.get("original_text") or "").strip()
    proposed = str(suggestion.get("proposed_text") or "").strip()
    if not section or not proposed:
        return False
    if original and len(original) > 1000:
        return False
    if len(proposed) > 1000:
        return False
    if original and count_section_markers(original) > 1:
        return False
    if count_section_markers(proposed) > 0:
        return False
    section_plain = normalize_plain_text(section)
    original_plain = normalize_plain_text(original)
    proposed_plain = normalize_plain_text(proposed)
    # Reject suggestions that are identical when normalized
    if original and original_plain == proposed_plain:
        return False
    # Reject suggestions that are almost identical (too small change)
    try:
        similarity = SequenceMatcher(None, original_plain, proposed_plain).ratio() if original_plain and proposed_plain else 0
        if similarity >= 0.92 and suggestion.get("generated_by") != "ollama":
            return False
    except Exception:
        pass
    if section_plain in {"chi sono", "profilo", "profilo professionale"}:
        blocked = ["contatti", "lingue", "hard skills", "soft skills", "formazione", "esperienze professionali"]
        if any(term in original_plain or term in proposed_plain for term in blocked):
            return False
    if section_plain == "hard skills":
        if len(proposed.split()) > 90 or any(term in proposed_plain for term in ["contatti", "telefono", "email", "linkedin"]):
            return False
    if "esperienze" in section_plain:
        if any(term in proposed_plain for term in ["contatti", "lingue", "hard skills", "soft skills", "formazione"]):
            return False
    return ResumeRewriter().is_safe_replacement(section, proposed)


def suggestion_targets_current_cv(suggestion: Dict, cv_text: str) -> bool:
    original = str(suggestion.get("original_text") or "").strip()
    section = str(suggestion.get("section") or "").strip()
    source_id = normalize_plain_text(str(suggestion.get("source_id") or suggestion.get("id") or ""))
    if not section or not cv_text.strip():
        return False

    if not original:
        section_plain = normalize_plain_text(section)
        appendable_sections = {
            "profilo",
            "progetti",
            "esperienze professionali",
            "esperienze",
            "formazione",
            "certificazioni",
            "lingue",
            "competenze tecniche",
            "hard skills",
            "soft skills",
            "attivita rilevanti",
        }
        return section_plain in appendable_sections or "user_additional_info" in source_id

    original_plain = normalize_plain_text(original)
    cv_plain = normalize_plain_text(cv_text)
    if original_plain and original_plain in cv_plain:
        return True

    section_key = _resume_section_key(canonical_edit_section_name(section) or section)
    section_text = extract_resume_sections(cv_text).get(section_key, "")
    original_tokens = tokenize_meaningful(original)
    section_tokens = tokenize_meaningful(section_text)
    if not original_tokens or not section_tokens:
        return False
    coverage = len(original_tokens.intersection(section_tokens)) / len(original_tokens)
    # Suggestions generated locally are conservative but often paraphrase the original.
    # Accept a lower overlap when the section clearly exists in the current CV.
    return coverage >= 0.55 or len(original_tokens.intersection(section_tokens)) >= 4


def is_additive_user_rewrite_source(source_id: str) -> bool:
    source = str(source_id or "")
    additive_markers = (
        "confirmed_",
        "user_box_",
        "user_additional_",
        "fallback_confirmed_",
    )
    if source.startswith(additive_markers):
        return True
    if source.startswith("consolidated:"):
        consolidated_ids = source.split(":", 1)[1].split("|")
        return any(item.startswith(additive_markers) for item in consolidated_ids)
    return False


def _structured_target_profile(evaluation: Dict, section_map: Optional[Dict[str, str]] = None) -> Dict[str, Any]:
    """Profilo target locale: ruolo, azienda, keyword e skill da usare per suggerimenti mirati."""
    section_map = section_map or {}
    target = evaluation.get("target") if isinstance(evaluation.get("target"), dict) else {}
    role = str(target.get("role") or evaluation.get("role") or "").strip()
    company = str(target.get("company") or evaluation.get("company") or "").strip()
    description = str(target.get("description") or evaluation.get("description") or "").strip()
    required_skills = str(target.get("required_skills") or evaluation.get("required_skills") or "").strip()

    role_family = infer_role_family(role, description, required_skills)
    library = ROLE_SKILL_LIBRARY.get(role_family) or infer_skill_library_from_role(role, description)
    cv_text = str(evaluation.get("cv_text") or "\n".join(section_map.values()))
    role_snapshot = role_keyword_snapshot(cv_text, role, description, required_skills)

    hard_skills = list(dict.fromkeys(library.get("hard_skills", [])))
    soft_skills = list(dict.fromkeys(library.get("soft_skills", [])))
    tools = list(dict.fromkeys([*library.get("tools", []), *library.get("programming_languages", [])]))

    return {
        "role": role,
        "company": company,
        "role_family": role_family,
        "hard_skills": hard_skills,
        "soft_skills": soft_skills,
        "tools": tools,
        "present_keywords": list(dict.fromkeys(role_snapshot.get("present", []))),
        "partially_present_keywords": list(dict.fromkeys(role_snapshot.get("partially_present", []))),
        "keywords_to_confirm": list(dict.fromkeys(role_snapshot.get("to_confirm", []))),
        "missing_keywords": evaluation.get("missing_keywords", []),
        "missing_hard_skills": evaluation.get("missing_hard_skills", []),
        "missing_soft_skills": evaluation.get("missing_soft_skills", []),
        "company_signals": [company] if company else [],
    }


def _first_structured_block(text: str, max_chars: int = 850) -> str:
    """Estrae un blocco applicabile da una sezione senza prendere più sezioni insieme."""
    clean = re.sub(r"[ \t]+", " ", text or "").strip()
    if not clean:
        return ""

    lines = [line.strip(" \t-•·") for line in clean.splitlines() if line.strip()]
    useful_lines = [
        line for line in lines
        if line
        and not canonical_edit_section_name(line)
        and len(line) >= 18
    ]
    if useful_lines:
        joined = "\n".join(useful_lines[:4]).strip()
        if 30 <= len(joined) <= max_chars:
            return joined

    chunks = [
        chunk.strip()
        for chunk in re.split(r"\n{2,}|(?<=[.!?])\s+(?=[A-ZÀ-ÖØ-Ý])", clean)
        if chunk.strip()
    ]
    for chunk in chunks:
        if 30 <= len(chunk) <= max_chars and not canonical_edit_section_name(chunk):
            return chunk

    return clean[:max_chars].strip()


def _infer_suggestion_impact(category: str, section: str = "") -> str:
    normalized = normalize_plain_text(f"{category} {section}")
    if any(term in normalized for term in ["profile", "profilo", "chi sono", "skills", "competenze", "hard skills"]):
        return "alto"
    if any(term in normalized for term in ["experience", "esperienze", "project", "progetti", "ats"]):
        return "medio"
    return "basso"


def _make_structured_suggestion(
    category: str,
    title: str,
    description: str,
    section: str,
    original_text: str,
    proposed_text: str,
    keywords_added: Optional[List[str]] = None,
    impact: Optional[str] = None,
    priority: int = 50,
) -> Optional[Dict]:
    """Crea un actionableEdit completo, validato e con campi utili al frontend."""
    original_text = re.sub(r"[ \t]+", " ", original_text or "").strip()
    proposed_text = re.sub(r"[ \t]+", " ", proposed_text or "").strip()

    if not section or not original_text or not proposed_text:
        return None

    suggestion = make_coach_suggestion(
        category,
        title,
        description,
        section=section,
        original_text=original_text,
        proposed_text=proposed_text,
        keywords_added=keywords_added or [],
        suggestion_type="actionableEdit",
        supported_by_cv=True,
    )
    suggestion["impact"] = impact or _infer_suggestion_impact(category, section)
    suggestion["priority"] = int(priority)
    suggestion["reason"] = description.strip()
    suggestion["requires_confirmation"] = False
    suggestion["supported_by_cv"] = True

    if not is_valid_actionable_suggestion(suggestion):
        return None
    return suggestion


def _dedupe_structured_suggestions(suggestions: List[Dict], limit: int = 8) -> List[Dict]:
    unique: List[Dict] = []
    seen_ids = set()
    seen_pairs = set()

    for suggestion in suggestions:
        if not isinstance(suggestion, dict):
            continue
        if suggestion.get("type") != "actionableEdit":
            continue
        section = str(suggestion.get("section") or "").strip()
        original = normalize_plain_text(str(suggestion.get("original_text") or ""))
        proposed = normalize_plain_text(str(suggestion.get("proposed_text") or ""))
        if not section or not original or not proposed or original == proposed:
            continue

        suggestion_id = str(suggestion.get("id") or "").strip()
        pair_key = (normalize_plain_text(section), original[:180], proposed[:180])
        if suggestion_id in seen_ids or pair_key in seen_pairs:
            continue

        seen_ids.add(suggestion_id)
        seen_pairs.add(pair_key)
        unique.append(suggestion)
        if len(unique) >= limit:
            break

    return unique


def _skills_from_existing_text(skills_text: str) -> List[str]:
    cleaned = re.sub(r"\b(competenze tecniche|hard skills|soft skills|skills)\s*[:\-]", " ", skills_text or "", flags=re.I)
    parts = [
        part.strip(" \t-•·:")
        for part in re.split(r"[,;|•·\n]+", cleaned)
        if part.strip(" \t-•·:")
    ]
    result: List[str] = []
    seen = set()
    for part in parts:
        normalized = normalize_plain_text(part)
        if not normalized or normalized in seen or len(part) > 70:
            continue
        seen.add(normalized)
        result.append(part)
    return result


def _reorganize_skills_text(skills_text: str) -> str:
    skills = _skills_from_existing_text(skills_text)
    if len(skills) < 3:
        return ""

    languages = []
    ai_data = []
    tools = []
    soft = []
    other = []

    language_terms = {"python", "java", "javascript", "typescript", "c++", "c#", "sql", "r", "html", "css"}
    ai_data_terms = {"machine learning", "ai", "artificial intelligence", "data", "analisi", "analytics", "nlp", "llm"}
    soft_terms = {"comunicazione", "collaborazione", "problem solving", "creativita", "creatività", "teamwork", "leadership"}

    for skill in skills:
        plain = normalize_plain_text(skill)
        if plain in language_terms or plain.replace(" ", "") in {"c++", "c#"}:
            languages.append(skill)
        elif any(term in plain for term in ai_data_terms):
            ai_data.append(skill)
        elif any(term in plain for term in soft_terms):
            soft.append(skill)
        elif any(term in plain for term in ["figma", "excel", "power bi", "tableau", "unity", "unreal", "git", "docker", "jira"]):
            tools.append(skill)
        else:
            other.append(skill)

    lines = []
    if languages:
        lines.append("Linguaggi: " + ", ".join(languages))
    if ai_data:
        lines.append("AI e dati: " + ", ".join(ai_data))
    if tools:
        lines.append("Strumenti: " + ", ".join(tools))
    if other:
        lines.append("Competenze tecniche: " + ", ".join(other))
    if soft:
        lines.append("Soft skills: " + ", ".join(soft))

    proposed = "\n".join(lines).strip()
    return proposed if proposed and normalize_plain_text(proposed) != normalize_plain_text(skills_text) else ""


def _profile_rewrite_candidate(profile: str, target_profile: Dict[str, Any]) -> str:
    profile = re.sub(r"\s+", " ", profile or "").strip()
    role = str(target_profile.get("role") or "").strip()
    company = str(target_profile.get("company") or "").strip()
    if not profile or not role:
        return ""

    plain_profile = normalize_plain_text(profile)
    role_plain = normalize_plain_text(role)
    company_phrase = f" presso {company}" if company and normalize_plain_text(company) not in plain_profile else ""

    if role_plain in plain_profile:
        return ""

    proposed = (
        f"{profile.rstrip('.')}. "
        f"Per il ruolo di {role}{company_phrase}, valorizzo il mio percorso tecnico e le competenze già presenti nel CV, "
        "orientandole alla candidatura in modo chiaro, concreto e coerente con le esperienze reali."
    )
    return proposed[:950].strip()


def _format_existing_section_as_bullets(value: str) -> str:
    clean = re.sub(r"[ \t]+", " ", value or "").strip()
    if not clean or clean.count("\n") >= 2:
        return ""

    sentences = [
        sentence.strip(" -•·")
        for sentence in re.split(r"(?<=[.!?])\s+", clean)
        if sentence.strip(" -•·")
    ]
    if len(sentences) < 2:
        return ""

    bullets = [f"- {sentence}" for sentence in sentences[:5]]
    proposed = "\n".join(bullets)
    return proposed if normalize_plain_text(proposed) != normalize_plain_text(value) else ""


def _compact_cv_text(value: str, max_chars: int = 900) -> str:
    clean = re.sub(r"[ \t]+", " ", value or "")
    clean = re.sub(r"\n{3,}", "\n\n", clean).strip()
    if len(clean) <= max_chars:
        return clean
    clipped = clean[:max_chars].rsplit(" ", 1)[0].strip()
    return clipped or clean[:max_chars].strip()


def _strip_skill_level_symbols(value: str) -> str:
    # Rimuove pallini/livelli grafici tipo ●●●○○, ma conserva skill e separatori.
    cleaned = re.sub(r"[●○■□▪▫◆◇★☆]{1,}", " ", value or "")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def _split_skill_items(value: str) -> List[str]:
    cleaned = _strip_skill_level_symbols(value)
    raw_parts = re.split(r"[,;|•·\n]+|\s{2,}", cleaned)
    items: List[str] = []
    seen = set()
    for part in raw_parts:
        item = re.sub(r"\s+", " ", part or "").strip(" -:•·")
        if not item:
            continue
        plain = normalize_plain_text(item)
        if (
            not plain
            or plain in {"hard skills", "soft skills", "competenze", "competenze tecniche"}
            or re.fullmatch(r"\d+", plain)
            or "@" in item
            or "linkedin" in plain
            or "http" in plain
            or "via " in plain
            or len(item) > 55
        ):
            continue
        if plain not in seen:
            seen.add(plain)
            items.append(item)
    return items[:18]


def _clean_section_fragment(value: str, section_key: str = "") -> str:
    """Taglia frammenti rumorosi dovuti a CV in tabelle/colonne."""
    if not value:
        return ""
    text = re.sub(r"\r\n?", "\n", value)
    text = re.sub(r"[ \t]+", " ", text)
    stop_markers = [
        "CONTATTI", "COMUNICAZIONE", "LINGUE", "SOFT SKILLS", "HARD SKILLS",
        "COMPETENZE", "FORMAZIONE", "ISTRUZIONE", "ESPERIENZA", "ESPERIENZE",
        "PROGETTI", "OBIETTIVO", "PROFILO", "PAGINA AGGIUNTIVA", "ATTIVITA RILEVANTI",
    ]
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    cleaned_lines: List[str] = []
    for index, line in enumerate(lines):
        plain = normalize_plain_text(line)
        upper = line.upper().strip(":")
        if index > 0 and upper in stop_markers:
            break
        if index > 0 and (
            re.search(r"[\w.+-]+@[\w.-]+\.\w+", line)
            or re.search(r"https?://|www\.|linkedin\.com", line, flags=re.I)
            or re.search(r"\+?\d[\d\s().-]{7,}", line)
            or plain.startswith("via ")
            or plain in {"pagina web personale", "linkedin profile"}
        ):
            break
        if section_key in {"hard_skills", "soft_skills"} and len(cleaned_lines) >= 12:
            break
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines).strip()


def _section_value(section_map: Dict[str, str], *keys: str) -> str:
    for key in keys:
        value = _clean_section_fragment(section_map.get(key, ""), key)
        if value:
            return value
    return ""


def _target_profile_from_evaluation(evaluation: Dict) -> Dict[str, Any]:
    target = evaluation.get("target") if isinstance(evaluation.get("target"), dict) else {}
    role = str(target.get("role") or evaluation.get("role") or "").strip()
    company = str(target.get("company") or evaluation.get("company") or "").strip()
    role_family = infer_role_family(role)
    suggested = evaluation.get("suggested_skills") if isinstance(evaluation.get("suggested_skills"), dict) else {}
    ats = evaluation.get("ats_analysis") if isinstance(evaluation.get("ats_analysis"), dict) else {}
    return {
        "role": role,
        "company": company,
        "role_family": role_family,
        "present_keywords": list(dict.fromkeys(
            [str(x).strip() for x in (evaluation.get("present_keywords") or []) if str(x).strip()]
            + [str(x).strip() for x in (ats.get("present_keywords") or ats.get("keywords_present") or []) if str(x).strip()]
        ))[:12],
        "missing_keywords": list(dict.fromkeys(
            [str(x).strip() for x in (evaluation.get("missing_keywords") or []) if str(x).strip()]
            + [str(x).strip() for x in (ats.get("missing_keywords") or ats.get("keywords_missing") or []) if str(x).strip()]
        ))[:12],
        "skills_to_confirm": suggested.get("confirmation_items", []) if isinstance(suggested.get("confirmation_items"), list) else [],
    }


def _make_structured_action(
    category: str,
    title: str,
    description: str,
    section: str,
    original_text: str,
    proposed_text: str,
    *,
    reason: str = "",
    impact: str = "medio",
    priority: int = 5,
    keywords_added: Optional[List[str]] = None,
) -> Dict:
    suggestion = make_coach_suggestion(
        category,
        title,
        description,
        section=section,
        original_text=_compact_cv_text(original_text, 950),
        proposed_text=_compact_cv_text(proposed_text, 950),
        keywords_added=keywords_added or [],
        supported_by_cv=True,
        suggestion_type="actionableEdit",
    )
    suggestion["reason"] = reason or description
    suggestion["impact"] = impact
    suggestion["priority"] = priority
    return suggestion


def _append_suggestion(
    suggestions: List[Dict[str, Any]],
    category: str,
    title: str,
    description: str,
    section: str,
    original_text: str,
    proposed_text: str,
    impact: str,
    priority: int,
    keywords_added: Optional[List[str]] = None,
    *,
    reason: str = "",
) -> None:
    suggestions.append(_make_structured_action(
        category,
        title,
        description,
        section,
        original_text,
        proposed_text,
        reason=reason or description,
        impact=impact,
        priority=priority,
        keywords_added=keywords_added or [],
    ))


def _profile_rewrite(profile: str, target_profile: Dict[str, Any], cv_text: str = "") -> str:
    text = re.sub(r"\s+", " ", profile or "").strip()
    role = str(target_profile.get("role") or "").strip()
    if not text:
        return ""
    if not role:
        return text
    return (
        f"{text}. "
        f"Per il ruolo di {role}, valorizzo il mio percorso tecnico e le competenze gi? presenti nel CV "
        "in modo chiaro, concreto e coerente con le esperienze reali."
    )


def _experience_rewrite(experience: str, target_profile: Dict[str, Any]) -> str:
    text = re.sub(r"\s+", " ", experience or "").strip()
    role = str(target_profile.get("role") or "").strip()
    if not text:
        return ""
    if role:
        return (
            f"Esperienza valorizzata per il ruolo target: {text}. "
            "Rendo le attività più leggibili con bullet chiari, mantenendo solo informazioni già presenti nel CV."
        )
    return (
        f"{text}. "
        "Rendo le attività più leggibili con bullet chiari, mantenendo solo informazioni già presenti nel CV."
    )


def _projects_rewrite(projects: str, target_profile: Dict[str, Any]) -> str:
    text = re.sub(r"\s+", " ", projects or "").strip()
    if not text:
        return ""
    return (
        f"Progetti riorganizzati in modo più leggibile: {text}. "
        "Metto in evidenza ruolo, attività e tecnologie già presenti senza aggiungere dettagli non documentati."
    )


def _education_rewrite(education: str, target_profile: Dict[str, Any], cv_text: str = "") -> str:
    text = re.sub(r"\s+", " ", education or "").strip()
    if not text:
        return ""
    return (
        f"{text}. "
        "Rendo la formazione più chiara e coerente con il profilo, senza forzare collegamenti non presenti nel CV."
    )


def _shorten_cv_text(text: str, max_length: int = 700) -> str:
    compact = re.sub(r"\s+", " ", text or "").strip()
    if not compact:
        return ""
    if len(compact) <= max_length:
        return compact

    shortened = compact[:max_length].rsplit(" ", 1)[0].strip()
    return (shortened or compact[:max_length].strip()) + "..."


def _unique_preserve_order(values: Iterable[str]) -> List[str]:
    result: List[str] = []
    seen = set()
    for value in values:
        item = re.sub(r"\s+", " ", str(value or "")).strip()
        plain = normalize_plain_text(item)
        if not plain or plain in seen:
            continue
        seen.add(plain)
        result.append(item)
    return result


def _extract_skill_terms(value: str) -> List[str]:
    return _unique_preserve_order(re.split(r"[,;|??\n]+", value or ""))


def _group_skill_terms(skill_terms: List[str]) -> str:
    return "\n".join(f"- {term}" for term in skill_terms[:12]).strip()


def build_generic_rewrite_fallbacks(section_map: Dict[str, str], role: str) -> List[Dict]:
    target_profile = {"role": role, "company": "", "role_family": infer_role_family(role)}
    cv_text = "\n".join(section_map.values())
    suggestions: List[Dict[str, Any]] = []

    profile = section_map.get("profile", "")
    if profile:
        _append_suggestion(
            suggestions,
            "profile",
            "Riscrivi il profilo in funzione del ruolo",
            "Trasforma il profilo in un testo mirato al ruolo, senza inventare esperienze o competenze.",
            "CHI SONO",
            profile,
            _profile_rewrite(profile, target_profile, cv_text),
            "alto",
            1,
            [role] if role else [],
        )

    hard_skills = section_map.get("hard_skills", "")
    skill_terms = _extract_skill_terms(hard_skills)
    grouped = _group_skill_terms(skill_terms)
    if hard_skills and grouped:
        _append_suggestion(
            suggestions,
            "skills",
            "Riorganizza le competenze tecniche già presenti",
            "Rende la sezione competenze più leggibile per recruiter e ATS, senza aggiungere skill non confermate.",
            "HARD SKILLS",
            hard_skills,
            grouped,
            "alto",
            2,
            skill_terms,
        )

    soft_skills = section_map.get("soft_skills", "")
    if soft_skills:
        soft_values = _unique_preserve_order(re.split(r"[,;|•·\n]+", soft_skills))
        if len(soft_values) >= 2:
            _append_suggestion(
                suggestions,
                "soft_skills",
                "Rendi più chiara la sezione soft skills",
                "Mantiene le soft skill reali e le presenta in modo più pulito.",
                "SOFT SKILLS",
                soft_skills,
                "Soft skills: " + ", ".join(soft_values[:8]),
                "medio",
                3,
                soft_values,
            )

    experience = section_map.get("experience", "")
    if experience:
        _append_suggestion(
            suggestions,
            "experience",
            "Valorizza l'esperienza più rilevante",
            "Riformula l'esperienza usando attività concrete e bullet leggibili, mantenendo i fatti presenti.",
            "ESPERIENZE PROFESSIONALI",
            experience,
            _experience_rewrite(experience, target_profile),
            "alto",
            4,
            [],
        )

    projects = section_map.get("projects") or section_map.get("progetti") or ""
    if projects:
        _append_suggestion(
            suggestions,
            "project",
            "Valorizza i progetti più coerenti",
            "Rende i progetti più leggibili e collegati al ruolo target senza inventare dettagli.",
            "PROGETTI",
            projects,
            _projects_rewrite(projects, target_profile),
            "medio",
            5,
            [],
        )

    education = section_map.get("education", "")
    if education:
        proposed = _education_rewrite(education, target_profile, cv_text)
        _append_suggestion(
            suggestions,
            "education",
            "Rendi la formazione più mirata",
            "Collega la formazione al target solo attraverso aree realmente presenti nel CV.",
            "FORMAZIONE",
            education,
            proposed,
            "medio",
            6,
            [],
        )

    return sorted(suggestions, key=lambda item: int(item.get("priority", 99)))[:8]





def build_coach_suggestions_from_evaluation(evaluation: Dict, allow_llm: bool = False) -> List[Dict]:
    """Genera sempre suggerimenti locali applicabili, senza dipendere dall'LLM.

    Prima prova il guard locale; se non trova nulla, usa fallback più permissivi
    costruiti dal testo del CV. Questo evita CV ottimizzati identici all'originale.
    """
    suggestions: List[Dict] = []
    cv_text = str(evaluation.get("cv_text") or "")

    def _is_allowed_coach_suggestion(item: Dict[str, Any]) -> bool:
        category = normalize_plain_text(str(item.get("category") or ""))
        if category in {"ats keywords", "keyword", "skills", "soft skills"}:
            return False
        return True

    try:
        from services.cv_optimizer.safe_cv_guard import build_structured_cv_suggestions

        suggestions = build_structured_cv_suggestions(evaluation)
        suggestions = [
            item for item in suggestions
            if is_valid_actionable_suggestion(item)
            and suggestion_targets_current_cv(item, cv_text)
            and _is_allowed_coach_suggestion(item)
        ]
        if suggestions:
            return suggestions[:8]
        print(
            "Analisi CV: il guard locale non ha prodotto suggerimenti da mostrare; "
            "provo il generatore locale deterministico."
        )
    except Exception as exc:
        print(f"Safe CV guard non disponibile: {exc}. Uso fallback locale applicabile.")

    if allow_llm:
        try:
            from services.cv_optimizer.skill_suggestions import build_skill_mini_shot_suggestions
            mini_suggestions = [
                item
                for item in build_skill_mini_shot_suggestions(evaluation)
                if _is_allowed_coach_suggestion(item)
            ]
            if mini_suggestions:
                return mini_suggestions[:5]
        except Exception as exc:
            print(f"Mini-shot skill suggestions non riuscito: {exc}")

    try:
        cv_text = str(evaluation.get("cv_text") or "")
        target = evaluation.get("target") if isinstance(evaluation.get("target"), dict) else {}
        role = str(target.get("role") or evaluation.get("role") or "").strip()
        section_map = _extract_sections_for_structured_suggestions(cv_text)
        fallback_suggestions = build_generic_rewrite_fallbacks(section_map, role)
        fallback_suggestions = [
            item for item in fallback_suggestions
            if is_valid_actionable_suggestion(item)
            and suggestion_targets_current_cv(item, cv_text)
            and _is_allowed_coach_suggestion(item)
        ]
        if fallback_suggestions:
            return fallback_suggestions[:8]
    except Exception as exc:
        print(f"Fallback suggerimenti CV non riuscito: {exc}")

    try:
        user_additional_data = evaluation.get("user_additional_data")
        if isinstance(user_additional_data, dict) and flatten_cv_support_data(user_additional_data):
            extra_instructions = build_additional_rewrite_instructions(
                user_additional_data,
                role,
                cv_text,
            )
            extra_suggestions = []
            for index, instruction in enumerate(extra_instructions):
                suggestion = make_coach_suggestion(
                    instruction.category or "extra_info",
                    "Inserisci le informazioni aggiuntive nel CV",
                    instruction.reason or "Trasforma le informazioni compilate altrove in testo CV utilizzabile.",
                    section=instruction.section,
                    original_text=instruction.original,
                    proposed_text=instruction.replacement,
                    supported_by_cv=True,
                    suggestion_type="actionableEdit",
                )
                if not suggestion:
                    continue
                suggestion["impact"] = "medio"
                suggestion["priority"] = 1 + index
                suggestion["source_id"] = instruction.source_id
                if is_valid_actionable_suggestion(suggestion) and suggestion_targets_current_cv(suggestion, cv_text):
                    extra_suggestions.append(suggestion)
            if extra_suggestions:
                return extra_suggestions[:8]
    except Exception as exc:
        print(f"Fallback suggerimenti dati extra non riuscito: {exc}")

    # Ultimo fallback: prendi un blocco reale del CV e rendilo più leggibile.
    try:
        cv_text = str(evaluation.get("cv_text") or "")
        target = evaluation.get("target") if isinstance(evaluation.get("target"), dict) else {}
        role = str(target.get("role") or evaluation.get("role") or "ruolo target").strip() or "ruolo target"
        section_map = _extract_sections_for_structured_suggestions(cv_text)
        for section_key, section_name in [
            ("profile", "CHI SONO"),
            ("experience", "ESPERIENZE PROFESSIONALI"),
            ("projects", "PROGETTI"),
            ("education", "FORMAZIONE"),
        ]:
            original = section_map.get(section_key, "")
            if not original.strip():
                continue
            proposed = (
                _experience_rewrite(original, {"role": role})
                if section_key in {"experience", "projects"}
                else _profile_rewrite(original, {"role": role, "company": ""}, cv_text)
            )
            if section_key == "education":
                proposed = _education_rewrite(original, {"role": role}, cv_text) or _shorten_cv_text(original, 700)
            elif section_key == "projects":
                proposed = _projects_rewrite(original, {"role": role, "company": ""})
            elif section_key == "experience":
                proposed = _experience_rewrite(original, {"role": role}) or _shorten_cv_text(original, 700)
            item = make_coach_suggestion(
                "experience" if section_key == "experience" else "profile",
                "Rendi più leggibile e mirata questa sezione",
                "Riformula esclusivamente fatti già presenti nel CV.",
                section=section_name,
                original_text=original,
                proposed_text=proposed,
                supported_by_cv=True,
                suggestion_type="actionableEdit",
            )
            if item and is_valid_actionable_suggestion(item) and suggestion_targets_current_cv(item, cv_text):
                item["impact"] = "medio"
                item["priority"] = 99
                return [item]
    except Exception as exc:
        print(f"Fallback minimo suggerimenti CV non riuscito: {exc}")

    # Se anche il fallback minimo non trova una sezione pulita, crea un suggerimento
    # estremamente conservativo sulla prima parte utile del CV.
    try:
        cv_text = str(evaluation.get("cv_text") or "")
        target = evaluation.get("target") if isinstance(evaluation.get("target"), dict) else {}
        role = str(target.get("role") or evaluation.get("role") or "ruolo target").strip() or "ruolo target"
        section_map = _extract_sections_for_structured_suggestions(cv_text)
        first_section_key = next((key for key in ["profile", "experience", "projects", "education"] if section_map.get(key, "").strip()), "")
        if first_section_key:
            original = section_map.get(first_section_key, "")
            proposed = _shorten_cv_text(original, 700)
            if first_section_key == "experience":
                proposed = _experience_rewrite(original, {"role": role}) or _shorten_cv_text(original, 700)
            elif first_section_key == "projects":
                proposed = _projects_rewrite(original, {"role": role, "company": ""})
            elif first_section_key == "education":
                proposed = _education_rewrite(original, {"role": role}, cv_text) or _shorten_cv_text(original, 700)
            else:
                proposed = _profile_rewrite(original, {"role": role, "company": ""}, cv_text) or _shorten_cv_text(original, 700)
            item = make_coach_suggestion(
                f"{first_section_key}-fallback",
                "Ottimizza questa sezione per il ruolo target",
                "Suggerimento conservativo basato su testo già presente nel CV.",
                section=(first_section_key.upper().replace("_", " ")),
                original_text=original,
                proposed_text=proposed or _shorten_cv_text(original, 500),
                supported_by_cv=True,
                suggestion_type="actionableEdit",
            )
            if item and is_valid_actionable_suggestion(item) and suggestion_targets_current_cv(item, cv_text):
                item["impact"] = "basso"
                item["priority"] = 100
                return [item]
    except Exception as exc:
        print(f"Fallback conservativo suggerimenti CV non riuscito: {exc}")

    # Ultima rete di sicurezza: se non riusciamo ancora a costruire un
    # suggerimento strutturato, proviamo con il primo blocco non vuoto del CV.
    try:
        cv_text = str(evaluation.get("cv_text") or "")
        target = evaluation.get("target") if isinstance(evaluation.get("target"), dict) else {}
        role = str(target.get("role") or evaluation.get("role") or "ruolo target").strip() or "ruolo target"
        fallback_block = next((block.strip() for block in re.split(r"\n\s*\n", cv_text) if block.strip()), "")
        if fallback_block:
            proposed = _shorten_cv_text(fallback_block, 700)
            item = make_coach_suggestion(
                "profile",
                "Valorizza questo blocco del CV",
                "Riscrive solo testo già presente per renderlo più chiaro e più mirato al ruolo target.",
                section="CHI SONO",
                original_text=fallback_block,
                proposed_text=_profile_rewrite(fallback_block, {"role": role, "company": ""}, cv_text) or proposed,
                supported_by_cv=True,
                suggestion_type="actionableEdit",
            )
            if item and is_valid_actionable_suggestion(item) and suggestion_targets_current_cv(item, cv_text):
                item["impact"] = "basso"
                item["priority"] = 101
                return [item]
    except Exception as exc:
        print(f"Fallback sicurezza suggerimenti CV non riuscito: {exc}")

    return []


def build_cv_job_suggestions(evaluation: Dict, allow_llm: bool = False) -> List[Dict]:
    from services.cv_optimizer.suggestions import build_cv_job_suggestions as build_suggestions

    return build_suggestions(evaluation, allow_llm=allow_llm)


def normalize_accepted_coach_suggestions(value: Any) -> List[Dict]:
    if not isinstance(value, list):
        return []

    accepted = []
    for index, item in enumerate(value):
        if not isinstance(item, dict):
            continue
        if item.get("type") != "actionableEdit":
            print(f"Suggerimento non applicabile ignorato: {item.get('id') or item.get('title') or index}")
            continue
        title = str(item.get("title") or item.get("category_label") or "Suggerimento accettato").strip()
        description = str(item.get("description") or item.get("action") or "").strip()
        proposed_text = str(item.get("proposed_text") or item.get("replacement") or "").strip()
        section = str(item.get("section") or "").strip()
        original_text = str(item.get("original_text") or item.get("original") or "").strip()
        if not section or not original_text or not proposed_text:
            print(f"Suggerimento applicabile incompleto ignorato: {item.get('id') or title}")
            continue
        if not is_valid_actionable_suggestion({
            **item,
            "type": "actionableEdit",
            "section": section,
            "original_text": original_text,
            "proposed_text": proposed_text,
        }):
            print(f"Suggerimento applicabile non sicuro ignorato: {item.get('id') or title} | {proposed_text[:180]}")
            continue
        suggestion = make_coach_suggestion(
            str(item.get("category") or "phrases"),
            title,
            description,
            str(item.get("action") or ""),
            bool(item.get("requires_confirmation")),
            section=section,
            original_text=original_text,
            proposed_text=proposed_text,
            supported_by_cv=item.get("supported_by_cv") is not False,
            keywords_added=item.get("keywords_added") if isinstance(item.get("keywords_added"), list) else [],
            suggestion_type="actionableEdit",
        )
        suggestion["id"] = str(item.get("id") or suggestion["id"]).strip()
        accepted.append(suggestion)
    return accepted[:30]


def build_accepted_suggestions_from_confirmed_skills(
    cv_text: str,
    user_additional_data: Optional[Dict[str, Any]],
    role: str = "",
) -> List[Dict]:
    suggestions: List[Dict] = []
    for instruction in build_confirmed_skill_rewrite_instructions(
        cv_text,
        user_additional_data or {},
        role,
    ):
        section = str(instruction.section or "").strip()
        proposed_text = str(instruction.replacement or "").strip()
        if not section or not proposed_text:
            continue
        title = "Skill confermata"
        description = str(instruction.reason or "Skill confermata dall'utente").strip()
        suggestion = {
            "id": str(instruction.source_id or "").strip()
            or re.sub(
                r"[^a-z0-9]+",
                "-",
                normalize_plain_text(f"{section}-{proposed_text}")[:90],
            ).strip("-"),
            "type": "actionableEdit",
            "category": "soft_skills" if normalize_plain_text(section) == "soft skills" else "skills",
            "category_label": "Soft Skills" if normalize_plain_text(section) == "soft skills" else "Competenze",
            "title": title,
            "message": description,
            "description": description,
            "action": "append" if not str(instruction.original or "").strip() else "replace",
            "section": section,
            "original_text": str(instruction.original or "").strip(),
            "proposed_text": proposed_text,
            "requires_confirmation": False,
            "supported_by_cv": True,
            "keywords_added": [],
            "source_id": str(instruction.source_id or "").strip(),
        }
        suggestions.append(suggestion)
    return suggestions[:30]


def build_fallback_cv_job_evaluation(
    cv_text: str,
    company: str,
    role: str,
    description: str,
    sources: List[Dict[str, str]],
    required_skills: str = "",
) -> Dict:
    cv_tokens = tokenize_meaningful(cv_text)
    role_tokens = tokenize_meaningful(role)
    description_tokens = tokenize_meaningful(description)
    scorecard = build_deterministic_cv_scorecard(
        cv_text, company, role, description, required_skills
    )
    completeness = scorecard["completeness_score"]
    role_match = scorecard["role_match_score"]
    company_fit = scorecard["company_fit_score"]
    clarity = scorecard["clarity_score"]
    professionalism = scorecard["professionalism_score"]

    required_skill_tokens = tokenize_meaningful(required_skills)
    relevant_found = sorted(list(cv_tokens.intersection(role_tokens.union(description_tokens).union(required_skill_tokens))))[:8]
    missing = filter_cv_keyword_list(sorted(list((role_tokens.union(description_tokens).union(required_skill_tokens)) - cv_tokens)))[:8]
    ats_analysis = scorecard["ats_analysis"]
    overall = scorecard["overall_score"]
    suggested_skills = build_role_skill_suggestions(cv_text, role, description, required_skills)
    questions_for_user = generate_cv_optimization_questions(cv_text, {
        "weaknesses": [],
        "missing_skills_for_role": missing,
    }, ats_analysis)
    strengths = [
        "Il CV contiene elementi utili per una prima valutazione della candidatura.",
        "La struttura include sezioni riconoscibili di un curriculum.",
        "Sono presenti informazioni confrontabili con ruolo e descrizione inseriti.",
    ]
    weaknesses = [
        f"Rendi più evidente la coerenza con il ruolo {role}." if role else "Rendi più chiaro l'obiettivo professionale.",
        "Valorizza solo competenze e parole chiave già presenti nel CV o confermate dall'utente.",
        "Rendi più specifiche le attività descritte, senza aggiungere risultati non documentati.",
    ]

    evaluation = {
        "target": {"company": company, "role": role},
        "overall_score": overall,
        "ats_score": ats_analysis["ats_score"],
        "job_match_score": role_match,
        "keyword_score": ats_analysis.get("keyword_score", ats_analysis["ats_score"]),
        "format_score": ats_analysis.get("format_score", completeness),
        "ats_analysis": ats_analysis,
        "suggested_skills": suggested_skills,
        "role_match_score": role_match,
        "company_fit_score": company_fit,
        "company_provided": scorecard["company_provided"],
        "clarity_score": clarity,
        "completeness_score": completeness,
        "professionalism_score": professionalism,
        "strengths": strengths,
        "weaknesses": weaknesses,
        "strong_points": strengths,
        "weak_points": weaknesses,
        "relevant_skills_found": relevant_found,
        "missing_skills_for_role": missing,
        "relevant_experiences": [
            "Evidenzia le esperienze piu vicine al ruolo indicato e collega ogni attivita a risultati concreti."
        ],
        "present_keywords": ats_analysis.get("present_keywords", ats_analysis.get("keywords_present", [])),
        "missing_keywords": ats_analysis.get("missing_keywords", ats_analysis.get("keywords_missing", [])),
        "missing_hard_skills": ats_analysis.get("missing_hard_skills", []),
        "missing_soft_skills": ats_analysis.get("missing_soft_skills", []),
        "sections_to_improve": ats_analysis.get("sections_to_improve", ats_analysis.get("missing_sections", [])),
        "suggestions": [
            f"Personalizza il profilo iniziale citando il ruolo {role}.",
            "Aggiungi risultati numerici alle esperienze, ad esempio metriche, obiettivi raggiunti o impatto.",
            "Rendi piu visibili le competenze richieste dall'annuncio nella sezione competenze.",
            f"Inserisci un riferimento chiaro al tipo di contesto aziendale di {company}.",
        ],
        "questions_for_user": questions_for_user,
        "cv_text": cv_text,
        "summary": "Il CV è valido ma va personalizzato su ruolo, azienda e parole chiave per essere competitivo.",
        "scoring_context": scorecard["scoring_context"],
    }
    evaluation["score_explanation"] = scorecard["score_explanation"]
    evaluation["coach_suggestions"] = []
    return evaluation


def normalize_cv_job_evaluation(result: Dict, fallback: Dict) -> Dict:
    result = result if isinstance(result, dict) else {}
    normalized = {}
    for key in [
        "overall_score", "ats_score", "role_match_score", "company_fit_score",
        "clarity_score", "completeness_score", "professionalism_score"
    ]:
        normalized[key] = clamp_score(fallback.get(key, 0))

    normalized["job_match_score"] = clamp_score(
        fallback.get("job_match_score", normalized["role_match_score"])
    )
    normalized["keyword_score"] = clamp_score(
        fallback.get("keyword_score", normalized["ats_score"])
    )
    normalized["format_score"] = clamp_score(
        fallback.get("format_score", normalized["completeness_score"])
    )
    normalized["company_provided"] = bool(fallback.get("company_provided"))
    normalized["score_explanation"] = fallback.get("score_explanation") or build_cv_score_explanation(normalized)

    if "strengths" not in result and isinstance(result.get("strong_points"), list):
        result["strengths"] = result["strong_points"]
    if "weaknesses" not in result and isinstance(result.get("weak_points"), list):
        result["weaknesses"] = result["weak_points"]

    list_fields = [
        "strengths", "weaknesses", "relevant_skills_found",
        "missing_skills_for_role", "relevant_experiences", "suggestions",
        "present_keywords", "missing_keywords", "missing_hard_skills",
        "missing_soft_skills", "sections_to_improve", "questions_for_user"
    ]
    for field in list_fields:
        value = result.get(field)
        normalized[field] = value if isinstance(value, list) and value else fallback.get(field, [])

    normalized["summary"] = result.get("summary") or fallback["summary"]
    normalized["target"] = fallback.get("target", {})
    normalized["cv_text"] = fallback.get("cv_text", "")
    normalized["scoring_context"] = fallback.get("scoring_context", {})
    normalized["suggested_skills"] = result.get("suggested_skills") if isinstance(result.get("suggested_skills"), dict) else fallback.get("suggested_skills", {})
    normalized["ats_analysis"] = dict(fallback.get("ats_analysis", {}))
    normalized["ats_analysis"]["keyword_score"] = normalized["ats_analysis"].get("keyword_score", normalized["keyword_score"])
    normalized["ats_analysis"]["format_score"] = normalized["ats_analysis"].get("format_score", normalized["format_score"])
    normalized["ats_analysis"]["present_keywords"] = normalized["ats_analysis"].get(
        "present_keywords",
        normalized["ats_analysis"].get("keywords_present", normalized["present_keywords"])
    )
    normalized["ats_analysis"]["missing_keywords"] = normalized["ats_analysis"].get(
        "missing_keywords",
        normalized["ats_analysis"].get("keywords_missing", normalized["missing_keywords"])
    )
    normalized["ats_analysis"]["keywords_present"] = normalized["ats_analysis"].get(
        "keywords_present",
        normalized["ats_analysis"].get("present_keywords", normalized["present_keywords"])
    )
    normalized["ats_analysis"]["keywords_missing"] = normalized["ats_analysis"].get(
        "keywords_missing",
        normalized["ats_analysis"].get("missing_keywords", normalized["missing_keywords"])
    )
    normalized["ats_analysis"]["missing_hard_skills"] = normalized["ats_analysis"].get("missing_hard_skills", normalized["missing_hard_skills"])
    normalized["ats_analysis"]["missing_soft_skills"] = normalized["ats_analysis"].get("missing_soft_skills", normalized["missing_soft_skills"])
    normalized["ats_analysis"]["keywords_partially_present"] = normalized["ats_analysis"].get(
        "keywords_partially_present",
        fallback.get("ats_analysis", {}).get("keywords_partially_present", [])
    )
    normalized["ats_analysis"]["keywords_to_confirm"] = normalized["ats_analysis"].get(
        "keywords_to_confirm",
        fallback.get("ats_analysis", {}).get("keywords_to_confirm", [])
    )
    normalized["ats_analysis"]["sections_to_improve"] = normalized["ats_analysis"].get(
        "sections_to_improve",
        normalized["ats_analysis"].get("missing_sections", normalized["sections_to_improve"])
    )
    normalized["coach_suggestions"] = []
    normalized["strong_points"] = normalized["strengths"]
    normalized["weak_points"] = normalized["weaknesses"]
    return normalized


def evaluate_cv_for_job(
    cv_text: str,
    company: str,
    role: str,
    description: str,
    link: str,
    sources: Optional[List[Dict[str, str]]] = None,
    required_skills: str = "",
) -> Dict:
    from services.cv_optimizer.evaluation import evaluate_cv_for_job as evaluate

    return evaluate(
        cv_text=cv_text,
        company=company,
        role=role,
        description=description,
        link=link,
        sources=sources,
        required_skills=required_skills,
    )


def build_lightweight_cv_evaluation_prompt(
    company: str,
    role: str,
    description: str,
    required_skills: str,
    link: str,
    sources_prompt: str,
    ats_analysis: Dict[str, Any],
    cv_text: str,
) -> str:
    return build_cv_analysis_prompt(
        company=company,
        role=role,
        description=description,
        required_skills=required_skills,
        link=link,
        sources_prompt=sources_prompt,
        ats_analysis=ats_analysis,
        cv_text=cv_text,
        compact=True,
    )


@app.post("/job/validate")
def validate_job_endpoint(data: JobValidationRequest):
    return validate_job_input(
        description=(data.description or "").strip(),
        company=(data.company or "").strip(),
        role=clean_job_role_title(data.role),
        link=(data.link or "").strip(),
        sector=(data.sector or "").strip(),
        required_skills=(data.required_skills or "").strip(),
    )


@app.post("/cv/analyze-for-job")
async def analyze_cv_for_job_endpoint(
    file: UploadFile = File(...),
    user_first_name: str = Form(""),
    user_last_name: str = Form(""),
    description: str = Form(""),
    company: str = Form(""),
    role: str = Form(""),
    role_level: str = Form(""),
    link: str = Form(""),
    sector: str = Form(""),
    required_skills: str = Form(""),
):
    file_bytes = await file.read()
    validation = validate_cv_content(file.filename or "", file_bytes, file.content_type)
    if not validation["is_cv"]:
        raise HTTPException(
            status_code=400,
            detail="Il file caricato non sembra essere un CV valido. Carica un curriculum contenente esperienze, formazione e competenze.",
        )

    cv_text = extract_text_from_file(file, file_bytes)
    cv_check = is_probably_cv(cv_text)
    if not cv_check["is_cv"]:
        raise HTTPException(
            status_code=400,
            detail="Il file caricato non sembra essere un CV valido. Carica un curriculum contenente esperienze, formazione e competenze.",
        )

    identity_check = check_cv_identity(cv_text, user_first_name, user_last_name)
    if identity_check["matches_user"] is False:
        raise HTTPException(status_code=400, detail=identity_check["message"])

    job_validation = validate_job_input(description, company, role, link, sector, required_skills)
    if not job_validation["is_valid"]:
        raise HTTPException(status_code=400, detail=job_validation)
    role = job_validation.get("normalized_role") or role
    company = job_validation.get("normalized_company") or company

    role_context = f"{role} ({role_level.strip()})" if role_level.strip() else role
    direct_job_link = job_validation.get("normalized_link") or link
    sources = search_job_context(company, role_context, direct_job_link) if TAVILY_API_KEY else filter_candidate_sources([
        {"title": f"Annuncio o pagina candidatura: {role_context}", "url": direct_job_link, "content": ""},
        *(job_validation.get("sources", []) or []),
    ])
    cv_evaluation = evaluate_cv_for_job(cv_text, company, role_context, description, link, sources, required_skills)
    cv_evaluation["sources"] = sources

    return {
        "is_valid_cv": True,
        "identity_check": identity_check,
        "job_validation": job_validation,
        "cv_evaluation": cv_evaluation,
        "cv_fingerprint": cv_content_fingerprint(cv_text),
        "warnings": job_validation.get("warnings", []),
    }


@app.post("/users/{user_id}/cv/analyze-for-job")
def analyze_saved_user_cv_for_job(
    user_id: int,
    data: JobValidationRequest,
    authorization: Optional[str] = Header(default=None),
):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()
    existing_user = fetch_user_by_id(cursor, user_id)
    if not existing_user:
        conn.close()
        raise HTTPException(status_code=404, detail="Utente non trovato.")

    cv_text = recover_saved_cv_text(cursor, existing_user)
    conn.commit()
    conn.close()

    if not existing_user[10] or not cv_text:
        raise HTTPException(status_code=400, detail="Carica un CV valido prima di avviare la valutazione.")

    cv_check = is_probably_cv(cv_text)
    warnings = []
    if not cv_check["is_cv"]:
        warnings.append(
            "Il CV salvato non supera uno dei controlli automatici di plausibilità, ma verrà comunque analizzato perché il caricamento iniziale è andato a buon fine."
        )

    public_user = user_to_response(existing_user)
    name_parts = (public_user.get("name") or "").split()
    identity_check = check_cv_identity(
        cv_text,
        name_parts[0] if name_parts else "",
        " ".join(name_parts[1:]) if len(name_parts) > 1 else "",
    )
    identity_warning = None
    if identity_check["matches_user"] is False:
        identity_warning = identity_check["message"]

    description = (data.description or "").strip()
    company = (data.company or "").strip()
    role = clean_job_role_title(data.role)
    role_level = (data.role_level or "").strip()
    link = (data.link or "").strip()
    sector = (data.sector or "").strip()
    required_skills = (data.required_skills or "").strip()
    job_validation = validate_job_input(description, company, role, link, sector, required_skills)
    if not job_validation["is_valid"]:
        raise HTTPException(status_code=400, detail=job_validation)
    role = job_validation.get("normalized_role") or role
    company = job_validation.get("normalized_company") or company
    if not job_validation["is_valid"]:
        warnings.append(job_validation.get("message") or "I dati del lavoro non sono completi o coerenti, ma l'analisi continuerà con i campi disponibili.")

    role_context = f"{role} ({role_level})" if role_level else role
    direct_job_link = job_validation.get("normalized_link") or link
    sources = search_job_context(company, role_context, direct_job_link) if TAVILY_API_KEY else filter_candidate_sources([
        {"title": f"Annuncio o pagina candidatura: {role_context}", "url": direct_job_link, "content": ""},
        *(job_validation.get("sources", []) or []),
    ])
    cv_evaluation = evaluate_cv_for_job(cv_text, company, role_context, description, link, sources, required_skills)
    cv_evaluation["sources"] = sources

    return {
        "is_valid_cv": True,
        "identity_check": identity_check,
        "identity_warning": identity_warning,
        "job_validation": job_validation,
        "cv_evaluation": cv_evaluation,
        "cv_fingerprint": cv_content_fingerprint(cv_text),
        "warnings": [*warnings, *(job_validation.get("warnings", []) or [])],
    }


@app.post("/users/{user_id}/cv-optimize")
def optimize_user_cv(
    user_id: int,
    data: CvOptimizationAnalysisRequest,
    authorization: Optional[str] = Header(default=None),
):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()
    existing_user = fetch_user_by_id(cursor, user_id)

    if not existing_user:
        conn.close()
        raise HTTPException(status_code=404, detail="Utente non trovato.")

    cv_text = recover_saved_cv_text(cursor, existing_user)
    conn.commit()

    if not existing_user[10] or not cv_text:
        conn.close()
        raise HTTPException(status_code=400, detail="Carica un CV valido prima di ottimizzarlo.")

    current_cv_fingerprint = cv_content_fingerprint(cv_text)
    request_cv_fingerprint = (data.cv_fingerprint or "").strip().lower()
    if not request_cv_fingerprint:
        conn.close()
        raise HTTPException(
            status_code=409,
            detail="L'analisi non è associata al CV corrente. Avvia una nuova analisi prima di ottimizzare.",
        )
    if not hmac.compare_digest(request_cv_fingerprint, current_cv_fingerprint):
        conn.close()
        raise HTTPException(
            status_code=409,
            detail="Il CV è cambiato dopo l'analisi. Avvia una nuova analisi per evitare di usare dati o suggerimenti del CV precedente.",
        )

    public_user = user_to_response(existing_user)
    public_user["cv_text"] = cv_text

    job_data = data.job_data or {}
    company = (job_data.get("company") or data.company or "Azienda Generica").strip() or "Azienda Generica"
    role = clean_job_role_title(job_data.get("role") or data.role or public_user.get("target_role"))
    role_level = (job_data.get("role_level") or data.role_level or "").strip()
    role_context = f"{role} ({role_level})" if role_level else role
    goal = (job_data.get("description") or data.goal or "").strip()
    job_link = normalize_public_profile_url(data.job_link or job_data.get("link"))
    required_skills = str(job_data.get("required_skills") or "").strip()
    expected_target_fingerprint = cv_analysis_target_fingerprint(
        role_context, company, goal, required_skills
    )
    evaluation_context = (
        data.cv_evaluation.get("scoring_context", {})
        if isinstance(data.cv_evaluation, dict)
        else {}
    )
    analyzed_target_fingerprint = str(
        evaluation_context.get("target_fingerprint") or ""
    ).strip().lower()
    if (
        not analyzed_target_fingerprint
        or not hmac.compare_digest(analyzed_target_fingerprint, expected_target_fingerprint)
    ):
        conn.close()
        raise HTTPException(
            status_code=409,
            detail=(
                "Ruolo, azienda o descrizione della candidatura sono cambiati dopo l'analisi. "
                "Avvia una nuova analisi prima di ottimizzare il CV."
            ),
        )
    raw_additional_data = dict(data.user_additional_data or {})
    if data.additionalInfo:
        if isinstance(data.additionalInfo, dict):
            raw_additional_data.update(data.additionalInfo)
        elif isinstance(data.additionalInfo, str):
            raw_additional_data["additional_notes"] = data.additionalInfo
    if data.answers and "adaptation_answers" not in raw_additional_data:
        raw_additional_data["adaptation_answers"] = data.answers
    if data.extraAnswers and "adaptation_answers" not in raw_additional_data:
        raw_additional_data["adaptation_answers"] = data.extraAnswers
    confirmed_skill_payload = []
    for payload in [
        data.acceptedSkillConfirmations,
        data.confirmedSkills,
    ]:
        if isinstance(payload, list):
            confirmed_skill_payload.extend(payload)
    if confirmed_skill_payload:
        raw_additional_data["confirmed_skills"] = confirmed_skill_payload
    # === DEBUG cv-optimize: dati extra ricevuti ===
    try:
        print("=" * 80)
        print("[CV-OPT DEBUG] raw_additional_data keys:", sorted(list(raw_additional_data.keys())))
        for _k, _v in raw_additional_data.items():
            _preview = str(_v)
            if len(_preview) > 300:
                _preview = _preview[:300] + "..."
            print(f"[CV-OPT DEBUG] raw_additional_data[{_k!r}] = {_preview}")
        print("=" * 80)
    except Exception as _dbg_exc:
        print(f"[CV-OPT DEBUG] errore nel logging raw_additional_data: {_dbg_exc}")

    user_additional_data, rejected_additional_fields = sanitize_cv_additional_data(raw_additional_data)
    allowed_skill_suggestions = build_role_skill_suggestions(
        cv_text,
        role_context,
        goal,
        required_skills,
    ).get("confirmation_items", [])
    confirmed_from_suggestions = filter_confirmed_skill_suggestions(
        cv_text,
        user_additional_data.get("confirmed_skills", []),
        allowed_skill_suggestions,
    )
    if confirmed_from_suggestions:
        user_additional_data["confirmed_skills"] = confirmed_from_suggestions
    elif user_additional_data.get("confirmed_skills"):
        print(
            "[CV-OPT DEBUG] confirmed_skills kept after sanitize: "
            "nessun match esatto con le suggestion correnti, uso payload confermato dall'utente."
        )

    # === DEBUG cv-optimize: risultati sanitize ===
    try:
        print("[CV-OPT DEBUG] user_additional_data keys (post-sanitize):", sorted(list((user_additional_data or {}).keys())))
        for _k, _v in (user_additional_data or {}).items():
            _preview = str(_v)
            if len(_preview) > 300:
                _preview = _preview[:300] + "..."
            print(f"[CV-OPT DEBUG] user_additional_data[{_k!r}] = {_preview}")
        print(f"[CV-OPT DEBUG] rejected_additional_fields = {rejected_additional_fields}")
        print("=" * 80)
    except Exception as _dbg_exc:
        print(f"[CV-OPT DEBUG] errore nel logging user_additional_data: {_dbg_exc}")

    accepted_suggestions = normalize_accepted_coach_suggestions(data.accepted_suggestions)
    rejected_suggestions = data.rejected_suggestions if isinstance(data.rejected_suggestions, list) else []
    rejected_suggestion_ids = [
        str(suggestion_id).strip()
        for suggestion_id in (data.rejected_suggestion_ids or [])
        if str(suggestion_id).strip()
    ]
    confirmed_skills = user_additional_data.get("confirmed_skills", []) if isinstance(user_additional_data.get("confirmed_skills"), list) else []
    confirmed_skill_suggestions = build_accepted_suggestions_from_confirmed_skills(
        cv_text,
        user_additional_data,
        role_context,
    )
    if confirmed_skill_suggestions:
        seen_suggestion_ids = {
            str(item.get("id") or "").strip()
            for item in accepted_suggestions
            if isinstance(item, dict)
        }
        for suggestion in confirmed_skill_suggestions:
            suggestion_id = str(suggestion.get("id") or "").strip()
            if suggestion_id and suggestion_id in seen_suggestion_ids:
                continue
            accepted_suggestions.append(suggestion)
            if suggestion_id:
                seen_suggestion_ids.add(suggestion_id)
    selected_ids = {
        str(suggestion_id).strip()
        for suggestion_id in ((data.selected_suggestion_ids or []) + (data.acceptedSuggestionIds or []))
        if str(suggestion_id).strip()
    }
    if selected_ids:
        accepted_suggestions = [
            suggestion for suggestion in accepted_suggestions
            if str(suggestion.get("id") or "").strip() in selected_ids
        ]
        if not accepted_suggestions:
            conn.close()
            raise HTTPException(
                status_code=400,
                detail="Nessuna modifica selezionata e applicabile: controlla i suggerimenti o riprova.",
            )

    if not accepted_suggestions and not flatten_cv_support_data(user_additional_data):
        conn.close()
        raise HTTPException(
            status_code=400,
            detail=(
                "Non ci sono modifiche accettate da applicare. "
                "Accetta almeno un suggerimento oppure aggiungi informazioni reali nella schermata Skill/Informazioni."
            ),
        )

    if rejected_additional_fields:
        conn.close()
        raise HTTPException(
            status_code=400,
            detail={
                "message": "Alcune informazioni inserite non sembrano coerenti o utili. Riscrivile con dettagli reali prima di procedere.",
                "rejected_fields": rejected_additional_fields,
            },
        )

    if not role or role.lower() == "da definire":
        conn.close()
        raise HTTPException(status_code=400, detail="Inserisci un ruolo target prima di ottimizzare il CV.")

    application_sources = job_data.get("application_sources") or []
    if isinstance(application_sources, str):
        application_sources = [{"title": "Fonti candidatura", "url": "", "content": application_sources}]
    elif not isinstance(application_sources, list):
        application_sources = []

    sources = filter_candidate_sources(application_sources) or (search_job_context(company, role_context, job_link) if TAVILY_API_KEY else [])
    analysis = data.strategic_analysis if isinstance(data.strategic_analysis, dict) else None
    if not analysis:
        analysis = analyze_cv_strategy(public_user, company, role_context, goal, job_link, sources)
    if sources and not analysis.get("sources"):
        analysis["sources"] = sources

    rewrite_result = build_resume_rewrite_result(
    cv_text=cv_text,
    company=company,
    role=role_context,
    goal=goal,
    accepted_suggestions=accepted_suggestions,
        user_additional_data=user_additional_data,
    )
    optimized_text = rewrite_result["optimized_text"]

    hallucination_warnings = detect_unsupported_optimized_claims(
        optimized_text,
        cv_text,
        user_additional_data,
        company,
        role_context,
        goal,
    )
    original_file_bytes = b""
    if existing_user[14]:
        try:
            original_file_bytes = base64.b64decode(existing_user[14], validate=True)
        except Exception:
            original_file_bytes = b""
    original_filename = (existing_user[10] or "").lower()

    export_service = ExportService()
    alternatives = []
    format_warnings = []
    applied_changes_count = 0
    quality_review: Dict[str, Any] = {
        "ready_to_send": False,
        "score": 0,
        "issues": [],
        "revisions": [],
    }

    if original_filename.endswith(".docx") and original_file_bytes:
        try:
            docx_pipeline = ResumeDocxOptimizationPipeline()
            rewrite_instructions = rewrite_result.get("instructions") or []

            docx_suggestions = [
                {
                    "suggestion_id": instruction.source_id or f"rewrite-instruction-{index + 1}",
                    "target_section": instruction.section,
                    "action": "append",
                    "old_text_hint": "",
                    "new_text": instruction.replacement,
                    "reason": instruction.reason,
                }
                for index, instruction in enumerate(rewrite_instructions)
                if (
                    instruction.replacement
                    and instruction.section
                    and is_additive_user_rewrite_source(instruction.source_id)
                )
            ]
            structured_instructions = docx_pipeline.generate_structured_instructions(
                cv_text=cv_text,
                role=role_context,
                company=company,
                goal=goal,
                accepted_suggestions=docx_suggestions,
                user_additional_data=user_additional_data,
                use_llm=False,
            )
            for i, instruction in enumerate(structured_instructions[:5]):
                print(
                    f"DEBUG DOCX instruction {i}: suggestion_id={instruction.suggestion_id}, "
                    f"source_field={getattr(instruction, 'source_field', '') or '-'}, "
                    f"target_section={instruction.target_section}, "
                    f"llm_target={getattr(instruction, 'llm_target_section', '') or '-'}, "
                    f"override_reason={getattr(instruction, 'section_override_reason', '') or '-'}, "
                    f"action={instruction.action}, old_text_hint={(instruction.old_text_hint or '')[:100]}"
                )

            apply_result = docx_pipeline.apply_instructions_to_docx(original_file_bytes, structured_instructions)
            file_bytes = apply_result.file_bytes
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            extension = "docx"
            applied_changes_count = len(apply_result.applied_ids)
            final_docx_text = apply_result.validation_report.get("final_text", "")
            # === DEBUG cv-optimize: stato applicazione DOCX ===
            try:
                print("[CV-OPT DEBUG] structured_instructions count =", len(structured_instructions))
                _section_counts: Dict[str, int] = {}
                for _inst in structured_instructions:
                    _section_counts[_inst.target_section] = _section_counts.get(_inst.target_section, 0) + 1
                print(f"[CV-OPT DEBUG] structured_instructions per sezione = {_section_counts}")
                print(f"[CV-OPT DEBUG] apply_result.applied_ids count = {len(apply_result.applied_ids)}")
                print(f"[CV-OPT DEBUG] apply_result.validation_report.status = {apply_result.validation_report.get('status')}")
                _changed = normalize_plain_text(final_docx_text) != normalize_plain_text(cv_text)
                print(f"[CV-OPT DEBUG] final_docx_text diverso da cv_text? {_changed}")
                _vr = apply_result.validation_report or {}
                for _vk, _vv in _vr.items():
                    if _vk == "final_text":
                        continue
                    _vp = str(_vv)
                    if len(_vp) > 400:
                        _vp = _vp[:400] + "..."
                    print(f"[CV-OPT DEBUG] validation_report[{_vk!r}] = {_vp}")
                print("=" * 80)
            except Exception as _dbg_exc:
                print(f"[CV-OPT DEBUG] errore nel logging DOCX apply: {_dbg_exc}")

            if normalize_plain_text(final_docx_text) == normalize_plain_text(cv_text):
                print("[CV-OPT DEBUG] 422 -> testo finale identico al CV originale")
                raise HTTPException(
                    status_code=422,
                    detail="Il motore non è riuscito ad applicare modifiche reali al DOCX. Il file originale non è stato salvato come ottimizzato.",
                )
            if apply_result.validation_report.get("status") == "failed":
                print("[CV-OPT DEBUG] 422 -> validation_report.status = failed")
                raise HTTPException(
                    status_code=422,
                    detail={
                        "message": "Il DOCX non contiene modifiche applicabili o non conserva correttamente la struttura originale.",
                        "validation_report": apply_result.validation_report,
                    },
                )
        except HTTPException:
            raise
        except Exception as exc:
            print(f"Errore nella pipeline DOCX strutturata: {exc}")
            raise HTTPException(status_code=500, detail="Impossibile ottimizzare il DOCX originale in modo sicuro.")

    final_cv_text = apply_result.validation_report.get("final_text", "") if original_filename.endswith(".docx") and original_file_bytes else optimized_text
    if not final_cv_text.strip():
        print("[CV-OPT DEBUG] 422 -> final_cv_text vuoto")
        raise HTTPException(status_code=422, detail="Il CV finale non contiene testo valido.")

    quality_review = review_generated_cv_quality(
        final_text=final_cv_text,
        original_cv_text=cv_text,
        role=role_context,
        company=company,
        accepted_instructions=rewrite_result.get("instructions") or [],
    )
    # === Quality review: blocchiamo SOLO se c'e' un problema reale ===
    # Il revisore LLM (qwen2.5:1.5b locale) tende a segnalare problemi minori
    # o inesistenti anche su CV ottimizzati correttamente. Per non bloccare
    # l'utente senza motivo, applichiamo la politica:
    #  - score < 50  -> blocco 422 (CV davvero scadente)
    #  - issues 'critical' presenti -> blocco 422
    #  - altrimenti il CV passa, le issues diventano warning informativi
    _qr_issues = quality_review.get("issues") if isinstance(quality_review.get("issues"), list) else []
    _qr_score_raw = quality_review.get("score") or 0
    try:
        _qr_score = int(_qr_score_raw)
    except (TypeError, ValueError):
        _qr_score = 0
    _has_critical = any(
        isinstance(_iss, dict) and str(_iss.get("severity", "")).lower() == "critical"
        for _iss in _qr_issues
    )
    _quality_blocking = _has_critical or _qr_score < 50

    if _quality_blocking:
        print(
            f"[CV-OPT] 422 quality_review bloccante: score={_qr_score}, "
            f"critical_issues={_has_critical}, total_issues={len(_qr_issues)}"
        )
        try:
            for _idx, _iss in enumerate(_qr_issues):
                if isinstance(_iss, dict):
                    print(
                        f"[CV-OPT] issue #{_idx} severity={_iss.get('severity')!r} "
                        f"section={_iss.get('section')!r} description={str(_iss.get('description'))[:300]!r}"
                    )
            _revs = quality_review.get("revisions") or []
            for _idx, _rev in enumerate(_revs):
                if isinstance(_rev, dict):
                    print(
                        f"[CV-OPT] revision #{_idx} section={_rev.get('section')!r} "
                        f"new_text_preview={str(_rev.get('new_text') or _rev.get('text') or '')[:200]!r}"
                    )
        except Exception as _dbg_exc:
            print(f"[CV-OPT] errore dump quality_review: {_dbg_exc}")
        raise HTTPException(
            status_code=422,
            detail={
                "message": "Il motore di ottimizzazione non ha prodotto un CV pronto da inviare.",
                "quality_review": quality_review,
            },
        )

    # Se non blocchiamo ma il revisore aveva detto ready_to_send=False,
    # registriamo lo skip e forziamo ready_to_send=True per il frontend,
    # conservando le issues come warning consultabili.
    if not quality_review.get("ready_to_send", False):
        print(
            f"[CV-OPT] quality_review non-bloccante: score={_qr_score}, "
            f"issues={len(_qr_issues)} - CV consegnato con warning"
        )
        quality_review["ready_to_send"] = True
        quality_review["warnings_only"] = True

    filename = get_target_optimized_cv_filename(public_user.get("name", "CV"), role_context, company, extension)
    file_base64 = base64.b64encode(file_bytes).decode("ascii") if file_bytes else ""
    docx_file = {
        "filename": filename,
        "content_type": content_type,
        "file_base64": file_base64,
    } if extension == "docx" else None
    pdf_file = {
        "filename": filename,
        "content_type": content_type,
        "file_base64": file_base64,
    } if extension == "pdf" else None
    for alternative in alternatives:
        alt_name = str(alternative.get("filename") or "").lower()
        if alt_name.endswith(".pdf"):
            pdf_file = alternative
        elif alt_name.endswith(".docx"):
            docx_file = alternative
    scoring_required_skills = required_skills
    initial_scorecard = build_deterministic_cv_scorecard(
        cv_text,
        company,
        role_context,
        goal,
        scoring_required_skills,
    )
    optimized_score_text = final_cv_text or optimized_text
    optimized_scorecard = build_deterministic_cv_scorecard(
        optimized_score_text,
        company,
        role_context,
        goal,
        scoring_required_skills,
    )
    score_comparison = compare_cv_scorecards(initial_scorecard, optimized_scorecard)
    analysis_score = optimized_scorecard["overall_score"]
    grouped_changes = rewrite_result.get("grouped_changes", {})
    skipped_change_details = []
    skipped_change_details.extend({"reason": warning} for warning in format_warnings)

    cursor.execute("""
    INSERT INTO optimized_cvs (
        user_id, filename, content_type, text, file_base64,
        target_role, target_company, job_description, analysis_score,
        selected_suggestions_json, rejected_suggestions_json, additional_info_json,
        confirmed_skills_json, generation_status, applied_changes_count, skipped_changes_json,
        docx_filename, docx_content_type, docx_file_base64,
        pdf_filename, pdf_content_type, pdf_file_base64
    )
    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        user_id, filename, content_type, optimized_text, file_base64,
        role, company, goal, analysis_score,
        json.dumps(accepted_suggestions, ensure_ascii=False),
        json.dumps({"ids": rejected_suggestion_ids, "suggestions": rejected_suggestions}, ensure_ascii=False),
        json.dumps(user_additional_data, ensure_ascii=False),
        json.dumps(confirmed_skills, ensure_ascii=False),
        "completed",
        applied_changes_count,
        json.dumps(skipped_change_details, ensure_ascii=False),
        docx_file.get("filename") if docx_file else None,
        docx_file.get("content_type") if docx_file else None,
        docx_file.get("file_base64") if docx_file else None,
        pdf_file.get("filename") if pdf_file else None,
        pdf_file.get("content_type") if pdf_file else None,
        pdf_file.get("file_base64") if pdf_file else None,
    ))
    optimized_cv_id = cursor.lastrowid
    conn.commit()
    conn.close()

    generated_at = datetime.utcnow().isoformat() + "Z"
    optimized_cv_payload = {
            "id": optimized_cv_id,
            "filename": filename,
            "file_name": filename,
            "content_type": content_type,
            "file_base64": file_base64,
            "text": optimized_text,
            "previewFinalCvContent": rewrite_result.get("previewFinalCvContent", {}),
            "download_url": f"/users/{user_id}/optimized-cvs/{optimized_cv_id}/file",
            "docx_url": f"/users/{user_id}/optimized-cvs/{optimized_cv_id}/file?format=docx" if docx_file else None,
            "pdf_url": f"/users/{user_id}/optimized-cvs/{optimized_cv_id}/file?format=pdf" if pdf_file else None,
            "alternatives": alternatives,
            "applied_changes_count": applied_changes_count,
            "skipped_changes": skipped_change_details,
            "generated_at": generated_at,
            "created_at": generated_at,
            "target_role": role,
            "role": role,
            "target_company": company,
            "company": company,
            "analysis_score": analysis_score,
            "score": analysis_score,
            "analysis": optimized_scorecard,
            "score_comparison": score_comparison,
            "has_docx": bool(docx_file),
            "has_pdf": bool(pdf_file),
            "generation_status": "completed",
            "quality_review": quality_review,
        }
    return {
        "success": True,
        "optimized_cv": optimized_cv_payload,
        "optimizedCv": optimized_cv_payload,
        "candidate_sources": sources,
        "analysis": analysis,
        "optimized_analysis": optimized_scorecard,
        "score_comparison": score_comparison,
        "hallucination_warnings": hallucination_warnings,
        "format_warnings": format_warnings,
        "quality_review": quality_review,
        "accepted_suggestions": accepted_suggestions,
        "rejected_suggestion_ids": rejected_suggestion_ids,
        "skipped_changes": skipped_change_details,
        "grouped_changes": grouped_changes,
        "appliedChanges": {
            "profile": grouped_changes.get("profile_updates", []),
            "hard_skills": grouped_changes.get("hard_skills_updates", []),
            "soft_skills": grouped_changes.get("soft_skills_updates", []),
            "experience": grouped_changes.get("experience_updates", []),
            "education": grouped_changes.get("education_updates", []),
            "projects": grouped_changes.get("projects_updates", []),
            "extra_sections": grouped_changes.get("extra_sections", []),
        },
        "applied_changes_count": applied_changes_count,
        "warnings": [*hallucination_warnings, *skipped_change_details],
        "message": "CV ottimizzato generato con alcune avvertenze." if format_warnings else "CV ottimizzato generato correttamente.",
    }


@app.get("/users/{user_id}/optimized-cvs/{optimized_cv_id}/file")
def download_optimized_cv(
    user_id: int,
    optimized_cv_id: int,
    format: Optional[str] = None,
    authorization: Optional[str] = Header(default=None),
):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("""
    SELECT filename, content_type, file_base64,
           docx_filename, docx_content_type, docx_file_base64,
           pdf_filename, pdf_content_type, pdf_file_base64
    FROM optimized_cvs
    WHERE id = ? AND user_id = ?
    """, (optimized_cv_id, user_id))
    row = cursor.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="CV ottimizzato non trovato.")

    filename, content_type, file_base64, docx_filename, docx_content_type, docx_file_base64, pdf_filename, pdf_content_type, pdf_file_base64 = row
    requested_format = (format or "").lower().strip(".")
    
    if requested_format == "docx" and docx_file_base64:
        filename, content_type, file_base64 = docx_filename, docx_content_type, docx_file_base64
    elif requested_format == "pdf" and pdf_file_base64:
        filename, content_type, file_base64 = pdf_filename, pdf_content_type, pdf_file_base64

    if not filename or not file_base64:
        raise HTTPException(status_code=500, detail="File CV ottimizzato non leggibile.")

    try:
        file_bytes = base64.b64decode(file_base64, validate=True)
    except Exception:
        raise HTTPException(status_code=500, detail="File CV ottimizzato non leggibile.")

    return Response(
        content=file_bytes,
        media_type=content_type or "application/octet-stream",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        },
    )


@app.get("/users/{user_id}/optimized-cvs/{optimized_cv_id}/download-docx")
def download_optimized_cv_docx(
    user_id: int,
    optimized_cv_id: int,
    authorization: Optional[str] = Header(default=None),
):
    return download_optimized_cv(user_id, optimized_cv_id, format="docx", authorization=authorization)


@app.get("/users/{user_id}/optimized-cvs/{optimized_cv_id}/download-pdf")
def download_optimized_cv_pdf(
    user_id: int,
    optimized_cv_id: int,
    authorization: Optional[str] = Header(default=None),
):
    return download_optimized_cv(user_id, optimized_cv_id, format="pdf", authorization=authorization)


@app.get("/users/{user_id}/optimized-cvs")
def list_optimized_cvs(user_id: int, authorization: Optional[str] = Header(default=None)):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM users WHERE id = ?", (user_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="Utente non trovato.")
    
    cursor.execute("""
    SELECT id, filename, content_type, created_at, target_role, target_company,
           job_description, analysis_score, docx_file_base64, pdf_file_base64, generation_status
    FROM optimized_cvs
    WHERE user_id = ?
    ORDER BY created_at DESC
    """, (user_id,))
    rows = cursor.fetchall()
    conn.close()

    return {
        "optimized_cvs": [
            {
                "id": row[0],
                "filename": row[1],
                "content_type": row[2],
                "created_at": row[3],
                "target_role": row[4],
                "target_company": row[5],
                "job_description": row[6],
                "analysis_score": row[7],
                "has_docx": bool(row[8]),
                "has_pdf": bool(row[9]),
                "generation_status": row[10] or "completed",
                "docx_download_url": f"/users/{user_id}/optimized-cvs/{row[0]}/file?format=docx" if row[8] else None,
                "pdf_download_url": f"/users/{user_id}/optimized-cvs/{row[0]}/file?format=pdf" if row[9] else None,
                "download_url": f"/users/{user_id}/optimized-cvs/{row[0]}/file",
            }
            for row in rows
        ]
    }


@app.delete("/users/{user_id}/optimized-cvs/{optimized_cv_id}")
def delete_optimized_cv(
    user_id: int,
    optimized_cv_id: int,
    authorization: Optional[str] = Header(default=None),
):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM optimized_cvs WHERE id = ? AND user_id = ?", (optimized_cv_id, user_id))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="CV ottimizzato non trovato.")
    
    cursor.execute("DELETE FROM optimized_cvs WHERE id = ? AND user_id = ?", (optimized_cv_id, user_id))
    conn.commit()
    conn.close()

    return {"message": "CV ottimizzato eliminato correttamente."}


@app.patch("/users/{user_id}/optimized-cvs/{optimized_cv_id}")
def update_optimized_cv(
    user_id: int,
    optimized_cv_id: int,
    update_data: Dict[str, str],
    authorization: Optional[str] = Header(default=None),
):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT filename FROM optimized_cvs WHERE id = ? AND user_id = ?", (optimized_cv_id, user_id))
    row = cursor.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="CV ottimizzato non trovato.")
    
    filename = update_data.get("filename", row[0])
    text = update_data.get("text")
    
    if text is not None:
        cursor.execute("""
        UPDATE optimized_cvs
        SET filename = ?, text = ?
        WHERE id = ? AND user_id = ?
        """, (filename, text, optimized_cv_id, user_id))
    else:
        cursor.execute("""
        UPDATE optimized_cvs
        SET filename = ?
        WHERE id = ? AND user_id = ?
        """, (filename, optimized_cv_id, user_id))
    
    conn.commit()
    conn.close()

    return {"message": "CV ottimizzato aggiornato correttamente.", "filename": filename}


# Legacy endpoint for backward compatibility
@app.get("/users/{user_id}/cv-optimized-file")
def download_optimized_cv_legacy(user_id: int, authorization: Optional[str] = Header(default=None)):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("""
    SELECT id, filename, content_type, file_base64
    FROM optimized_cvs
    WHERE user_id = ?
    ORDER BY created_at DESC
    LIMIT 1
    """, (user_id,))
    row = cursor.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="Genera prima il CV ottimizzato.")

    optimized_cv_id, filename, content_type, file_base64 = row
    if not filename or not file_base64:
        raise HTTPException(status_code=500, detail="File CV ottimizzato non leggibile.")

    try:
        file_bytes = base64.b64decode(file_base64, validate=True)
    except Exception:
        raise HTTPException(status_code=500, detail="File CV ottimizzato non leggibile.")

    return Response(
        content=file_bytes,
        media_type=content_type or "application/octet-stream",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        },
    )


def is_visual_service_unavailable_error(exc: Exception) -> bool:
    message = str(exc).lower()
    if isinstance(exc, requests.exceptions.RequestException):
        return True
    if isinstance(exc, (MaxRetryError, NewConnectionError)):
        return True
    if "httpconnectionpool" in message or "connection" in message or "connessione" in message:
        return True
    return False


def detect_disallowed_cv_content(text: str) -> Dict[str, Any]:
    normalized = normalize_for_cv_detection(text)
    disallowed_patterns = [
        r"\bporn(?:o|ografia)?\b",
        r"\bnud(?:o|ità)\b",
        r"\bnaked\b",
        r"\bsex\b",
        r"\bsesso\b",
        r"\berotic(?:o|a)?\b",
        r"\bviol(?:en[za]|ento)\b",
        r"\b(?:sangue|gore|ferit[oa]?|omicid[io]|assass|arma|pistola|coltello|bomba|terror(?:ismo|ista)?)\b",
        r"\b(?:drugs?|droga|cocaina|eroina|hashish|marijuana|spaccio)\b",
        r"\b(?:razz(?:is(?:ta|mo)|ismo)|odio|hate|genocidio)\b",
        r"\b(?:cazzo|merda|stronzo|stronza|troia|puttana|coglione|vaffanculo|figa|fottiti|frocio|negro|bastardo)\b",
    ]
    for pattern in disallowed_patterns:
        if re.search(pattern, normalized):
            return {
                "blocked": True,
                "reason": "Il CV contiene contenuti non idonei alle linee guida.",
            }
    return {"blocked": False}


def is_gibberish_cv_text(text: str) -> bool:
    words = text.split()
    if not words:
        return True
    
    vowels = "aeiouyAEIOUY"
    weird_words = 0
    total_words = 0
    
    for word in words:
        clean_word = "".join(ch for ch in word if ch.isalpha())
        if not clean_word:
            continue
            
        total_words += 1
        
        if len(clean_word) >= 5 and not any(v in clean_word for v in vowels):
            weird_words += 1
            continue
            
        consecutive_consonants = 0
        max_consecutive_consonants = 0
        for ch in clean_word:
            if ch.isalpha() and ch not in vowels:
                consecutive_consonants += 1
                max_consecutive_consonants = max(max_consecutive_consonants, consecutive_consonants)
            else:
                consecutive_consonants = 0
                
        if max_consecutive_consonants >= 5:
            weird_words += 1
            continue

    if total_words == 0:
        return True
        
    weird_ratio = weird_words / total_words
    return weird_ratio >= 0.4

def validate_cv_content(filename: str, file_bytes: bytes, content_type: Optional[str] = None) -> Dict:
    filename = filename.strip()
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    allowed_extensions = {"pdf", "docx"}

    if not filename or extension not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Carica un file PDF o DOCX.")

    if len(file_bytes) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Il CV non puo superare 5MB.")

    debug = {
        "filename": filename,
        "content_type": content_type,
        "file_size": len(file_bytes),
        "extension": extension,
        "text_length": 0,
        "extraction_method": "",
    }

    if not file_bytes:
        return {
            "is_cv": False,
            "confidence": 0,
            "reason": "Il file e vuoto.",
            "detected_sections": [],
            "visual_validation": {
                "status": "no_content",
                "blocked": False,
            },
            "debug": debug,
        }

    extracted_text, extraction_method = extract_text_from_file_bytes(file_bytes, filename)
    normalized_text = clean_extracted_text(extracted_text)
    debug["text_length"] = len(normalized_text)
    debug["extraction_method"] = extraction_method

    if len(normalized_text) == 0:
        print(
            f"CV validation rejected: file={filename}, extraction_method={extraction_method}, text_length=0, reason=no_text"
        )
        return {
            "is_cv": False,
            "confidence": 0,
            "reason": "Non riesco a estrarre testo dal file. Potrebbe essere un PDF scannerizzato o composto da immagini.",
            "detected_sections": [],
            "visual_validation": {
                "status": "no_text",
                "blocked": False,
            },
            "debug": debug,
        }

    print(
        f"CV validation inspected: file={filename}, extension={extension}, extraction_method={extraction_method}, text_length={len(normalized_text)}"
    )

    disallowed = detect_disallowed_cv_content(normalized_text)
    if disallowed["blocked"]:
        print(f"CV validation rejected: disallowed_content=True, reason={disallowed['reason']}")
        return {
            "is_cv": False,
            "confidence": 0,
            "reason": disallowed["reason"],
            "detected_sections": [],
            "visual_validation": {
                "status": "disallowed_content",
                "blocked": False,
            },
            "debug": debug,
        }

    if is_gibberish_cv_text(normalized_text):
        return {
            "is_cv": False,
            "confidence": 0,
            "reason": "Il testo sembra contenere parole casuali o non ha senso logico.",
            "detected_sections": [],
            "visual_validation": {
                "status": "gibberish_text",
                "blocked": False,
            },
            "debug": debug,
        }

    heuristic = analyze_cv_heuristics(normalized_text)
    heuristic_score = heuristic["score"]

    if len(normalized_text) < 50 and heuristic_score < 35:
        return {
            "is_cv": False,
            "confidence": clamp_score(heuristic_score),
            "reason": "Il testo estratto e molto breve e non contiene abbastanza elementi tipici di un curriculum.",
            "detected_sections": heuristic["detected_sections"],
            "visual_validation": {
                "status": "short_text",
                "blocked": False,
            },
            "debug": debug,
        }

    strong_cv_sections = {"contatti", "formazione", "esperienze professionali", "competenze", "lingue"}
    strong_section_count = len(strong_cv_sections.intersection(set(heuristic["detected_sections"])))

    if heuristic_score >= 45:
        is_cv = True
        confidence = heuristic_score
        reason = heuristic["reason"]
    elif heuristic_score >= 35 and strong_section_count >= 2:
        is_cv = True
        confidence = heuristic_score
        reason = heuristic["reason"]
    else:
        is_cv = False
        confidence = heuristic_score
        reason = "Il documento e leggibile, ma non contiene abbastanza elementi tipici di un curriculum."

    if not is_cv:
        print(
            f"CV validation rejected: file={filename}, heuristic_score={heuristic_score}, detected_sections={heuristic['detected_sections']}, reason={reason}"
        )
        return {
            "is_cv": False,
            "confidence": clamp_score(confidence),
            "reason": reason,
            "detected_sections": heuristic["detected_sections"],
            "visual_validation": {
                "status": "skipped",
                "blocked": False,
            },
            "debug": debug,
        }

    visual_validation = {
        "status": "not_analyzed",
        "blocked": False,
    }
    visual_warning = None
    try:
        visual_validation = validate_cv_images(
            filename,
            file_bytes,
            analyze_embedded_cv_image,
        )
    except Exception as exc:
        print(
            f"Controllo immagini non completato perché il servizio visuale non e disponibile: {exc}"
        )
        visual_validation = {
            "status": "analysis_failed",
            "image_count": 0,
            "analyzed_count": 0,
            "blocked": False,
            "message": "Controllo immagini non completato perché il servizio visuale non e disponibile.",
        }
        visual_warning = (
            "Il CV e stato letto correttamente. Il controllo automatico delle immagini non e disponibile al momento, ma il caricamento puo proseguire."
        )

    if visual_validation.get("blocked"):
        categories = ", ".join(visual_validation.get("blocked_categories") or [])
        print(f"CV validation rejected: visual_blocked=True, categories={categories}")
        return {
            "is_cv": False,
            "confidence": 0,
            "reason": (
                "Il CV contiene immagini o contenuti non idonei alle linee guida"
                + (f": {categories}." if categories else ".")
            ),
            "detected_sections": heuristic["detected_sections"],
            "visual_validation": visual_validation,
            "debug": debug,
        }

    return {
        "is_cv": True,
        "confidence": clamp_score(confidence),
        "reason": visual_warning or reason,
        "detected_sections": heuristic["detected_sections"],
        "visual_validation": visual_validation,
        "debug": debug,
    }


@app.post("/validate-cv-file")
async def validate_cv_file(file: UploadFile = File(...)):
    file_bytes = await file.read()
    return validate_cv_content(file.filename or "", file_bytes, file.content_type)


@app.post("/users/{user_id}/cv")
def upload_user_cv(
    user_id: int,
    data: UserCvUpload,
    authorization: Optional[str] = Header(default=None),
):
    require_user_session(user_id, authorization)
    filename = data.filename.strip()
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    allowed_extensions = {"pdf", "docx"}

    if not filename or extension not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Carica un file PDF o DOCX.")

    if data.size > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Il CV non puo superare 5MB.")

    if not data.file_base64:
        raise HTTPException(status_code=400, detail="Il contenuto del CV non e disponibile.")

    try:
        file_bytes = base64.b64decode(data.file_base64, validate=True)
    except Exception:
        raise HTTPException(status_code=400, detail="Il contenuto del CV non e valido.")

    if len(file_bytes) != data.size:
        raise HTTPException(status_code=400, detail="Il file ricevuto non corrisponde al CV selezionato.")

    validation = validate_cv_content(filename, file_bytes, data.content_type)
    if not validation["is_cv"]:
        raise HTTPException(
            status_code=400,
            detail=f"Il file caricato non sembra essere un CV. {validation['reason']}",
        )

    extracted_text, _extraction_method = extract_text_from_file_bytes(file_bytes, filename)

    conn = get_connection()
    cursor = conn.cursor()
    existing_user = fetch_user_by_id(cursor, user_id)
    if not existing_user:
        conn.close()
        raise HTTPException(status_code=404, detail="Utente non trovato.")

    profile_name_parts = (existing_user[1] or "").split()
    first_name = profile_name_parts[0] if profile_name_parts else ""
    last_name = " ".join(profile_name_parts[1:]) if len(profile_name_parts) > 1 else ""
    identity_check = check_cv_identity(extracted_text, first_name, last_name)
    if identity_check["matches_user"] is False:
        conn.close()
        raise HTTPException(status_code=400, detail=identity_check["message"])

    cursor.execute("""
    UPDATE users
    SET cv_filename = ?,
        cv_content_type = ?,
        cv_size = ?,
        cv_text = ?,
        cv_file_base64 = ?,
        cv_uploaded_at = CURRENT_TIMESTAMP,
        education = CASE WHEN education IS NULL OR trim(education) = '' THEN 'Da CV caricato' ELSE education END,
        target_role = CASE WHEN target_role IS NULL OR trim(target_role) = '' THEN 'Da definire' ELSE target_role END,
        sector = CASE WHEN sector IS NULL OR trim(sector) = '' THEN 'Da definire' ELSE sector END,
        experience_level = CASE WHEN experience_level IS NULL OR trim(experience_level) = '' THEN 'Junior' ELSE experience_level END,
        interview_language = CASE WHEN interview_language IS NULL OR trim(interview_language) = '' THEN 'Italiano' ELSE interview_language END
    WHERE id = ?
    """, (
        filename,
        data.content_type,
        data.size,
        clean_extracted_text(extracted_text or data.text or "")[:20000],
        data.file_base64 or "",
        user_id,
    ))

    conn.commit()
    user = fetch_user_by_id(cursor, user_id)
    conn.close()

    return {
        "user": user_to_response(user),
        "message": "CV salvato nel profilo.",
        "identity_check": identity_check,
    }


@app.get("/users/{user_id}/cv-file")
def get_user_cv_file(user_id: int, authorization: Optional[str] = Header(default=None)):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()
    user = fetch_user_by_id(cursor, user_id)
    conn.close()

    if not user:
        raise HTTPException(status_code=404, detail="Utente non trovato.")

    if not user[10] or not user[14]:
        raise HTTPException(status_code=404, detail="CV non trovato.")

    return {
        "filename": user[10],
        "content_type": user[11] or "application/octet-stream",
        "size": user[12],
        "text": user[13] or "",
        "file_base64": user[14],
        "uploaded_at": user[15],
    }

@app.post("/users/{user_id}/linkedin-profile")
async def upload_linkedin_profile(
    user_id: int,
    file: UploadFile = File(...),
    authorization: Optional[str] = Header(default=None),
):
    require_user_session(user_id, authorization)
    filename = (file.filename or "").strip()
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    allowed_extensions = {"pdf", "docx"}

    if not filename or extension not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Carica l'esportazione LinkedIn in formato PDF o DOCX.")

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Il file LinkedIn e vuoto.")

    if len(file_bytes) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="L'esportazione LinkedIn non puo superare 5MB.")

    extracted_text, _extraction_method = extract_text_from_file_bytes(file_bytes, filename)
    normalized_text = clean_extracted_text(extracted_text)
    if len(normalized_text) < 80:
        raise HTTPException(
            status_code=400,
            detail="Non riesco a leggere abbastanza testo dal file LinkedIn. Esporta il profilo come PDF oppure usa un file DOCX.",
        )

    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM users WHERE id = ?", (user_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="Utente non trovato.")

    cursor.execute("""
    UPDATE users
    SET linkedin_profile_filename = ?,
        linkedin_profile_text = ?,
        digital_analysis_json = NULL
    WHERE id = ?
    """, (filename, normalized_text[:30000], user_id))
    conn.commit()
    user = fetch_user_by_id(cursor, user_id)
    conn.close()

    return {
        "user": user_to_response(user),
        "message": "Esportazione LinkedIn caricata correttamente.",
    }


@app.delete("/users/{user_id}/linkedin-profile")
def delete_linkedin_profile(user_id: int, authorization: Optional[str] = Header(default=None)):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM users WHERE id = ?", (user_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="Utente non trovato.")

    cursor.execute("""
    UPDATE users
    SET linkedin_profile_filename = NULL,
        linkedin_profile_text = NULL,
        digital_analysis_json = NULL
    WHERE id = ?
    """, (user_id,))
    conn.commit()
    user = fetch_user_by_id(cursor, user_id)
    conn.close()

    return {
        "user": user_to_response(user),
        "message": "Esportazione LinkedIn rimossa.",
    }


@app.get("/users/{user_id}/official-profiles")
def get_user_official_profiles(user_id: int, authorization: Optional[str] = Header(default=None)):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()
    existing_user = fetch_user_by_id(cursor, user_id)
    conn.close()

    if not existing_user:
        raise HTTPException(status_code=404, detail="Utente non trovato.")

    public_user = user_to_response(existing_user)
    public_user["linkedin_profile_text"] = existing_user[21] or ""
    official_sources = build_official_profile_sources(public_user)

    return {
        "sources": official_sources,
        "source_count": len(official_sources),
        "capabilities": OFFICIAL_PROFILE_CAPABILITIES,
        "message": (
            "Fonti OAuth/API ufficiali standardizzate recuperate."
            if official_sources
            else "Nessun profilo ufficiale collegato via OAuth/API."
        ),
    }


@app.post("/users/{user_id}/cv-optimization-analysis")
def analyze_user_cv_for_optimization(
    user_id: int,
    data: CvOptimizationAnalysisRequest,
    authorization: Optional[str] = Header(default=None),
):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()
    existing_user = fetch_user_by_id(cursor, user_id)
    conn.close()

    if not existing_user:
        raise HTTPException(status_code=404, detail="Utente non trovato.")

    if not existing_user[10]:
        raise HTTPException(status_code=400, detail="Carica un CV prima di avviare l'ottimizzazione.")

    public_user = user_to_response(existing_user)
    public_user["cv_text"] = existing_user[13] or ""

    company = (data.company or "Azienda Generica").strip() or "Azienda Generica"
    role = (data.role or public_user.get("target_role") or "Ruolo da definire").strip() or "Ruolo da definire"
    role_level = (data.role_level or "").strip()
    role_context = f"{role} ({role_level})" if role_level else role
    goal = (data.goal or "").strip()
    job_link = normalize_public_profile_url(data.job_link)
    sources = search_job_context(company, role_context, job_link)
    analysis = analyze_cv_strategy(public_user, company, role_context, goal, job_link, sources)

    return {
        "analysis": analysis,
        "message": "Analisi strategica CV completata.",
    }


@app.put("/users/{user_id}/digital-presence")
def update_digital_presence(
    user_id: int,
    data: DigitalPresenceUpdate,
    authorization: Optional[str] = Header(default=None),
):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()
    existing_user = fetch_user_by_id(cursor, user_id)
    if not existing_user:
        conn.close()
        raise HTTPException(status_code=404, detail="Utente non trovato.")

    public_user = user_to_response(existing_user)
    public_user["cv_text"] = recover_saved_cv_text(cursor, existing_user)
    public_user["linkedin_profile_text"] = existing_user[21] or ""
    public_user["target_role"] = clean_job_role_title(
        data.target_role or public_user.get("target_role")
    )
    public_user["linkedin_url"] = normalize_linkedin_profile_url(data.linkedin_url)
    public_user["portfolio_url"] = (data.portfolio_url or "").strip()
    instagram_handle = normalize_instagram_handle(data.instagram_handle)
    public_user["instagram_handle"] = f"@{instagram_handle}" if instagram_handle else ""
    public_user["visual_media_analysis"] = analyze_public_social_media(public_user)

    sources = search_public_profile_signals(public_user, data)
    digital_analysis = analyze_digital_profile(public_user, sources)
    previous_analysis = json.loads(existing_user[19]) if existing_user[19] else {}
    previous_evidence = previous_analysis.get("analysis_evidence", {})
    previous_profile_analyses = dict(previous_evidence.get("visual_media_analyses", {}))
    previous_text_analyses = dict(previous_evidence.get("social_text_analyses", {}))
    same_target_role = (
        normalize_plain_text(previous_evidence.get("target_role"))
        == normalize_plain_text(public_user.get("target_role"))
    )
    same_instagram_profile = (
        normalize_instagram_handle(existing_user[18])
        == normalize_instagram_handle(public_user["instagram_handle"])
    )
    if not same_instagram_profile:
        previous_profile_analyses.pop("instagram", None)
        previous_text_analyses.pop("instagram", None)
    elif not same_target_role:
        for profile_type, analysis in previous_text_analyses.items():
            ocr = analysis.get("ocr") or {}
            analysis["evaluation"] = evaluate_social_profile_text(
                ocr.get("extracted_text", ""),
                profile_type,
                public_user,
            )
    if previous_profile_analyses or previous_text_analyses:
        evidence = digital_analysis.setdefault("analysis_evidence", {})
        evidence["visual_media_analyses"] = previous_profile_analyses
        evidence["social_text_analyses"] = previous_text_analyses
        evidence["social_screenshot_batches"] = list(previous_evidence.get("social_screenshot_batches", []))
        evidence["profile_screenshots_analyzed"] = previous_evidence.get(
            "profile_screenshots_analyzed", sorted(previous_profile_analyses)
        )
        evidence["instagram_media_analyzed"] = (
            previous_evidence.get("instagram_media_analyzed", False)
            if same_instagram_profile
            else False
        )
        evidence["instagram_bio_analyzed"] = (
            previous_evidence.get("instagram_bio_analyzed", False)
            if same_instagram_profile
            else False
        )
        evidence["instagram_identity_check"] = (
            previous_evidence.get("instagram_identity_check", {})
            if same_instagram_profile
            else {}
        )
        evidence["visual_score_adjustment"] = int(
            previous_evidence.get("visual_score_adjustment", 0) or 0
        )
        if previous_evidence.get("visual_media_analysis"):
            evidence["visual_media_analysis"] = previous_evidence["visual_media_analysis"]
        digital_analysis["score"] = compute_digital_presence_score(evidence)

        preserved_titles = {"foto e contenuti pubblici", "bio e testo profilo"}
        preserved_findings = [
            finding
            for finding in previous_analysis.get("findings", [])
            if str(finding.get("title", "")).strip().lower() in preserved_titles
        ]
        if preserved_findings:
            digital_analysis["findings"] = [
                finding
                for finding in digital_analysis.get("findings", [])
                if str(finding.get("title", "")).strip().lower() not in preserved_titles
            ]
            digital_analysis["findings"].extend(preserved_findings)
    digital_analysis_json = json.dumps(digital_analysis, ensure_ascii=False)

    cursor.execute("""
    UPDATE users
    SET linkedin_url = ?,
        portfolio_url = ?,
        instagram_handle = ?,
        target_role = ?,
        digital_analysis_json = ?
    WHERE id = ?
    """, (
        public_user["linkedin_url"],
        public_user["portfolio_url"],
        public_user["instagram_handle"],
        public_user["target_role"],
        digital_analysis_json,
        user_id,
    ))

    conn.commit()
    user = fetch_user_by_id(cursor, user_id)
    conn.close()

    return {
        "user": user_to_response(user),
        "analysis": digital_analysis,
        "message": "Presenza digitale salvata.",
    }


@app.post("/users/{user_id}/social-screenshots")
async def analyze_social_screenshots(
    user_id: int,
    profile_type: str = Form("instagram"),
    instagram_handle: str = Form(""),
    files: List[UploadFile] = File(...),
    authorization: Optional[str] = Header(default=None),
):
    require_user_session(user_id, authorization)
    profile_type = profile_type.strip().lower()
    if profile_type not in VISUAL_PROFILE_LABELS:
        raise HTTPException(status_code=400, detail="Tipo di profilo non valido.")
    if not files or len(files) > 8:
        raise HTTPException(status_code=400, detail="Carica da 1 a 8 screenshot.")

    image_inputs = []
    for file in files:
        content_type = (file.content_type or "").lower()
        if content_type not in {"image/jpeg", "image/png", "image/webp"}:
            raise HTTPException(status_code=400, detail="Usa screenshot JPG, PNG o WEBP.")
        content = await file.read()
        if not content or len(content) > 5 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="Ogni screenshot deve pesare al massimo 5 MB.")
        encoded = base64.b64encode(content).decode("ascii")
        image_inputs.append({
            "type": "image_url",
            "image_url": {"url": f"data:{content_type};base64,{encoded}"},
        })

    ocr_analysis = await run_in_threadpool(
        extract_social_screenshot_texts,
        image_inputs,
    )
    content_classification = classify_social_screenshot_text(ocr_analysis.get("extracted_text", ""))
    if not content_classification["valid"]:
        raise HTTPException(status_code=400, detail=content_classification["reason"])
    screenshot_analysis = await run_in_threadpool(
        moderate_visual_inputs,
        image_inputs,
        "uploaded_screenshots",
        len(image_inputs),
    )
    if screenshot_analysis.get("status") == "rate_limited":
        raise HTTPException(status_code=429, detail=screenshot_analysis["message"])
    if screenshot_analysis.get("status") in {"provider_not_configured", "provider_unavailable"}:
        raise HTTPException(status_code=503, detail=screenshot_analysis["message"])

    conn = get_connection()
    cursor = conn.cursor()
    user = fetch_user_by_id(cursor, user_id)
    if not user:
        conn.close()
        raise HTTPException(status_code=404, detail="Utente non trovato.")
    if instagram_handle.strip():
        normalized_instagram = normalize_instagram_handle(instagram_handle)
        cursor.execute(
            "UPDATE users SET instagram_handle = ? WHERE id = ?",
            (f"@{normalized_instagram}" if normalized_instagram else "", user_id),
        )
        conn.commit()
        user = fetch_user_by_id(cursor, user_id)
    profile_text_analysis = evaluate_social_profile_text(
        ocr_analysis.get("extracted_text", ""),
        profile_type,
        user_to_response(user),
    )

    digital_analysis = json.loads(user[19]) if user[19] else {
        "score": 0,
        "headline": "Analisi screenshot completata",
        "summary": screenshot_analysis["message"],
        "findings": [],
        "sources": [],
        "analysis_evidence": {},
    }
    evidence = digital_analysis.setdefault("analysis_evidence", {})
    profile_batches = evidence.setdefault("social_screenshot_batches", [])
    profile_analyses = evidence.setdefault("visual_media_analyses", {})
    screenshot_analysis["profile_type"] = profile_type
    screenshot_analysis["profile_label"] = VISUAL_PROFILE_LABELS[profile_type]
    existing_batch_index = next(
        (
            index
            for index, item in enumerate(profile_batches)
            if item.get("profile_type") == profile_type and item.get("profile_label") == VISUAL_PROFILE_LABELS[profile_type]
        ),
        None,
    )
    batch_payload = {
        "profile_type": profile_type,
        "profile_label": VISUAL_PROFILE_LABELS[profile_type],
        "valid": True,
        "classification": content_classification,
        "ocr": ocr_analysis,
        "visual_analysis": screenshot_analysis,
        "text_analysis": profile_text_analysis,
        "flagged_count": int(screenshot_analysis.get("flagged_count", 0) or 0),
        "sensitive_flagged_count": int(screenshot_analysis.get("sensitive_flagged_count", 0) or 0),
        "analyzed_count": int(screenshot_analysis.get("analyzed_count", 0) or 0),
    }
    if existing_batch_index is None:
        profile_batches.append(batch_payload)
    else:
        profile_batches[existing_batch_index] = batch_payload
    profile_analyses[profile_type] = {
        **screenshot_analysis,
        "profile_type": profile_type,
        "profile_label": VISUAL_PROFILE_LABELS[profile_type],
        "batch_count": sum(1 for item in profile_batches if item.get("profile_type") == profile_type),
    }
    visual_score_adjustment = calculate_social_screenshot_score_adjustment(profile_batches)
    evidence["visual_score_adjustment"] = visual_score_adjustment
    evidence["visual_media_analysis"] = profile_analyses[profile_type]
    evidence["instagram_media_analyzed"] = any(
        item.get("profile_type") == "instagram" and int(item.get("analyzed_count", 0) or 0) > 0
        for item in profile_batches
    )
    evidence["profile_screenshots_analyzed"] = sorted(
        {
            item.get("profile_type")
            for item in profile_batches
            if item.get("profile_type")
        }
    )
    evidence["screenshots_summary"] = summarize_screenshot_evidence(evidence)
    text_analyses = evidence.setdefault("social_text_analyses", {})
    previous_text_entry = text_analyses.get(profile_type, {})
    history = list(previous_text_entry.get("history", []))
    history.append({
        "ocr": ocr_analysis,
        "evaluation": profile_text_analysis,
        "classification": content_classification,
    })
    text_analyses[profile_type] = {
        "ocr": ocr_analysis,
        "evaluation": profile_text_analysis,
        "classification": content_classification,
        "history": history[-6:],
    }
    evidence["instagram_bio_analyzed"] = bool(
        any(
            item.get("ocr", {}).get("extracted_text")
            for item in (text_analyses.get("instagram", {}) or {}).get("history", [])
        )
        or text_analyses.get("instagram", {}).get("ocr", {}).get("extracted_text")
    )
    digital_analysis["score"] = compute_digital_presence_score(evidence)
    findings = digital_analysis.setdefault("findings", [])
    media_finding = next(
        (
            finding
            for finding in findings
            if "foto" in str(finding.get("title", "")).lower()
            or "contenuti pubblici" in str(finding.get("title", "")).lower()
        ),
        None,
    )
    if not media_finding:
        media_finding = {"title": "Foto e contenuti pubblici", "coach_tip": ""}
        findings.append(media_finding)
    media_finding["status"] = (
        "warning"
        if any(analysis.get("flagged_count", 0) for analysis in profile_analyses.values())
        else "success"
    )
    media_finding["description"] = describe_profile_screenshot_analyses(profile_analyses)
    media_finding["coach_tip"] = (
        "Rivedi manualmente i contenuti intimi o sensibili segnalati prima di candidarti."
        if any(analysis.get("flagged_count", 0) for analysis in profile_analyses.values())
        else "Gli screenshot non mostrano contenuti sensibili evidenti e ricevono un piccolo bonus nel punteggio."
    )
    text_finding = next(
        (
            finding
            for finding in findings
            if str(finding.get("title", "")).lower() == "bio e testo profilo"
        ),
        None,
    )
    if not text_finding:
        text_finding = {"title": "Bio e testo profilo"}
        findings.append(text_finding)
    text_finding["status"] = (
        "success" if profile_text_analysis.get("status") == "aligned" else "warning"
    )
    text_finding["description"] = (
        f"{ocr_analysis.get('message', '')} {profile_text_analysis.get('message', '')}"
    ).strip()
    text_finding["coach_tip"] = " ".join(profile_text_analysis.get("suggestions") or []) or (
        "Mantieni la bio sintetica, verificabile e coerente con il ruolo target."
    )
    score_finding = next(
        (
            finding
            for finding in findings
            if "punteggio" in str(finding.get("title", "")).lower() or "score" in str(finding.get("title", "")).lower()
        ),
        None,
    )
    if not score_finding:
        score_finding = {"title": "Punteggio digitale"}
        findings.append(score_finding)
    score_finding["status"] = "success"
    score_finding["description"] = (
        f"Punteggio digitale aggiornato dopo l'aggiunta di {len(image_inputs)} screenshot validi. "
        f"Contributo screenshot: {visual_score_adjustment:+d}."
    )
    score_finding["coach_tip"] = (
        "Aggiungi screenshot validi di profili o piattaforme professionali per aggiornare il punteggio in modo cumulativo."
    )

    cursor.execute(
        "UPDATE users SET digital_analysis_json = ? WHERE id = ?",
        (json.dumps(digital_analysis, ensure_ascii=False), user_id),
    )
    conn.commit()
    updated_user = fetch_user_by_id(cursor, user_id)
    conn.close()
    return {
        "user": user_to_response(updated_user),
        "analysis": digital_analysis,
        "message": (
            f"{VISUAL_PROFILE_LABELS[profile_type]}: {screenshot_analysis['message']} "
            f"{ocr_analysis.get('message', '')} "
            f"Punteggio digitale aggiornato: {digital_analysis['score']}% "
            f"(screenshot: {visual_score_adjustment:+d})."
        ),
    }


# =========================
# ENDPOINT RICERCA WEB
# =========================

@app.post("/search-interview-questions")
def search_interview_questions(data: SearchInterviewQuestionsRequest):
    sources = search_web_interview_questions(
        company=data.company,
        role=data.role,
        interview_type=data.interview_type,
        language=data.language
    )

    return {
        "query": build_search_query(
            data.company,
            data.role,
            data.interview_type,
            data.language
        ),
        "sources": sources
    }


# =========================
# ENDPOINT GENERA 10 DOMANDE
# =========================

@app.post("/generate-question")
def generate_question(data: GenerateQuestionRequest, authorization: Optional[str] = Header(default=None)):
    require_user_session(data.user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT 
        id,
        name,
        education,
        target_role,
        sector,
        experience_level,
        interview_language
    FROM users
    WHERE id = ?
    """, (data.user_id,))

    user = cursor.fetchone()

    if not user:
        conn.close()
        raise HTTPException(
            status_code=404,
            detail="Utente non trovato. Prima devi creare il profilo con POST /users."
        )

    user_id, name, education, target_role, sector, experience_level, interview_language = user

    company = (data.company or "Azienda Generica").strip() or "Azienda Generica"
    personalized_goal = (data.goal or "").strip()
    job_link = (data.job_link or "").strip()
    role_for_questions = (data.role or target_role or "Ruolo da definire").strip()
    question_mode = data.question_mode or "web"

    sources = []

    if question_mode in ["web", "mixed"]:
        sources = search_web_interview_questions(
            company=company,
            role=role_for_questions,
            interview_type=data.interview_type,
            language=interview_language
        )

    if not sources and question_mode in ["web", "mixed"]:
        print("Nessuna fonte web trovata. Procedo con generazione AI basata sul profilo.")

    sources_text = sources_to_prompt(sources)

    question_type_instructions = get_question_type_instructions(
        interview_type=data.interview_type,
        role=role_for_questions,
        company=company
    )

    difficulty_instructions = get_difficulty_instructions(data.difficulty)

    if question_mode == "ai":
        prompt = f"""
Sei un recruiter esperto.

Devi generare 10 domande di colloquio per questo candidato.

Profilo candidato:
- Nome: {name}
- Percorso di studi: {education}
- Ruolo target: {role_for_questions}
- Settore: {sector}
- Livello esperienza: {experience_level}
- Lingua colloquio: {interview_language}
- Azienda target: {company}
- Obiettivo dichiarato dal candidato: {personalized_goal or "Non specificato"}
- Link annuncio o azienda: {job_link or "Non specificato"}
- Tipo colloquio: {data.interview_type}
- Difficoltà: {data.difficulty}

Istruzioni sulla tipologia scelta:
{question_type_instructions}

Istruzioni sul livello di difficoltà:
{difficulty_instructions}

Le 10 domande devono simulare un colloquio realistico.
Devono essere diverse tra loro e progressive.
Devono rispettare rigorosamente la tipologia scelta: {data.interview_type}.

Regole:
- Non aggiungere testo prima o dopo il JSON.
- Non scrivere frasi introduttive come "Ecco le 10 domande".
- Scrivi esattamente 10 domande.
- Ogni elemento della lista deve essere una domanda vera e deve finire con il punto interrogativo.
- Le domande devono essere coerenti con obiettivo, ruolo, azienda, eventuale link, livello e difficoltà.
- Non numerare le domande dentro il testo.

Restituisci SOLO un JSON valido.
La struttura deve essere ESATTAMENTE questa:

{{
  "questions": [
    "Prima domanda?",
    "Seconda domanda?",
    "Terza domanda?",
    "Quarta domanda?",
    "Quinta domanda?",
    "Sesta domanda?",
    "Settima domanda?",
    "Ottava domanda?",
    "Nona domanda?",
    "Decima domanda?"
  ]
}}
"""
    else:
        prompt = f"""
Sei un recruiter esperto.

Devi generare 10 domande di colloquio realistiche per un candidato, ispirandoti ai risultati web trovati.

Profilo candidato:
- Nome: {name}
- Percorso di studi: {education}
- Ruolo target: {role_for_questions}
- Settore: {sector}
- Livello esperienza: {experience_level}
- Lingua colloquio: {interview_language}
- Azienda target: {company}
- Obiettivo dichiarato dal candidato: {personalized_goal or "Non specificato"}
- Link annuncio o azienda: {job_link or "Non specificato"}
- Tipo colloquio: {data.interview_type}
- Difficoltà: {data.difficulty}

Risultati web trovati:
{sources_text}

Istruzioni sulla tipologia scelta:
{question_type_instructions}

Istruzioni sul livello di difficoltà:
{difficulty_instructions}

Le 10 domande devono simulare un colloquio realistico.
Devono essere diverse tra loro e progressive.
Devono rispettare rigorosamente la tipologia scelta: {data.interview_type}.

Regole:
- Non copiare frasi lunghe dalle fonti.
- Non dire "secondo la fonte".
- Genera domande originali ma coerenti con i temi ricorrenti nei risultati.
- Se i risultati sono poco pertinenti, usa comunque il profilo candidato e il ruolo target.
- Non aggiungere testo prima o dopo il JSON.
- Non scrivere frasi introduttive come "Ecco le 10 domande".
- Scrivi esattamente 10 domande.
- Ogni elemento della lista deve essere una domanda vera e deve finire con il punto interrogativo.
- Le domande devono essere coerenti con obiettivo, ruolo, azienda, eventuale link, livello e difficoltà.
- Non numerare le domande dentro il testo.

Restituisci SOLO un JSON valido.
La struttura deve essere ESATTAMENTE questa:

{{
  "questions": [
    "Prima domanda?",
    "Seconda domanda?",
    "Terza domanda?",
    "Quarta domanda?",
    "Quinta domanda?",
    "Sesta domanda?",
    "Settima domanda?",
    "Ottava domanda?",
    "Nona domanda?",
    "Decima domanda?"
  ]
}}
"""

    try:
        raw_questions = call_groq(prompt, temperature=0.3, max_tokens=1400)

        questions_list = extract_questions_list(raw_questions)

    except Exception as e:
        conn.close()
        raise e

    if data.interview_type == "conoscitive_motivazionali":
        fallback_questions = [
            "Parlami di te e del tuo percorso.",
            "Perché ti interessa questa posizione?",
            "Cosa sai della nostra azienda?",
            "Quali sono i tuoi obiettivi professionali?",
            "Come ti vedi tra cinque anni?",
            "Quali sono i tuoi punti di forza?",
            "Qual è un aspetto su cui vorresti migliorare?",
            "Raccontami una situazione in cui hai lavorato in gruppo.",
            "Cosa ti aspetti da questa esperienza lavorativa?",
            "Perché pensi di essere adatto a questo ruolo?"
        ]

    elif data.interview_type == "tecniche":
        fallback_questions = [
            f"Quali competenze tecniche ritieni fondamentali per il ruolo di {role_for_questions}?",
            "Raccontami un progetto tecnico che hai svolto e quali problemi hai incontrato.",
            "Come affronteresti un problema tecnico che non sai risolvere subito?",
            "Quali strumenti o tecnologie conosci che potrebbero essere utili per questo ruolo?",
            "Come verificheresti la correttezza del tuo lavoro tecnico?",
            "Descrivi un caso in cui hai dovuto analizzare dati, codice, requisiti o informazioni tecniche.",
            "Come spiegheresti un concetto tecnico complesso a una persona non tecnica?",
            "Qual è una competenza tecnica che vuoi migliorare?",
            "Come ti organizzi quando devi completare un task tecnico entro una scadenza?",
            "Quale esperienza ti ha aiutato di più a sviluppare competenze utili per questo ruolo?"
        ]

    elif data.interview_type == "logica":
        fallback_questions = [
            f"Immagina di lavorare come {role_for_questions}: hai tre attività urgenti, risorse limitate e informazioni incomplete. Come decideresti da cosa partire e perché?",
            "Quante palline da ping pong servirebbero, secondo te, per riempire la stanza in cui ti trovi? Spiega il ragionamento.",
            "Completa la serie numerica e spiega la regola: 2, 6, 12, 20, 30, ?",
            "Qual è l’angolo tra la lancetta dell’ora e quella dei minuti alle tre e quindici? Spiega il procedimento.",
            f"Se {company} dovesse stimare quanti utenti usano un suo servizio in un giorno, quali ipotesi faresti?",
            "Completa la serie di lettere e spiega la logica: A, C, F, J, O, ?",
            "Quanti Big Mac vengono venduti da McDonald’s ogni anno negli Stati Uniti? Non serve il numero esatto: spiega come lo stimeresti.",
            "Una procedura funziona nel 90% dei casi, ma fallisce nel restante 10%. Come analizzeresti il problema?",
            "Se un cliente ti fornisse dati contraddittori, come ragioneresti per capire quale informazione è più affidabile?",
            "Quante birre vengono consumate in media in un anno in Italia? Spiega quali dati useresti per stimarlo."
        ]

    else:
        fallback_questions = [
            "Parlami brevemente di te e del tuo percorso.",
            "Perché ti interessa questa posizione?",
            "Quali competenze pensi di poter portare in questo ruolo?",
            "Raccontami un progetto o un’esperienza rilevante per questa posizione.",
            "Qual è stata una difficoltà che hai affrontato e come l’hai gestita?",
            "Come ti organizzi quando hai una scadenza importante?",
            "Descrivi una situazione in cui hai lavorato in team.",
            "Qual è un tuo punto di forza e come lo useresti in questo ruolo?",
            "Qual è un aspetto su cui vuoi migliorare professionalmente?",
            "Hai qualche domanda sull’azienda o sul ruolo?"
        ]

    while len(questions_list) < 10:
        questions_list.append(fallback_questions[len(questions_list)])

    questions_list = questions_list[:10]

    sources_json = json.dumps(sources, ensure_ascii=False)

    cursor.execute("""
    INSERT INTO interview_sessions (
        user_id,
        interview_type,
        difficulty,
        company,
        role,
        question_mode
    )
    VALUES (?, ?, ?, ?, ?, ?)
    """, (
        user_id,
        data.interview_type,
        data.difficulty,
        company,
        role_for_questions,
        question_mode
    ))

    session_id = cursor.lastrowid
    saved_questions = []

    for question_text in questions_list:
        cursor.execute("""
        INSERT INTO questions (
            session_id,
            question_text,
            category,
            difficulty,
            company,
            question_mode,
            sources_json
        )
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (
            session_id,
            question_text,
            data.interview_type,
            data.difficulty,
            company,
            question_mode,
            sources_json
        ))

        question_id = cursor.lastrowid

        save_sources_for_question(cursor, question_id, sources)

        saved_questions.append({
            "question_id": question_id,
            "question": question_text
        })

    conn.commit()
    conn.close()

    print("10 domande salvate e restituite al frontend.")

    return {
        "session_id": session_id,
        "questions": saved_questions,
        "question_id": saved_questions[0]["question_id"],
        "question": saved_questions[0]["question"],
        "company": company,
        "question_mode": question_mode
    }


# =========================
# ENDPOINT VALUTA RISPOSTA
# =========================

@app.post("/evaluate-answer")
def evaluate_answer(data: EvaluateAnswerRequest, authorization: Optional[str] = Header(default=None)):
    print("Endpoint /evaluate-answer chiamato.")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT 
        q.id,
        q.question_text,
        q.session_id,
        q.company,
        q.question_mode,
        q.sources_json,
        s.user_id,
        s.interview_type,
        s.difficulty,
        u.name,
        u.education,
        u.target_role,
        u.sector,
        u.experience_level,
        u.interview_language
    FROM questions q
    JOIN interview_sessions s ON q.session_id = s.id
    JOIN users u ON s.user_id = u.id
    WHERE q.id = ?
    """, (data.question_id,))

    row = cursor.fetchone()

    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Domanda non trovata")

    (
        question_id,
        question_text,
        session_id,
        company,
        question_mode,
        sources_json,
        user_id,
        interview_type,
        difficulty,
        name,
        education,
        target_role,
        sector,
        experience_level,
        interview_language
    ) = row

    require_user_session(user_id, authorization)

    speech_metrics_json = get_speech_metrics_json(data.speech_metrics)

    zero_reason = get_zero_answer_reason(data.answer, question_text)

    if zero_reason:
        zero_result = build_zero_feedback(
            reason=zero_reason,
            question=question_text,
            interview_type=interview_type
        )

        save_answer_result(
            cursor=cursor,
            question_id=question_id,
            user_answer=data.answer,
            result=zero_result,
            speech_metrics_json=speech_metrics_json
        )

        cursor.execute("""
            UPDATE interview_sessions
            SET total_score = (
                SELECT CAST(ROUND(AVG(total_score)) AS INTEGER)
                FROM answers
                WHERE question_id IN (SELECT id FROM questions WHERE session_id = ?)
            )
            WHERE id = ?
        """, (session_id, session_id))

        conn.commit()
        conn.close()

        return zero_result

    has_speech_metrics = data.speech_metrics is not None

    speech_info = ""

    if has_speech_metrics:
        speech_info = f"""
Dati sul parlato rilevati dal frontend:
- Durata risposta: {data.speech_metrics.duration_seconds} secondi
- Numero parole: {data.speech_metrics.words_count}
- Parole al minuto: {data.speech_metrics.words_per_minute}
- Numero parole riempitive: {data.speech_metrics.filler_words_count}
- Parole riempitive rilevate: {data.speech_metrics.filler_words}

Valuta anche il modo di parlare considerando:
- ritmo;
- chiarezza espositiva;
- parole riempitive;
- sicurezza percepita;
- sintesi;
- capacità di organizzare il discorso.
"""

    prompt = f"""
Sei un coach esperto di colloqui di lavoro.

Devi valutare la risposta di un candidato.

Profilo candidato:
- Nome: {name}
- Percorso di studi: {education}
- Ruolo target: {target_role}
- Settore: {sector}
- Livello esperienza: {experience_level}
- Lingua colloquio: {interview_language}
- Azienda target: {company}
- Tipo colloquio: {interview_type}
- Difficoltà: {difficulty}

Domanda del colloquio:
{question_text}

Risposta del candidato:
{data.answer}

{speech_info}

ISTRUZIONE PRIORITARIA:
Prima di assegnare qualsiasi punteggio, devi decidere internamente se la risposta è una vera risposta alla domanda oppure no.

La classificazione interna può essere solo:
- VALID_ANSWER
- NON_ANSWER

NON devi inserire questa classificazione nel JSON finale.

Una risposta è VALID_ANSWER solo se il candidato prova concretamente a rispondere alla domanda del colloquio.

Una risposta è NON_ANSWER se il candidato non fornisce un vero tentativo di risposta alla domanda posta.

Classifica sempre come NON_ANSWER qualsiasi risposta che:
- faccia una domanda invece di rispondere;
- chieda come viene valutata la risposta;
- chieda quali criteri vengono usati per valutare;
- chieda un voto, un punteggio, un giudizio o una correzione;
- chieda la risposta corretta;
- chieda aiuto, suggerimenti o indicazioni su cosa dire;
- parli del sistema di valutazione, del codice, della logica del progetto o dell'AI invece di rispondere;
- sia una frase di rinuncia, dubbio o incertezza;
- sia evasiva, fuori tema o non pertinente;
- sia troppo generica e non collegata alla domanda originale;
- sia una richiesta rivolta al recruiter, all'intervistatore o all'assistente;
- non dimostri alcuna competenza, esperienza, ragionamento o contenuto utile rispetto alla domanda.

Sono NON_ANSWER anche frasi grammaticalmente corrette, educate o comprensibili, se non rispondono alla domanda originale.

Esempi di NON_ANSWER:
- "che criterio usi per valutare le mie risposte?"
- "come mi valuti?"
- "quanto mi dai?"
- "che voto mi daresti?"
- "qual è la risposta corretta?"
- "cosa dovrei rispondere?"
- "mi aiuti a rispondere?"
- "puoi farmi un esempio?"
- "puoi spiegarmi la domanda?"
- "non lo so"
- "boh"
- "non saprei"
- "non ho idea"
- "dipende"

REGOLA PRIORITARIA ASSOLUTA:
Se la risposta viene classificata internamente come NON_ANSWER, devi assegnare obbligatoriamente 0 a TUTTI i punteggi:

clarity_score = 0
completeness_score = 0
relevance_score = 0
professionalism_score = 0
synthesis_score = 0
speaking_score = 0
total_score = 0

Non assegnare mai punteggi parziali a una risposta classificata come NON_ANSWER.

In caso di NON_ANSWER:
- feedback deve spiegare che il candidato non ha risposto alla domanda posta;
- improved_answer deve contenere una risposta modello corretta alla domanda originale;
- speaking_feedback deve spiegare che la risposta non è una risposta valida alla domanda;
- solution_explanation deve essere compilata solo se la domanda originale è di logica o richiede un ragionamento passo passo, altrimenti deve essere una stringa vuota.

Valuta la risposta sui seguenti criteri:
1. Chiarezza
2. Completezza
3. Pertinenza rispetto al ruolo e all'azienda
4. Professionalità
5. Sintesi
6. Modo di parlare, solo se sono presenti metriche vocali

Restituisci SOLO un JSON valido, senza testo prima e senza testo dopo.

La struttura deve essere ESATTAMENTE questa:

{{
  "clarity_score": 0,
  "completeness_score": 0,
  "relevance_score": 0,
  "professionalism_score": 0,
  "synthesis_score": 0,
  "speaking_score": 0,
  "total_score": 0,
  "feedback": "feedback dettagliato ma comprensibile",
  "improved_answer": "risposta migliorata o risposta modello alla domanda",
  "speaking_feedback": "feedback sul modo di parlare",
  "solution_explanation": "soluzione spiegata se la domanda è di logica, altrimenti stringa vuota"
}}

Regole di valutazione:
- Tutti i punteggi devono essere numeri interi compresi tra 0 e 100.
- Se la risposta è classificata internamente come NON_ANSWER, tutti i punteggi devono essere 0.
- Se la risposta è incomprensibile, casuale, composta da lettere senza senso o non risponde minimamente alla domanda, assegna 0 a tutti i punteggi.
- Se la risposta è "boh", "non lo so", "chi lo sa", "non saprei" o simili, assegna 0 a tutti i punteggi.
- Se la risposta è completamente fuori tema rispetto alla domanda, assegna 0 a tutti i punteggi.
- Se la risposta fa domande al recruiter, all'intervistatore, all'assistente o all'AI invece di rispondere, assegna 0 a tutti i punteggi.
- Se la risposta chiede informazioni sulla valutazione, sui criteri, sul punteggio, sul codice o sulla logica del sistema, assegna 0 a tutti i punteggi.
- Se la risposta è molto vaga ma contiene almeno un minimo tentativo reale di risposta, assegna punteggi bassi, tra 5 e 20.
- Se la risposta usa un lessico povero, confuso o poco professionale, abbassa clarity_score e professionalism_score.
- Se la risposta non contiene esempi, motivazioni o contenuti concreti, abbassa completeness_score.
- Non premiare una risposta solo perché è lunga: deve essere comprensibile, pertinente e utile.
- Se la risposta è molto prolissa, divaga o contiene molte informazioni non rilevanti, abbassa clarity_score e synthesis_score.
- improved_answer deve essere SEMPRE valorizzata.
- Se il candidato non risponde alla domanda, improved_answer deve contenere una risposta modello corretta alla domanda originale.
- Il punteggio massimo può essere raggiunto anche con risposte corrette ma formulate in modo diverso dalla risposta modello.
- Se la domanda è tecnica, improved_answer deve mostrare una risposta tecnica corretta ma credibile per il livello del candidato.
- Se la domanda è conoscitiva o motivazionale, improved_answer deve mostrare una risposta naturale, professionale e adatta a un colloquio.
- Se la domanda è di logica, improved_answer deve contenere una possibile risposta ragionata e solution_explanation deve contenere la soluzione spiegata passo passo.
- Se la risposta è errata ma la domanda è di logica, spiega comunque il ragionamento corretto in solution_explanation.
- Se sono presenti metriche vocali e la risposta è valida, devi obbligatoriamente valutare il modo di parlare usando quei dati.
- Se sono presenti metriche vocali e la risposta è valida, speaking_score deve essere maggiore di 0.
- Se sono presenti metriche vocali, NON devi mai scrivere che le metriche vocali mancano.
- Nel campo speaking_feedback devi commentare ritmo, chiarezza espositiva, velocità del parlato, parole riempitive e sicurezza percepita.
- Se NON ci sono metriche vocali e la risposta è valida, speaking_score deve essere 0 e speaking_feedback deve dire che la risposta è stata valutata solo sul testo.
- Se la risposta non è valida, speaking_score deve essere 0 anche se sono presenti metriche vocali.
- Non inventare esperienze personali non presenti nella risposta.
- La risposta migliorata deve essere naturale, credibile e coerente con la domanda originale.
- Restituisci solo JSON valido.
"""

    try:
        raw_output = call_groq(prompt, temperature=0.4, max_tokens=1600)
        result = extract_json(raw_output)
    except Exception as e:
        conn.close()
        raise HTTPException(
            status_code=500,
            detail=f"Errore nella valutazione della risposta: {str(e)}"
        )

    clarity_score = clamp_score(result.get("clarity_score", 0))
    completeness_score = clamp_score(result.get("completeness_score", 0))
    relevance_score = clamp_score(result.get("relevance_score", 0))
    professionalism_score = clamp_score(result.get("professionalism_score", 0))
    synthesis_score = clamp_score(result.get("synthesis_score", 0))

    all_textual_scores_are_zero = all(score == 0 for score in [
        clarity_score,
        completeness_score,
        relevance_score,
        professionalism_score,
        synthesis_score
    ])

    if has_speech_metrics:
        speaking_score = clamp_score(result.get("speaking_score", 0))

        # Non trasformare mai lo speaking_score in 60 se il contenuto è già stato valutato 0.
        if all_textual_scores_are_zero:
            speaking_score = 0
        elif speaking_score == 0:
            speaking_score = 60
    else:
        speaking_score = 0

    total_score = compute_total_score(
        clarity_score,
        completeness_score,
        relevance_score,
        professionalism_score,
        synthesis_score,
        speaking_score
    )

    # Se tutti i punteggi testuali sono zero, il totale deve restare zero.
    if all_textual_scores_are_zero:
        total_score = 0
        speaking_score = 0

    feedback = result.get("feedback", "")
    improved_answer = result.get("improved_answer", "")
    solution_explanation = result.get("solution_explanation", "")

    if not improved_answer:
        improved_answer = (
            f"Una risposta migliore alla domanda '{question_text}' dovrebbe rispondere direttamente, "
            "essere chiara, pertinente e includere almeno un esempio o una motivazione concreta."
        )

    if has_speech_metrics:
        speaking_feedback = result.get("speaking_feedback", "")

        invalid_speaking_feedback = (
            not speaking_feedback
            or "non contiene metriche vocali" in speaking_feedback.lower()
            or "non ci sono metriche vocali" in speaking_feedback.lower()
            or "non sono presenti metriche vocali" in speaking_feedback.lower()
            or "valutata solo sul testo" in speaking_feedback.lower()
            or "solo sul contenuto testuale" in speaking_feedback.lower()
        )

        if all_textual_scores_are_zero:
            speaking_feedback = "Il modo di parlare non viene valutato perché la risposta non è valida rispetto alla domanda posta."
        elif invalid_speaking_feedback:
            speaking_feedback = build_speaking_feedback_from_metrics(data.speech_metrics)
    else:
        speaking_feedback = (
            "La risposta è stata valutata solo sul contenuto testuale perché non sono presenti dati sul parlato."
        )

    final_result = {
        "clarity_score": clarity_score,
        "completeness_score": completeness_score,
        "relevance_score": relevance_score,
        "professionalism_score": professionalism_score,
        "synthesis_score": synthesis_score,
        "speaking_score": speaking_score,
        "total_score": total_score,
        "feedback": feedback,
        "improved_answer": improved_answer,
        "speaking_feedback": speaking_feedback,
        "solution_explanation": solution_explanation
    }

    save_answer_result(
        cursor=cursor,
        question_id=question_id,
        user_answer=data.answer,
        result=final_result,
        speech_metrics_json=speech_metrics_json
    )

    cursor.execute("""
        UPDATE interview_sessions
        SET total_score = (
            SELECT CAST(ROUND(AVG(total_score)) AS INTEGER)
            FROM answers
            WHERE question_id IN (SELECT id FROM questions WHERE session_id = ?)
        )
        WHERE id = ?
    """, (session_id, session_id))

    conn.commit()
    conn.close()

    return final_result


# =========================
# ENDPOINT STORICO
# =========================

@app.get("/history/{user_id}")
def get_history(user_id: int, authorization: Optional[str] = Header(default=None)):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT 
        s.id,
        s.interview_type,
        s.difficulty,
        s.company,
        s.role,
        s.question_mode,
        s.total_score,
        s.created_at,
        q.question_text,
        a.user_answer,
        a.clarity_score,
        a.completeness_score,
        a.relevance_score,
        a.professionalism_score,
        a.synthesis_score,
        a.speaking_score,
        a.feedback,
        a.improved_answer,
        a.speaking_feedback,
        a.solution_explanation
    FROM interview_sessions s
    JOIN questions q ON q.session_id = s.id
    LEFT JOIN answers a ON a.question_id = q.id
    WHERE s.user_id = ?
    ORDER BY s.created_at DESC
    """, (user_id,))

    rows = cursor.fetchall()
    conn.close()

    history = []

    for row in rows:
        history.append({
            "session_id": row[0],
            "interview_type": row[1],
            "difficulty": row[2],
            "company": row[3],
            "role": row[4],
            "question_mode": row[5],
            "total_score": row[6],
            "created_at": row[7],
            "question": row[8],
            "user_answer": row[9],
            "clarity_score": row[10],
            "completeness_score": row[11],
            "relevance_score": row[12],
            "professionalism_score": row[13],
            "synthesis_score": row[14],
            "speaking_score": row[15],
            "feedback": row[16],
            "improved_answer": row[17],
            "speaking_feedback": row[18],
            "solution_explanation": row[19]
        })

    return history


# =========================
# ENDPOINT PROGRESSI
# =========================

@app.get("/progress/{user_id}")
def get_progress(user_id: int, authorization: Optional[str] = Header(default=None)):
    require_user_session(user_id, authorization)
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT 
        COUNT(a.id),
        AVG(a.total_score),
        AVG(a.clarity_score),
        AVG(a.completeness_score),
        AVG(a.relevance_score),
        AVG(a.professionalism_score),
        AVG(a.synthesis_score),
        AVG(a.speaking_score)
    FROM answers a
    JOIN questions q ON a.question_id = q.id
    JOIN interview_sessions s ON q.session_id = s.id
    WHERE s.user_id = ?
    """, (user_id,))

    row = cursor.fetchone()
    conn.close()

    if not row or row[0] == 0:
        return {
            "message": "Non ci sono ancora risposte valutate per questo utente.",
            "total_answers": 0
        }

    return {
        "total_answers": row[0],
        "average_total_score": round(row[1], 2) if row[1] is not None else 0,
        "average_clarity_score": round(row[2], 2) if row[2] is not None else 0,
        "average_completeness_score": round(row[3], 2) if row[3] is not None else 0,
        "average_relevance_score": round(row[4], 2) if row[4] is not None else 0,
        "average_professionalism_score": round(row[5], 2) if row[5] is not None else 0,
        "average_synthesis_score": round(row[6], 2) if row[6] is not None else 0,
        "average_speaking_score": round(row[7], 2) if row[7] is not None else 0
    }


# =========================
# ENDPOINT FONTI DOMANDA
# =========================

@app.get("/question-sources/{question_id}")
def get_question_sources(question_id: int, authorization: Optional[str] = Header(default=None)):
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("""
    SELECT 
        ws.id,
        ws.title,
        ws.url,
        ws.content,
        ws.created_at,
        s.user_id
    FROM web_sources ws
    JOIN question_web_sources qws ON ws.id = qws.source_id
    JOIN questions q ON q.id = qws.question_id
    JOIN interview_sessions s ON s.id = q.session_id
    WHERE qws.question_id = ?
    ORDER BY ws.created_at DESC
    """, (question_id,))

    rows = cursor.fetchall()
    conn.close()

    if rows:
        require_user_session(rows[0][5], authorization)

    sources = []

    for row in rows:
        sources.append({
            "source_id": row[0],
            "title": row[1],
            "url": row[2],
            "content": row[3],
            "created_at": row[4]
        })

    return {
        "question_id": question_id,
        "sources": sources
    }
@app.post("/debug-cv-read")
async def debug_cv_read(file: UploadFile = File(...)):
    if not DEBUG_MODE:
        raise HTTPException(status_code=404, detail="Endpoint debug non disponibile.")

    from io import BytesIO
    import fitz
    from pypdf import PdfReader

    file_bytes = await file.read()

    result = {
        "filename": file.filename,
        "content_type": file.content_type,
        "file_size": len(file_bytes),
        "pymupdf_text_length": 0,
        "pymupdf_preview": "",
        "pypdf_text_length": 0,
        "pypdf_preview": "",
    }

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []

        for page in doc:
            page_text = page.get_text("text")
            if page_text:
                text_parts.append(page_text)

        doc.close()

        pymupdf_text = "\n".join(text_parts).strip()
        result["pymupdf_text_length"] = len(pymupdf_text)
        result["pymupdf_preview"] = pymupdf_text[:1000]

    except Exception as e:
        result["pymupdf_error"] = str(e)

    try:
        reader = PdfReader(BytesIO(file_bytes))
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        pypdf_text = "\n".join(text_parts).strip()
        result["pypdf_text_length"] = len(pypdf_text)
        result["pypdf_preview"] = pypdf_text[:1000]

    except Exception as e:
        result["pypdf_error"] = str(e)

    return result
