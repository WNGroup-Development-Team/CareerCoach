import io
import unittest
from unittest.mock import patch

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt

from main import (
    build_additional_rewrite_instructions,
    build_confirmed_skill_rewrite_instructions,
    call_rewrite_llm,
    is_additive_user_rewrite_source,
    canonical_edit_section_name,
    extract_resume_sections,
    fallback_skill_detail_instruction,
    infer_extra_content_section,
)
from services.cv_optimizer.rewrite import build_resume_rewrite_result
from services.cv_optimizer.structured_cv_engine import build_optimized_cv_text
from services.cv_optimizer.structured_cv_engine import _call_copywriting_llm
from services.cv_optimizer.pipeline import (
    DocxPreserver,
    ResumeDocxOptimizationPipeline,
    ResumeParser,
    ResumeRewriter,
    RewriteInstruction,
    StructuredRewriteInstruction,
    canonical_section,
    is_section_heading,
)
from services.cv_optimizer.section_catalog import canonical_section_key


class ResumeParserLayoutTests(unittest.TestCase):
    def test_parse_text_recovers_inline_section_headings(self):
        text = (
            "Mario Rossi CHI SONO Profilo breve orientato ai dati "
            "HARD SKILLS Python SQL Power BI FORMAZIONE Laurea in Informatica "
            "ESPERIENZE PROFESSIONALI Analisi dati in progetto universitario"
        )

        sections = {section.name: section.text for section in ResumeParser().parse_text(text)}

        self.assertIn("profilo", sections)
        self.assertIn("hard_skills", sections)
        self.assertIn("formazione", sections)
        self.assertIn("esperienze", sections)
        self.assertEqual(sections["hard_skills"], "Python SQL Power BI")

    def test_common_heading_variants_share_the_same_section(self):
        variants = {
            "Esperienze lavorative": "experience",
            "ESPERIENZE PROFESSIONALI:": "experience",
            "Employment History": "experience",
            "Percorso accademico": "education",
            "ISTRUZIONE E FORMAZIONE": "education",
            "Conoscenze tecniche": "hard_skills",
            "Competenze personali": "soft_skills",
            "CAPACITÀ": "soft_skills",
            "Conoscenze linguistiche": "languages",
            "LANGUAGE": "languages",
            "IT SKILLS": "hard_skills",
            "Corsi e certificazioni": "certifications",
            "PUBBLICAZIONI": "publications",
            "Progetti accademici": "projects",
            "Additional Information": "projects",
            "FORMAZIONE 2023–2025": "education",
            "COMUNICAZIONE": "languages",
            "Pagina Web personale": "contacts",
            "Informazioni di contatto": "contacts",
        }

        for heading, expected in variants.items():
            with self.subTest(heading=heading):
                self.assertEqual(canonical_section_key(heading), expected)

    def test_resume_parser_accepts_work_experience_synonym(self):
        sections = {
            section.name: section.text
            for section in ResumeParser().parse_text(
                "Mario Rossi\nESPERIENZE LAVORATIVE\nAnalista dati presso Acme"
            )
        }

        self.assertEqual(
            sections["esperienze"],
            "Analista dati presso Acme",
        )

    def test_colon_terminated_subheading_stops_previous_section(self):
        document = Document()
        document.add_paragraph("Courses and Certifications:")
        document.add_paragraph("Professional certification.")
        document.add_paragraph("Voluntary activities:")
        document.add_paragraph("Community volunteering.")

        contexts = ResumeDocxOptimizationPipeline()._paragraph_contexts(document)

        self.assertEqual(contexts[1].section, "certificazioni")
        self.assertEqual(contexts[2].section, "voluntary activities")
        self.assertEqual(contexts[3].section, "voluntary activities")

    def test_uppercase_name_and_date_are_not_section_headings(self):
        self.assertFalse(is_section_heading("ALESSIO MARINUCCI"))
        self.assertFalse(is_section_heading("2025-IN CORSO"))
        self.assertTrue(is_section_heading("FORMAZIONE 2023–2025"))

    def test_legacy_parser_uses_the_shared_heading_catalog(self):
        self.assertEqual(
            canonical_edit_section_name("Esperienze lavorative"),
            "ESPERIENZE PROFESSIONALI",
        )
        sections = extract_resume_sections(
            "Mario Rossi\nESPERIENZE LAVORATIVE\nAnalista dati presso Acme\n"
            "PERCORSO ACCADEMICO\nLaurea in Informatica"
        )

        self.assertEqual(sections["experience"], "Analista dati presso Acme")
        self.assertEqual(sections["education"], "Laurea in Informatica")


class DynamicAdditionalContentTests(unittest.TestCase):
    def test_language_level_is_routed_to_languages(self):
        self.assertEqual(
            infer_extra_content_section("Ho preso il B2 in Inglese"),
            ("LINGUE", "languages"),
        )

    @patch("main.build_professional_extra_text")
    def test_short_project_is_kept_as_title_and_description(self, rewrite_mock):
        rewrite_mock.return_value = "Partecipazione a un progetto di imprenditorialita digitale."

        instructions = build_additional_rewrite_instructions(
            {"projects": "Ho lavorato in un progetto di imprenditorialita digitale"},
            role="Data Analyst",
            cv_text="PROFILO\nCandidato junior.",
        )

        project = next(item for item in instructions if item.section == "PROGETTI")
        self.assertEqual(len(project.replacement.splitlines()), 2)
        self.assertTrue(project.replacement.splitlines()[1].endswith("."))

    def test_section_formatter_supports_all_dynamic_sections(self):
        pipeline = ResumeDocxOptimizationPipeline()

        self.assertEqual(
            pipeline._format_section_text("lingue", "Inglese B2\nFrancese A2", ""),
            "Inglese B2\nFrancese A2",
        )
        self.assertIn(
            "Progetto personale",
            pipeline._format_section_text(
                "progetti",
                "Partecipazione a un progetto di imprenditorialita digitale.",
                "",
            ),
        )
        self.assertEqual(
            pipeline._format_section_text("certificazioni", "Certificazione AWS.", ""),
            "Certificazione AWS",
        )

    def test_clean_user_fact_is_not_discarded_as_duplicate_note(self):
        pipeline = ResumeDocxOptimizationPipeline()
        instructions = pipeline._sanitize_structured_instructions(
            [
                StructuredRewriteInstruction(
                    suggestion_id="language-b2",
                    target_section="LINGUE",
                    action="append",
                    old_text_hint="",
                    new_text="Inglese B2",
                ),
                StructuredRewriteInstruction(
                    suggestion_id="cert-aws",
                    target_section="CERTIFICAZIONI",
                    action="append",
                    old_text_hint="",
                    new_text="Certificazione AWS",
                ),
            ],
            cv_text="PROFILO\nCandidato junior.",
            user_additional_data={
                "languages": "Inglese B2",
                "certifications": "Certificazione AWS",
            },
        )

        self.assertEqual(
            {item.suggestion_id for item in instructions},
            {"language-b2", "cert-aws"},
        )

    def test_user_example_routes_each_fact_without_llm_reinterpretation(self):
        instructions = build_additional_rewrite_instructions(
            {
                "experiences": (
                    "Ho lavorato in Datewave dove ho realizzato un chatbot "
                    "per la spedizione di pacchi"
                ),
                "projects": (
                    "progetto di ML sulla predizione del possibile vincitore "
                    "dei mondiali e progetto di Ingegneria dei Dati per un "
                    "sistema di Information retrieval"
                ),
                "certifications": (
                    "Ho preso la certificazione B2 di Inglese e C1 di francese."
                ),
            },
            role="Data Analyst",
            cv_text="ESPERIENZE LAVORATIVE\nEsperienza precedente.",
        )

        by_section = {}
        for instruction in instructions:
            by_section.setdefault(instruction.section, []).append(instruction)

        self.assertIn("ESPERIENZE PROFESSIONALI", by_section)
        self.assertIn("Datewave", by_section["ESPERIENZE PROFESSIONALI"][0].replacement)
        self.assertIn("PROGETTI", by_section)
        self.assertEqual(len(by_section["PROGETTI"]), 2)
        self.assertIn("Machine Learning", by_section["PROGETTI"][0].replacement)
        self.assertIn("Data Engineering", by_section["PROGETTI"][1].replacement)
        self.assertEqual(
            by_section["LINGUE"][0].replacement,
            "Inglese B2\nFrancese C1",
        )
        self.assertEqual(by_section["LINGUE"][0].original, "")

    @patch("main.call_rewrite_llm")
    def test_project_paragraph_stays_coherent_and_language_level_is_exact(
        self,
        rewrite_mock,
    ):
        rewrite_mock.return_value = {
            "replacement": (
                "Realizzazione di progetti universitari di analisi dati. "
                "Utilizzo di SQL e Python per interrogazione, analisi e reportistica."
            )
        }
        instructions = build_additional_rewrite_instructions(
            {
                "projects": (
                    "Ho realizzato progetti universitari legati all'analisi dei dati. "
                    "In particolare, ho lavorato con SQL e Python per analisi e report."
                ),
                "certifications": "Ho preso il B2 in Inglese",
            },
            role="Data Analyst",
            cv_text="PROGETTI\nProgetto esistente.",
        )

        projects = [item for item in instructions if item.section == "PROGETTI"]
        languages = [item for item in instructions if item.section == "LINGUE"]

        self.assertEqual(len(projects), 1)
        self.assertIn("SQL", projects[0].replacement)
        self.assertEqual(languages[0].replacement, "Inglese B2")

    @patch("main.call_rewrite_llm")
    def test_every_frontend_additional_box_produces_a_cv_instruction(
        self,
        rewrite_mock,
    ):
        rewrite_mock.return_value = {"replacement": "Testo professionale confermato."}
        fields = {
            "experiences": "Ho lavorato presso Acme occupandomi di assistenza clienti.",
            "technical_skills": "Ho utilizzato Python per elaborare dataset.",
            "soft_skills": "Ho coordinato il lavoro con un gruppo universitario.",
            "projects": "Progetto di analisi dati con SQL e Python.",
            "measurable_results": "Ho ridotto i tempi di elaborazione del 20%.",
            "certifications": "Certificazione Google Data Analytics.",
            "tools": "Ho utilizzato Power BI per creare dashboard.",
            "company_role_notes": "Interesse per il ruolo Data Analyst in Google.",
            "additional_notes": "Attivita di volontariato nella gestione di eventi.",
        }

        for field_name, value in fields.items():
            with self.subTest(field=field_name):
                instructions = build_additional_rewrite_instructions(
                    {field_name: value},
                    role="Data Analyst",
                    cv_text="PROFILO\nProfilo esistente.",
                )
                self.assertTrue(
                    instructions,
                    f"La box {field_name} non ha prodotto istruzioni per il CV.",
                )

    @patch("main.call_structured_llm")
    def test_rewrite_llm_does_not_fallback_to_ollama(self, structured_mock):
        structured_mock.return_value = {"instructions": []}

        call_rewrite_llm("prompt", context="test")

        self.assertEqual(
            structured_mock.call_args.kwargs["preferred_order"],
            ["gemini"],
        )

    @patch("main.call_structured_llm")
    def test_cv_analysis_helpers_are_gemini_only(self, structured_mock):
        from main import call_analysis_llm, call_lightweight_analysis_llm

        structured_mock.return_value = {}
        call_analysis_llm("prompt")
        call_lightweight_analysis_llm("prompt")

        self.assertEqual(
            [call.kwargs["preferred_order"] for call in structured_mock.call_args_list],
            [["gemini"], ["gemini"]],
        )

    @patch("main.call_gemini")
    @patch("main.call_ollama")
    def test_structured_copywriting_uses_gemini_and_never_ollama(
        self,
        ollama_mock,
        gemini_mock,
    ):
        gemini_mock.return_value = {"instructions": []}

        result = _call_copywriting_llm("prompt")

        self.assertEqual(result, {"instructions": []})
        gemini_mock.assert_called_once()
        ollama_mock.assert_not_called()

    def test_confirmed_skill_examples_are_professionalized_locally(self):
        kpi = fallback_skill_detail_instruction(
            {
                "name": "KPI",
                "category": "hard_skill",
                "detail": (
                    "Usati in progetti universitari di analisi dati per "
                    "valutare risultati e monitorare indicatori."
                ),
            },
            0,
            "",
        )
        analytics = fallback_skill_detail_instruction(
            {
                "name": "Google Analytics",
                "category": "hard_skill",
                "detail": "l'ho usato in un corso di Google",
            },
            1,
            "",
        )

        self.assertEqual(kpi.section, "PROGETTI")
        self.assertIn("KPI", kpi.replacement)
        self.assertEqual(analytics.section, "FORMAZIONE")
        self.assertNotIn("l'ho usato", analytics.replacement.lower())
        self.assertIn("Google Analytics", analytics.replacement)

    @patch("main.CV_REWRITE_LLM_ENABLED", False)
    def test_carol_confirmed_content_is_routed_without_losing_project_evidence(
        self,
    ):
        user_data = {
            "confirmed_skills": [
                {
                    "name": "KPI",
                    "category": "hard_skill",
                    "detail": (
                        "Usata in progetti universitari di analisi dati per "
                        "definire e monitorare indicatori utili alla valutazione "
                        "dei risultati, confrontando metriche e performance."
                    ),
                }
            ],
            "experiences": (
                "ho lavorato presso Poste Italiane dove ho realizzato una "
                "chatbot che risponde a domande relative a documentazione aziendale"
            ),
            "projects": (
                "Ho lavorato su progetti universitari di Data Analysis e Machine "
                "Learning in cui ho analizzato dati, definito indicatori di "
                "performance e interpretato metriche tramite report e tabelle."
            ),
            "measurable_results": (
                "Ho contribuito alla produzione di analisi piu chiare e "
                "confrontabili, metriche quantitative per valutare risultati, "
                "individuare criticita e presentare conclusioni in modo strutturato."
            ),
        }

        instructions = [
            *build_confirmed_skill_rewrite_instructions(
                "HARD SKILLS\nPython\nC++",
                user_data,
                "Data Analyst",
            ),
            *build_additional_rewrite_instructions(
                user_data,
                "Data Analyst",
                "HARD SKILLS\nPython\nC++",
            ),
        ]
        projects = [
            item for item in instructions
            if item.section == "PROGETTI"
        ]
        experiences = [
            item for item in instructions
            if item.section == "ESPERIENZE PROFESSIONALI"
        ]

        self.assertGreaterEqual(len(projects), 3)
        project_text = "\n".join(item.replacement for item in projects)
        self.assertIn("KPI", project_text)
        self.assertIn("Machine Learning", project_text)
        self.assertIn("analisi piu chiare", project_text.lower())
        self.assertEqual(len(experiences), 1)
        self.assertIn("Presso Poste Italiane", experiences[0].replacement)
        self.assertNotIn("Esperienza presso", experiences[0].replacement)

    def test_accepted_skill_cards_are_all_converted_to_cv_instructions(self):
        instructions = build_confirmed_skill_rewrite_instructions(
            "HARD SKILLS\nPython\nSOFT SKILLS\nProblem solving",
            {
                "confirmed_skills": [
                    {
                        "id": "hard-card",
                        "name": "SQL",
                        "category": "hard_skill",
                        "detail": "",
                        "target_section": "HARD SKILLS",
                    },
                    {
                        "id": "soft-card",
                        "name": "Coordinamento team",
                        "category": "soft_skill",
                        "detail": (
                            "Dimostrato durante un progetto universitario "
                            "organizzando attivita e scadenze."
                        ),
                        "target_section": "SOFT SKILLS",
                    },
                    {
                        "id": "keyword-card",
                        "name": "KPI",
                        "category": "keyword",
                        "detail": (
                            "Usati in progetti universitari di analisi dati "
                            "per monitorare i risultati."
                        ),
                        "target_section": "COMPETENZE TECNICHE",
                    },
                ]
            },
            role="Data Analyst",
        )

        replacements = "\n".join(item.replacement for item in instructions)
        self.assertIn("SQL", replacements)
        self.assertIn("Coordinamento team", replacements)
        self.assertIn("KPI", replacements)
        self.assertTrue(
            any(item.category == "soft_skills" for item in instructions)
        )
        self.assertTrue(
            any(item.source_id.startswith("confirmed_skill_detail_") for item in instructions)
        )

    def test_only_box_and_confirmation_sources_are_additive(self):
        self.assertTrue(is_additive_user_rewrite_source("user_box_projects_0"))
        self.assertTrue(is_additive_user_rewrite_source("confirmed_hard_skills"))
        self.assertTrue(
            is_additive_user_rewrite_source(
                "consolidated:coach-edit|user_additional_answer_0_0"
            )
        )
        self.assertFalse(is_additive_user_rewrite_source("coach-suggestion-1"))
        self.assertFalse(is_additive_user_rewrite_source("llm_instruction"))


class ResumeRewriterTests(unittest.TestCase):
    def test_applies_generated_instructions_to_optimized_text(self):
        original = (
            "PROFILO\n"
            "Analista con esperienza in reportistica.\n\n"
            "COMPETENZE\n"
            "Excel e SQL"
        )

        optimized = ResumeRewriter().apply_to_text(
            original,
            [
                RewriteInstruction(
                    section="profilo",
                    original="Analista con esperienza in reportistica.",
                    replacement="Analista orientato ai dati con esperienza in reportistica.",
                    source_id="generated-profile",
                )
            ],
        )

        self.assertIn(
            "Analista orientato ai dati con esperienza in reportistica.",
            optimized,
        )
        self.assertNotEqual(optimized, original)

    def test_build_optimized_cv_text_applies_accepted_edits(self):
        original = (
            "SILVIA MUCCI\n"
            "CHI SONO\n"
            "Sono una studentessa magistrale in Ingegneria Informatica.\n\n"
            "HARD SKILLS\n"
            "Python SQL Java\n\n"
            "ESPERIENZE PROFESSIONALI\n"
            "Supporto all'analisi legislativa tramite LLM."
        )

        accepted = [
            {
                "id": "profile-1",
                "type": "actionableEdit",
                "section": "CHI SONO",
                "original_text": "Sono una studentessa magistrale in Ingegneria Informatica.",
                "proposed_text": "Sono una studentessa magistrale in Ingegneria Informatica orientata all'analisi dei dati e allo sviluppo di soluzioni digitali.",
            },
            {
                "id": "skills-1",
                "type": "actionableEdit",
                "section": "HARD SKILLS",
                "original_text": "Python SQL Java",
                "proposed_text": "Python, SQL, Java, Machine Learning",
            },
        ]

        optimized = build_optimized_cv_text(
            original,
            accepted,
            user_additional_data={},
            role="Data Analyst",
            company="Google",
            use_llm=False,
        )

        self.assertNotEqual(optimized, original)
        self.assertIn("orientata all'analisi dei dati", optimized)
        self.assertIn("Machine Learning", optimized)

    def test_build_optimized_cv_text_preserves_experience_rewrite_when_merge_is_sparse(self):
        original = (
            "Luca Zerella\n"
            "ESPERIENZE PROFESSIONALI\n"
            "Utilizzato in esame di basi di dati per analisi e manipolazione di dati relazionali.\n\n"
            "PROGETTI\n"
            "Progetto Machine Learning\n"
            "Implementazione di pipeline ML per analisi predittiva.\n"
        )

        accepted = [
            {
                "id": "experience-1",
                "type": "actionableEdit",
                "section": "ESPERIENZE PROFESSIONALI",
                "original_text": "Utilizzato in esame di basi di dati per analisi e manipolazione di dati relazionali.",
                "proposed_text": (
                    "Esperienza orientata ad analisi dati, automazione e valutazione dei risultati per il ruolo di Data Analyst:\n"
                    "- Utilizzato in esame di basi di dati per analisi e manipolazione di dati relazionali."
                ),
            }
        ]

        optimized = build_optimized_cv_text(
            original,
            accepted,
            user_additional_data={},
            role="Data Analyst",
            company="",
            use_llm=False,
        )

        self.assertNotIn("Esperienza orientata ad analisi dati", optimized)
        self.assertNotIn("Utilizzato in esame di basi di dati", optimized)
        self.assertIn("PROGETTI", optimized)

    def test_build_resume_rewrite_result_converts_informal_notes_and_deduplicates_skills(self):
        with patch("main.CV_REWRITE_LLM_ENABLED", False):
            result = build_resume_rewrite_result(
                cv_text=(
                    "PROFILO\n"
                    "Studente magistrale in Ingegneria Informatica.\n\n"
                    "HARD SKILLS\n"
                    "Python, SQL, Data visualization\n\n"
                    "PROGETTI\n"
                    "Progetto Machine Learning\n"
                    "Implementazione di pipeline ML per analisi predittiva.\n"
                ),
                company="",
                role="Data Analyst",
                goal="Ottimizzazione CV",
                accepted_suggestions=[],
                user_additional_data={
                    "adaptation_answers": [
                        {
                            "question": "Hai usato database a lezione?",
                            "answer": "Database l'ho visto all'esame di basi dati.",
                            "category": "experience",
                        },
                        {
                            "question": "Hai usato data visualization?",
                            "answer": "Ho applicato la data visualization in progetti di analisi dati, creando grafici e dashboard.",
                            "category": "project",
                        },
                    ],
                    "confirmed_skills": [
                        {
                            "id": "dup-sql",
                            "name": "SQL",
                            "target_section": "HARD SKILLS",
                            "category": "hard_skill",
                        },
                        {
                            "id": "dup-sql-2",
                            "name": "SQL",
                            "target_section": "HARD SKILLS",
                            "category": "hard_skill",
                        },
                        {
                            "id": "database-relazionali",
                            "name": "Database relazionali",
                            "target_section": "HARD SKILLS",
                            "category": "hard_skill",
                        },
                    ],
                },
            )

        text = result["optimized_text"]
        lowered = text.lower()
        self.assertNotIn("l'ho visto", lowered)
        self.assertNotIn("ho applicato", lowered)
        self.assertIn("data visualization", lowered)
        self.assertIn("database relazionali", lowered)
        self.assertEqual(lowered.count("sql"), 1)
        self.assertTrue("competenze tecniche" in lowered or "progetti" in lowered)
        self.assertNotIn("esperienza orientata", lowered)

    def test_resume_rewrite_result_includes_additional_screen_inputs(self):
        with patch("main.CV_REWRITE_LLM_ENABLED", False):
            result = build_resume_rewrite_result(
                cv_text=(
                    "PROFILO\n"
                    "Studente magistrale in Ingegneria Informatica.\n\n"
                    "HARD SKILLS\n"
                    "Python, SQL\n"
                ),
                company="Poste Italiane",
                role="Project Manager",
                goal="Ottimizzazione CV",
                accepted_suggestions=[],
                user_additional_data={
                    "adaptation_answers": [
                        {
                            "question": "Hai gestito attività o progetti?",
                            "answer": "Ho coordinato un progetto universitario con divisione attività e scadenze.",
                            "category": "experience",
                        }
                    ],
                    "confirmed_skills": [
                        {
                            "id": "pm-jira",
                            "name": "Jira",
                            "detail": "Usato per organizzare attività e priorità.",
                            "target_section": "HARD SKILLS",
                            "type": "skillConfirmation",
                        }
                    ],
                },
            )

        text = result["optimized_text"]
        self.assertIn("progetto universitario", text.lower())
        self.assertIn("jira", text.lower())
        self.assertTrue(result["instructions"])


class DocxPreserverLayoutTests(unittest.TestCase):
    def _docx_bytes(self, document: Document) -> bytes:
        output = io.BytesIO()
        document.save(output)
        return output.getvalue()

    def _add_textbox(self, document: Document, lines):
        text_xml = "".join(
            f"<w:p><w:r><w:t>{line}</w:t></w:r></w:p>"
            for line in lines
        )
        shape = parse_xml(
            f"""
            <w:r {nsdecls('w')} xmlns:v="urn:schemas-microsoft-com:vml">
              <w:pict>
                <v:shape style="width:180pt;height:300pt">
                  <v:textbox>
                    <w:txbxContent>{text_xml}</w:txbxContent>
                  </v:textbox>
                </v:shape>
              </w:pict>
            </w:r>
            """
        )
        document.add_paragraph()._p.append(shape)

    def test_profile_rewrite_can_mention_skills_in_a_normal_sentence(self):
        pipeline = ResumeDocxOptimizationPipeline()
        instruction = StructuredRewriteInstruction(
            suggestion_id="profile-data-analyst",
            target_section="profilo",
            action="replace",
            old_text_hint="Studentessa magistrale in Ingegneria Informatica.",
            new_text=(
                "Studentessa magistrale orientata al ruolo di Data Analyst, "
                "con hard skills in Python e SQL e attenzione ai dettagli."
            ),
            items=[],
            reason="Profilo più mirato.",
            confidence=0.8,
        )

        self.assertTrue(pipeline._is_safe_target_text(instruction))

    def test_appends_skill_inside_existing_textbox_section(self):
        document = Document()
        document.add_paragraph("PROFILO")
        document.add_paragraph("Profilo professionale.")
        self._add_textbox(document, [
            "CONTATTI",
            "Email: user@example.com",
            "HARD SKILLS",
            "Python",
            "SOFT SKILLS",
            "Problem solving",
        ])

        result = ResumeDocxOptimizationPipeline().apply_instructions_to_docx(
            self._docx_bytes(document),
            [StructuredRewriteInstruction(
                suggestion_id="textbox-hard-skill",
                target_section="HARD SKILLS",
                action="append",
                old_text_hint="",
                new_text="SQL",
                items=["SQL"],
                reason="Skill confermata.",
                confidence=1.0,
            )],
        )

        final_document = Document(io.BytesIO(result.file_bytes))
        contexts = ResumeDocxOptimizationPipeline()._paragraph_contexts(final_document)
        hard_skill_text = [
            context.paragraph.text
            for context in contexts
            if context.section == "hard_skills"
            and context.paragraph.text.strip()
            and context.paragraph.text.strip() != "HARD SKILLS"
        ]
        self.assertIn("Python", hard_skill_text)
        self.assertIn("SQL", hard_skill_text)
        self.assertEqual(
            sum(1 for context in contexts if context.paragraph.text.strip() == "HARD SKILLS"),
            1,
        )

    def test_inline_table_inherits_the_preceding_section(self):
        document = Document()
        document.add_paragraph("COMPETENZE TECNICHE")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Linguaggi"
        table.cell(0, 1).text = "Python, Java"
        table.cell(1, 0).text = "Framework"
        table.cell(1, 1).text = "TensorFlow"
        document.add_paragraph("LINGUE")
        document.add_paragraph("Italiano, Inglese")

        contexts = ResumeDocxOptimizationPipeline()._paragraph_contexts(document)
        section_by_text = {
            context.paragraph.text.strip(): context.section
            for context in contexts
            if context.paragraph.text.strip()
        }

        self.assertEqual(section_by_text["Python, Java"], "hard_skills")
        self.assertEqual(section_by_text["TensorFlow"], "hard_skills")
        self.assertEqual(section_by_text["Italiano, Inglese"], "lingue")

    def test_contact_cells_do_not_inherit_hard_skills(self):
        document = Document()
        table = document.add_table(rows=4, cols=1)
        table.cell(0, 0).text = "HARD SKILLS"
        table.cell(0, 0).add_paragraph("Python")
        table.cell(1, 0).text = "Via Roma 10, Milano"
        table.cell(2, 0).text = "user@example.com"
        table.cell(3, 0).text = "LinkedIn: linkedin.com/in/example"
        table.cell(3, 0).add_paragraph("Pagina Web personale:")
        table.cell(3, 0).add_paragraph("https://example.com")

        contexts = ResumeDocxOptimizationPipeline()._paragraph_contexts(document)
        section_by_text = {
            context.paragraph.text.strip(): context.section
            for context in contexts
            if context.paragraph.text.strip()
        }

        self.assertEqual(section_by_text["Python"], "hard_skills")
        self.assertEqual(section_by_text["Via Roma 10, Milano"], "contatti")
        self.assertEqual(section_by_text["user@example.com"], "contatti")
        self.assertEqual(section_by_text["LinkedIn: linkedin.com/in/example"], "contatti")
        self.assertEqual(section_by_text["https://example.com"], "contatti")

    def test_profile_rewrite_rejects_an_injected_section_heading(self):
        pipeline = ResumeDocxOptimizationPipeline()
        instruction = StructuredRewriteInstruction(
            suggestion_id="profile-injection",
            target_section="profilo",
            action="replace",
            old_text_hint="Profilo originale.",
            new_text="Profilo aggiornato.\nHARD SKILLS\nPython, SQL",
            items=[],
            reason="Test sicurezza.",
            confidence=0.8,
        )

        self.assertFalse(pipeline._is_safe_target_text(instruction))

    def test_structured_instruction_guard_removes_unsupported_pm_content(self):
        cv_text = (
            "PROFILO\n"
            "Studente magistrale in Ingegneria Informatica.\n"
            "HARD SKILLS\n"
            "Python, Java\n"
            "SOFT SKILLS\n"
            "Problem solving, organizzazione\n"
            "PROGETTI\n"
            "Progetto Software\n"
            "Sviluppo di un'applicazione Java con lavoro in team.\n"
        )
        pipeline = ResumeDocxOptimizationPipeline()
        instructions = pipeline.generate_structured_instructions(
            cv_text=cv_text,
            role="Project Manager",
            company="",
            goal="",
            accepted_suggestions=[
                {
                    "id": "invented-experience",
                    "section": "ESPERIENZE PROFESSIONALI",
                    "original_text": "",
                    "proposed_text": "Esperienza professionale come Project Manager.",
                },
                {
                    "id": "unsupported-skills",
                    "section": "HARD SKILLS",
                    "original_text": "Python, Java",
                    "proposed_text": "Python, Java, Gestione budget, Microsoft Project, Monday.com, Notion",
                },
                {
                    "id": "projects-with-note",
                    "section": "PROGETTI",
                    "original_text": "Progetto Software\nSviluppo di un'applicazione Java con lavoro in team.",
                    "proposed_text": (
                        "Progetto Software\n"
                        "Sviluppo di un'applicazione Java con lavoro in team.\n"
                        "Usata in progetto di sviluppo software per allineare requisiti e aspettative."
                    ),
                },
            ],
            user_additional_data={
                "confirmed_skills": [
                    {
                        "name": "Gestione requisiti",
                        "detail": "Usata in progetto di sviluppo software per allineare requisiti e aspettative.",
                    }
                ]
            },
            use_llm=False,
        )

        joined = "\n".join(item.new_text for item in instructions)
        targets = {canonical_section(item.target_section) for item in instructions}
        self.assertNotIn("esperienze", targets)
        self.assertNotIn("Gestione budget", joined)
        self.assertNotIn("Microsoft Project", joined)
        self.assertNotIn("Monday.com", joined)
        self.assertNotIn("Notion", joined)
        self.assertNotIn("Usata in progetto", joined)
        self.assertIn("Progetto Software", joined)

    def test_project_rewrite_keeps_title_description_pairs(self):
        original = (
            "PROFILO\nStudente magistrale.\n"
            "PROGETTI\n"
            "Progetto Machine Learning\n"
            "Implementazione di pipeline di analisi predittiva con feature engineering.\n"
            "Progetto Big Data\n"
            "Elaborazione distribuita di dataset con Hadoop, Hive e Spark.\n"
            "Progetto Deep Learning\n"
            "Sviluppo e addestramento di reti neurali per classificazione.\n"
        )
        accepted = [{
            "id": "projects-clean",
            "type": "actionableEdit",
            "section": "PROGETTI",
            "original_text": (
                "Progetto Machine Learning\n"
                "Implementazione di pipeline di analisi predittiva con feature engineering.\n"
                "Progetto Big Data\n"
                "Elaborazione distribuita di dataset con Hadoop, Hive e Spark.\n"
                "Progetto Deep Learning\n"
                "Sviluppo e addestramento di reti neurali per classificazione."
            ),
            "proposed_text": (
                "Progetto Machine Learning\n"
                "Implementazione di pipeline di analisi predittiva con feature engineering.\n"
                "Progetto Big Data\n"
                "Elaborazione distribuita di dataset con Hadoop, Hive e Spark.\n"
                "Progetto Deep Learning\n"
                "Sviluppo e addestramento di reti neurali per classificazione."
            ),
        }]

        optimized = build_optimized_cv_text(
            original,
            accepted,
            user_additional_data={},
            role="Project Manager",
            use_llm=False,
        )

        self.assertIn("Progetto Machine Learning\nImplementazione", optimized)
        self.assertIn("Progetto Big Data\nElaborazione", optimized)
        self.assertIn("Progetto Deep Learning\nSviluppo", optimized)

    def test_structured_pipeline_applies_profile_and_skill_updates_together(self):
        document = Document()
        document.add_paragraph("CHI SONO")
        document.add_paragraph(
            "Sono una studentessa magistrale in Ingegneria Informatica."
        )
        document.add_paragraph("HARD SKILLS")
        document.add_paragraph("Python, SQL, Java")
        document.add_paragraph("SOFT SKILLS")
        document.add_paragraph("Creatività, flessibilità, problem solving")

        instructions = [
            StructuredRewriteInstruction(
                suggestion_id="profile-data-analyst",
                target_section="profilo",
                action="replace",
                old_text_hint="Sono una studentessa magistrale in Ingegneria Informatica.",
                new_text=(
                    "Studentessa magistrale orientata al ruolo di Data Analyst, "
                    "con hard skills in Python e SQL e approccio analitico."
                ),
                items=[],
                reason="Profilo più mirato.",
                confidence=0.8,
            ),
            StructuredRewriteInstruction(
                suggestion_id="confirmed-hard-skills",
                target_section="HARD SKILLS",
                action="replace",
                old_text_hint="Python, SQL, Java",
                new_text="Python, SQL, Java, analisi dei dati",
                items=[],
                reason="Competenze confermate.",
                confidence=0.8,
            ),
            StructuredRewriteInstruction(
                suggestion_id="confirmed-soft-skills",
                target_section="SOFT SKILLS",
                action="replace",
                old_text_hint="Creatività, flessibilità, problem solving",
                new_text="Problem solving, attenzione ai dettagli, flessibilità",
                items=[],
                reason="Soft skill confermate.",
                confidence=0.8,
            ),
        ]

        result = ResumeDocxOptimizationPipeline().apply_instructions_to_docx(
            self._docx_bytes(document),
            instructions,
        )

        self.assertEqual(result.validation_report["status"], "applied")
        self.assertEqual(len(result.applied_ids), 3)
        self.assertIn("Studentessa magistrale orientata", result.validation_report["final_text"])
        self.assertIn("CHI SONO", result.validation_report["final_text"])

    def test_structured_pipeline_copies_direct_run_format_to_inserted_lines(self):
        document = Document()
        heading = document.add_paragraph("ESPERIENZE PROFESSIONALI")
        heading.runs[0].font.name = "Montserrat"
        heading.runs[0].font.bold = True
        body = document.add_paragraph("Attività originale.")
        body.runs[0].font.name = "Open Sans"
        body.runs[0].font.size = Pt(10)
        body.paragraph_format.space_after = Pt(12)

        instruction = StructuredRewriteInstruction(
            suggestion_id="formatted-experience",
            target_section="ESPERIENZE PROFESSIONALI",
            action="replace",
            old_text_hint="Attività originale.",
            new_text="Prima attività aggiornata.\nSeconda attività aggiornata.",
            items=[],
            reason="Test formattazione.",
            confidence=0.8,
        )

        result = ResumeDocxOptimizationPipeline().apply_instructions_to_docx(
            self._docx_bytes(document),
            [instruction],
        )
        updated = Document(io.BytesIO(result.file_bytes))
        inserted = next(
            paragraph
            for paragraph in updated.paragraphs
            if paragraph.text == "Seconda attività aggiornata."
        )

        self.assertEqual(inserted.runs[0].font.name, "Open Sans")
        self.assertEqual(inserted.runs[0].font.size.pt, 10)
        self.assertEqual(inserted.paragraph_format.space_after.pt, 12)

    def test_structured_pipeline_preserves_heading_format_when_appending_section(self):
        document = Document()
        heading = document.add_paragraph("FORMAZIONE")
        heading.runs[0].font.name = "Montserrat"
        heading.runs[0].font.size = Pt(16)
        heading.paragraph_format.space_after = Pt(6)
        body = document.add_paragraph("Laurea")
        body.runs[0].font.name = "Open Sans"
        body.runs[0].font.size = Pt(10)

        instruction = StructuredRewriteInstruction(
            suggestion_id="append-projects",
            target_section="PROGETTI",
            action="replace",
            old_text_hint="",
            new_text="Progetto uno\nProgetto due",
            items=[],
            reason="Aggiunta sezione.",
            confidence=0.8,
        )

        result = ResumeDocxOptimizationPipeline().apply_instructions_to_docx(
            self._docx_bytes(document),
            [instruction],
        )
        updated = Document(io.BytesIO(result.file_bytes))
        project_heading = next(paragraph for paragraph in updated.paragraphs if paragraph.text == "PROGETTI")
        project_body = next(paragraph for paragraph in updated.paragraphs if paragraph.text == "Progetto uno")

        self.assertTrue(project_heading.runs[0].bold)

    def test_structured_pipeline_formats_project_titles_and_descriptions(self):
        document = Document()
        document.add_paragraph("PROGETTI")
        document.add_paragraph("Progetto precedente")
        document.add_paragraph("Descrizione precedente.")

        instruction = StructuredRewriteInstruction(
            suggestion_id="projects-readable",
            target_section="PROGETTI",
            action="replace",
            old_text_hint="Progetto precedente\nDescrizione precedente.",
            new_text=(
                "Progetto Machine Learning\n"
                "Implementazione di una pipeline di analisi predittiva.\n"
                "Progetto Big Data\n"
                "Elaborazione distribuita di dataset con Hadoop e Spark."
            ),
            items=[],
            reason="Rende i progetti leggibili.",
            confidence=0.9,
        )

        result = ResumeDocxOptimizationPipeline().apply_instructions_to_docx(
            self._docx_bytes(document),
            [instruction],
        )
        updated = Document(io.BytesIO(result.file_bytes))
        projects = {
            paragraph.text: paragraph
            for paragraph in updated.paragraphs
            if paragraph.text.startswith(("Progetto ", "Implementazione ", "Elaborazione "))
        }

        self.assertTrue(projects["Progetto Machine Learning"].runs[0].bold)
        self.assertFalse(projects["Implementazione di una pipeline di analisi predittiva."].runs[0].bold)
        self.assertTrue(projects["Progetto Big Data"].paragraph_format.keep_with_next)

    def test_structured_pipeline_harmonizes_arial_mt_without_replacing_custom_fonts(self):
        document = Document()
        arial_body = document.add_paragraph("Testo principale")
        arial_body.runs[0].font.name = "Arial MT"
        custom_heading = document.add_paragraph("PROGETTI")
        custom_heading.runs[0].font.name = "Montserrat"
        document.add_paragraph("Progetto Software")
        document.add_paragraph("Sviluppo di una applicazione.")

        result = ResumeDocxOptimizationPipeline().apply_instructions_to_docx(
            self._docx_bytes(document),
            [],
        )
        updated = Document(io.BytesIO(result.file_bytes))

        self.assertEqual(updated.paragraphs[0].runs[0].font.name, "Arial")
        self.assertEqual(updated.paragraphs[1].runs[0].font.name, "Montserrat")

    def test_structured_pipeline_preserves_empty_input_section_heading(self):
        document = Document()
        document.add_paragraph("COMPETENZE TECNICHE")
        document.add_paragraph("")
        document.add_paragraph("PROGETTI")
        document.add_paragraph("Progetto Software")
        document.add_paragraph("Sviluppo di una applicazione Java.")

        result = ResumeDocxOptimizationPipeline().apply_instructions_to_docx(
            self._docx_bytes(document),
            [],
        )
        updated = Document(io.BytesIO(result.file_bytes))
        texts = [paragraph.text for paragraph in updated.paragraphs if paragraph.text.strip()]

        self.assertIn("COMPETENZE TECNICHE", texts)
        self.assertIn("PROGETTI", texts)

    def test_structured_pipeline_minimizes_trailing_paragraph_after_table(self):
        document = Document()
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "HARD SKILLS"
        table.rows[0].height = Pt(700)
        document.add_paragraph("")

        result = ResumeDocxOptimizationPipeline().apply_instructions_to_docx(
            self._docx_bytes(document),
            [],
        )
        updated = Document(io.BytesIO(result.file_bytes))
        trailing = updated.paragraphs[-1]

        self.assertEqual(trailing.paragraph_format.space_after.pt, 0)
        self.assertEqual(trailing.runs[0].font.size.pt, 1)
        self.assertTrue(trailing.runs[0].font.hidden)
        row_xml = updated.tables[0].rows[0]._tr.xml
        self.assertNotIn("trHeight", row_xml)

    def test_structured_pipeline_collapses_multiple_trailing_blank_paragraphs(self):
        document = Document()
        document.add_paragraph("ESPERIENZE PROFESSIONALI")
        document.add_paragraph("Attività")
        document.add_paragraph("")
        document.add_paragraph("")

        result = ResumeDocxOptimizationPipeline().apply_instructions_to_docx(
            self._docx_bytes(document),
            [],
        )
        updated = Document(io.BytesIO(result.file_bytes))
        blank_trailing = [paragraph for paragraph in updated.paragraphs if not paragraph.text.strip()]

        self.assertLessEqual(len(blank_trailing), 1)

    def test_structured_pipeline_prefers_body_anchor_over_sidebar_table(self):
        document = Document()
        sidebar = document.add_table(rows=1, cols=1)
        sidebar.cell(0, 0).text = "CONTATTI"
        sidebar.cell(0, 0).add_paragraph("Telefono")
        document.add_paragraph("CHI SONO")
        document.add_paragraph("Profilo esistente")
        document.add_paragraph("FORMAZIONE")
        document.add_paragraph("Laurea")

        instruction = StructuredRewriteInstruction(
            suggestion_id="append-experience",
            target_section="ESPERIENZE PROFESSIONALI",
            action="replace",
            old_text_hint="",
            new_text="Esperienza uno\nEsperienza due",
            items=[],
            reason="Aggiunta esperienza.",
            confidence=0.8,
        )

        result = ResumeDocxOptimizationPipeline().apply_instructions_to_docx(
            self._docx_bytes(document),
            [instruction],
        )
        updated = Document(io.BytesIO(result.file_bytes))
        body_texts = [paragraph.text for paragraph in updated.paragraphs if paragraph.text.strip()]
        sidebar_texts = [paragraph.text for paragraph in updated.tables[0].cell(0, 0).paragraphs if paragraph.text.strip()]

        self.assertIn("Esperienza uno", body_texts)
        self.assertNotIn("Esperienza uno", sidebar_texts)

    def test_replaces_text_inside_table_section(self):
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "CHI SONO"
        table.cell(0, 1).text = "Profilo originale"
        table.cell(1, 0).text = "HARD SKILLS"
        table.cell(1, 1).text = "Python, SQL"

        updated_bytes, applied = DocxPreserver().apply(
            self._docx_bytes(document),
            [
                RewriteInstruction(
                    section="HARD SKILLS",
                    original="Python, SQL",
                    replacement="Python · SQL · Power BI",
                    category="skills",
                    source_id="test",
                )
            ],
        )

        updated = Document(io.BytesIO(updated_bytes))
        self.assertEqual(applied, 1)
        self.assertEqual(updated.tables[0].cell(1, 1).text, "Python · SQL · Power BI")
        self.assertEqual(updated.tables[0].cell(0, 1).text, "Profilo originale")

    def test_does_not_replace_same_text_from_wrong_table_section(self):
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "CHI SONO"
        table.cell(0, 1).text = "Python, SQL"
        table.cell(1, 0).text = "HARD SKILLS"
        table.cell(1, 1).text = "Python, SQL"

        updated_bytes, applied = DocxPreserver().apply(
            self._docx_bytes(document),
            [
                RewriteInstruction(
                    section="HARD SKILLS",
                    original="Python, SQL",
                    replacement="Python · SQL · Power BI",
                    category="skills",
                    source_id="test",
                )
            ],
        )

        updated = Document(io.BytesIO(updated_bytes))
        self.assertEqual(applied, 1)
        self.assertEqual(updated.tables[0].cell(0, 1).text, "Python, SQL")
        self.assertEqual(updated.tables[0].cell(1, 1).text, "Python · SQL · Power BI")

    def test_table_column_with_own_headings_does_not_inherit_soft_skills(self):
        document = Document()
        table = document.add_table(rows=1, cols=2)
        left = table.cell(0, 0)
        left.text = "SOFT SKILLS"
        left.add_paragraph("Problem solving")
        right = table.cell(0, 1)
        right.text = "SILVIA MUCCI\nIngegnere Informatico"
        right.add_paragraph("CHI SONO")
        right.add_paragraph("Profilo professionale.")

        contexts = ResumeDocxOptimizationPipeline()._paragraph_contexts(document)
        name_context = next(
            context
            for context in contexts
            if "SILVIA MUCCI" in (context.paragraph.text or "")
        )

        self.assertEqual(name_context.section, "intestazione")

    def test_validation_rejects_education_content_inside_hard_skills(self):
        document = Document()
        document.add_paragraph("HARD SKILLS")
        document.add_paragraph("Python · SQL")
        document.add_paragraph("Università degli Studi | 2019-2024 Laurea")

        pipeline = ResumeDocxOptimizationPipeline()
        warnings = pipeline._skill_section_contamination_warnings(
            self._docx_bytes(document)
        )

        self.assertTrue(warnings)

    def test_validation_rejects_narrative_fragments_and_names_inside_skills(self):
        document = Document()
        document.add_paragraph("HARD SKILLS")
        document.add_paragraph("Python · SQL")
        document.add_paragraph("e approccio analitico alla risoluzione di problemi complessi.")
        document.add_paragraph("SOFT SKILLS")
        document.add_paragraph("Problem solving · SILVIA MUCCI Ingegnere Informatico")

        warnings = ResumeDocxOptimizationPipeline()._skill_section_contamination_warnings(
            self._docx_bytes(document)
        )

        self.assertEqual(len(warnings), 2)

    def test_appends_to_existing_section_without_page_break(self):
        document = Document()
        document.add_paragraph("PROGETTI")
        document.add_paragraph("Progetto esistente")

        updated_bytes, applied = DocxPreserver().apply(
            self._docx_bytes(document),
            [
                RewriteInstruction(
                    section="PROGETTI",
                    original="",
                    replacement="Nuovo progetto confermato",
                    category="project",
                    source_id="test-project",
                )
            ],
        )

        updated = Document(io.BytesIO(updated_bytes))
        texts = [paragraph.text for paragraph in updated.paragraphs]
        self.assertEqual(applied, 1)
        self.assertEqual(texts, ["PROGETTI", "Progetto esistente", "Nuovo progetto confermato"])

    def test_structured_pipeline_appends_without_deleting_existing_section(self):
        document = Document()
        document.add_paragraph("ESPERIENZE PROFESSIONALI")
        document.add_paragraph("Esperienza originale presso Azienda Uno.")
        document.add_paragraph("Attivita originale da conservare.")
        document.add_paragraph("FORMAZIONE")
        document.add_paragraph("Laurea originale.")

        pipeline = ResumeDocxOptimizationPipeline()
        instructions = pipeline.generate_structured_instructions(
            cv_text=(
                "ESPERIENZE PROFESSIONALI\n"
                "Esperienza originale presso Azienda Uno.\n"
                "Attivita originale da conservare.\n"
                "FORMAZIONE\nLaurea originale."
            ),
            role="Data Engineer",
            company="",
            goal="",
            accepted_suggestions=[{
                "suggestion_id": "user-box-experience",
                "target_section": "ESPERIENZE PROFESSIONALI",
                "old_text_hint": "",
                "new_text": "Data Engineer presso Poste Italiane, 2020-2024.",
            }],
            user_additional_data={
                "experiences": "Ho lavorato dal 2020 al 2024 presso Poste Italiane come Data Engineer."
            },
            use_llm=False,
        )

        self.assertEqual(instructions[0].action, "append")
        result = pipeline.apply_instructions_to_docx(
            self._docx_bytes(document),
            instructions,
        )
        updated = Document(io.BytesIO(result.file_bytes))
        texts = [paragraph.text for paragraph in updated.paragraphs if paragraph.text]

        self.assertEqual(result.validation_report["status"], "applied")
        self.assertEqual(
            texts,
            [
                "ESPERIENZE PROFESSIONALI",
                "Esperienza originale presso Azienda Uno.",
                "Attivita originale da conservare.",
                "Data Engineer presso Poste Italiane, 2020-2024.",
                "FORMAZIONE",
                "Laurea originale.",
            ],
        )

    def test_appends_inside_existing_table_section(self):
        document = Document()
        table = document.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "PROGETTI"
        cell.add_paragraph("Progetto esistente")

        updated_bytes, applied = DocxPreserver().apply(
            self._docx_bytes(document),
            [
                RewriteInstruction(
                    section="PROGETTI",
                    original="",
                    replacement="Nuovo progetto confermato",
                    category="project",
                    source_id="test-table-project",
                )
            ],
        )

        updated = Document(io.BytesIO(updated_bytes))
        texts = [paragraph.text for paragraph in updated.tables[0].cell(0, 0).paragraphs]
        self.assertEqual(applied, 1)
        self.assertEqual(texts, ["PROGETTI", "Progetto esistente", "Nuovo progetto confermato"])

    def test_removes_redundant_profile_paragraph_after_rewrite(self):
        document = Document()
        document.add_paragraph("CHI SONO")
        document.add_paragraph("Profilo originale.")
        document.add_paragraph("Motivata a crescere professionalmente.")
        document.add_paragraph("FORMAZIONE")
        document.add_paragraph("Laurea")

        updated_bytes, applied = DocxPreserver().apply(
            self._docx_bytes(document),
            [
                RewriteInstruction(
                    section="CHI SONO",
                    original="Profilo originale.",
                    replacement="Profilo aggiornato. Motivata a crescere professionalmente.",
                    category="profile",
                    source_id="test-profile",
                )
            ],
        )

        updated = Document(io.BytesIO(updated_bytes))
        texts = [paragraph.text for paragraph in updated.paragraphs]
        self.assertEqual(applied, 1)
        self.assertEqual(texts[1], "Profilo aggiornato. Motivata a crescere professionalmente.")
        self.assertEqual(texts[2], "")
        self.assertEqual(texts[3], "FORMAZIONE")

    def test_replaces_original_text_split_across_multiple_paragraphs(self):
        document = Document()
        document.add_paragraph("ESPERIENZE PROFESSIONALI")
        document.add_paragraph("Prima parte dell'esperienza con attività principali.")
        document.add_paragraph("Seconda parte con strumenti e risultati.")
        document.add_paragraph("FORMAZIONE")
        document.add_paragraph("Corso di studio")

        replacement = "Esperienza riscritta in modo completo e professionale."
        updated_bytes, applied = DocxPreserver().apply(
            self._docx_bytes(document),
            [
                RewriteInstruction(
                    section="ESPERIENZE PROFESSIONALI",
                    original=(
                        "Prima parte dell'esperienza con attività principali. "
                        "Seconda parte con strumenti e risultati."
                    ),
                    replacement=replacement,
                    category="experience",
                    source_id="split-experience",
                )
            ],
        )

        updated = Document(io.BytesIO(updated_bytes))
        texts = [paragraph.text for paragraph in updated.paragraphs]
        self.assertEqual(applied, 1)
        self.assertEqual(texts[1], replacement)
        self.assertEqual(texts[2], "")
        self.assertEqual(texts[3], "FORMAZIONE")

    def test_distributes_multiline_section_rewrite_across_existing_paragraphs(self):
        document = Document()
        document.add_paragraph("ESPERIENZE PROFESSIONALI")
        first = document.add_paragraph("Attività iniziale.")
        first.runs[0].font.name = "Open Sans"
        first.paragraph_format.space_after = Pt(9)
        second = document.add_paragraph("Risultato iniziale.")
        second.runs[0].font.name = "Open Sans"
        document.add_paragraph("FORMAZIONE")

        replacement = "Attività aggiornata.\nRisultato aggiornato.\nStrumento confermato."
        updated_bytes, applied = DocxPreserver().apply(
            self._docx_bytes(document),
            [
                RewriteInstruction(
                    section="ESPERIENZE PROFESSIONALI",
                    original="Attività iniziale. Risultato iniziale.",
                    replacement=replacement,
                    category="experience",
                    source_id="multiline-experience",
                )
            ],
        )

        updated = Document(io.BytesIO(updated_bytes))
        texts = [paragraph.text for paragraph in updated.paragraphs]
        self.assertEqual(applied, 1)
        self.assertEqual(
            texts,
            [
                "ESPERIENZE PROFESSIONALI",
                "Attività aggiornata.",
                "Risultato aggiornato.",
                "Strumento confermato.",
                "FORMAZIONE",
            ],
        )
        self.assertEqual(updated.paragraphs[3].runs[0].font.name, "Open Sans")

    def test_appends_accepted_change_when_target_section_has_no_editable_body(self):
        document = Document()
        document.add_paragraph("PROGETTI")

        updated_bytes, applied = DocxPreserver().apply(
            self._docx_bytes(document),
            [
                RewriteInstruction(
                    section="PROGETTI",
                    original="Testo non più individuabile",
                    replacement="Progetto confermato dall'utente.",
                    category="project",
                    source_id="missing-original",
                )
            ],
        )

        updated = Document(io.BytesIO(updated_bytes))
        texts = [paragraph.text for paragraph in updated.paragraphs if paragraph.text]
        self.assertEqual(applied, 1)
        self.assertEqual(texts, ["PROGETTI", "Progetto confermato dall'utente."])

    def test_appends_new_skill_section_without_new_page(self):
        document = Document()
        document.add_paragraph("FORMAZIONE")
        document.add_paragraph("Laurea in Economia")

        updated_bytes, applied = DocxPreserver().apply(
            self._docx_bytes(document),
            [
                RewriteInstruction(
                    section="COMPETENZE",
                    original="",
                    replacement="Python",
                    category="skills",
                    source_id="skill-end",
                )
            ],
        )

        updated = Document(io.BytesIO(updated_bytes))
        texts = [paragraph.text for paragraph in updated.paragraphs if paragraph.text]
        self.assertEqual(applied, 1)
        self.assertEqual(texts[-2:], ["COMPETENZE", "Python"])

    def test_inserts_confirmed_section_when_document_has_no_known_anchor(self):
        document = Document()
        document.add_paragraph("Curriculum vitae")
        document.add_paragraph("Contenuto con struttura personalizzata")
        pipeline = ResumeDocxOptimizationPipeline()
        instruction = StructuredRewriteInstruction(
            suggestion_id="confirmed-language",
            target_section="LINGUE",
            action="append",
            old_text_hint="",
            new_text="Inglese B2",
            items=["Inglese B2"],
        )

        result = pipeline.apply_instructions_to_docx(
            self._docx_bytes(document),
            [instruction],
        )

        updated = Document(io.BytesIO(result.file_bytes))
        texts = [paragraph.text for paragraph in updated.paragraphs if paragraph.text]
        self.assertEqual(result.failed_ids, [])
        self.assertIn("confirmed-language", result.applied_ids)
        self.assertEqual(texts[-2:], ["LINGUE", "Inglese B2"])

    def test_replacing_one_experience_does_not_delete_other_existing_entries(self):
        document = Document()
        document.add_paragraph("ESPERIENZE PROFESSIONALI")
        document.add_paragraph("Esperienza A da aggiornare")
        document.add_paragraph("Esperienza B da conservare")
        pipeline = ResumeDocxOptimizationPipeline()
        instruction = StructuredRewriteInstruction(
            suggestion_id="rewrite-one-experience",
            target_section="ESPERIENZE PROFESSIONALI",
            action="replace",
            old_text_hint="Esperienza A da aggiornare",
            new_text="Esperienza A aggiornata",
            items=[],
        )

        result = pipeline.apply_instructions_to_docx(
            self._docx_bytes(document),
            [instruction],
        )

        updated = Document(io.BytesIO(result.file_bytes))
        texts = [paragraph.text for paragraph in updated.paragraphs if paragraph.text]
        self.assertIn("Esperienza A aggiornata", texts)
        self.assertIn("Esperienza B da conservare", texts)

    def test_skips_role_like_confirmed_skill(self):
        instructions = build_confirmed_skill_rewrite_instructions(
            "PROFILO\nCandidato interessato a Project Manager (stage)",
            {"confirmed_skills": [{"name": "Project Manager (stage)", "category": "keyword"}]},
            "Project Manager (stage)"
        )
        self.assertEqual(instructions, [])

    def test_reuses_new_section_for_multiple_additions(self):
        document = Document()
        document.add_paragraph("FORMAZIONE")
        document.add_paragraph("Laurea")

        updated_bytes, applied = DocxPreserver().apply(
            self._docx_bytes(document),
            [
                RewriteInstruction(
                    section="PROGETTI",
                    original="",
                    replacement="Primo progetto",
                    category="project",
                    source_id="first-project",
                ),
                RewriteInstruction(
                    section="PROGETTI",
                    original="",
                    replacement="Secondo progetto",
                    category="project",
                    source_id="second-project",
                ),
            ],
        )

        updated = Document(io.BytesIO(updated_bytes))
        texts = [paragraph.text for paragraph in updated.paragraphs if paragraph.text]
        self.assertEqual(applied, 2)
        self.assertEqual(texts.count("PROGETTI"), 1)
        self.assertEqual(texts[-2:], ["Primo progetto", "Secondo progetto"])

    def test_new_page_section_copies_direct_font_formatting(self):
        document = Document()
        heading = document.add_paragraph("FORMAZIONE")
        heading.runs[0].font.name = "Montserrat"
        heading.runs[0].font.bold = True
        body = document.add_paragraph("Laurea")
        body.runs[0].font.name = "Open Sans"

        updated_bytes, applied = DocxPreserver().apply(
            self._docx_bytes(document),
            [
                RewriteInstruction(
                    section="PROGETTI",
                    original="",
                    replacement="Progetto confermato",
                    category="project",
                    source_id="formatted-project",
                )
            ],
        )

        updated = Document(io.BytesIO(updated_bytes))
        project_heading = next(paragraph for paragraph in updated.paragraphs if paragraph.text == "PROGETTI")
        project_body = next(paragraph for paragraph in updated.paragraphs if paragraph.text == "Progetto confermato")
        self.assertEqual(applied, 1)
        self.assertEqual(project_heading.runs[0].font.name, "Montserrat")
        self.assertTrue(project_heading.runs[0].font.bold)
        self.assertEqual(project_body.runs[0].font.name, "Open Sans")


if __name__ == "__main__":
    unittest.main()
