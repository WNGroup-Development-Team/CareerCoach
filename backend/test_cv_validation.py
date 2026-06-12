import io
import unittest
import zipfile
from unittest.mock import patch

import requests
from docx import Document

import main


def create_docx_with_image(text_lines):
    png_bytes = (
        b"\x89PNG\r\n\x1a\n"
        b"\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
        b"\x00\x00\x00\nIDATx\x9cc\xf8\x0f\x00\x01\x01\x01\x00\x18\xdd\x8d\xb6\x00\x00\x00\x00IEND\xaeB`\x82"
    )

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            """<?xml version='1.0' encoding='UTF-8'?>
<Types xmlns='http://schemas.openxmlformats.org/package/2006/content-types'>
<Default Extension='rels' ContentType='application/vnd.openxmlformats-package.relationships+xml'/>
<Default Extension='xml' ContentType='application/xml'/>
<Override PartName='/word/document.xml' ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'/>
<Override PartName='/word/media/image1.png' ContentType='image/png'/>
</Types>
""",
        )
        archive.writestr(
            "_rels/.rels",
            """<?xml version='1.0' encoding='UTF-8'?>
<Relationships xmlns='http://schemas.openxmlformats.org/package/2006/relationships'>
<Relationship Id='rId1' Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument' Target='word/document.xml'/>
</Relationships>
""",
        )
        archive.writestr(
            "word/_rels/document.xml.rels",
            """<?xml version='1.0' encoding='UTF-8'?>
<Relationships xmlns='http://schemas.openxmlformats.org/package/2006/relationships'>
<Relationship Id='rIdImage1' Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/image' Target='media/image1.png'/>
</Relationships>
""",
        )
        archive.writestr(
            "word/document.xml",
            """<?xml version='1.0' encoding='UTF-8'?>
<w:document xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main' xmlns:v='urn:schemas-microsoft-com:vml'>
<w:body>
{paragraphs}
</w:body>
</w:document>
""".replace(
                "{paragraphs}",
                "\n".join(f"<w:p><w:r><w:t>{line}</w:t></w:r></w:p>" for line in text_lines),
            ),
        )
        archive.writestr("word/media/image1.png", png_bytes)

    output.seek(0)
    return output.getvalue()


def create_docx_with_table(name, email):
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Nome"
    table.cell(0, 1).text = name
    table.cell(1, 0).text = "Email"
    table.cell(1, 1).text = email
    doc.add_paragraph("Esperienze professionali")
    doc.add_paragraph("Formazione")
    doc.add_paragraph("Competenze")
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def create_docx_with_header(name, email):
    doc = Document()
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.add_paragraph()
    header_para.add_run(name)
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph("Esperienze professionali")
    doc.add_paragraph("Formazione")
    doc.add_paragraph("Competenze")
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def create_docx_with_textbox(name, email):
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            """<?xml version='1.0' encoding='UTF-8'?>
<Types xmlns='http://schemas.openxmlformats.org/package/2006/content-types'>
<Default Extension='rels' ContentType='application/vnd.openxmlformats-package.relationships+xml'/>
<Default Extension='xml' ContentType='application/xml'/>
<Override PartName='/word/document.xml' ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'/>
</Types>
""",
        )
        archive.writestr(
            "_rels/.rels",
            """<?xml version='1.0' encoding='UTF-8'?>
<Relationships xmlns='http://schemas.openxmlformats.org/package/2006/relationships'>
<Relationship Id='rId1' Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument' Target='word/document.xml'/>
</Relationships>
""",
        )
        archive.writestr(
            "word/document.xml",
            f"""<?xml version='1.0' encoding='UTF-8'?>
<w:document xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main' xmlns:wps='http://schemas.openxmlformats.org/drawingml/2006/wordprocessingShape' xmlns:v='urn:schemas-microsoft-com:vml'>
<w:body>
<w:p><w:r><w:t>Email: {email}</w:t></w:r></w:p>
<w:p><w:r><w:pict><v:shape style='width:100pt;height:30pt'><v:textbox><w:txbxContent><w:p><w:r><w:t>{name}</w:t></w:r></w:p></w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>
<w:p><w:r><w:t>Esperienze professionali</w:t></w:r></w:p>
<w:p><w:r><w:t>Formazione</w:t></w:r></w:p>
<w:p><w:r><w:t>Competenze</w:t></w:r></w:p>
</w:body>
</w:document>
""",
        )
    output.seek(0)
    return output.getvalue()


class CvValidationTests(unittest.TestCase):
    def test_valid_docx_cv_with_images_is_accepted_when_visual_service_unavailable(self):
        docx_bytes = create_docx_with_image(
            [
                "Luca Rossi",
                "Email: luca.rossi@example.com",
                "Telefono: +39 333 1234567",
                "Esperienze professionali",
                "Formazione",
                "Competenze",
            ]
        )

        with patch("main.validate_cv_images", side_effect=requests.exceptions.ConnectionError(
            "HTTPConnectionPool(host='127.0.0.1', port=11434): Max retries exceeded with url: /api/chat"
        )):
            result = main.validate_cv_content("curriculum.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        self.assertTrue(result["is_cv"])
        self.assertIn("controllo automatico delle immagini non e disponibile", result["reason"].lower())
        self.assertEqual(result["visual_validation"]["status"], "analysis_failed")
        self.assertFalse(result["visual_validation"]["blocked"])

    def test_valid_cv_accepted_when_visual_service_down_for_txt(self):
        text_bytes = b"Mario Bianchi\nEmail: mario.bianchi@example.com\nTelefono: +39 345 9876543\nEsperienze professionali\nFormazione\nCompetenze\nLingue\n"

        with patch("main.validate_cv_images", side_effect=requests.exceptions.ConnectionError(
            "HTTPConnectionPool(host='127.0.0.1', port=11434): Max retries exceeded with url: /api/chat"
        )):
            result = main.validate_cv_content("cv.txt", text_bytes, "text/plain")

        self.assertTrue(result["is_cv"])
        self.assertIn("controllo automatico delle immagini non e disponibile", result["reason"].lower())
        self.assertEqual(result["visual_validation"]["status"], "analysis_failed")
        self.assertFalse(result["visual_validation"]["blocked"])

    def test_file_without_text_and_structure_is_rejected(self):
        result = main.validate_cv_content("not_a_cv.txt", b"just a note without curriculum structure", "text/plain")

        self.assertFalse(result["is_cv"])
        self.assertIn("non contiene abbastanza elementi tipici", result["reason"].lower())

    def test_profile_name_and_cv_name_different_format_are_compatible(self):
        result = main.check_cv_identity(
            "ROSSI LUCA\nEmail: luca.rossi@example.com\n",
            "Luca",
            "Rossi",
        )

        self.assertTrue(result["matches_user"])
        self.assertIn("rossi luca", result["detected_name"].lower())

    def test_images_unverifiable_produce_warning_not_error(self):
        text_bytes = b"Anna Verdi\nEmail: anna.verdi@example.com\nTelefono: +39 333 4445556\nFormazione\nEsperienze professionali\nCompetenze\nLingue\n"

        with patch("main.validate_cv_images", side_effect=requests.exceptions.ConnectionError(
            "HTTPConnectionPool(host='127.0.0.1', port=11434): Max retries exceeded with url: /api/chat"
        )):
            result = main.validate_cv_content("cv.docx", create_docx_with_image(["Anna Verdi", "Email: anna.verdi@example.com", "Esperienze professionali"]), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        self.assertTrue(result["is_cv"])
        self.assertIn("controllo automatico delle immagini non e disponibile", result["reason"].lower())
        self.assertFalse(result["visual_validation"]["blocked"])

    def test_valid_cv_without_images_is_accepted(self):
        result = main.validate_cv_content(
            "cv.txt",
            b"Giorgia Neri\nEmail: giorgia.neri@example.com\nTelefono: +39 345 1112223\nEsperienze professionali\nFormazione\nCompetenze\nLingue\n",
            "text/plain",
        )

        self.assertTrue(result["is_cv"])
        self.assertEqual(result["visual_validation"]["status"], "no_images")
        self.assertFalse(result["visual_validation"]["blocked"])

    def test_valid_cv_with_normal_profile_photo_is_accepted(self):
        docx_bytes = create_docx_with_image(
            [
                "Francesca Bianchi",
                "Email: francesca.bianchi@example.com",
                "Telefono: +39 333 5556667",
                "Esperienze professionali",
                "Formazione",
                "Competenze",
            ]
        )

        with patch("main.validate_cv_images", return_value={
            "status": "completed",
            "image_count": 1,
            "analyzed_count": 1,
            "blocked": False,
            "blocked_categories": [],
        }):
            result = main.validate_cv_content("cv.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        self.assertTrue(result["is_cv"])
        self.assertFalse(result["visual_validation"]["blocked"])

    def test_cv_with_name_mismatch_is_rejected(self):
        result = main.check_cv_identity(
            "Mario Rossi\nEmail: mario.rossi@example.com\n",
            "Luca",
            "Verdi",
        )

        self.assertFalse(result["matches_user"])
        self.assertIn("nome presente nel cv non corrisponde", result["message"].lower())

    def test_cv_text_with_disallowed_content_is_rejected(self):
        result = main.validate_cv_content(
            "cv.txt",
            b"Laura Gialli\nEmail: laura.gialli@example.com\nCompetenze\nProgetto su droga e violenza\n",
            "text/plain",
        )

        self.assertFalse(result["is_cv"])
        self.assertIn("non idonei", result["reason"].lower())

    def test_cv_with_image_blocked_is_rejected(self):
        docx_bytes = create_docx_with_image(
            [
                "Elena Neri",
                "Email: elena.neri@example.com",
                "Telefono: +39 333 7778889",
                "Esperienze professionali",
                "Formazione",
                "Competenze",
            ]
        )

        with patch("main.validate_cv_images", return_value={
            "status": "completed",
            "image_count": 1,
            "analyzed_count": 1,
            "blocked": True,
            "blocked_categories": ["nudità"],
        }):
            result = main.validate_cv_content("cv.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        self.assertFalse(result["is_cv"])
        self.assertIn("contiene immagini o contenuti non idonei", result["reason"].lower())
        self.assertTrue(result["visual_validation"]["blocked"])

    def test_docx_cv_with_name_in_table_is_accepted(self):
        docx_bytes = create_docx_with_table("Luca Rossi", "luca.rossi@example.com")
        text, method = main.extract_text_from_file_bytes(docx_bytes, "cv.docx")
        self.assertIn("luca rossi", text.lower())
        self.assertIn("luca.rossi@example.com", text.lower())
        self.assertIn("nome", text.lower())
        self.assertIn("email", text.lower())
        self.assertIn("esperienze professionali", text.lower())
        self.assertIn(method, {"docx", "docx_xml"})

        with patch("main.validate_cv_images", return_value={
            "status": "no_images",
            "image_count": 0,
            "analyzed_count": 0,
            "blocked": False,
            "blocked_categories": [],
        }):
            result = main.validate_cv_content("cv.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        self.assertTrue(result["is_cv"])
        self.assertFalse(result["visual_validation"]["blocked"])

    def test_docx_cv_with_name_in_header_is_accepted(self):
        docx_bytes = create_docx_with_header("Marta Bianchi", "marta.bianchi@example.com")
        text, method = main.extract_text_from_file_bytes(docx_bytes, "cv.docx")
        self.assertIn("marta bianchi", text.lower())
        self.assertIn("email: marta.bianchi@example.com", text.lower())
        self.assertIn(method, {"docx", "docx_xml"})

        with patch("main.validate_cv_images", return_value={
            "status": "no_images",
            "image_count": 0,
            "analyzed_count": 0,
            "blocked": False,
            "blocked_categories": [],
        }):
            result = main.validate_cv_content("cv.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        self.assertTrue(result["is_cv"])
        self.assertFalse(result["visual_validation"]["blocked"])

    def test_docx_cv_with_name_in_textbox_is_accepted(self):
        docx_bytes = create_docx_with_textbox("Marco Neri", "marco.neri@example.com")
        text, method = main.extract_text_from_file_bytes(docx_bytes, "cv.docx")
        self.assertIn("marco neri", text.lower())
        self.assertIn("email: marco.neri@example.com", text.lower())
        self.assertIn(method, {"docx", "docx_xml"})

        with patch("main.validate_cv_images", return_value={
            "status": "no_images",
            "image_count": 0,
            "analyzed_count": 0,
            "blocked": False,
            "blocked_categories": [],
        }):
            result = main.validate_cv_content("cv.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        self.assertTrue(result["is_cv"])
        self.assertFalse(result["visual_validation"]["blocked"])

    def test_quality_review_treats_semantically_present_replacement_as_kept(self):
        final_text = "Profilo aggiornato con orientamento analitico.\nHARD SKILLS\nPython, SQL, data analysis\n"
        instructions = [
            main.RewriteInstruction(
                section="HARD SKILLS",
                original="Python, SQL",
                replacement="Python, SQL, data analysis",
                category="skills",
                source_id="skills-update",
            )
        ]

        review = main.review_generated_cv_quality_locally(
            final_text=final_text,
            accepted_instructions=instructions,
        )

        self.assertTrue(review["ready_to_send"])
        self.assertEqual(review["issues"], [])


if __name__ == "__main__":
    unittest.main()
